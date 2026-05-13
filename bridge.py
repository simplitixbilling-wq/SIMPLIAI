"""Pywebview JS API bridge — exposes backend to the web frontend."""

import gc
import base64
import csv
import concurrent.futures
import hashlib
import hmac
import io
import json
import os
import re
import subprocess
import sys
import tempfile
import threading
import textwrap
import time
import uuid
import difflib
from datetime import datetime, timezone
from pathlib import Path
from urllib.parse import parse_qs, unquote, urlparse
import pandas as pd

import psutil

from database import ChatDatabase, migrate_json_to_sqlite
from plugin_manager import PluginManager
from mcp_manager import MCPManager
from rag_manager import RAGManager
from utils import app_data_path
from memory_optimizer import MemoryOptimizer
from metrics import get_metrics

try:
    from llama_cpp import Llama
except ImportError:
    Llama = None

import requests  # For Ollama health check


# ── Ollama HTTP backend (drop-in replacement for Llama) ─────────
class OllamaModel:
    """Wraps Ollama's REST API to match the llama_cpp.Llama interface
    so the rest of Bridge works unchanged."""

    def __init__(self, model_name: str, base_url: str = "http://127.0.0.1:11434"):
        self.model_name = model_name
        self.base_url = base_url.rstrip("/")
        self._n_ctx = 65536  # Default context window (increased from 8192 to support larger models)
        # Warm the model so first inference isn't slow
        try:
            import urllib.request, json as _json
            req = urllib.request.Request(
                f"{self.base_url}/api/generate",
                data=_json.dumps({"model": self.model_name, "prompt": "hi",
                                  "options": {"num_predict": 1}}).encode(),
                headers={"Content-Type": "application/json"},
                method="POST",
            )
            urllib.request.urlopen(req, timeout=120).read()
        except Exception:
            pass  # Model will load on first real call

    # --- Compatibility helpers ---
    def n_ctx(self):
        return self._n_ctx

    def _post(self, endpoint: str, payload: dict, stream: bool = False):
        import urllib.request, json as _json
        data = _json.dumps(payload).encode()
        req = urllib.request.Request(
            f"{self.base_url}{endpoint}",
            data=data,
            headers={"Content-Type": "application/json"},
            method="POST",
        )
        resp = urllib.request.urlopen(req, timeout=600)
        if stream:
            return resp  # caller reads line-by-line
        return _json.loads(resp.read().decode())

    # --- create_completion (non-streaming) ---
    def create_completion(self, prompt, **kwargs):
        payload = {
            "model": self.model_name,
            "prompt": prompt,
            "stream": False,
            "options": {
                "num_predict": kwargs.get("max_tokens", 512),
                "temperature": kwargs.get("temperature", 0.25),
                "top_p": kwargs.get("top_p", 0.8),
                "top_k": kwargs.get("top_k", 50),
                "repeat_penalty": kwargs.get("repeat_penalty", 1.1),
            },
        }
        stop = kwargs.get("stop")
        if stop:
            payload["options"]["stop"] = stop
        resp = self._post("/api/generate", payload)
        return {
            "choices": [{"text": resp.get("response", ""), "finish_reason": "stop"}],
            "usage": {
                "prompt_tokens": resp.get("prompt_eval_count", 0),
                "completion_tokens": resp.get("eval_count", 0),
            },
        }

    # --- __call__ (streaming, matches Llama().__call__) ---
    def __call__(self, prompt, **kwargs):
        import json as _json
        payload = {
            "model": self.model_name,
            "prompt": prompt,
            "stream": True,
            "options": {
                "num_predict": kwargs.get("max_tokens", 512),
                "temperature": kwargs.get("temperature", 0.25),
                "top_p": kwargs.get("top_p", 0.8),
                "top_k": kwargs.get("top_k", 50),
                "repeat_penalty": kwargs.get("repeat_penalty", 1.1),
            },
        }
        stop = kwargs.get("stop")
        if stop:
            payload["options"]["stop"] = stop
        resp = self._post("/api/generate", payload, stream=True)
        for line in resp:
            if not line:
                continue
            try:
                obj = _json.loads(line.decode())
            except Exception:
                continue
            token = obj.get("response", "")
            done = obj.get("done", False)
            yield {
                "choices": [{
                    "text": token,
                    "finish_reason": "stop" if done else None,
                }]
            }
            if done:
                break

    # --- create_chat_completion (for vision/multimodal) ---
    def create_chat_completion(self, messages=None, **kwargs):
        import json as _json, base64 as _b64
        ollama_msgs = []
        for m in (messages or []):
            role = m.get("role", "user")
            content = m.get("content", "")
            images = []
            if isinstance(content, list):
                # Multimodal content blocks
                text_parts = []
                for block in content:
                    if block.get("type") == "text":
                        text_parts.append(block["text"])
                    elif block.get("type") == "image_url":
                        url = block["image_url"]["url"]
                        if url.startswith("data:"):
                            b64 = url.split(",", 1)[1] if "," in url else url
                            images.append(b64)
                content = "\n".join(text_parts)
            msg = {"role": role, "content": content}
            if images:
                msg["images"] = images
            ollama_msgs.append(msg)

        payload = {
            "model": self.model_name,
            "messages": ollama_msgs,
            "stream": False,
            "options": {"num_predict": kwargs.get("max_tokens", 512)},
        }
        resp = self._post("/api/chat", payload)
        msg = resp.get("message", {})
        return {
            "choices": [{
                "message": {"role": msg.get("role", "assistant"),
                            "content": msg.get("content", "")},
                "finish_reason": "stop",
            }],
            "usage": {
                "prompt_tokens": resp.get("prompt_eval_count", 0),
                "completion_tokens": resp.get("eval_count", 0),
            },
        }

    @staticmethod
    def is_available(base_url: str = "http://127.0.0.1:11434") -> bool:
        """Check if Ollama server is running."""
        try:
            import urllib.request
            resp = urllib.request.urlopen(f"{base_url}/api/tags", timeout=3)
            return resp.status == 200
        except Exception:
            return False

    @staticmethod
    def list_models(base_url: str = "http://127.0.0.1:11434") -> list:
        """Return list of models available in Ollama."""
        try:
            import urllib.request, json as _json
            resp = urllib.request.urlopen(f"{base_url}/api/tags", timeout=5)
            data = _json.loads(resp.read().decode())
            models = []
            for m in data.get("models", []):
                name = m.get("name", "")
                size_gb = round(m.get("size", 0) / (1024 ** 3), 1)
                models.append({"name": name, "size_gb": size_gb})
            return models
        except Exception:
            return []


# Combined pattern to strip think blocks from both Gemma 3 (<think>...</think>)
# and Gemma 4 (<|channel>thought...<channel|>) model output
_RE_THINK = re.compile(r"<think>.*?</think>|<\|channel>thought.*?<channel\|>", re.DOTALL)
_RE_THINK_INCOMPLETE = re.compile(
    r"<think>(?:(?!</think>).)*$|<\|channel>thought(?:(?!<channel\|>).)*$", re.DOTALL)


class Bridge:
    """Exposed to JavaScript via pywebview window.api."""

    TRIAL_DAYS = 30
    PASSKEY_ENV = "SIMPLIAI_PASSKEY"
    DEFAULT_PASSKEY = "SIMPLIAI-FULL-ACCESS"

    def __init__(self):
        # ── State ──────────────────────────────────────────────
        # Chat state
        self.chats: dict = {}
        self.current_chat_id: str | None = None
        self.message_history: list = []
        self.chat_counter: int = 0

        # Model state
        self.model = None
        self.model_path: str | None = None
        self.model_map: dict = {}
        self.model_config: dict = {}
        self.model_configs: dict = {}
        self.actual_n_ctx: int = 65536  # Actual context window detected from loaded model

        # Generation
        self.generation_in_progress: bool = False
        self.stop_generation_flag: bool = False

        # RAG
        self.rag_manager: RAGManager | None = None
        self.current_rag_database: str | None = None
        self.chat_rag_settings: dict = {}
        self.uploaded_content: str | None = None
        self.uploaded_file_name: str | None = None
        self.uploaded_file_path: str | None = None
        self.uploaded_pages: list[str] | None = None  # page-level chunks for large docs
        self.uploaded_files: list[dict] = []  # multi-file: [{name, path, content, pages}]
        self.chat_file_state: dict = {}  # per-chat uploaded file state
        self._pdf_cache_dir = os.path.join(app_data_path(), "pdf_extract_cache")
        os.makedirs(self._pdf_cache_dir, exist_ok=True)
        self._agent_context_trimmed: bool = False
        self.temp_rag_db_name: str | None = None
        self.last_rag_hits: int = 0

        # Settings
        self.app_settings: dict = {}
        self.chat_system_prompts: dict = {}
        self.web_search_enabled: bool = False
        self._web_search_cache: dict = {}       # session cache: query → (context, sources)
        self.current_theme: str = "Dark"
        self.attached_image: str | None = None
        self.plugin_commands: dict = {}  # '/cmd' -> {handler, plugin, description}

        # System
        self.system_ram: int = round(psutil.virtual_memory().total / (1024**3))
        self.gpu_info: dict = {}
        self.config: dict = {}
        self._torch_module = None
        self._torch_unavailable = False
        self.model_lock = threading.Lock()  # Protect model from concurrent access
        self._tts_lock = threading.Lock()
        self._tts_engine = None
        self._tts_active = False
        self._tts_stop_event = threading.Event()
        self._piper_tmp_path = None

        # Database
        db_path = app_data_path("chats.db")
        self.chat_db = ChatDatabase(db_path)
        self.plugin_manager = PluginManager(app_data_path("plugins"))
        self.mcp_manager = MCPManager(self.chat_db)
        if migrate_json_to_sqlite(self.chat_db):
            print("[MIGRATE] Imported legacy JSON -> chats.db")

        # Dirty flag: set True whenever chat data changes; cleared by auto-save
        self._chats_dirty: bool = False

        # ── Initialize ─────────────────────────────────────────
        self.gpu_info = self._detect_gpu_info()
        self._build_config()
        self._load_model_configs()
        self._load_app_settings()

        # ── Memory optimizer (virtual RAM + watchdog) ──────────
        self.mem_optimizer = MemoryOptimizer(
            system_ram_gb=self.system_ram,
            gpu_info=self.gpu_info,
            emit_fn=lambda ev, d: self._emit(ev, d),
        )
        self.mem_optimizer.start_watchdog(interval=3.0)
        self._initialize_activation_state()
        self._load_chats()
        self._init_rag()
        self._load_system_prompts()
        self._load_plugins()
        self._init_tool_permission()
        self._start_auto_save_loop()

    # ── JS helper: emit event to frontend ──────────────────────

    def _emit(self, event: str, data=None):
        """Send event to JS frontend."""
        try:
            import webview
            if not webview.windows:
                return
            window = webview.windows[0]
        except Exception:
            return
        payload = json.dumps(data) if data is not None else "null"
        try:
            window.evaluate_js(
                f"window.dispatchEvent(new CustomEvent('{event}',"
                f" {{detail: {payload}}}));"
            )
        except Exception:
            pass

    def _status(self, text: str):
        """Emit app status text for top-right status bar in web UI."""
        self._emit("app_status", {"text": text})

    # ── AI Tool-Calling Infrastructure ─────────────────────────

    _TOOL_CALL_RE = re.compile(
        r"<tool_call>\s*(\{.*?\})\s*</tool_call>", re.DOTALL)

    # Read-only tools that don't need user permission
    _SAFE_TOOLS = {"/ls", "/tree", "/find", "/size", "/info"}

    def _init_tool_permission(self):
        """Initialise tool-permission synchronisation primitives."""
        self._tool_permission_event = threading.Event()
        self._tool_permission_granted = False

    def _build_tool_definitions(self) -> str:
        """Build tool description block for the system prompt from registered plugins."""
        if not self.plugin_commands:
            return ""
        tools = []
        for cmd, entry in self.plugin_commands.items():
            desc = entry.get("description", "")
            tools.append(f'  - {cmd}: {desc}')

        # If user has uploaded a file, tell the AI about it
        file_hint = ""
        if self.uploaded_files:
            if len(self.uploaded_files) == 1:
                f0 = self.uploaded_files[0]
                ext = os.path.splitext(f0["name"])[1].lower()
                file_hint = (
                    f"\n\nUPLOADED FILE INFO:\n"
                    f"  Name: {f0['name']}\n"
                    f"  Path: {f0['path']}\n"
                    f"When using /code, these variables are pre-set:\n"
                    f"  UPLOADED_FILE = original file path (use for CSV/Excel/text: pd.read_csv(UPLOADED_FILE))\n"
                    f"  UPLOADED_TEXT = path to extracted text version (use for PDFs/scanned docs: open(UPLOADED_TEXT).read())\n"
                    f"  UPLOADED_NAME = original filename\n"
                    f"For CSV/Excel: read UPLOADED_FILE directly with pandas.\n"
                    f"For PDFs (including scanned): read UPLOADED_TEXT which has all extracted/OCR text.\n"
                    f"Always save output to a file (e.g. df.to_csv('result.csv', index=False)).\n"
                )
            else:
                lines = ["\n\nUPLOADED FILES INFO:"]
                for i, uf in enumerate(self.uploaded_files):
                    lines.append(f"  File {i+1}: {uf['name']} → {uf['path']}")
                lines.append(
                    "When using /code, these variables are pre-set:\n"
                    "  UPLOADED_FILE = path of first uploaded file\n"
                    "  UPLOADED_TEXT = path to extracted text of first file\n"
                    "  UPLOADED_NAME = first filename\n"
                    "For multiple files, use the paths listed above directly.\n"
                    "Always save output to a file (e.g. df.to_csv('result.csv', index=False)).\n"
                )
                file_hint = "\n".join(lines)
        elif self.uploaded_file_path and self.uploaded_file_name:
            ext = os.path.splitext(self.uploaded_file_name)[1].lower()
            file_hint = (
                f"\n\nUPLOADED FILE INFO:\n"
                f"  Name: {self.uploaded_file_name}\n"
                f"  Path: {self.uploaded_file_path}\n"
                f"When using /code, these variables are pre-set:\n"
                f"  UPLOADED_FILE = original file path (use for CSV/Excel/text: pd.read_csv(UPLOADED_FILE))\n"
                f"  UPLOADED_TEXT = path to extracted text version (use for PDFs/scanned docs: open(UPLOADED_TEXT).read())\n"
                f"  UPLOADED_NAME = original filename\n"
                f"For CSV/Excel: read UPLOADED_FILE directly with pandas.\n"
                f"For PDFs (including scanned): read UPLOADED_TEXT which has all extracted/OCR text.\n"
                f"Always save output to a file (e.g. df.to_csv('result.csv', index=False)).\n"
            )

        mcp_defs = self.mcp_manager.build_tool_definitions() if hasattr(self, 'mcp_manager') else ""

        # Use real user home path so the LLM doesn't hallucinate usernames
        _user_home = os.path.expanduser("~").replace("\\", "\\\\")

        return (
            "You have access to local tools. When the user asks you to perform a file or system operation, "
            "call the appropriate tool by emitting EXACTLY this format (no extra text around it):\n"
            "<tool_call>{\"tool\": \"/command\", \"args\": \"full command with arguments\"}</tool_call>\n\n"
            "Available tools:\n"
            + "\n".join(tools) + "\n\n"
            f"The current user's home directory is: {_user_home}\n\n"
            "Examples:\n"
            f'  User: "list files in my Downloads" → <tool_call>{{"tool": "/ls", "args": "/ls {_user_home}\\\\Downloads"}}</tool_call>\n'
            f'  User: "sort the Downloads folder" → <tool_call>{{"tool": "/sort", "args": "/sort {_user_home}\\\\Downloads"}}</tool_call>\n'
            f'  User: "find PDFs in Documents" → <tool_call>{{"tool": "/find", "args": "/find {_user_home}\\\\Documents *.pdf"}}</tool_call>\n'
            '  User: "how big is C drive" → <tool_call>{"tool": "/size", "args": "/size C:\\\\"}</tool_call>\n'
            '  User: "what is 15% of 4800?" → <tool_call>{"tool": "/code", "args": "/code print(4800 * 0.15)"}</tool_call>\n'
            '  User: "generate fibonacci up to 100" → <tool_call>{"tool": "/code", "args": "/code\\na, b = 0, 1\\nwhile a <= 100:\\n    print(a)\\n    a, b = b, a + b"}</tool_call>\n'
            '  User: "extract this PDF to CSV" → <tool_call>{"tool": "/code", "args": "/code\\nimport pandas as pd\\ntext = open(UPLOADED_TEXT).read()\\nlines = [l.strip() for l in text.splitlines() if l.strip()]\\n# parse lines into structured data...\\ndf = pd.DataFrame(data)\\ndf.to_csv(\'extracted.csv\', index=False)\\nprint(df.to_string())"}</tool_call>\n\n'
            "Rules:\n"
            "- ONLY use <tool_call> when the user clearly wants a computation, file/system action, or data task.\n"
            "- For normal questions, just answer directly without tool calls.\n"
            "- For /code: write COMPLETE Python scripts. Use print(), pandas, numpy, matplotlib as needed.\n"
            "- After the tool result is provided, summarize it naturally for the user.\n"
            + file_hint
            + mcp_defs
        )

    def _parse_tool_call(self, text: str):
        """Parse a <tool_call> block from AI response. Returns (cmd, args_str) or None."""
        m = self._TOOL_CALL_RE.search(text)
        if not m:
            return None
        try:
            payload = json.loads(m.group(1))
            tool = payload.get("tool", "").strip()
            args = payload.get("args", "")
            # MCP tool: "mcp:server:tool_name"
            if tool.startswith("mcp:"):
                return (tool, args if isinstance(args, dict) else {})
            # Local plugin tool
            if isinstance(args, str):
                args = args.strip()
            else:
                args = json.dumps(args)
            if tool and tool in self.plugin_commands:
                return (tool, args)
        except (json.JSONDecodeError, AttributeError):
            pass
        return None

    def _execute_tool_call(self, cmd, args_str) -> str:
        """Execute a plugin command or MCP tool and return the result text."""
        # MCP tool: "mcp:server:tool_name"
        if isinstance(cmd, str) and cmd.startswith("mcp:"):
            parts = cmd.split(":", 2)
            if len(parts) != 3:
                return f"Invalid MCP tool format: {cmd}  (expected mcp:server:tool)"
            server_name, tool_name = parts[1], parts[2]
            arguments = args_str if isinstance(args_str, dict) else {}
            result = self.mcp_manager.call_tool(server_name, tool_name, arguments)
            if "error" in result:
                return f"MCP tool error: {result['error']}"
            return result.get("result", str(result))

        # Local plugin tool
        entry = self.plugin_commands.get(cmd)
        if not entry:
            return f"Unknown tool: {cmd}"
        handler = entry.get("handler")
        if not callable(handler):
            return f"Tool {cmd} has no callable handler"
        try:
            result = handler(self, args_str)
            if isinstance(result, dict):
                return result.get("content", str(result))
            return str(result)
        except Exception as e:
            return f"Tool error: {e}"

    def _extract_rag_sources(self, chunks, max_sources: int = 4):
        """Extract unique source file names + snippets from retrieved RAG chunks."""
        if not chunks:
            return []

        sources = []
        seen = set()
        for chunk in chunks:
            for match in re.findall(r"---\s*File:\s*(.+?)\s*---", chunk):
                name = os.path.basename(str(match).strip())
                if name and name not in seen:
                    seen.add(name)
                    # Extract a short snippet from the chunk text
                    snippet = re.sub(r"---\s*File:\s*.+?\s*---", "", chunk).strip()
                    snippet = snippet[:120].replace("\n", " ").strip()
                    if len(snippet) > 117:
                        snippet = snippet[:117] + "..."
                    sources.append({"name": name, "snippet": snippet})
                    if len(sources) >= max_sources:
                        return sources
        return sources

    def _resolve_search_result_url(self, raw_url: str) -> str:
        if not raw_url:
            return ""
        if raw_url.startswith("//"):
            raw_url = "https:" + raw_url
        try:
            parsed = urlparse(raw_url)
            if "duckduckgo.com" in parsed.netloc:
                target = parse_qs(parsed.query).get("uddg", [""])[0]
                if target:
                    return unquote(target)
        except Exception:
            pass
        return raw_url

    def _parse_ddg_lite(self, soup, query: str, num_results: int = 8) -> tuple:
        """Parse DuckDuckGo Lite HTML (table-based layout) as fallback."""
        lines, sources = [], []
        try:
            result_links = soup.find_all("a", class_="result-link")
            # Build list of (title, url) pairs
            pairs = []
            for a in result_links:
                title = a.get_text(strip=True)
                href = self._resolve_search_result_url(a.get("href", ""))
                if title and href:
                    pairs.append((title, href))
            # Get snippets from result-snippet class
            snippets_els = soup.find_all("td", class_="result-snippet")
            for i, (title, url) in enumerate(pairs[:num_results]):
                snippet = ""
                if i < len(snippets_els):
                    snippet = " ".join(snippets_els[i].get_text(strip=True).split())[:280]
                sources.append({
                    "title": title,
                    "url": url,
                    "snippet": snippet,
                    "excerpt": "",
                })
            # Fetch excerpts in parallel
            return self._build_search_context(sources, query)
        except Exception:
            pass
        context = "\n".join(lines) if lines else "No results found."
        return context, sources

    # ── Web search helpers ────────────────────────────────────

    @staticmethod
    def _is_safe_url(url: str) -> bool:
        """SSRF protection — block internal / private network URLs."""
        try:
            import ipaddress
            from urllib.parse import urlparse as _up
            parsed = _up(url)
            host = parsed.hostname or ""
            if not host:
                return False
            # Block obviously internal hostnames
            if host in ("localhost", "127.0.0.1", "0.0.0.0", "::1", "[::1]"):
                return False
            if host.endswith(".local") or host.endswith(".internal"):
                return False
            # Resolve and check private IP ranges
            import socket
            try:
                resolved = socket.getaddrinfo(host, None, socket.AF_UNSPEC,
                                              socket.SOCK_STREAM)
                for _, _, _, _, addr in resolved:
                    ip = ipaddress.ip_address(addr[0])
                    if ip.is_private or ip.is_loopback or ip.is_reserved:
                        return False
            except (socket.gaierror, ValueError):
                pass  # DNS failure — allow, will fail on actual request
            return True
        except Exception:
            return False

    @staticmethod
    def _truncate_at_sentence(text: str, max_chars: int) -> str:
        """Truncate text at the nearest sentence boundary before max_chars."""
        if len(text) <= max_chars:
            return text
        # Look for sentence-ending punctuation near the limit
        truncated = text[:max_chars]
        # Search backwards for sentence boundary (.!? followed by space or end)
        for i in range(len(truncated) - 1, max(0, len(truncated) - 80), -1):
            if truncated[i] in ".!?" and (i + 1 >= len(truncated) or truncated[i + 1] in " \n\t"):
                return truncated[: i + 1]
        # Fallback: break at last space to avoid mid-word cut
        last_space = truncated.rfind(" ", max(0, len(truncated) - 60))
        if last_space > 0:
            return truncated[:last_space] + "…"
        return truncated

    def _fetch_web_excerpt(self, url: str, query: str, max_chars: int = 800) -> str:
        """Fetch the most relevant excerpt from a web page.
        Includes SSRF protection, relevance scoring, and sentence-boundary truncation."""
        if not url or not url.startswith(("http://", "https://")):
            return ""
        if not self._is_safe_url(url):
            print(f"[WEB] Blocked SSRF attempt: {url}")
            return ""
        try:
            import requests
            from bs4 import BeautifulSoup

            response = requests.get(
                url,
                headers={
                    "User-Agent": (
                        "Mozilla/5.0 (Windows NT 10.0; Win64; x64) "
                        "AppleWebKit/537.36 Chrome/125.0 Safari/537.36"
                    ),
                    "Accept": "text/html,application/xhtml+xml",
                    "Accept-Language": "en-US,en;q=0.9",
                },
                timeout=8,
                allow_redirects=True,
            )
            response.raise_for_status()
            soup = BeautifulSoup(response.text, "html.parser")
            # Remove boilerplate
            for tag in soup(["script", "style", "nav", "footer", "header",
                             "aside", "noscript", "form", "iframe"]):
                tag.decompose()

            query_terms = [t.lower() for t in re.findall(r"\b\w{3,}\b", query)]

            # Score candidates — paragraphs, list items, table rows, headings
            candidates = []
            for block in soup.find_all(["p", "li", "td", "th", "article",
                                        "main", "section", "h1", "h2", "h3",
                                        "h4", "blockquote", "pre", "dd"]):
                text = " ".join(block.get_text(" ", strip=True).split())
                if len(text) < 30:
                    continue
                score = sum(1 for t in query_terms if t in text.lower())
                # Bonus for longer, more substantial paragraphs
                length_bonus = min(2, len(text) // 200)
                candidates.append((score + length_bonus, text))

            if not candidates:
                page_text = " ".join(soup.get_text(" ", strip=True).split())
                return self._truncate_at_sentence(page_text, max_chars)

            candidates.sort(key=lambda x: (x[0], len(x[1])), reverse=True)

            # Collect top-scored paragraphs up to max_chars
            excerpt_parts = []
            total = 0
            for _score, text in candidates[:5]:  # up to 5 paragraphs
                remaining = max_chars - total
                if remaining < 50:
                    break
                chunk = self._truncate_at_sentence(text, remaining)
                if chunk:
                    excerpt_parts.append(chunk)
                    total += len(chunk) + 2
            return " ".join(excerpt_parts)[:max_chars]
        except requests.exceptions.Timeout:
            print(f"[WEB] Timeout fetching excerpt: {url}")
            return ""
        except requests.exceptions.ConnectionError:
            print(f"[WEB] Connection error: {url}")
            return ""
        except Exception as e:
            print(f"[WEB] Excerpt error for {url}: {type(e).__name__}")
            return ""

    # ── URL scraping for Agent mode ────────────────────────────

    _URL_PATTERN = re.compile(
        r'https?://[^\s<>"\')\]]+', re.IGNORECASE)

    def scrape_url(self, url: str) -> dict:
        """Scrape a URL and return structured content (text + tables).
        Exposed as a public API method for the frontend."""
        if not url or not url.strip():
            return {"error": "No URL provided"}
        url = url.strip()
        if not url.startswith(("http://", "https://")):
            url = "https://" + url
        if not self._is_safe_url(url):
            return {"error": "URL blocked for security reasons (internal/private address)"}
        try:
            result = self._scrape_page(url)
            return result
        except Exception as e:
            return {"error": f"Scraping failed: {type(e).__name__}: {e}"}

    def _scrape_page(self, url: str, max_chars: int = 0) -> dict:
        """Full page scraper — extracts text, tables, and metadata.
        If max_chars is 0, uses a dynamic limit based on context window."""
        import requests
        from bs4 import BeautifulSoup

        if max_chars <= 0:
            n_ctx = getattr(self, "actual_n_ctx", 2048)
            max_chars = max(8000, int(n_ctx * 3.5 * 0.50))

        headers = {
            "User-Agent": (
                "Mozilla/5.0 (Windows NT 10.0; Win64; x64) "
                "AppleWebKit/537.36 (KHTML, like Gecko) "
                "Chrome/125.0.0.0 Safari/537.36"
            ),
            "Accept": "text/html,application/xhtml+xml,application/xml;q=0.9,*/*;q=0.8",
            "Accept-Language": "en-US,en;q=0.9",
        }
        resp = requests.get(url, headers=headers, timeout=15, allow_redirects=True)
        resp.raise_for_status()

        soup = BeautifulSoup(resp.text, "html.parser")

        # Extract page title
        title = ""
        title_tag = soup.find("title")
        if title_tag:
            title = title_tag.get_text(strip=True)

        # Remove boilerplate
        for tag in soup(["script", "style", "nav", "footer", "header",
                         "aside", "noscript", "form", "iframe", "svg",
                         "button", "input"]):
            tag.decompose()

        # ── Extract tables ──
        tables_md = []
        for table in soup.find_all("table"):
            rows = []
            for tr in table.find_all("tr"):
                cells = []
                for td in tr.find_all(["td", "th"]):
                    cell_text = " ".join(td.get_text(" ", strip=True).split())
                    cells.append(cell_text)
                if cells:
                    rows.append(cells)
            if rows and len(rows) >= 2:
                # Convert to markdown table
                # Normalize column count
                max_cols = max(len(r) for r in rows)
                md_lines = []
                for i, row in enumerate(rows):
                    padded = row + [""] * (max_cols - len(row))
                    md_lines.append("| " + " | ".join(padded) + " |")
                    if i == 0:
                        md_lines.append("| " + " | ".join(["---"] * max_cols) + " |")
                tables_md.append("\n".join(md_lines))
            # Remove table from soup so it doesn't appear in body text too
            table.decompose()

        # ── Extract main body text ──
        body_parts = []
        for block in soup.find_all(["p", "li", "h1", "h2", "h3", "h4",
                                     "h5", "h6", "blockquote", "pre",
                                     "article", "main", "section", "dd"]):
            text = " ".join(block.get_text(" ", strip=True).split())
            if len(text) >= 20:
                # Add markdown heading markers for headings
                tag_name = block.name
                if tag_name in ("h1", "h2", "h3", "h4", "h5", "h6"):
                    level = int(tag_name[1])
                    text = "#" * level + " " + text
                elif tag_name == "li":
                    text = "• " + text
                body_parts.append(text)

        body_text = "\n".join(body_parts)

        # ── Combine tables + body text within budget ──
        tables_text = ""
        if tables_md:
            tables_text = "\n\n### Extracted Tables\n\n" + "\n\n".join(tables_md)

        combined = ""
        if tables_text:
            # Tables first (they're usually what the user wants)
            combined = tables_text[:max_chars]
            remaining = max_chars - len(combined)
            if remaining > 500 and body_text:
                combined += "\n\n### Page Content\n\n" + body_text[:remaining]
        else:
            combined = body_text[:max_chars]

        # Truncate at sentence boundary
        if len(combined) >= max_chars:
            combined = self._truncate_at_sentence(combined, max_chars)

        result = {
            "ok": True,
            "url": url,
            "title": title,
            "content": combined,
            "tables_count": len(tables_md),
            "text_length": len(combined),
        }
        print(f"[SCRAPE] {url} -> {len(combined)} chars, "
              f"{len(tables_md)} tables, title='{title[:60]}'")
        return result

    @staticmethod
    def _extract_urls_from_text(text: str) -> list:
        """Extract all http/https URLs from user text."""
        if not text:
            return []
        return Bridge._URL_PATTERN.findall(text)

    @staticmethod
    def _strip_urls_from_text(text: str) -> str:
        """Remove URLs from text, leaving the rest of the user's message."""
        cleaned = Bridge._URL_PATTERN.sub("", text)
        return " ".join(cleaned.split()).strip()

    # ── System detection ───────────────────────────────────────

    def _get_torch_module(self):
        if self._torch_unavailable:
            return None
        if self._torch_module is not None:
            return self._torch_module
        try:
            import torch
            self._torch_module = torch
            return torch
        except Exception:
            self._torch_unavailable = True
            return None

    def _detect_gpu_info(self):
        # NVIDIA
        try:
            r = subprocess.run(
                ["nvidia-smi", "--query-gpu=memory.total",
                 "--format=csv,noheader,nounits"],
                capture_output=True, text=True, timeout=5)
            if r.returncode == 0 and r.stdout.strip():
                vram_mb = int(r.stdout.strip().splitlines()[0].strip())
                return {"type": "NVIDIA", "backend": "cuda",
                        "vram": round(vram_mb / 1024, 2)}
        except Exception:
            pass
        # Apple Metal
        if sys.platform == "darwin":
            try:
                import platform as _p
                if "arm" in _p.machine().lower():
                    total = round(psutil.virtual_memory().total / (1024**3), 1)
                    return {"type": "APPLE_METAL", "backend": "metal",
                            "vram": total}
                sp = subprocess.run(
                    ["system_profiler", "SPDisplaysDataType"],
                    capture_output=True, text=True, timeout=5)
                if sp.returncode == 0 and "Metal" in sp.stdout:
                    return {"type": "APPLE_METAL", "backend": "metal",
                            "vram": 0}
            except Exception:
                pass
        # AMD
        try:
            if sys.platform == "win32":
                w = subprocess.run(
                    ["wmic", "path", "win32_videocontroller", "get",
                     "Name,AdapterRAM", "/format:csv"],
                    capture_output=True, text=True, timeout=5)
                if w.returncode == 0:
                    for line in w.stdout.strip().splitlines():
                        parts = line.strip().split(",")
                        if len(parts) >= 3:
                            name = parts[2].strip().lower()
                            if "amd" in name or "radeon" in name:
                                try:
                                    vram = round(int(parts[1].strip()) / (1024**3), 1)
                                except (ValueError, TypeError):
                                    vram = 0
                                return {"type": "AMD", "backend": "vulkan",
                                        "vram": max(vram, 0)}
        except Exception:
            pass
        return {"type": "CPU", "backend": "cpu", "vram": 0}

    def _build_config(self):
        gpu_type = self.gpu_info.get("type", "CPU")
        vram = self.gpu_info.get("vram", 0)
        if gpu_type == "NVIDIA":
            gl = -1 if vram >= 4 else (20 if vram >= 2 else 8)
            self.config = {"mode": "GPU", "n_gpu_layers": gl,
                           "n_threads": 4, "max_tokens": 2048}
        elif gpu_type == "APPLE_METAL":
            self.config = {"mode": "GPU (Metal)", "n_gpu_layers": -1,
                           "n_threads": max(4, (os.cpu_count() or 8) - 2),
                           "max_tokens": 2048}
        elif gpu_type == "AMD":
            gl = -1 if vram >= 4 else (20 if vram >= 2 else 8)
            self.config = {"mode": "GPU (Vulkan)", "n_gpu_layers": gl,
                           "n_threads": max(4, (os.cpu_count() or 8) - 2),
                           "max_tokens": 2048}
        elif self.system_ram <= 8:
            self.config = {"mode": "LOW_RAM", "n_gpu_layers": 0,
                           "n_threads": 4, "max_tokens": 800}
        elif self.system_ram <= 16:
            self.config = {"mode": "CPU", "n_gpu_layers": 0,
                           "n_threads": max(8, (os.cpu_count() or 8) - 2),
                           "max_tokens": 1024}
        else:
            self.config = {"mode": "CPU_HIGH", "n_gpu_layers": 0,
                           "n_threads": max(10, (os.cpu_count() or 10) - 2),
                           "max_tokens": 2048}

    # ── Data loading ───────────────────────────────────────────

    def _load_model_configs(self):
        self.model_configs = self.chat_db.get_kv("model_configs", {})

    def _load_app_settings(self):
        self.app_settings = self.chat_db.get_kv("app_settings", {})
        # Restore persisted theme on startup.
        saved_theme = str(self.app_settings.get("theme", self.current_theme)).strip()
        if saved_theme:
            self.current_theme = saved_theme[0].upper() + saved_theme[1:].lower()

    def _is_ai_debug_enabled(self) -> bool:
        """Return True when verbose AI generation diagnostics are enabled.

        Toggle sources:
        - app settings key: debug_ai_generation (bool/1/true/on)
        - env var: SIMPLE_AI_DEBUG_GEN=1
        """
        raw = self.app_settings.get("debug_ai_generation", False)
        if isinstance(raw, str):
            raw = raw.strip().lower() in ("1", "true", "yes", "on")
        env_on = str(os.environ.get("SIMPLE_AI_DEBUG_GEN", "")).strip().lower() in (
            "1", "true", "yes", "on"
        )
        return bool(raw) or env_on

    def _debug_ai_generation(self, stage: str, payload: dict | None = None):
        """Print structured debug diagnostics for model generation when enabled."""
        if not self._is_ai_debug_enabled():
            return
        data = payload or {}
        try:
            # Keep logs readable and avoid dumping full prompts.
            safe = dict(data)
            if "prompt" in safe:
                p = str(safe.get("prompt") or "")
                safe["prompt_chars"] = len(p)
                safe["prompt_head"] = p[:220]
                safe["prompt_tail"] = p[-220:] if len(p) > 220 else p
                safe.pop("prompt", None)
            print(f"[AI-DEBUG] {stage}: {json.dumps(safe, ensure_ascii=False, default=str)}")
        except Exception as dbg_err:
            print(f"[AI-DEBUG] {stage}: <debug log error: {dbg_err}>")

    def _utc_now_iso(self) -> str:
        return datetime.now(timezone.utc).replace(microsecond=0).isoformat()

    def _parse_iso_datetime(self, value):
        try:
            raw = str(value or "").strip()
            if not raw:
                return None
            return datetime.fromisoformat(raw.replace("Z", "+00:00"))
        except Exception:
            return None

    def _activation_secret(self) -> str:
        return str(os.environ.get(self.PASSKEY_ENV, self.DEFAULT_PASSKEY)).strip()

    def _activation_machine_hint(self) -> str:
        return f"{uuid.getnode()}:{os.environ.get('COMPUTERNAME', '')}"

    def _activation_system_code(self) -> str:
        """Short stable code derived from local machine identity."""
        raw = self._activation_machine_hint().encode("utf-8")
        return hashlib.sha256(raw).hexdigest()[:12].upper()

    def _build_machine_bound_key(self, system_code: str) -> str:
        secret = self._activation_secret().encode("utf-8")
        msg = str(system_code or "").strip().upper().encode("utf-8")
        digest = hmac.new(secret, msg, hashlib.sha256).hexdigest()[:24].upper()
        return f"{str(system_code).strip().upper()}-{digest}"

    def _is_valid_activation_key(self, entered_key: str) -> bool:
        expected = self._build_machine_bound_key(self._activation_system_code())
        provided = str(entered_key or "").strip().upper()
        return bool(provided and hmac.compare_digest(provided, expected))

    def _activation_signature(self, first_opened_at: str, activated_at: str) -> str:
        payload = f"{first_opened_at}|{activated_at}|{self._activation_machine_hint()}|{self._activation_secret()}"
        return hashlib.sha256(payload.encode("utf-8")).hexdigest()

    def _activation_store_path(self) -> str:
        """Persistent activation file in user profile (survives app folder delete)."""
        base = os.environ.get("LOCALAPPDATA") or os.path.expanduser("~")
        folder = os.path.join(base, "SIMPLIAI")
        os.makedirs(folder, exist_ok=True)
        return os.path.join(folder, "activation_state.json")

    def _load_activation_store(self) -> dict:
        path = self._activation_store_path()
        try:
            if not os.path.exists(path):
                return {}
            with open(path, "r", encoding="utf-8") as f:
                data = json.load(f)
            return data if isinstance(data, dict) else {}
        except Exception:
            return {}

    def _save_activation_store(self, payload: dict):
        path = self._activation_store_path()
        try:
            with open(path, "w", encoding="utf-8") as f:
                json.dump(payload or {}, f, ensure_ascii=False, indent=2)
        except Exception as e:
            print(f"[ACTIVATION] Could not write activation store: {e}")

    def _initialize_activation_state(self):
        """Persist first-open timestamp and validate existing activation record."""
        changed = False
        if not isinstance(self.app_settings, dict):
            self.app_settings = {}
            changed = True

        # Prefer user-profile activation state so deleting app folder won't reset trial.
        persisted = self._load_activation_store()

        first_opened_at = str(persisted.get("trial_first_opened_at", "")).strip()
        if not first_opened_at:
            first_opened_at = str(self.app_settings.get("trial_first_opened_at", "")).strip()
        if not first_opened_at:
            first_opened_at = self._utc_now_iso()
            changed = True
        self.app_settings["trial_first_opened_at"] = first_opened_at

        activation = persisted.get("activation", {}) if isinstance(persisted.get("activation", {}), dict) else {}
        if not activation:
            activation = self.app_settings.get("activation", {})
        if not isinstance(activation, dict):
            activation = {"activated": False}
            changed = True
        elif activation.get("activated"):
            first_opened = first_opened_at
            activated_at = str(activation.get("activated_at", "")).strip()
            sig = str(activation.get("sig", "")).strip()
            expected = self._activation_signature(first_opened, activated_at) if first_opened and activated_at else ""
            if not expected or not sig or not hmac.compare_digest(sig, expected):
                activation = {"activated": False}
                changed = True

        self.app_settings["activation"] = activation

        persisted_payload = {
            "trial_first_opened_at": first_opened_at,
            "activation": activation,
        }
        if persisted_payload != persisted:
            self._save_activation_store(persisted_payload)

        if changed:
            self.chat_db.set_kv("app_settings", self.app_settings)

    def _get_activation_status(self) -> dict:
        first_opened_at = str(self.app_settings.get("trial_first_opened_at", "")).strip()
        first_dt = self._parse_iso_datetime(first_opened_at)
        if not first_dt:
            first_opened_at = self._utc_now_iso()
            self.app_settings["trial_first_opened_at"] = first_opened_at
            self.chat_db.set_kv("app_settings", self.app_settings)
            first_dt = self._parse_iso_datetime(first_opened_at)

        now_dt = datetime.now(timezone.utc)
        days_used = max(0, (now_dt - first_dt).days) if first_dt else 0
        is_trial_active = days_used < self.TRIAL_DAYS
        days_left = max(0, self.TRIAL_DAYS - days_used)

        activation = self.app_settings.get("activation", {})
        is_activated = False
        activated_at = ""
        if isinstance(activation, dict) and activation.get("activated"):
            activated_at = str(activation.get("activated_at", "")).strip()
            sig = str(activation.get("sig", "")).strip()
            expected = self._activation_signature(first_opened_at, activated_at) if first_opened_at and activated_at else ""
            is_activated = bool(expected and sig and hmac.compare_digest(sig, expected))

        return {
            "trial_days_total": self.TRIAL_DAYS,
            "first_opened_at": first_opened_at,
            "days_used": days_used,
            "days_left": days_left,
            "is_trial_active": is_trial_active,
            "is_activated": is_activated,
            "activated_at": activated_at,
            "requires_passkey": (not is_activated and not is_trial_active),
            "system_code": self._activation_system_code(),
        }

    def _has_full_access(self) -> bool:
        status = self._get_activation_status()
        return bool(status.get("is_trial_active") or status.get("is_activated"))

    def get_activation_status(self):
        """Return current trial/license status for frontend activation UI."""
        return self._get_activation_status()

    def activate_full_access(self, passkey: str):
        """Activate full access after trial expiration using passkey."""
        provided = str(passkey or "").strip().upper()
        if not provided:
            return {"ok": False, "error": "Passkey is required"}
        if not self._is_valid_activation_key(provided):
            return {"ok": False, "error": "Invalid passkey"}

        first_opened_at = str(self.app_settings.get("trial_first_opened_at", "")).strip() or self._utc_now_iso()
        activated_at = self._utc_now_iso()
        self.app_settings["trial_first_opened_at"] = first_opened_at
        self.app_settings["activation"] = {
            "activated": True,
            "activated_at": activated_at,
            "sig": self._activation_signature(first_opened_at, activated_at),
        }
        self.chat_db.set_kv("app_settings", self.app_settings)
        self._save_activation_store({
            "trial_first_opened_at": first_opened_at,
            "activation": self.app_settings["activation"],
        })
        return {"ok": True, "status": self._get_activation_status()}

    def _load_chats(self):
        self.chats = self.chat_db.load_all_chats()
        if self.chats:
            ids = self.chat_db.sorted_chat_ids()
            nums = []
            for cid in self.chats:
                m = re.search(r"(\d+)$", cid)
                if m:
                    nums.append(int(m.group(1)))
            self.chat_counter = max(nums) if nums else len(self.chats)
        else:
            self.chat_counter = 0

    def _init_rag(self):
        try:
            self.rag_manager = RAGManager(app_data_path("rag_databases"))
        except Exception as e:
            print(f"[RAG] Init error: {e}")

    def _load_system_prompts(self):
        self.chat_system_prompts = self.chat_db.get_all_meta("system_prompt")

    def _load_plugins(self):
        try:
            self.plugin_manager.load_all(self)
        except Exception as e:
            print(f"[PLUGINS] Error: {e}")

    def register_plugin_command(self, command: str, handler, plugin_name: str = "", description: str = ""):
        """Register a plugin command handler (e.g. '/hello')."""
        cmd = str(command or "").strip().lower()
        if not cmd:
            return {"error": "Command is empty"}
        if not cmd.startswith("/"):
            cmd = "/" + cmd
        if not callable(handler):
            return {"error": "Handler must be callable"}

        self.plugin_commands[cmd] = {
            "handler": handler,
            "plugin": str(plugin_name or "").strip(),
            "description": str(description or "").strip(),
        }
        return {"ok": True, "command": cmd}

    def unregister_plugin_command(self, command: str):
        """Unregister a plugin command handler."""
        cmd = str(command or "").strip().lower()
        if not cmd.startswith("/"):
            cmd = "/" + cmd
        self.plugin_commands.pop(cmd, None)
        return {"ok": True, "command": cmd}

    def _run_plugin_command(self, text: str) -> dict:
        """Run plugin command if text starts with a registered command.

        Returns {'handled': bool, 'result': ...}.
        """
        txt = str(text or "").strip()
        if not txt.startswith("/"):
            return {"handled": False}

        cmd = txt.split(None, 1)[0].lower()
        entry = self.plugin_commands.get(cmd)
        if not entry:
            return {"handled": False}

        handler = entry.get("handler")
        try:
            result = handler(self, txt)
        except Exception as e:
            content = f"Plugin command {cmd} failed: {e}"
            self._emit("message_added", {"role": "assistant", "content": content})
            self.message_history.append({"role": "assistant", "content": content})
            self.chats[self.current_chat_id] = self.message_history
            self._save_current_chat()
            self._status(f"Plugin error: {cmd}")
            return {"handled": True, "error": str(e)}

        # Handler can return:
        # - str: assistant content
        # - dict with 'content' and optional metadata
        # - dict with handled=False to continue normal routing
        # - any other value: treated as handled with no auto message
        if isinstance(result, dict) and result.get("handled") is False:
            return {"handled": False, "result": result}

        content = None
        if isinstance(result, str):
            content = result.strip()
        elif isinstance(result, dict):
            content = str(result.get("content", "") or "").strip()

        if content:
            self._emit("message_added", {"role": "assistant", "content": content})
            self.message_history.append({"role": "assistant", "content": content})
            self.chats[self.current_chat_id] = self.message_history
            self._save_current_chat()

        self._status(f"Plugin command: {cmd}")
        return {"handled": True, "result": result}

    def execute_python_snippet(self, code: str, timeout: int = 8, allowed_imports=None) -> dict:
        """Execute a Python snippet in a restricted subprocess.

        Intended for plugin command handlers that need lightweight computation.
        Blocks filesystem/network/process access and enforces a timeout.
        """
        snippet = str(code or "").strip()
        if not snippet:
            return {"error": "Empty code"}

        if len(snippet) > 12000:
            return {"error": "Code too long (max 12000 chars)"}

        # Fast deny-list to prevent obvious unsafe operations.
        blocked_patterns = [
            r"\bimport\s+os\b",
            r"\bimport\s+sys\b",
            r"\bimport\s+subprocess\b",
            r"\bimport\s+socket\b",
            r"\bimport\s+shutil\b",
            r"\bimport\s+pathlib\b",
            r"\bopen\s*\(",
            r"\bexec\s*\(",
            r"\beval\s*\(",
            r"__import__",
            r"os\.system",
            r"subprocess\.",
            r"socket\.",
        ]
        lowered = snippet.lower()
        for pat in blocked_patterns:
            if re.search(pat, lowered):
                return {"error": "Unsafe code blocked by policy"}

        whitelist = allowed_imports or ["math", "statistics", "random", "re", "json"]
        safe_imports = [str(x).strip() for x in whitelist if str(x).strip()]

        runner = r'''
import json
import traceback
import importlib
import builtins as _builtins
import sys

payload = json.loads(sys.stdin.read())
code = payload.get("code", "")
allowed = set(payload.get("allowed_imports", []))

def _safe_import(name, globals=None, locals=None, fromlist=(), level=0):
    root = (name or "").split(".", 1)[0]
    if root not in allowed:
        raise ImportError(f"Import blocked: {root}")
    return importlib.import_module(name)

safe_builtins = {
    "abs": _builtins.abs,
    "all": _builtins.all,
    "any": _builtins.any,
    "bool": _builtins.bool,
    "dict": _builtins.dict,
    "enumerate": _builtins.enumerate,
    "float": _builtins.float,
    "int": _builtins.int,
    "len": _builtins.len,
    "list": _builtins.list,
    "max": _builtins.max,
    "min": _builtins.min,
    "pow": _builtins.pow,
    "print": _builtins.print,
    "range": _builtins.range,
    "round": _builtins.round,
    "set": _builtins.set,
    "sorted": _builtins.sorted,
    "str": _builtins.str,
    "sum": _builtins.sum,
    "tuple": _builtins.tuple,
    "zip": _builtins.zip,
    "__import__": _safe_import,
}

scope = {"__builtins__": safe_builtins}
try:
    exec(code, scope, scope)
except Exception:
    traceback.print_exc()
    raise
'''

        try:
            payload = json.dumps({"code": snippet, "allowed_imports": safe_imports})
            proc = subprocess.run(
                [sys.executable, "-I", "-c", runner],
                input=payload,
                text=True,
                capture_output=True,
                timeout=max(1, int(timeout)),
            )
        except subprocess.TimeoutExpired:
            return {"error": f"Execution timed out after {timeout}s"}
        except Exception as e:
            return {"error": f"Execution failed: {e}"}

        stdout = (proc.stdout or "").strip()
        stderr = (proc.stderr or "").strip()
        if len(stdout) > 4000:
            stdout = stdout[:4000] + "\n...[truncated]"
        if len(stderr) > 4000:
            stderr = stderr[:4000] + "\n...[truncated]"

        if proc.returncode != 0:
            return {
                "error": "Python execution failed",
                "stdout": stdout,
                "stderr": stderr,
                "returncode": proc.returncode,
            }

        return {
            "ok": True,
            "stdout": stdout,
            "stderr": stderr,
            "returncode": proc.returncode,
        }

    def _start_auto_save_loop(self):
        def _loop():
            while True:
                try:
                    if self._chats_dirty:
                        self.chat_db.save_all_chats(self.chats)
                        self._chats_dirty = False
                except Exception:
                    pass
                time.sleep(60)

        threading.Thread(target=_loop, daemon=True).start()

    # ── Model scanning ─────────────────────────────────────────

    def get_models(self):
        """Return list of available models (local GGUF + Ollama)."""
        models_dir = Path(app_data_path("models"))
        models_dir.mkdir(exist_ok=True)
        files = list(models_dir.glob("*.gguf"))
        self.model_map = {}
        result = []
        for f in files:
            name = f.stem
            m = re.search(r"(Q\d[_A-Z0-9]*)", name, re.IGNORECASE)
            quant = m.group(1).upper() if m else ""
            try:
                size_gb = round(f.stat().st_size / (1024 ** 3), 1)
            except Exception:
                size_gb = 0
            label = f"{name} ({quant} · {size_gb}GB)" if quant else f"{name} ({size_gb}GB)"
            self.model_map[label] = str(f)
            result.append({"label": label, "path": str(f),
                           "quant": quant, "size_gb": size_gb,
                           "backend": "llama_cpp"})

        # Ollama models (if server is running)
        ollama_models = OllamaModel.list_models()
        for om in ollama_models:
            name = om["name"]
            size_gb = om["size_gb"]
            label = f"[Ollama] {name} ({size_gb}GB)"
            self.model_map[label] = f"ollama://{name}"
            result.append({"label": label, "path": f"ollama://{name}",
                           "quant": "", "size_gb": size_gb,
                           "backend": "ollama"})
        return result

    def select_model(self, label: str):
        """Set the active model path by display label."""
        self.model_path = self.model_map.get(label)
        if self.model_path:
            self._status(f"Selected model: {Path(self.model_path).name}")
        return {"ok": True, "path": self.model_path}

    # ── Model loading ──────────────────────────────────────────

    def load_model(self, *_args, **_kwargs):
        """Load the selected model. Returns immediately; sends events."""
        if not self._has_full_access():
            return {"error": "Trial expired. Enter passkey to activate full access."}
        if not self.model_path:
            return {"error": "No model selected"}
        if self.model is not None:
            return {"error": "Model already loaded. Unload first."}
        self._status("Loading model...")
        threading.Thread(target=self._load_model_thread, daemon=True).start()
        return {"status": "loading"}

    def _load_model_thread(self):
        """Load model in background thread.
        
        Context Window Control:
        - User controls n_ctx via UI: Model Settings → Context Window selector
        - Options: default, 1024, 2048, 4096, 8192, 16384, 32768, 65536, or "Max"
        - Default values based on available RAM (65536 for 32GB+, down to 8192 for <8GB)
        - "Max" (-1) uses model's native trained context (auto-detected)
        - mmap=True enables SSD-backed virtual RAM for safe high context windows
        """
        model_name = Path(self.model_path).stem if not self.model_path.startswith("ollama://") else self.model_path.replace("ollama://", "")
        self._emit("model_status", {"text": f"Loading {model_name}...", "progress": 0})
        try:
            start = time.time()

            # Progress simulation
            for p in range(0, 90, 10):
                self._emit("model_status", {"text": f"Loading {model_name}... {p}%",
                                            "progress": p})
                time.sleep(0.2)

            # ── Ollama backend ──
            if self.model_path.startswith("ollama://"):
                ollama_name = self.model_path.replace("ollama://", "")
                if not OllamaModel.is_available():
                    raise RuntimeError("Ollama server not running. Start it with 'ollama serve'.")
                new_model = OllamaModel(ollama_name)
                n_ctx = new_model.n_ctx()
                with self.model_lock:
                    self.model = new_model
                    self.actual_n_ctx = n_ctx
                    self._active_backend = "ollama"
                    print(f"[MODEL] Loaded Ollama model: {ollama_name}, n_ctx={n_ctx}")

            # ── Local GGUF backend ──
            else:
                if Llama is None:
                    raise RuntimeError("llama-cpp-python not installed")

                filename = Path(self.model_path).name
                per_model = self.model_configs.get(filename, {})

                # Determine n_ctx (context window) with user override support:
                # 1. Start with RAM-based defaults (user can override in UI)
                # 2. Check for per-model user settings (saved from UI)
                # 3. Support "max" (-1) to use model's native trained context
                
                # RAM-based defaults with mmap/SSD-backed virtual memory:
                if self.system_ram >= 32:
                    n_ctx = 65536    # 64K for high-end systems (was 16384)
                elif self.system_ram >= 16:
                    n_ctx = 32768    # 32K for mid-range (was 8192)
                elif self.system_ram >= 8:
                    n_ctx = 16384    # 16K for 8GB+ (was 4096)
                else:
                    n_ctx = 8192     # 8K minimum for constrained systems (was 2048)
                
                # User can override via UI → Model Settings → Context Window
                user_n_ctx = per_model.get("n_ctx", None)
                if user_n_ctx == -1:
                    # "Max" — use model's native context limit (0 = auto-detect in llama.cpp)
                    n_ctx = 0
                    print(f"[MODEL] User selected 'Max' context — using model's native limit (auto-detect)")
                elif user_n_ctx and int(user_n_ctx) > 0:
                    n_ctx = int(user_n_ctx)
                    print(f"[MODEL] User selected n_ctx={n_ctx} from UI settings")
                n_threads = per_model.get("n_threads", self.config["n_threads"])

                # Detect mmproj/clip file for vision support
                clip_path = None
                model_dir = os.path.dirname(self.model_path)
                try:
                    for f in os.listdir(model_dir):
                        lower = f.lower()
                        if "mmproj" in lower or "clip" in lower:
                            clip_path = os.path.join(model_dir, f)
                            break
                except Exception:
                    pass

                # ── Memory-optimised params (mmap + n_batch only) ──
                opt = self.mem_optimizer.optimal_llama_params(
                    base_n_ctx=n_ctx,
                    base_n_threads=n_threads,
                    base_n_gpu_layers=self.config["n_gpu_layers"],
                )
                llama_kwargs = dict(
                    model_path=self.model_path,
                    **opt,
                )
                if clip_path:
                    llama_kwargs["chat_handler"] = self._make_clip_handler(clip_path)

                try:
                    new_model = Llama(**llama_kwargs, flash_attn=True)
                except TypeError:
                    llama_kwargs.pop("flash_attn", None)
                    new_model = Llama(**llama_kwargs)

                with self.model_lock:
                    self.model = new_model
                    self.actual_n_ctx = n_ctx
                    self._active_backend = "llama_cpp"
                    print(f"[MODEL] Loaded GGUF: n_ctx={n_ctx}, mmap=True (SSD-backed virtual RAM enabled)")
                
                try:
                    nc = getattr(self.model, "n_ctx", None)
                    if callable(nc):
                        nc = nc()
                    if isinstance(nc, (int, float)) and nc > 0:
                        self.actual_n_ctx = int(nc)
                        print(f"[MODEL] Model's actual n_ctx: {int(nc)}")
                except Exception:
                    pass

            filename = model_name
            per_model = self.model_configs.get(filename, {}) if not self.model_path.startswith("ollama://") else {}
            self.model_config = {
                "temperature": per_model.get("temperature", 0.25),
                "top_p": 0.8, "repeat_penalty": 1.1,
                "max_initial_tokens": 250, "max_context_tokens": 1200,
            }

            load_time = round(time.time() - start, 2)
            mem_stage = self.mem_optimizer.stage if hasattr(self, 'mem_optimizer') else "normal"
            self._emit("model_loaded", {
                "name": filename, "load_time": load_time,
                "n_ctx": self.actual_n_ctx, "progress": 100,
                "backend": getattr(self, '_active_backend', 'llama_cpp'),
                "memory_stage": mem_stage})
            self._status(f"Loaded in {load_time}s | ctx={self.actual_n_ctx} | mem={mem_stage}")

        except Exception as e:
            self._emit("model_error", {"error": str(e)})
            self._status(f"Model load error: {e}")
            self.model = None

    def unload_model(self):
        """Unload current model and reclaim memory aggressively."""
        with self.model_lock:
            if self.model is not None:
                del self.model
                self.model = None
                self._active_backend = None
                gc.collect()
                if self.gpu_info.get("backend") == "cuda":
                    torch = self._get_torch_module()
                    if torch and torch.cuda.is_available():
                        torch.cuda.empty_cache()
                self.mem_optimizer.force_gc()
                self._status("Model unloaded")
        return {"ok": True}

    def get_memory_status(self):
        """Return current memory state and optimizer status (JS-callable)."""
        return self.mem_optimizer.snapshot()

    def get_pagefile_advice(self):
        """Check system page file config and return recommendations (JS-callable)."""
        return self.mem_optimizer.pagefile_advice()

    def get_model_status(self):
        """Return current model state."""
        if self.model is None:
            return {"loaded": False}
        name = self.model_path
        if name and name.startswith("ollama://"):
            name = name.replace("ollama://", "")
        elif name:
            name = Path(name).name
        return {"loaded": True, "name": name,
                "n_ctx": self.actual_n_ctx,
                "backend": getattr(self, '_active_backend', 'llama_cpp')}

    # ── Chat CRUD ──────────────────────────────────────────────

    def get_chats(self):
        """Return sorted list of chat ids."""
        try:
            ids = self.chat_db.sorted_chat_ids()
        except Exception:
            ids = list(self.chats.keys())
        return ids

    def _save_chat_file_state(self):
        """Save current uploaded file info for the active chat.
        Only stores path + name (NOT full content) to avoid holding 250-page
        PDFs in memory per-chat.  Content is re-extracted on restore."""
        if self.current_chat_id:
            self.chat_file_state[self.current_chat_id] = {
                "name": self.uploaded_file_name,
                "path": self.uploaded_file_path,
                "files": [{"name": f["name"], "path": f["path"]} for f in self.uploaded_files],
            }

    def _restore_chat_file_state(self, chat_id: str):
        """Restore uploaded file info for a chat (or clear if none).
        Re-extracts content from disk path to avoid holding stale data in RAM."""
        state = self.chat_file_state.get(chat_id, {})
        self.uploaded_file_name = state.get("name")
        self.uploaded_file_path = state.get("path")
        self.uploaded_content = None
        self.uploaded_pages = None
        self.uploaded_files = []
        # Restore multi-file list
        saved_files = state.get("files", [])
        for finfo in saved_files:
            fpath = finfo.get("path")
            fname = finfo.get("name")
            if fpath and os.path.isfile(fpath):
                self.uploaded_files.append({
                    "name": fname, "path": fpath,
                    "content": None, "pages": None,
                })
        # Re-extract content from file path if it still exists
        if self.uploaded_file_path and os.path.isfile(self.uploaded_file_path):
            ext = os.path.splitext(self.uploaded_file_path)[1].lower()
            if ext == ".pdf":
                # Async extraction — don't block chat switch for large PDFs
                threading.Thread(
                    target=self._load_pdf_background,
                    args=(self.uploaded_file_path,),
                    daemon=True,
                ).start()
            elif self._is_tabular_file(self.uploaded_file_name or ""):
                self.uploaded_content = self._build_tabular_preview_fast(
                    self.uploaded_file_path)
            else:
                self.uploaded_content = self._extract_text_from_file(
                    self.uploaded_file_path)
        elif self.uploaded_file_name:
            # File path gone — clear stale reference
            self.uploaded_file_name = None
            self.uploaded_file_path = None

    def new_chat(self):
        """Create a new empty chat."""
        self._save_chat_file_state()
        self.chat_counter += 1
        chat_id = f"Chat {self.chat_counter}"
        self.chats[chat_id] = []
        self.message_history = []
        self.current_chat_id = chat_id
        # New chat has no file attached
        self.uploaded_content = None
        self.uploaded_file_name = None
        self.uploaded_file_path = None
        self.uploaded_pages = None
        self.uploaded_files = []
        self.chat_db.save_chat(chat_id, [])
        # Clean up temp files from previous agent processing
        self._cleanup_agent_temp()
        return {"chat_id": chat_id}

    def load_chat(self, chat_id: str):
        """Switch to an existing chat."""
        if chat_id not in self.chats:
            return {"error": "Chat not found"}
        self._save_chat_file_state()
        self.current_chat_id = chat_id
        self.message_history = list(self.chats[chat_id])
        # Load RAG setting
        self.current_rag_database = self.chat_rag_settings.get(chat_id)
        # Restore per-chat file state
        self._restore_chat_file_state(chat_id)
        return {
            "chat_id": chat_id,
            "messages": self.message_history,
            "attached_file": self.uploaded_file_name,
            "attached_files": [f["name"] for f in self.uploaded_files],
        }

    def delete_chat(self, chat_id: str):
        """Delete a chat."""
        self.chats.pop(chat_id, None)
        self.chat_rag_settings.pop(chat_id, None)
        self.chat_system_prompts.pop(chat_id, None)
        self.chat_file_state.pop(chat_id, None)
        self.chat_db.delete_chat(chat_id)
        if self.current_chat_id == chat_id:
            self.current_chat_id = None
            self.message_history = []
        return {"ok": True}

    def rename_chat(self, old_id: str, new_id: str):
        """Rename a chat."""
        if old_id not in self.chats:
            return {"error": "Chat not found"}
        if new_id in self.chats:
            return {"error": "Name already taken"}
        self.chats[new_id] = self.chats.pop(old_id)
        self.chat_db.rename_chat(old_id, new_id)
        # Migrate settings
        if old_id in self.chat_rag_settings:
            self.chat_rag_settings[new_id] = self.chat_rag_settings.pop(old_id)
        if old_id in self.chat_system_prompts:
            self.chat_system_prompts[new_id] = self.chat_system_prompts.pop(old_id)
        if old_id in self.chat_file_state:
            self.chat_file_state[new_id] = self.chat_file_state.pop(old_id)
        if self.current_chat_id == old_id:
            self.current_chat_id = new_id
        return {"ok": True, "chat_id": new_id}

    def branch_chat(self, branch_at_index: int | None = None):
        """Create a branched chat from current chat history up to index."""
        if not self.current_chat_id or not self.message_history:
            return {"error": "No active chat to branch"}

        if branch_at_index is None:
            branch_at_index = len(self.message_history) - 1

        try:
            idx = max(0, min(int(branch_at_index), len(self.message_history) - 1))
        except Exception:
            idx = len(self.message_history) - 1

        branched = [dict(m) for m in self.message_history[: idx + 1]]

        # Next chat id: Chat N+1
        max_n = 0
        for cid in self.chats.keys():
            m = re.search(r"^Chat\s+(\d+)$", cid)
            if m:
                max_n = max(max_n, int(m.group(1)))
        new_id = f"Chat {max_n + 1}"

        self.chats[new_id] = branched
        self.chat_db.save_chat(new_id, branched)

        if self.current_chat_id in self.chat_rag_settings:
            self.chat_rag_settings[new_id] = self.chat_rag_settings[self.current_chat_id]
            self.chat_db.set_meta(new_id, "rag_db", self.chat_rag_settings[new_id])
        if self.current_chat_id in self.chat_system_prompts:
            self.chat_system_prompts[new_id] = self.chat_system_prompts[self.current_chat_id]
            self.chat_db.set_meta(new_id, "system_prompt", self.chat_system_prompts[new_id])

        self.current_chat_id = new_id
        self.message_history = list(branched)
        self._emit("chat_branched", {"chat_id": new_id})
        return {"ok": True, "chat_id": new_id, "messages": branched}

    def _get_response_format(self, request_options: dict | None = None) -> str:
        if request_options and request_options.get("response_format") == "table":
            return "table"
        value = str(self.app_settings.get("response_format", "normal")).strip().lower()
        return "table" if value == "table" else "normal"

    def _get_response_format_instruction(self, request_options: dict | None = None) -> str:
        if self._get_response_format(request_options) != "table":
            return ""
        return (
            "Format the final answer as a markdown table whenever possible. "
            "If the answer is a single fact, use a two-column table with headers 'Field' and 'Value'. "
            "If a table needs context, add at most one short sentence before the table."
        )

    def _resolve_rag_mention(self, text: str) -> tuple[str, str | None]:
        """Parse @rag_name from message text.

        Returns (cleaned_text, matched_db_name | None).
        Supports: @name, @"name with spaces", @'name with spaces'.
        """
        if not self.rag_manager or "@" not in text:
            return text, None
        db_names = self.rag_manager.list_databases()
        if not db_names:
            return text, None

        # Try quoted mentions first: @"My DB" or @'My DB'
        m = re.search(r'@(["\'])(.+?)\1', text)
        if m:
            candidate = m.group(2).strip()
            for db in db_names:
                if db.lower() == candidate.lower():
                    cleaned = text[:m.start()].rstrip() + " " + text[m.end():].lstrip()
                    return cleaned.strip(), db
            return text, None

        # Try unquoted @name  (match longest db name first to avoid partial hits)
        for db in sorted(db_names, key=len, reverse=True):
            pattern = re.compile(r'@' + re.escape(db) + r'(?=\s|$|[.,!?;:])', re.IGNORECASE)
            match = pattern.search(text)
            if match:
                cleaned = text[:match.start()].rstrip() + " " + text[match.end():].lstrip()
                return cleaned.strip(), db
        return text, None

    def _parse_prompt_options(self, text: str) -> dict:
        normalized = " ".join((text or "").lower().split())
        options = {"response_format": "normal", "export_format": None}

        table_markers = (
            "table format", "tabular", "in a table", "as a table",
            "markdown table", "table with columns", "columns and rows"
        )
        if any(marker in normalized for marker in table_markers):
            options["response_format"] = "table"

        export_patterns = {
            "xlsx": (
                r"\b(export|save|download|create|generate|give|provide|convert|make|turn)\b.{0,30}\b(excel|xlsx|spreadsheet)\b",
                r"\b(excel|xlsx|spreadsheet)\b.{0,20}\b(file|sheet|download|export|save|format)\b",
                r"\bto\s+(excel|xlsx|spreadsheet)\b",
            ),
            "csv": (
                r"\b(export|save|download|create|generate|give|provide|convert|make|turn)\b.{0,30}\bcsv\b",
                r"\bcsv\b.{0,20}\b(file|download|export|save|format)\b",
                r"\bto\s+csv\b",
            ),
            "docx": (
                r"\b(export|save|download|create|generate|give|provide|convert|make|turn)\b.{0,30}\b(docx?|word|ms\s*word|microsoft\s*word)\b",
                r"\b(docx?|word|ms\s*word|microsoft\s*word)\b.{0,20}\b(file|document|download|export|save|format)\b",
                r"\bto\s+(docx?|word|ms\s*word|microsoft\s*word)\b",
            ),
            "pdf": (
                r"\b(export|save|download|create|generate|give|provide|convert|make|turn)\b.{0,30}\bpdf\b",
                r"\bpdf\b.{0,20}\b(file|document|download|export|save|format)\b",
                r"\bto\s+pdf\b",
            ),
            "txt": (
                r"\b(export|save|download|create|generate|give|provide|convert|make|turn)\b.{0,30}\b(txt|text file|plain text)\b",
                r"\b(txt|text file|plain text)\b.{0,20}\b(file|download|export|save|format)\b",
                r"\bto\s+(txt|text\s*file|plain\s*text)\b",
            ),
        }
        for export_format, patterns in export_patterns.items():
            if any(re.search(pattern, normalized) for pattern in patterns):
                options["export_format"] = export_format
                break

        if options["export_format"] in {"csv", "xlsx"}:
            options["response_format"] = "table"

        return options

    def _build_chat_execution_guidance(self, text: str) -> str:
        """Return intent-aware guidance so chat answers include best in-app execution path."""
        q = " ".join((text or "").lower().split())
        if not q:
            return ""

        recon_markers = (
            "recon", "reconciliation", "mismatch", "difference", "compare", "2a", "books", "gst"
        )
        pdf_markers = (
            "pdf", "agreement", "contract", "clause", "document", "policy", "instruction"
        )
        sql_markers = (
            "sql", "duckdb", "query", "join", "group by", "with "
        )

        wants_recon = any(m in q for m in recon_markers)
        wants_pdf = any(m in q for m in pdf_markers)
        wants_sql = any(m in q for m in sql_markers)

        if wants_recon and (wants_sql or "query" in q or "logic" in q):
            return (
                "CHAT RESPONSE STYLE (must include both):\n"
                "1) Best Way In This App: Recommend Agent tab with tabular pipeline (DuckDB SQL execution), "
                "upload files, and if table aliases are user-friendly (Books/2A), note mapping to loaded df_ tables.\n"
                "2) Solution: Provide one executable DuckDB SQL query template and mention expected output shape.\n"
                "Keep it practical and execution-first."
            )

        if wants_pdf:
            return (
                "CHAT RESPONSE STYLE (must include both):\n"
                "1) Best Way In This App: Recommend Agent tab, upload PDF, provide explicit instruction checks/thresholds, "
                "and note instruction-compliance validation retries.\n"
                "2) Solution: Provide a ready-to-use instruction template with measurable pass/fail checks and expected output table columns.\n"
                "Do not give generic advice only; provide an actionable template."
            )

        if wants_recon:
            return (
                "CHAT RESPONSE STYLE (must include both):\n"
                "1) Best Way In This App: route to Agent tab + tabular analysis.\n"
                "2) Solution: provide concrete reconciliation steps and optional SQL starter."
            )

        return (
            "CHAT RESPONSE STYLE:\n"
            "Briefly suggest the best in-app path (Chat vs Agent, tabular vs document), then answer the user question directly."
        )

    def _sanitize_export_name(self, value: str) -> str:
        safe = re.sub(r'[<>:"/\\|?*]+', '_', (value or '').strip())
        safe = safe.strip(' ._')
        return safe or 'export'

    def _extract_markdown_table(self, text: str):
        lines = [line.strip() for line in (text or '').splitlines() if line.strip()]
        table_lines = [line for line in lines if line.count('|') >= 2]
        if len(table_lines) < 2:
            return None, None

        rows = []
        for line in table_lines:
            cells = [cell.strip() for cell in line.strip().strip('|').split('|')]
            if len(cells) >= 2:
                rows.append(cells)

        if len(rows) < 2:
            return None, None

        is_separator = all(re.fullmatch(r'[:\- ]+', cell or '') for cell in rows[1])
        if is_separator:
            headers = rows[0]
            data_rows = rows[2:]
        else:
            # No separator row — first row is still the header
            headers = rows[0]
            data_rows = rows[1:]

        if not data_rows:
            return None, None

        width = len(headers)
        normalized = []
        for row in data_rows:
            current = list(row[:width])
            if len(current) < width:
                current.extend([''] * (width - len(current)))
            normalized.append(current)
        return headers, normalized

    def _get_export_payload(self, chat_id: str, message_index: int | None = None):
        msgs = self.chats.get(chat_id, [])
        if message_index is not None:
            if message_index < 0 or message_index >= len(msgs):
                raise IndexError("Message index out of range")
            msg = msgs[message_index]
            role = msg.get("role", "assistant")
            content = msg.get("content", "")
            headers, table_rows = self._extract_markdown_table(content)
            return {
                "title": f"{chat_id}_{role}_{message_index + 1}",
                "text": content if role == "assistant" else f"[{role.upper()}]\n{content}",
                "rows": [{"role": role, "content": content}],
                "table_headers": headers,
                "table_rows": table_rows,
            }

        rows = []
        for msg in msgs:
            rows.append({
                "role": msg.get("role", "assistant"),
                "content": msg.get("content", ""),
            })
        return {
            "title": chat_id,
            "text": self.export_chat(chat_id),
            "rows": rows,
            "table_headers": None,
            "table_rows": None,
        }

    def _write_pdf_export(self, path: str, title: str, text: str):
        """Export to PDF with proper table formatting using reportlab."""
        try:
            from reportlab.lib.pagesizes import letter, A4
            from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
            from reportlab.lib.units import inch
            from reportlab.platypus import SimpleDocTemplate, Table, TableStyle, Paragraph, Spacer
            from reportlab.lib import colors
            import re
            
            # Create PDF
            doc = SimpleDocTemplate(path, pagesize=letter, topMargin=0.5*inch, bottomMargin=0.5*inch)
            story = []
            styles = getSampleStyleSheet()
            
            # Custom styles
            title_style = ParagraphStyle(
                'CustomTitle',
                parent=styles['Heading1'],
                fontSize=16,
                textColor=colors.HexColor('#1a1a1a'),
                spaceAfter=0.3*inch,
            )
            
            body_style = ParagraphStyle(
                'CustomBody',
                parent=styles['BodyText'],
                fontSize=10,
                leading=12,
            )
            
            # Add title
            story.append(Paragraph(title, title_style))
            story.append(Spacer(1, 0.2*inch))
            
            # Parse content for tables and text
            table_pattern = r'\|[^\n]*\|(?:\n\|[^\n]*\|)*'
            parts = re.split(f'({table_pattern})', text or '')
            
            for part in parts:
                if not part or not part.strip():
                    continue
                
                # Check if this is a table
                if part.strip().startswith('|') and '\n' in part:
                    lines = part.strip().split('\n')
                    lines = [l.strip() for l in lines if l.strip().startswith('|')]
                    
                    if len(lines) >= 2:
                        try:
                            # Skip separator if exists
                            data_start = 1
                            if re.match(r'^[\|\s\-:]+$', lines[1]):
                                data_start = 2
                            
                            def get_cells(line):
                                return [c.strip() for c in line.split('|') if c.strip() and not re.match(r'^-+$', c.strip())]
                            
                            headers = get_cells(lines[0])
                            rows = [get_cells(lines[i]) for i in range(data_start, len(lines))]
                            rows = [r for r in rows if len(r) > 0]
                            
                            if headers and rows:
                                # Build table data
                                table_data = [headers] + rows
                                
                                # Limit table size for performance
                                table_data = table_data[:51]  # Max 50 rows + header
                                max_cols = 8
                                table_data = [[cell[:40] for cell in row[:max_cols]] for row in table_data]
                                
                                # Create table
                                table = Table(table_data, colWidths=[7.5*inch/len(headers) for _ in headers])
                                table.setStyle(TableStyle([
                                    ('BACKGROUND', (0, 0), (-1, 0), colors.HexColor('#e8e8e8')),
                                    ('TEXTCOLOR', (0, 0), (-1, 0), colors.black),
                                    ('ALIGN', (0, 0), (-1, -1), 'LEFT'),
                                    ('FONTNAME', (0, 0), (-1, 0), 'Helvetica-Bold'),
                                    ('FONTSIZE', (0, 0), (-1, 0), 10),
                                    ('BOTTOMPADDING', (0, 0), (-1, 0), 12),
                                    ('BACKGROUND', (0, 1), (-1, -1), colors.white),
                                    ('GRID', (0, 0), (-1, -1), 0.5, colors.black),
                                    ('FONTSIZE', (0, 1), (-1, -1), 9),
                                    ('ROWBACKGROUNDS', (0, 1), (-1, -1), [colors.white, colors.HexColor('#f5f5f5')]),
                                ]))
                                
                                story.append(table)
                                story.append(Spacer(1, 0.2*inch))
                                continue
                        except Exception as e:
                            print(f"[PDF] Table render error: {e}")
                
                # Regular text
                for paragraph in part.strip().split('\n'):
                    if paragraph.strip():
                        story.append(Paragraph(paragraph, body_style))
                story.append(Spacer(1, 0.1*inch))
            
            # Build PDF
            doc.build(story)
            
        except ImportError:
            # Fallback if reportlab not installed
            print("[PDF] reportlab not available, using simple text export")
            self._write_pdf_simple(path, title, text)
        except Exception as e:
            print(f"[PDF] Export error: {e}")
            self._write_pdf_simple(path, title, text)

    def _write_pdf_simple(self, path: str, title: str, text: str):
        """Fallback: Simple PDF export with formatted text."""
        import fitz
        
        doc = fitz.open()
        page = doc.new_page()
        left = 42
        top = 46
        y = top
        max_y = page.rect.height - 42
        line_height = 14
        
        def write_line(content: str, font_size: int = 11):
            nonlocal page, y, max_y
            if y > max_y:
                page = doc.new_page()
                y = top
                max_y = page.rect.height - 42
            
            try:
                safe = str(content).encode('latin-1', errors='replace').decode('latin-1')
                page.insert_text((left, y), safe, fontsize=font_size)
            except Exception:
                pass
            
            y += line_height if font_size <= 11 else line_height + 4
        
        # Title
        write_line(title, font_size=14)
        y += 6
        
        # Content
        for line in (text or '').split('\n'):
            if line.strip():
                wrapped = textwrap.wrap(line, width=95) or ['']
                for wrapped_line in wrapped:
                    write_line(wrapped_line)
            y += 3
        
        doc.save(path)
        doc.close()

    def _write_docx_export(self, path: str, title: str, text: str,
                           rows: list[dict], table_headers, table_rows):
        from docx import Document

        doc = Document()
        doc.add_heading(title, level=1)

        if table_headers and table_rows:
            table = doc.add_table(rows=len(table_rows) + 1, cols=len(table_headers))
            table.style = 'Table Grid'
            for col, header in enumerate(table_headers):
                table.cell(0, col).text = str(header)
            for row_index, row in enumerate(table_rows, start=1):
                for col, value in enumerate(row):
                    table.cell(row_index, col).text = str(value)
        elif len(rows) > 1:
            for row in rows:
                role = str(row.get('role', 'assistant')).capitalize()
                doc.add_paragraph(f"[{role}]")
                doc.add_paragraph(str(row.get('content', '')))
        else:
            for paragraph in (text or '').splitlines():
                doc.add_paragraph(paragraph)

        doc.save(path)

    def _parse_structured_lines(self, text: str):
        """Try to extract key-value pairs or list items from AI response text.

        Returns (headers, rows) if structured data found, else (None, None).
        """
        lines = [l.strip() for l in (text or "").splitlines() if l.strip()]
        if not lines:
            return None, None

        # 1) Key: Value pairs  (e.g. "Name: John", "Amount: 500")
        kv_pairs = []
        kv_pattern = re.compile(r"^([\w\s/().#-]{2,40})\s*[:=]\s*(.+)$")
        for line in lines:
            clean = re.sub(r"^[\-\*\d.)+]+\s*", "", line)  # strip bullet/number prefix
            m = kv_pattern.match(clean)
            if m:
                kv_pairs.append((m.group(1).strip(), m.group(2).strip()))
        if len(kv_pairs) >= 3:
            return ["Field", "Value"], [[k, v] for k, v in kv_pairs]

        # 2) Bullet / numbered list items
        list_items = []
        for line in lines:
            clean = re.sub(r"^[\-\*•]\s*", "", line)
            clean = re.sub(r"^\d+[.)]+\s*", "", clean)
            if clean != line:  # line had a list prefix
                list_items.append(clean)
        if len(list_items) >= 2:
            return ["Item"], [[item] for item in list_items]

        return None, None

    def _write_csv_export(self, path: str, rows: list[dict], table_headers, table_rows):
        with open(path, 'w', encoding='utf-8', newline='') as handle:
            writer = csv.writer(handle)
            if table_headers and table_rows:
                writer.writerow(table_headers)
                writer.writerows(table_rows)
                return
            # Try to parse structured data from the assistant response
            for row in rows:
                if row.get('role') == 'assistant':
                    parsed_h, parsed_r = self._parse_structured_lines(row.get('content', ''))
                    if parsed_h and parsed_r:
                        writer.writerow(parsed_h)
                        writer.writerows(parsed_r)
                        return
            # Final fallback — export as role/content
            writer.writerow(["role", "content"])
            for row in rows:
                writer.writerow([row.get('role', ''), row.get('content', '')])

    def _write_xlsx_export(self, path: str, rows: list[dict], table_headers, table_rows):
        import pandas as pd

        if table_headers and table_rows:
            frame = pd.DataFrame(table_rows, columns=table_headers)
        else:
            # Try to parse structured data from the assistant response
            for row in rows:
                if row.get('role') == 'assistant':
                    parsed_h, parsed_r = self._parse_structured_lines(row.get('content', ''))
                    if parsed_h and parsed_r:
                        frame = pd.DataFrame(parsed_r, columns=parsed_h)
                        break
            else:
                frame = pd.DataFrame(rows)
        frame.to_excel(path, index=False)

    def _write_wav_export(self, path: str, text: str):
        """Export text content as WAV using pyttsx3."""
        try:
            import pyttsx3
        except Exception:
            raise RuntimeError("pyttsx3 is not available for WAV export")

        clean = re.sub(r"[#*`_~\[\]()]+", "", str(text or "")).strip()
        clean = re.sub(r"\n+", ". ", clean)
        if not clean:
            clean = "No content available for audio export."

        engine = pyttsx3.init()
        try:
            self._apply_selected_tts_voice(engine)
        except Exception:
            pass
        engine.save_to_file(clean[:12000], path)
        engine.runAndWait()

    def _write_export_file(self, path: str, export_format: str, payload: dict):
        if export_format == "txt":
            with open(path, "w", encoding="utf-8") as handle:
                handle.write(payload["text"])
            return
        if export_format == "pdf":
            self._write_pdf_export(path, payload["title"], payload["text"])
            return
        if export_format == "docx":
            self._write_docx_export(
                path,
                payload["title"],
                payload["text"],
                payload["rows"],
                payload["table_headers"],
                payload["table_rows"],
            )
            return
        if export_format == "csv":
            self._write_csv_export(path, payload["rows"], payload["table_headers"], payload["table_rows"])
            return
        if export_format == "xlsx":
            self._write_xlsx_export(path, payload["rows"], payload["table_headers"], payload["table_rows"])
            return
        if export_format == "wav":
            self._write_wav_export(path, payload["text"])
            return
        raise ValueError(f"Unsupported export format: {export_format}")

    def _auto_export_response(self, chat_id: str, message_index: int, export_format: str | None):
        if not export_format:
            return None

        payload = self._get_export_payload(chat_id, message_index)

        # When the user asked for CSV/XLSX/DOCX export but the AI response
        # has no structured table, and a file is currently uploaded, try to
        # export the uploaded file content directly.
        if (
            export_format in ("csv", "xlsx", "docx")
            and not payload.get("table_headers")
            and self.uploaded_file_path
        ):
            try:
                uploaded_exported = self._export_uploaded_file_as(
                    export_format, payload.get("title", chat_id)
                )
                if uploaded_exported:
                    self._emit("export_ready", {"path": uploaded_exported, "format": export_format, "chat_id": chat_id})
                    self._status(f"Exported uploaded file to {uploaded_exported}")
                    return uploaded_exported
            except Exception:
                pass  # fall through to default export

        exports_dir = app_data_path("exports")
        os.makedirs(exports_dir, exist_ok=True)
        timestamp = time.strftime("%Y%m%d_%H%M%S")
        filename = f"{self._sanitize_export_name(payload['title'])}_{timestamp}.{export_format}"
        path = os.path.join(exports_dir, filename)
        self._write_export_file(path, export_format, payload)
        self._emit("export_ready", {"path": path, "format": export_format, "chat_id": chat_id})
        self._status(f"Exported reply to {path}")
        return path

    def _export_uploaded_file_as(self, export_format: str, title: str) -> str | None:
        """Convert the uploaded file (PDF/text/tabular) to CSV or XLSX directly."""
        import pandas as pd

        src = self.uploaded_file_path
        if not src or not os.path.isfile(src):
            return None

        ext = os.path.splitext(self.uploaded_file_name or "")[1].lower()
        df = None

        # Tabular files → read directly
        if ext in (".csv", ".xlsx", ".xls", ".tsv"):
            try:
                df = self._load_tabular_df(src)
            except Exception:
                pass

        # PDF → use PyMuPDF table extraction for structured data
        if df is None and ext == ".pdf":
            try:
                import fitz
                doc = fitz.open(src)
                all_frames = []
                for page in doc:
                    tables = page.find_tables()
                    for tab in tables.tables:
                        page_df = tab.to_pandas()
                        if not page_df.empty:
                            # Clean newlines in cell values
                            for col in page_df.columns:
                                page_df[col] = page_df[col].astype(str).str.replace(
                                    r"\n", " ", regex=True
                                ).str.strip()
                            all_frames.append(page_df)
                doc.close()
                if all_frames:
                    # Concatenate all page tables (same headers)
                    df = pd.concat(all_frames, ignore_index=True)
                    # Clean newlines in column names too
                    df.columns = [str(c).replace("\n", " ").strip() for c in df.columns]
            except Exception:
                pass

        # PDF / text fallback → parse from extracted text
        if df is None:
            raw = None
            # 1) Pre-extracted text file (set during upload flow)
            text_path = getattr(self, "uploaded_text_path", None)
            if text_path and os.path.isfile(text_path):
                with open(text_path, "r", encoding="utf-8", errors="replace") as fh:
                    raw = fh.read()
            # 2) Plain text files
            elif ext in (".txt", ".md", ".log"):
                with open(src, "r", encoding="utf-8", errors="replace") as fh:
                    raw = fh.read()
            # 3) Other formats — extract text directly
            else:
                try:
                    raw = self._extract_text_from_file(src)
                except Exception:
                    pass

            if raw:
                # Try markdown table extraction first
                headers, rows = self._extract_markdown_table(raw)
                if headers and rows:
                    df = pd.DataFrame(rows, columns=headers)
                else:
                    # Fallback: split text into lines as a single-column frame
                    lines = [l.strip() for l in raw.splitlines() if l.strip()]
                    if lines:
                        df = pd.DataFrame(lines, columns=["content"])

        if df is None or df.empty:
            return None

        exports_dir = app_data_path("exports")
        os.makedirs(exports_dir, exist_ok=True)
        timestamp = time.strftime("%Y%m%d_%H%M%S")
        safe_name = self._sanitize_export_name(
            os.path.splitext(self.uploaded_file_name or title)[0]
        )
        out_path = os.path.join(exports_dir, f"{safe_name}_{timestamp}.{export_format}")

        if export_format == "csv":
            df.to_csv(out_path, index=False)
        elif export_format == "xlsx":
            df.to_excel(out_path, index=False)
        elif export_format == "docx":
            from docx import Document as DocxDocument
            doc = DocxDocument()
            doc.add_heading(safe_name, level=1)
            table = doc.add_table(rows=len(df) + 1, cols=len(df.columns))
            table.style = "Table Grid"
            for col_i, col_name in enumerate(df.columns):
                table.cell(0, col_i).text = str(col_name)
            for row_i, (_, row) in enumerate(df.iterrows(), start=1):
                for col_i, val in enumerate(row):
                    table.cell(row_i, col_i).text = str(val)
            doc.save(out_path)
        else:
            return None
        return out_path

    def export_chat(self, chat_id: str):
        """Return chat as formatted text for export."""
        msgs = self.chats.get(chat_id, [])
        lines = [f"=== {chat_id} ===\n"]
        for m in msgs:
            role = "[You]" if m["role"] == "user" else "[AI]"
            lines.append(f"{role}\n{m['content']}\n")
        return "\n".join(lines)

    def export_chat_dialog(self, chat_id: str | None = None,
                           export_format: str = "txt",
                           message_index: int | None = None):
        """Open save dialog and export a chat or a single reply."""
        try:
            import webview
            if chat_id is None:
                chat_id = self.current_chat_id
            if not chat_id or chat_id not in self.chats:
                return {"error": "No chat to export"}
            if not webview.windows:
                return {"error": "Window not ready"}

            export_format = str(export_format or "txt").strip().lower()
            # Normalize UI/input variants: "WAV", ".wav", "audio/wav", etc.
            normalized_fmt = re.sub(r"[^a-z0-9]+", "", export_format)
            fmt_aliases = {
                "text": "txt",
                "plaintext": "txt",
                "word": "docx",
                "excel": "xlsx",
                "wave": "wav",
                "audiowav": "wav",
            }
            if normalized_fmt in fmt_aliases:
                export_format = fmt_aliases[normalized_fmt]
            elif normalized_fmt in {"txt", "pdf", "docx", "csv", "xlsx", "wav"}:
                export_format = normalized_fmt
            labels = {
                "txt": "Text files (*.txt)",
                "pdf": "PDF files (*.pdf)",
                "docx": "Word documents (*.docx)",
                "csv": "CSV files (*.csv)",
                "xlsx": "Excel files (*.xlsx)",
                "wav": "WAV audio (*.wav)",
            }
            if export_format not in labels:
                return {"error": f"Unsupported export format: {export_format}"}

            payload = self._get_export_payload(chat_id, message_index)
            print(f"[EXPORT] Opening save dialog for format: {export_format}")
            file_path = webview.windows[0].create_file_dialog(
                webview.FileDialog.SAVE,
                save_filename=f"{self._sanitize_export_name(payload['title'])}.{export_format}",
                file_types=(labels[export_format],),
            )
            print(f"[EXPORT] File dialog returned: {file_path}")
            if not file_path:
                print("[EXPORT] User cancelled - returning selected:False")
                return {"selected": False}
            path = file_path if isinstance(file_path, str) else file_path[0]
            print(f"[EXPORT] Writing to: {path}")
            self._write_export_file(path, export_format, payload)
            print(f"[EXPORT] Success: {path}")
            result = {"ok": True, "path": path}
            print(f"[EXPORT] Returning result: {result}")
            return result
        except IndexError as e:
            print(f"[EXPORT] IndexError: {e}")
            return {"error": str(e)}
        except Exception as e:
            import traceback
            print(f"[EXPORT] Exception: {e}")
            traceback.print_exc()
            return {"error": str(e)}

    # ── Instruction Templates ──────────────────────────────────

    def get_instruction_templates(self):
        """Return saved instruction templates."""
        templates_file = os.path.join(app_data_path(), "instruction_templates.json")
        try:
            if os.path.exists(templates_file):
                with open(templates_file, 'r', encoding='utf-8') as f:
                    return json.load(f)
        except Exception as e:
            print(f"[TEMPLATES] Load error: {e}")
        return {}

    def save_instruction_template(self, name: str, instructions: str):
        """Save an instruction template."""
        if not name or not name.strip():
            return {"error": "Template name required"}
        if not instructions or not instructions.strip():
            return {"error": "Instructions required"}
        templates_file = os.path.join(app_data_path(), "instruction_templates.json")
        templates = self.get_instruction_templates()
        templates[name.strip()] = instructions.strip()
        try:
            with open(templates_file, 'w', encoding='utf-8') as f:
                json.dump(templates, f, indent=2)
            return {"ok": True}
        except Exception as e:
            return {"error": str(e)}

    def delete_instruction_template(self, name: str):
        """Delete an instruction template."""
        templates_file = os.path.join(app_data_path(), "instruction_templates.json")
        templates = self.get_instruction_templates()
        if name in templates:
            del templates[name]
            with open(templates_file, 'w', encoding='utf-8') as f:
                json.dump(templates, f, indent=2)
        return {"ok": True}

    def rewrite_instruction(self, fields: dict):
        """Use the loaded LLM to convert agent instructions into execution-ready form."""
        if self.model is None:
            return {"error": "Load a model first"}
        role = (fields.get("role") or "").strip()
        task = (fields.get("task") or "").strip()
        steps = (fields.get("steps") or "").strip()

        def _format_steps_text(value: str) -> str:
            """Normalize steps to numbered sections with '*' bullets."""
            text = (value or "").replace("\r\n", "\n")
            lines = [ln.strip() for ln in text.split("\n") if ln.strip()]
            if not lines:
                return ""

            out = []
            section_idx = 0
            for ln in lines:
                if re.match(r"^\d+\.\s*", ln):
                    section_idx += 1
                    section = re.sub(r"^\d+\.\s*", "", ln).strip()
                    if not section.endswith(":"):
                        section += ":"
                    out.append(f"{section_idx}. {section}")
                    continue

                if ln.startswith(("*", "-", "•")):
                    out.append(f"* {ln.lstrip('*-• ').strip()}")
                    continue

                if section_idx == 0:
                    section_idx = 1
                    out.append("1. Steps:")
                out.append(f"* {ln}")

            return "\n".join(out)

        if not role and not task and not steps:
            return {"error": "Nothing to rewrite"}

        prompt_parts = [
            "Rewrite the following agent instruction fields into an execution-ready task. "
            "Do more than grammar correction: convert vague wording into precise, operational steps an AI agent can execute. "
            "Keep the original intent, but make the instructions concrete, explicit, and implementation-ready.\n"
            "Return ONLY a JSON object with keys \"role\", \"task\", \"steps\" "
            "(steps as a single string). No markdown fences.\n"
            "REWRITE RULES:\n"
            "- Preserve the user's goal and domain (learning, support, coding, finance, etc.).\n"
            "- Keep the same section intent and labels when present (e.g., Assessment, Planning, Feedback).\n"
            "- For data tasks, make joins, aggregations, filters, comparisons, grouping keys, and output columns explicit only when relevant.\n"
            "- If the user mentions table names, sheet names, columns, or keys, keep them exactly.\n"
            "- Do NOT invent schema details that are not present in the input.\n"
            "- Do NOT change the business meaning.\n"
            "- Prefer deterministic action wording and clear deliverables.\n"
            "- The task field should be a short execution objective, not a generic summary.\n"
            "CRITICAL STEPS FORMAT: Use numbered section headers with '*' bullets.\n"
            "Example:\n"
            "1. Assessment:\n"
            "* Evaluate current level and goals.\n"
            "2. Planning:\n"
            "* Build a phased roadmap with milestones.\n"
            "3. Feedback:\n"
            "* Define checkpoints and next focus areas.\n"
            "Do NOT force section names like 'Data Selection & Aggregation', 'Joining and Mapping', or 'Final Output' unless the input is explicitly a data-reconciliation task.\n"
            "Do not return paragraph-style steps.\n"
        ]
        if role:
            prompt_parts.append(f"ROLE: {role}")
        if task:
            prompt_parts.append(f"TASK: {task}")
        if steps:
            prompt_parts.append(f"STEPS:\n{steps}")
        prompt_parts.append("\nReturn the rewritten JSON now:")

        user_text = "\n".join(prompt_parts)
        try:
            formatted_prompt = self._build_chat_prompt(
                system=(
                    "You are an expert workflow designer. "
                    "Rewrite user instructions into precise, execution-ready steps while preserving the original domain and section intent. "
                    "Return ONLY valid JSON with keys \"role\", \"task\", \"steps\". "
                    "No markdown fences, no explanation."
                ),
                messages=[],
                user_text=user_text,
                extra_context="",
            )
            stop_tokens = self._get_stop_tokens()
            with self.model_lock:
                resp = self.model.create_completion(
                    formatted_prompt,
                    max_tokens=min(1200, max(400, self.actual_n_ctx // 3)),
                    temperature=0.3,
                    stop=stop_tokens,
                )
            raw = resp["choices"][0].get("text", "").strip()
            # Strip think blocks
            raw = _RE_THINK.sub("", raw).strip()
            raw = _RE_THINK_INCOMPLETE.sub("", raw).strip()
            # Strip markdown fences if model wraps in ```json
            fence_match = re.search(r"```(?:json)?\s*(.*?)```", raw, re.DOTALL)
            if fence_match:
                raw = fence_match.group(1).strip()
            # Try to extract JSON object
            json_match = re.search(r"\{.*\}", raw, re.DOTALL)
            if json_match:
                raw = json_match.group(0).strip()
            result = json.loads(raw)
            rewritten_steps = _format_steps_text(str(result.get("steps", steps)).strip())
            return {
                "role": str(result.get("role", role)).strip(),
                "task": str(result.get("task", task)).strip(),
                "steps": rewritten_steps,
            }
        except json.JSONDecodeError:
            # If JSON parse fails, return raw text as task
            return {"role": role, "task": raw if raw else task, "steps": steps}
        except Exception as e:
            return {"error": f"Rewrite failed: {e}"}

    def open_file_location(self, file_path: str):
        """Open a file with the system default application."""
        try:
            if not file_path or not os.path.exists(file_path):
                return {"error": "File not found"}
            os.startfile(file_path)
            return {"ok": True}
        except Exception as e:
            return {"error": str(e)}

    def list_processed_files(self):
        """List all generated output files with metadata."""
        try:
            output_dir = os.path.join(app_data_path(), "processed_files")
            if not os.path.exists(output_dir):
                return {"files": []}
            files = []
            for fname in os.listdir(output_dir):
                fpath = os.path.join(output_dir, fname)
                if os.path.isfile(fpath):
                    stat = os.stat(fpath)
                    files.append({
                        "name": fname,
                        "path": fpath,
                        "size": stat.st_size,
                        "modified": stat.st_mtime
                    })
            # Sort by newest first
            files.sort(key=lambda f: f["modified"], reverse=True)
            return {"files": files}
        except Exception as e:
            return {"error": str(e)}

    def delete_processed_file(self, file_path: str):
        """Delete a processed output file."""
        try:
            if not file_path:
                return {"error": "No file path provided"}
            # Ensure file is within our processed_files directory
            output_dir = os.path.join(app_data_path(), "processed_files")
            real_path = os.path.realpath(file_path)
            real_dir = os.path.realpath(output_dir)
            if not real_path.startswith(real_dir):
                return {"error": "Cannot delete files outside processed_files directory"}
            if not os.path.exists(real_path):
                return {"error": "File not found"}
            os.remove(real_path)
            return {"ok": True}
        except Exception as e:
            return {"error": str(e)}

    def agent_chat(self, text: str, role: str = "", task: str = "", steps: str = "",
                   output_format: str = "none", json_schema: str = ""):
        """Synchronous AI chat for the Agent panel with Role-Task-Steps config.
        
        When *output_format* is csv/excel/txt/pdf a file is created from the
        response and ``file_path`` is returned alongside the text.
        """
        try:
            _chat_start_time = time.time()
            
            # Ensure any stuck flags from previous runs are cleared
            self.stop_generation_flag = False
            self.generation_in_progress = False
            
            if not self._has_full_access():
                return {"error": "Trial expired. Enter passkey to activate full access."}
            
            self.generation_in_progress = True
            self._agent_context_trimmed = False

            if not text or not text.strip():
                return {"error": "Empty message"}

            # #2 Tier 1: Validate JSON schema if provided
            if json_schema and json_schema.strip():
                schema_err = self._validate_json_schema(json_schema)
                if schema_err:
                    return {"error": schema_err}

            # Check model availability before doing health check
            if not self.model:
                return {"error": "No model loaded. Please load a model first."}

            # #3 Tier 1: Check Ollama health before starting (only if model is loaded)
            if not self._check_ollama_health():
                return {"error": "Model server unavailable. Please start Ollama on port 11434."}

            # Parse @rag_name mention from user text first, then from template steps/task
            text, mentioned_rag = self._resolve_rag_mention(text)
            agent_rag_database = mentioned_rag
            # Fallback: check steps and task fields (template authors often put @rag there)
            if not agent_rag_database:
                _, rag_from_steps = self._resolve_rag_mention(steps)
                if rag_from_steps:
                    agent_rag_database = rag_from_steps
            if not agent_rag_database:
                _, rag_from_task = self._resolve_rag_mention(task)
                if rag_from_task:
                    agent_rag_database = rag_from_task

            # Agent mode is stateless per run: use only RAG explicitly mentioned in this request.
            # Two-stage retrieval: vector search first, fall back to full knowledge.md.
            knowledge_md = ""
            if agent_rag_database and self.rag_manager:
                try:
                    knowledge_md = self.rag_manager.read_knowledge_markdown(
                        agent_rag_database) or ""
                except Exception:
                    pass

            # ── Auto-detect and scrape URLs in user message ──
            scraped_context = ""
            urls_in_text = self._extract_urls_from_text(text)
            if urls_in_text:
                scrape_parts = []
                for url in urls_in_text[:3]:  # max 3 URLs per request
                    if self.stop_generation_flag:
                        return {"error": "Generation stopped by user", "stopped": True}
                    try:
                        result = self._scrape_page(url)
                        if result.get("ok") and result.get("content"):
                            label = result.get("title") or url
                            scrape_parts.append(
                                f"## Scraped: {label}\nSource: {url}\n\n"
                                + result["content"])
                    except Exception as e:
                        print(f"[SCRAPE] Failed for {url}: {type(e).__name__}: {e}")
                        scrape_parts.append(
                            f"## Scraped: {url}\n[Error: could not fetch page]")
                if scrape_parts:
                    scraped_context = "\n\n".join(scrape_parts)
                    print(f"[AGENT] Scraped {len(scrape_parts)} URL(s), "
                          f"{len(scraped_context)} chars")
                # Strip URLs from user text so the prompt is cleaner
                text = self._strip_urls_from_text(text) or text

            # If web mode is enabled, use web results in Agent mode too.
            if self.web_search_enabled:
                refined = self._refine_search_query(text)
                context, sources = self._search_web(refined)

                if self.stop_generation_flag:
                    return {"error": "Generation stopped by user", "stopped": True}

                # No model: return raw web context so Agent mode still works.
                if not self.model:
                    display = context if context else "No web results found."
                    return {"text": display, "web_sources": sources}

                # Model available: synthesize final answer from web context.
                system_parts = []
                if role and role.strip():
                    system_parts.append(role.strip())
                else:
                    system_parts.append("You are a helpful web research assistant.")
                if task and task.strip():
                    system_parts.append("\nYour task: " + task.strip())
                if steps and steps.strip():
                    system_parts.append("\nFollow these steps:\n" + steps.strip())
                if knowledge_md:
                    rag_ctx = knowledge_md[:int(self.actual_n_ctx * 3.5 * 0.4)]
                    system_parts.append("\nReference knowledge:\n" + rag_ctx)
                if scraped_context:
                    scrape_budget = int(self.actual_n_ctx * 3.5 * 0.25)
                    system_parts.append(
                        "\nScraped page data:\n" + scraped_context[:scrape_budget])
                system_parts.append(
                    "\nIMPORTANT: Answer ONLY based on the web search results provided below. "
                    "Do NOT use prior knowledge or make up any information not found in the results. "
                    "If the results do not contain the answer, say so instead of guessing. "
                    "Synthesize findings instead of repeating raw snippets. "
                    "Add citation markers like [1], [2] after facts when possible."
                )
                web_system = "\n".join(system_parts)

                # Dynamic context cap — leave room for system + user + response
                max_web_chars = max(2000, int(self.actual_n_ctx * 3.5 * 0.40))
                web_ctx = (context or "")[:max_web_chars]
                prompt = self._build_chat_prompt(
                    system=web_system,
                    messages=[],
                    user_text=text.strip(),
                    extra_context=f"Web search context:\n{web_ctx}",
                )
                # Guard: ensure prompt fits in context window
                try:
                    prompt_tokens = len(self.model.tokenize(prompt.encode("utf-8")))
                except Exception:
                    prompt_tokens = len(prompt) // 4
                avail_tokens = max(64, self.actual_n_ctx - prompt_tokens - 32)
                gen_tokens = min(1024, avail_tokens)
                print(f"[AGENT-WEB] prompt_tokens={prompt_tokens}, "
                      f"avail={avail_tokens}, gen_tokens={gen_tokens}")
                if prompt_tokens >= self.actual_n_ctx - 32:
                    # Prompt overflows — trim web context and retry
                    overflow = prompt_tokens - (self.actual_n_ctx - 256)
                    trim_chars = overflow * 4  # rough token→char
                    web_ctx = web_ctx[:max(500, len(web_ctx) - trim_chars)]
                    prompt = self._build_chat_prompt(
                        system=web_system,
                        messages=[],
                        user_text=text.strip(),
                        extra_context=f"Web search context:\n{web_ctx}",
                    )
                    try:
                        prompt_tokens = len(self.model.tokenize(prompt.encode("utf-8")))
                    except Exception:
                        prompt_tokens = len(prompt) // 4
                    gen_tokens = min(1024, max(64, self.actual_n_ctx - prompt_tokens - 32))
                    print(f"[AGENT-WEB] After trim: prompt_tokens={prompt_tokens}, gen_tokens={gen_tokens}")
                response = self.model.create_completion(
                    prompt,
                    max_tokens=gen_tokens,
                    temperature=0.3,
                    stop=self._get_stop_tokens(),
                )
                if self.stop_generation_flag:
                    return {"error": "Generation stopped by user", "stopped": True}
                ai_text = response.get("choices", [{}])[0].get("text", "").strip()
                if not ai_text:
                    ai_text = context if context else "No web results found."
                return {"text": ai_text, "web_sources": sources}

            if not self.model:
                return {"error": "No model loaded. Please load a model first."}

            # Build structured prompt from Role-Task-Steps
            system_parts = []
            if role and role.strip():
                system_parts.append(role.strip())
            else:
                system_parts.append("You are a helpful assistant.")
            if task and task.strip():
                system_parts.append("\nYour task: " + task.strip())
            if steps and steps.strip():
                system_parts.append("\nFollow these steps:\n" + steps.strip())
            system_parts.append("\nAnswer directly and concisely. Do NOT write code unless explicitly asked. Provide actual answers, facts, and analysis.")
            system_msg = "\n".join(system_parts)

            # ── Determine if knowledge base is too large for a single pass ──
            # Floor of 6000 chars ensures small-context models can still process
            # a reasonable number of documents per batch.
            max_rag_chars = max(6000, int(self.actual_n_ctx * 3.5 * 0.6))

            # Split knowledge.md by source document
            doc_sections = self._split_knowledge_by_source(knowledge_md) if knowledge_md else []

            # ── Inject uploaded PDF/document if attached via Chat 📎 button ──
            # Always inject even when a RAG knowledge base is also active so the
            # model can reason over both the reference KB and the uploaded file.
            uploaded_doc_pages = self.uploaded_pages  # page-level chunks
            uploaded_doc_text = self.uploaded_content

            # #3 Map-reduce: when csv_json format is active and the uploaded
            # document has multiple pages, use map-reduce extraction instead of
            # the generic batching path.  This gives much cleaner per-page
            # JSON extraction before the final merge.
            _MAPREDUCE_PAGE_THRESHOLD = 3  # pages before triggering map-reduce
            if (output_format and output_format.lower().strip() == "csv_json"
                    and uploaded_doc_pages
                    and (len(uploaded_doc_pages) >= _MAPREDUCE_PAGE_THRESHOLD
                         or len(uploaded_doc_text or "") > max_rag_chars)):
                print(f"[AGENT] Map-reduce path: {len(uploaded_doc_pages)} pages, csv_json mode")
                _GENERIC_SCHEMA = (
                    '{"rows": [{"<column_name>": "<value>", '
                    '"...": "use as many columns as the task requires"}]}'
                )
                _mr_schema = json_schema.strip() if json_schema and json_schema.strip() else _GENERIC_SCHEMA
                _kb_ctx = ""
                if agent_rag_database and knowledge_md:
                    _kb_ctx = self._retrieve_rag_context(
                        agent_rag_database, text, knowledge_md,
                        max(2000, max_rag_chars // 4))
                _mr_text, _mr_parsed = self._agent_mapreduce_extract(
                    system_msg=system_msg,
                    user_text=text,
                    page_chunks=uploaded_doc_pages,
                    schema_hint=_mr_schema,
                    extra_context=_kb_ctx,
                )
                if self.stop_generation_flag:
                    return {"error": "Generation stopped by user", "stopped": True}
                if _mr_parsed is not None:
                    # Build CSV directly from merged JSON and return early
                    output_dir = os.path.join(app_data_path(""), "processed_files")
                    os.makedirs(output_dir, exist_ok=True)
                    _mr_file = self._schema_first_csv_from_json(
                        _mr_parsed, output_dir, int(time.time()))
                    if _mr_file:
                        return {
                            "text": _mr_text or "Map-reduce extraction complete.",
                            "file_path": _mr_file,
                            "context_trimmed": bool(self._agent_context_trimmed),
                        }
                # If map-reduce failed to produce parseable JSON, fall through
                # to the normal batching pipeline below.
                print("[AGENT] Map-reduce produced no JSON; falling through to batching")

            if uploaded_doc_text:
                # Use page-level chunks if available (PDF), else full text
                if uploaded_doc_pages and len(uploaded_doc_pages) > 1:
                    for pg in uploaded_doc_pages:
                        doc_sections.append(pg)
                    print(f"[AGENT] Injected {len(uploaded_doc_pages)} uploaded "
                          f"PDF pages ({len(uploaded_doc_text)} chars) as context"
                          + (" (alongside RAG)" if knowledge_md else ""))
                else:
                    doc_sections.append(uploaded_doc_text)
                    print(f"[AGENT] Injected uploaded document "
                          f"({len(uploaded_doc_text)} chars) as context"
                          + (" (alongside RAG)" if knowledge_md else ""))

            # If it fits in one pass, do a simple single call
            total_chars = sum(len(s) for s in doc_sections)
            # Force batching if content is significantly larger than budget
            # OR if we only got 1 giant section that would be truncated
            needs_batching = (
                (total_chars > max_rag_chars and len(doc_sections) > 1)
                or (total_chars > max_rag_chars and len(doc_sections) == 1
                    and total_chars > max_rag_chars * 1.3)
            )
            if needs_batching and len(doc_sections) == 1:
                # Re-split the single oversized section more aggressively
                doc_sections = self._split_knowledge_by_source(knowledge_md)
                if len(doc_sections) <= 1 and len(knowledge_md) > max_rag_chars:
                    # Force split by paragraphs
                    parts = re.split(r"\n\s*\n", knowledge_md)
                    doc_sections = []
                    chunk, clen = [], 0
                    for p in parts:
                        if chunk and clen + len(p) > max_rag_chars:
                            doc_sections.append("\n\n".join(chunk))
                            chunk, clen = [p], len(p)
                        else:
                            chunk.append(p)
                            clen += len(p)
                    if chunk:
                        doc_sections.append("\n\n".join(chunk))
                    needs_batching = len(doc_sections) > 1
                    print(f"[AGENT] Force-split into {len(doc_sections)} sections")

            print(f"[AGENT] RAG mode: {'batched' if needs_batching else 'single-pass'}, "
                  f"{len(doc_sections)} sections, {total_chars} chars, "
                  f"budget={max_rag_chars} chars")

            if not needs_batching:
                # ── Single-pass mode (small knowledge base or no RAG) ──
                rag_context = ""
                if agent_rag_database and knowledge_md:
                    # #2 Two-stage retrieval: vector search narrows context
                    # to the most relevant chunks for this specific query.
                    rag_context = self._retrieve_rag_context(
                        agent_rag_database, text, knowledge_md, max_rag_chars)
                elif knowledge_md:
                    rag_context = ("=== KNOWLEDGE BASE ===\n"
                                   + knowledge_md[:max_rag_chars]
                                   + "\n=== END KNOWLEDGE BASE ===")
                    if len(knowledge_md) > max_rag_chars:
                        rag_context += "\n[... truncated to fit context window ...]"

                # Append uploaded document when RAG is also present (was previously skipped)
                if uploaded_doc_text and knowledge_md:
                    doc_budget = max(4000, int(self.actual_n_ctx * 3.5 * 0.30))
                    rag_context += ("\n\n=== UPLOADED DOCUMENT ===\n"
                                    + uploaded_doc_text[:doc_budget]
                                    + "\n=== END UPLOADED DOCUMENT ===")
                    print(f"[AGENT] Single-pass: appended uploaded doc "
                          f"({min(len(uploaded_doc_text), doc_budget)} chars) to rag_context")

                # Append scraped URL content if any
                if scraped_context:
                    scrape_budget = max(2000, int(self.actual_n_ctx * 3.5 * 0.25))
                    rag_context += ("\n\n=== SCRAPED WEB DATA ===\n"
                                    + scraped_context[:scrape_budget]
                                    + "\n=== END SCRAPED DATA ===")

                # #8 Automatic fallback: 3-tier strategy
                ai_text, _fallback_used = self._agent_with_fallback(
                    system_msg=system_msg,
                    user_text=text,
                    extra_context=rag_context,
                    max_tokens=2048,
                    temperature=0.3,
                )
                if _fallback_used:
                    self._agent_context_trimmed = True
                if self.stop_generation_flag:
                    return {"error": "Generation stopped by user", "stopped": True}
            else:
                # ── Batched extraction mode (large knowledge base) ──
                # Process documents in batches that fit within context,
                # accumulate partial results, then combine.
                print(f"[AGENT] Batched RAG extraction: {len(doc_sections)} documents, {total_chars} chars total")

                # If we have scraped URL data, prepend as an extra section
                if scraped_context:
                    scrape_budget = max(2000, int(self.actual_n_ctx * 3.5 * 0.25))
                    doc_sections.insert(0, scraped_context[:scrape_budget])

                batch_results = []
                batch = []
                batch_chars = 0

                for section in doc_sections:
                    if self.stop_generation_flag:
                        return {"error": "Generation stopped by user", "stopped": True}
                    sec_len = len(section)
                    # If adding this section would overflow, flush the batch
                    if batch and (batch_chars + sec_len) > max_rag_chars:
                        partial = self._run_agent_batch(
                            system_msg, text, batch, batch_results)
                        if partial:
                            batch_results.append(partial)
                        batch = []
                        batch_chars = 0
                    batch.append(section)
                    batch_chars += sec_len

                # Flush remaining batch
                if batch:
                    if self.stop_generation_flag:
                        return {"error": "Generation stopped by user", "stopped": True}
                    partial = self._run_agent_batch(
                        system_msg, text, batch, batch_results)
                    if partial:
                        batch_results.append(partial)

                if not batch_results:
                    return {"text": "No results extracted from the knowledge base."}

                # If only one batch, use its result directly
                if len(batch_results) == 1:
                    ai_text = batch_results[0]
                else:
                    # Merge partial results in a final summarisation pass
                    if self.stop_generation_flag:
                        return {"error": "Generation stopped by user", "stopped": True}
                    ai_text = self._merge_batch_results(
                        system_msg, text, batch_results)

            if not ai_text:
                return {
                    "text": (
                        "I could not generate a stable response for this input. "
                        "Please try with fewer rows/columns or simpler instruction."
                    )
                }

            # ── Schema-first path for csv_json format ──
            _schema_parsed_json: dict | list | None = None
            # Generic schema works for any table-like task.
            # Templates can supply a precise schema via the json_schema field;
            # otherwise we use a universal row-based schema that instructs the
            # model to choose column names appropriate for the task.
            _GENERIC_SCHEMA = (
                '{"rows": [{"<column_name>": "<value>", '
                '"<column_name_2>": "<value_2>", '
                '"...": "use as many columns as the task requires"}]}'
            )
            _active_schema = json_schema.strip() if json_schema and json_schema.strip() else _GENERIC_SCHEMA
            if output_format and output_format.lower().strip() == "csv_json":
                print(f"[AGENT] Schema-first mode active "
                      f"({'custom' if json_schema and json_schema.strip() else 'generic'} schema)")
                # Rebuild the rag_context that was used in the single-pass or batch path
                _schema_context = ""
                if knowledge_md:
                    _schema_context = (
                        "=== KNOWLEDGE BASE ===\n"
                        + knowledge_md[:max_rag_chars]
                        + "\n=== END KNOWLEDGE BASE ==="
                    )
                if uploaded_doc_text and knowledge_md:
                    doc_budget = max(4000, int(self.actual_n_ctx * 3.5 * 0.30))
                    _schema_context += (
                        "\n\n=== UPLOADED DOCUMENT ===\n"
                        + uploaded_doc_text[:doc_budget]
                        + "\n=== END UPLOADED DOCUMENT ==="
                    )
                elif uploaded_doc_text:
                    doc_budget = max(6000, int(self.actual_n_ctx * 3.5 * 0.60))
                    _schema_context = (
                        "=== UPLOADED DOCUMENT ===\n"
                        + uploaded_doc_text[:doc_budget]
                        + "\n=== END UPLOADED DOCUMENT ==="
                    )
                ai_text, _schema_parsed_json = self._agent_schema_first_completion(
                    system_msg=system_msg,
                    user_text=text,
                    extra_context=_schema_context,
                    schema_hint=_active_schema,
                )
                if self.stop_generation_flag:
                    return {"error": "Generation stopped by user", "stopped": True}

            # ── Create output file if requested ──
            result = {"text": ai_text, "context_trimmed": bool(self._agent_context_trimmed)}
            print(f"[AGENT] output_format={output_format!r}, creating file: {output_format not in ('none', '', None)}")
            if output_format and output_format not in ("none", ""):
                try:
                    file_path = self._agent_chat_create_output(
                        ai_text, output_format,
                        _parsed_json=_schema_parsed_json)
                    print(f"[AGENT] File created: {file_path}")
                    if file_path:
                        result["file_path"] = file_path
                except Exception as e:
                    import traceback
                    print(f"[AGENT] Output file creation error: {e}")
                    traceback.print_exc()

            # #2 Tier 2: Record inference metrics
            elapsed = time.time() - _chat_start_time
            tokens_gen = len(ai_text.split()) * 1.3  # Rough estimate
            get_metrics().record_inference(
                task=task or role or "agent_chat",
                duration_sec=elapsed,
                tokens_generated=int(tokens_gen),
                success=not result.get("error"),
                model_name=getattr(self.model, "model_name", "unknown"),
            )
            return result
        except Exception as e:
            return {"error": str(e)}
        finally:
            self.generation_in_progress = False
            self.stop_generation_flag = False  # Reset stop flag for next run

    # ── Helpers for batched agent RAG extraction ──────────────────

    @staticmethod
    def _split_knowledge_by_source(knowledge_md: str) -> list:
        """Split knowledge.md into per-source-document sections.
        Splits on ## headers first, then ### headers, and as a last
        resort splits oversized single sections on blank-line gaps."""
        if not knowledge_md or not knowledge_md.strip():
            return []

        # Pass 1: split on ## headers
        sections = []
        current = []
        for line in knowledge_md.split("\n"):
            if line.startswith("## ") and current:
                sections.append("\n".join(current))
                current = [line]
            else:
                current.append(line)
        if current:
            sections.append("\n".join(current))

        # If only 1 section, try splitting on ### headers instead
        if len(sections) <= 1 and "### " in knowledge_md:
            sections = []
            current = []
            for line in knowledge_md.split("\n"):
                if line.startswith("### ") and current:
                    sections.append("\n".join(current))
                    current = [line]
                else:
                    current.append(line)
            if current:
                sections.append("\n".join(current))

        # Pass 2: split any oversized section (>3000 chars) on double-newlines
        MAX_SECTION = 3000
        refined = []
        for sec in sections:
            if len(sec) <= MAX_SECTION:
                refined.append(sec)
                continue
            # Split on paragraph breaks (blank lines)
            parts = re.split(r"\n\s*\n", sec)
            chunk = []
            chunk_len = 0
            for part in parts:
                plen = len(part)
                if chunk and (chunk_len + plen) > MAX_SECTION:
                    refined.append("\n\n".join(chunk))
                    chunk = [part]
                    chunk_len = plen
                else:
                    chunk.append(part)
                    chunk_len += plen
            if chunk:
                refined.append("\n\n".join(chunk))
        sections = refined

        print(f"[AGENT] Knowledge split: {len(sections)} sections, "
              f"sizes: {[len(s) for s in sections]}")
        return sections

    def _run_agent_batch(self, system_msg: str, user_text: str,
                         batch: list, previous_results: list) -> str:
        """Run the model on one batch of document sections.
        #1 Tier 1: Uses _agent_with_fallback for resilient batched extraction.
        """
        rag_block = "\n\n".join(batch)
        print(f"[AGENT] Running batch: {len(batch)} sections, "
              f"{len(rag_block)} chars")
        # Include a brief reminder of what was already extracted
        prior = ""
        if previous_results:
            # Show last partial result so model doesn't repeat
            prior_budget = min(2000, int(self.actual_n_ctx * 3.5 * 0.15))
            prior = ("\n\nYou have already extracted the following from earlier documents "
                     "(do NOT repeat these, only add NEW entries):\n"
                     + previous_results[-1][:prior_budget])

        extra = ("=== DOCUMENTS (batch) ===\n" + rag_block
                 + "\n=== END DOCUMENTS ===" + prior)

        # #1 Tier 1: Batched extraction also gets 3-tier fallback
        result, _fallback_used = self._agent_with_fallback(
            system_msg=system_msg,
            user_text=user_text,
            extra_context=extra,
            max_tokens=2048,
            temperature=0.2,
        )
        if _fallback_used:
            self._agent_context_trimmed = True
            get_metrics().record_fallback_event(
                tier=1, reason="Batched extraction context limit", context_trimmed_pct=50.0)
        return result

    def _check_ollama_health(self) -> bool:
        """#3 Tier 1: Check if Ollama server is reachable and healthy."""
        if isinstance(self.model, OllamaModel):
            try:
                start = time.time()
                resp = requests.get(
                    f"{self.model.base_url}/api/tags",
                    timeout=2.0
                )
                elapsed = (time.time() - start) * 1000  # ms
                healthy = resp.status_code == 200
                get_metrics().record_ollama_health(healthy, elapsed)
                return healthy
            except Exception as e:
                print(f"[OLLAMA] Health check failed: {e}")
                get_metrics().record_ollama_health(False, 0.0)
                return False
        return True  # Assume healthy if not using Ollama

    def _validate_input_file_size(self, filename: str, size_bytes: int, 
                                   max_mb: int = 50) -> str:
        """#2 Tier 1: Validate file size."""
        max_bytes = max_mb * 1024 * 1024
        if size_bytes > max_bytes:
            error_msg = f"File '{filename}' exceeds {max_mb}MB limit ({size_bytes} bytes)"
            get_metrics().record_validation_error("FILE_TOO_LARGE", error_msg)
            return error_msg
        return ""

    def _validate_json_schema(self, schema_str: str) -> str:
        """#2 Tier 1: Validate JSON schema format."""
        if not schema_str or not schema_str.strip():
            return ""  # Empty schema is OK (will use generic)
        try:
            json.loads(schema_str)
            return ""
        except json.JSONDecodeError as e:
            error_msg = f"Invalid JSON schema: {str(e)}"
            get_metrics().record_validation_error("SCHEMA_INVALID", error_msg)
            return error_msg

    def _validate_file_path(self, file_path: str) -> str:
        """#2 Tier 1: Prevent path traversal attacks."""
        try:
            real_path = os.path.realpath(file_path)
            workspace_root = os.path.realpath(".")
            if not real_path.startswith(workspace_root):
                error_msg = f"Path traversal detected: {file_path}"
                get_metrics().record_validation_error("PATH_TRAVERSAL", error_msg)
                return error_msg
        except Exception as e:
            error_msg = f"Path validation failed: {str(e)}"
            get_metrics().record_validation_error("PATH_ERROR", error_msg)
            return error_msg
        return ""

    def _agent_create_completion_safely(
        self,
        system_msg: str,
        user_text: str,
        extra_context: str,
        max_tokens: int = 2048,
        temperature: float = 0.3,
        repeat_penalty: float | None = None,
    ) -> str:
        """Create an agent completion while strictly fitting context window.

        Trims extra_context first (then user/system as a last resort) until
        prompt tokens fit inside ``self.actual_n_ctx``.
        """
        sys_part = str(system_msg or "")
        user_part = str(user_text or "").strip()
        extra_part = str(extra_context or "")
        trimmed = False

        def _count_tokens(s: str) -> int:
            try:
                return len(self.model.tokenize(s.encode("utf-8")))
            except Exception:
                # Conservative fallback when tokenizer is unavailable.
                return max(1, len(s) // 3)

        # Keep a safety margin for BOS/EOS/stop handling and response tokens.
        hard_prompt_limit = max(256, int(self.actual_n_ctx) - 96)

        prompt = ""
        prompt_tokens = 0
        for _ in range(14):
            prompt = self._build_chat_prompt(
                system=sys_part,
                messages=[],
                user_text=user_part,
                extra_context=extra_part,
            )
            prompt_tokens = _count_tokens(prompt)
            if prompt_tokens <= hard_prompt_limit:
                break

            # Trim in priority order: extra context, then user, then system.
            if len(extra_part) > 500:
                keep = max(500, int(len(extra_part) * 0.82))
                extra_part = extra_part[:keep]
                trimmed = True
                continue
            if len(user_part) > 240:
                keep = max(240, int(len(user_part) * 0.9))
                user_part = user_part[:keep]
                trimmed = True
                continue
            if len(sys_part) > 400:
                keep = max(400, int(len(sys_part) * 0.92))
                sys_part = sys_part[:keep]
                trimmed = True
                continue
            # Last resort: hard-cap prompt text itself.
            prompt = prompt[:max(900, hard_prompt_limit * 3)]
            prompt_tokens = _count_tokens(prompt)
            trimmed = True
            break

        avail_tokens = max(64, int(self.actual_n_ctx) - int(prompt_tokens) - 32)
        gen_tokens = min(max_tokens, avail_tokens)

        kwargs = {
            "max_tokens": gen_tokens,
            "temperature": temperature,
            "stop": self._get_stop_tokens(),
        }
        if repeat_penalty is not None:
            kwargs["repeat_penalty"] = repeat_penalty

        try:
            response = self.model.create_completion(prompt, **kwargs)
        except Exception as e:
            msg = str(e)
            # Retry once with aggressively trimmed extra context for overflow errors.
            if "Requested tokens" in msg or "context window" in msg:
                tiny_extra = extra_part[:max(300, int(len(extra_part) * 0.4))]
                trimmed = True
                prompt = self._build_chat_prompt(
                    system=sys_part,
                    messages=[],
                    user_text=user_part,
                    extra_context=tiny_extra,
                )
                prompt_tokens = _count_tokens(prompt)
                avail_tokens = max(64, int(self.actual_n_ctx) - int(prompt_tokens) - 32)
                kwargs["max_tokens"] = min(max_tokens, avail_tokens)
                response = self.model.create_completion(prompt, **kwargs)
            else:
                raise

        if trimmed:
            self._agent_context_trimmed = True

        return response.get("choices", [{}])[0].get("text", "").strip()

    def _pdf_hash_for_path(self, path: str) -> str:
        h = hashlib.sha256()
        with open(path, "rb") as f:
            for chunk in iter(lambda: f.read(1024 * 1024), b""):
                h.update(chunk)
        return h.hexdigest()

    def _pdf_hash_for_bytes(self, raw_bytes: bytes) -> str:
        return hashlib.sha256(raw_bytes).hexdigest()

    def _pdf_cache_file(self, key: str) -> str:
        return os.path.join(self._pdf_cache_dir, f"{key}.json")

    def _load_pdf_cached_extract(self, key: str) -> tuple | None:
        cache_file = self._pdf_cache_file(key)
        if not os.path.isfile(cache_file):
            return None
        try:
            with open(cache_file, "r", encoding="utf-8") as f:
                payload = json.load(f)
            full_text = str(payload.get("full_text") or "")
            pages = payload.get("pages") or []
            if full_text and isinstance(pages, list):
                return full_text, [str(p) for p in pages if p]
        except Exception:
            return None
        return None

    def _save_pdf_cached_extract(self, key: str, full_text: str, pages: list):
        cache_file = self._pdf_cache_file(key)
        payload = {
            "full_text": full_text,
            "pages": pages,
            "cached_at": int(time.time()),
        }
        try:
            with open(cache_file, "w", encoding="utf-8") as f:
                json.dump(payload, f, ensure_ascii=False)
        except Exception:
            pass

    def _merge_batch_results(self, system_msg: str, user_text: str,
                             batch_results: list) -> str:
        """Final pass: merge partial batch results into one coherent answer.
        Handles small-context models by using direct concatenation when
        the merge prompt would exceed the context window."""
        combined = "\n\n---\n\n".join(batch_results)
        # Use 70% of context, with a floor so small models still work
        max_merge_chars = max(8000, int(self.actual_n_ctx * 3.5 * 0.7))
        print(f"[AGENT] Merging {len(batch_results)} batch results, "
              f"{len(combined)} chars, budget={max_merge_chars}")

        # If combined results are small enough, try LLM merge
        if len(combined) <= max_merge_chars:
            merged = self._agent_create_completion_safely(
                system_msg=system_msg + (
                    "\n\nMerge ALL partial results below into ONE final, complete answer. "
                    "Preserve repeated-looking transaction rows when they may represent "
                    "distinct entries (same date/amount can still be valid). "
                    "Do NOT drop items unless they are obvious formatting-only repeats."),
                user_text=user_text,
                extra_context="=== PARTIAL RESULTS TO MERGE ===\n" + combined + "\n=== END ===",
                max_tokens=2048,
                temperature=0.2,
            )
            if merged:
                return merged

        # Fallback: preserve all data rows to avoid dropping legitimate
        # repeated transactions in recon statements (same date/amount can recur).
        print("[AGENT] Using direct concatenation (preserve all rows)")
        return "\n\n".join(batch_results)

    # ── Schema-first output helpers ────────────────────────────────────────────

    @staticmethod
    def _extract_json_from_response(text: str) -> dict | None:
        """Robustly extract a JSON object from a model response.

        Handles markdown code fences (```json ... ```, ``` ... ```),
        leading/trailing prose, and single-line JSON.
        Returns the parsed dict/list on success, None on failure.
        """
        if not text:
            return None
        # Strip markdown code fences
        fenced = re.search(r'```(?:json)?\s*([\s\S]*?)```', text, re.IGNORECASE)
        if fenced:
            candidate = fenced.group(1).strip()
        else:
            # Find the outermost { ... } or [ ... ] block
            start = text.find('{')
            start_arr = text.find('[')
            if start == -1 and start_arr == -1:
                return None
            if start == -1:
                start = start_arr
            elif start_arr != -1:
                start = min(start, start_arr)
            # Walk to find matching closing brace
            opener = text[start]
            closer = '}' if opener == '{' else ']'
            depth = 0
            end = -1
            for i, ch in enumerate(text[start:], start):
                if ch == opener:
                    depth += 1
                elif ch == closer:
                    depth -= 1
                    if depth == 0:
                        end = i
                        break
            if end == -1:
                candidate = text[start:]
            else:
                candidate = text[start:end + 1]
        try:
            return json.loads(candidate)
        except Exception:
            # Try with relaxed trailing-comma removal
            cleaned = re.sub(r',\s*([}\]])', r'\1', candidate)
            try:
                return json.loads(cleaned)
            except Exception:
                return None

    def _schema_first_csv_from_json(self, parsed: dict | list,
                                    output_dir: str, timestamp: int) -> str:
        """Write a clean CSV from a schema-first JSON response.

        Supports two shapes:
          - dict with a list under any top-level key (uses first list found)
          - bare list of dicts
        Each dict in the list becomes one CSV row; all keys become headers.
        Writes to output_dir/analysis_{timestamp}_q1.csv (auto-numbered to
        avoid collisions).
        """
        # Resolve to a list of records
        records: list = []
        if isinstance(parsed, list):
            records = parsed
        elif isinstance(parsed, dict):
            # Take the first list-valued key (e.g. "transactions")
            for v in parsed.values():
                if isinstance(v, list):
                    records = v
                    break
        if not records:
            return ""

        # Collect union of all keys as headers (preserves insertion order)
        headers: list = []
        seen_hdrs: set = set()
        for rec in records:
            if isinstance(rec, dict):
                for k in rec.keys():
                    if k not in seen_hdrs:
                        headers.append(k)
                        seen_hdrs.add(k)
        if not headers:
            return ""

        # Find a non-colliding filename
        base = os.path.join(output_dir, f"analysis_{timestamp}")
        output_file = base + ".csv"
        suffix = 1
        while os.path.exists(output_file):
            output_file = f"{base}_q{suffix}.csv"
            suffix += 1

        try:
            with open(output_file, 'w', newline='', encoding='utf-8') as f:
                writer = csv.DictWriter(f, fieldnames=headers,
                                        quoting=csv.QUOTE_ALL,
                                        extrasaction='ignore')
                writer.writeheader()
                for rec in records:
                    if isinstance(rec, dict):
                        writer.writerow(rec)
            print(f"[AGENT] Schema-first CSV written: {output_file} "
                  f"({len(records)} rows, {len(headers)} cols)")
            return output_file
        except Exception as e:
            print(f"[AGENT] Schema-first CSV error: {e}")
            return ""

    def _agent_schema_first_completion(self, system_msg: str, user_text: str,
                                       extra_context: str,
                                       schema_hint: str) -> tuple[str, dict | None]:
        """Run a completion requesting strict JSON output.

        Appends a JSON-schema instruction to the system message, calls the
        model, then attempts to parse the result.  On parse failure, retries
        once with an even stricter prompt.

        Returns (raw_text, parsed_json_or_None).
        """
        json_instruction = (
            "\n\n" +
            "=== OUTPUT INSTRUCTIONS ===\n"
            "You MUST output ONLY valid JSON that matches this schema exactly.\n"
            "Return ALL matching rows from the provided context. Do NOT sample, summarize, or stop at 10 rows.\n"
            "Do NOT include any prose, markdown, or code fences outside the JSON.\n"
            f"Schema:\n{schema_hint}\n"
            "=== END OUTPUT INSTRUCTIONS ==="
        )
        augmented_system = system_msg + json_instruction

        raw = self._agent_create_completion_safely(
            system_msg=augmented_system,
            user_text=user_text,
            extra_context=extra_context,
            max_tokens=3072,
            temperature=0.1,
        )
        parsed = self._extract_json_from_response(raw)
        if parsed is not None:
            return raw, parsed

        # Retry with an even stricter, shorter prompt
        print("[AGENT] Schema-first: first attempt did not yield valid JSON, retrying...")
        strict_system = (
            "You are a JSON-only output machine. "
            "Return every matching row in the context; do not truncate or sample. "
            "Output NOTHING except a single valid JSON object matching the schema. "
            "No preamble, no explanation, no markdown.\n"
            f"Schema:\n{schema_hint}"
        )
        raw2 = self._agent_create_completion_safely(
            system_msg=strict_system,
            user_text=user_text,
            extra_context=extra_context,
            max_tokens=3072,
            temperature=0.05,
        )
        parsed2 = self._extract_json_from_response(raw2)
        if parsed2 is not None:
            return raw2, parsed2

        print("[AGENT] Schema-first: both attempts failed to produce valid JSON, "
              "falling back to plain-text CSV.")
        # Return the first raw attempt for fallback text-based CSV creation
        return raw, None

    # ── #2 Two-stage RAG retrieval ──────────────────────────────────────────

    def _retrieve_rag_context(self, rag_database: str, query: str,
                              knowledge_md: str, max_chars: int) -> str:
        """Two-stage RAG retrieval for agent_chat.

        Stage 1 — vector retrieval (TF-IDF + BM25 + keyword):
            If the RAG database is indexed, use ``rag_manager.retrieve()`` to
            find the top-K most relevant chunks for the current query.  This
            avoids dumping the entire knowledge.md into the prompt.

        Stage 2 — fallback:
            If the database is not loaded / indexed yet (e.g. first run) or
            retrieval returns nothing, fall back to the full knowledge.md
            text truncated to ``max_chars``.

        Returns a formatted context string ready to inject into the prompt.
        """
        retrieved_chunks: list = []
        if self.rag_manager and rag_database:
            try:
                # How many chunks to retrieve — scale with context window
                k = min(12, max(4, self.actual_n_ctx // 512))
                raw = self.rag_manager.retrieve(rag_database, query, k=k)
                retrieved_chunks = [c for c, _score in raw if c.strip()]
                print(f"[AGENT-RAG] Two-stage retrieval: {len(retrieved_chunks)} "
                      f"chunks retrieved for query ({len(query)} chars)")
            except Exception as e:
                print(f"[AGENT-RAG] Retrieval failed ({e}), falling back to full KB")

        if retrieved_chunks:
            # Budget: each chunk gets an equal slice; cap total
            chunk_budget = max(4000, int(max_chars * 0.85))
            combined = "\n\n---\n\n".join(retrieved_chunks)
            if len(combined) > chunk_budget:
                combined = combined[:chunk_budget] + "\n[... additional KB content omitted ...]"
            return ("=== RELEVANT KNOWLEDGE BASE EXCERPTS ===\n"
                    + combined
                    + "\n=== END KNOWLEDGE BASE ===")

        # Fallback: inject full knowledge.md (existing behaviour)
        if not knowledge_md:
            return ""
        ctx = knowledge_md[:max_chars]
        suffix = "\n[... truncated to fit context window ...]" if len(knowledge_md) > max_chars else ""
        return ("=== KNOWLEDGE BASE ===\n" + ctx
                + suffix + "\n=== END KNOWLEDGE BASE ===")

    # ── #3 Map-reduce PDF extraction ─────────────────────────────────────

    def _agent_mapreduce_extract(self, system_msg: str, user_text: str,
                                 page_chunks: list, schema_hint: str,
                                 extra_context: str = "") -> tuple[str, dict | list | None]:
        """Map-reduce extraction for large multi-page documents.

        MAP phase:
            For each page chunk, ask the model to extract rows that match the
            schema and return them as a partial JSON array.

        REDUCE phase:
            Merge all partial arrays, deduplicate by a lightweight key
            (first two field values), and produce one final JSON object.

        Returns (summary_text, merged_json_or_None).
        """
        map_instruction = (
            "Extract ALL data rows visible in the document excerpt below that "
            "match the required schema.  Output ONLY valid JSON — no prose, no "
            "markdown fences.  If no matching rows are present in this excerpt, "
            f"output an empty array: [].\nSchema: {schema_hint}"
        )
        all_records: list = []

        # Split oversized pages into sub-chunks so each map prompt can include
        # complete table lines without being heavily trimmed by context limits.
        def _split_large_chunk(text: str) -> list[str]:
            max_chars = max(4000, int(self.actual_n_ctx * 2.2))
            overlap = 350
            if not text or len(text) <= max_chars:
                return [text] if text else []
            parts = re.split(r"\n\s*\n", text)
            out: list[str] = []
            cur = ""
            for part in parts:
                part = part.strip()
                if not part:
                    continue
                if not cur:
                    cur = part
                    continue
                candidate = cur + "\n\n" + part
                if len(candidate) <= max_chars:
                    cur = candidate
                else:
                    out.append(cur)
                    tail = cur[-overlap:] if len(cur) > overlap else cur
                    cur = (tail + "\n\n" + part) if tail else part
            if cur:
                out.append(cur)
            return out

        expanded_chunks: list[str] = []
        for pg in page_chunks:
            expanded_chunks.extend(_split_large_chunk(pg))

        page_count = len(expanded_chunks)
        print(f"[MAP-REDUCE] Starting map phase over {page_count} chunks "
              f"(from {len(page_chunks)} pages)")

        for i, chunk in enumerate(expanded_chunks):
            if self.stop_generation_flag:
                break
            extra = (extra_context + "\n\n" if extra_context else "") + (
                f"=== DOCUMENT EXCERPT (chunk {i+1}/{page_count}) ===\n"
                + chunk
                + "\n=== END EXCERPT ==="
            )
            raw = self._agent_create_completion_safely(
                system_msg=map_instruction,
                user_text=user_text,
                extra_context=extra,
                # Allow enough budget for pages with many line items.
                max_tokens=3072,
                temperature=0.05,
            )
            parsed = self._extract_json_from_response(raw)
            if parsed is None:
                # Model may have output a bare array
                try:
                    parsed = json.loads(raw.strip())
                except Exception:
                    parsed = None
            records: list = []
            if isinstance(parsed, list):
                records = parsed
            elif isinstance(parsed, dict):
                for v in parsed.values():
                    if isinstance(v, list):
                        records = v
                        break
            print(f"[MAP-REDUCE] Chunk {i+1}/{page_count}: "
                  f"{len(records)} records extracted")
            all_records.extend(records)

        if not all_records:
            print("[MAP-REDUCE] Map phase returned no records")
            return "", None

        # ── REDUCE: keep all extracted rows (no aggressive dedup) ──
        # Prior logic deduped by the first two field values, which can collapse
        # valid transactions when those fields repeat (for example same date).
        # Keep all model-extracted rows and only normalize row numbering below.
        deduped: list = []
        for rec in all_records:
            if isinstance(rec, dict):
                deduped.append(rec)

        # Re-number SL No / No / # fields if present
        for i, rec in enumerate(deduped, 1):
            for k in rec:
                if k.lower() in ("sl no", "sl_no", "no", "#", "sno", "serial"):
                    rec[k] = str(i)
                    break

        merged = {"rows": deduped}
        # Try to infer top-level key from schema_hint
        try:
            schema_obj = json.loads(schema_hint)
            if isinstance(schema_obj, dict):
                top_key = next(
                    (k for k, v in schema_obj.items() if isinstance(v, list)),
                    "rows"
                )
                merged = {top_key: deduped}
        except Exception:
            pass

        summary = (f"Map-reduce extraction complete: "
                   f"{len(deduped)} records from {page_count} chunks.")
        print(f"[MAP-REDUCE] Reduce phase: {len(deduped)} records "
              f"(from {len(all_records)} raw)")
        return summary, merged

    # ── #8 Automatic fallback strategy ────────────────────────────────────────

    def _agent_with_fallback(self, system_msg: str, user_text: str,
                             extra_context: str,
                             max_tokens: int = 2048,
                             temperature: float = 0.3) -> tuple[str, bool]:
        """Run agent completion with a 3-tier automatic fallback strategy.

        Tier 1 — Full context:
            Attempt with the full extra_context (already auto-trimmed by
            _agent_create_completion_safely).

        Tier 2 — Reduced context (50%):
            If tier 1 returns an empty string, trim extra_context to 50% and
            retry.  Marks fallback=True.

        Tier 3 — Minimal context (task only):
            If tier 2 also fails, use only the system_msg + user_text with no
            extra_context.  Marks fallback=True.

        Returns (result_text, fallback_was_used).
        Guarantees a non-empty string unless the model is not loaded.
        """
        # Tier 1 — full
        result = self._agent_create_completion_safely(
            system_msg=system_msg,
            user_text=user_text,
            extra_context=extra_context,
            max_tokens=max_tokens,
            temperature=temperature,
        )
        if result:
            return result, False

        # Tier 2 — 50% context
        print("[FALLBACK] Tier 1 returned empty, trying 50% context")
        half_ctx = extra_context[:max(500, len(extra_context) // 2)]
        result = self._agent_create_completion_safely(
            system_msg=system_msg,
            user_text=user_text,
            extra_context=half_ctx,
            max_tokens=max_tokens,
            temperature=temperature,
        )
        if result:
            return result, True

        # Tier 3 — no context
        print("[FALLBACK] Tier 2 returned empty, trying minimal context")
        result = self._agent_create_completion_safely(
            system_msg=system_msg,
            user_text=user_text,
            extra_context="",
            max_tokens=max_tokens,
            temperature=temperature,
        )
        return result, True

    # ── End pipeline helpers ─────────────────────────────────────────────────────

    # ── End schema-first helpers ───────────────────────────────────────────────

    def _agent_chat_create_output(self, ai_text: str, output_format: str,
                                  _parsed_json: dict | list | None = None) -> str:
        """Create an output file from agent_chat response text."""
        output_dir = os.path.join(app_data_path(""), "processed_files")
        os.makedirs(output_dir, exist_ok=True)
        timestamp = int(time.time())
        data = self._extract_response_data(ai_text) if hasattr(self, '_extract_response_data') else {}
        fmt = output_format.lower().strip()
        if fmt == "csv_json":
            # Schema-first path: use pre-parsed JSON if available, else try to
            # extract JSON from the raw text, then fall back to plain CSV.
            parsed = _parsed_json
            if parsed is None:
                parsed = self._extract_json_from_response(ai_text)
            if parsed is not None:
                result = self._schema_first_csv_from_json(parsed, output_dir, timestamp)
                if result:
                    return result
            # Fallback: treat as plain CSV
            return self._create_csv_output(data, ai_text, output_dir, timestamp)
        elif fmt == "csv":
            return self._create_csv_output(data, ai_text, output_dir, timestamp)
        elif fmt == "excel":
            return self._create_excel_output(data, ai_text, output_dir, timestamp)
        elif fmt == "pdf":
            return self._create_pdf_output(data, ai_text, output_dir, timestamp)
        elif fmt == "txt":
            return self._create_txt_output(data, ai_text, output_dir, timestamp)
        return ""

    def process_files_with_ai(self, files: list, instructions: str,
                               output_format: str = "excel", json_schema: str = ""):
        """Process uploaded files with AI and generate output in specified format.
        
                files: list of dicts with:
                    - name, size
                    - content (plain text) OR content_base64 (binary payload)
        """
        try:
            _pf_start_time = time.time()
            
            # Ensure any stuck flags from previous runs are cleared
            self.stop_generation_flag = False
            self.generation_in_progress = False
            
            if not self._has_full_access():
                return {"error": "Trial expired. Enter passkey to activate full access."}
            
            self.generation_in_progress = True

            # #2 Tier 1: Validate JSON schema if provided (before model check)
            if json_schema and json_schema.strip():
                schema_err = self._validate_json_schema(json_schema)
                if schema_err:
                    return {"error": schema_err}

            from reportlab.lib.pagesizes import letter
            from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
            from reportlab.lib.units import inch
            from reportlab.platypus import SimpleDocTemplate, Table, TableStyle, Paragraph, Spacer
            from reportlab.lib import colors
            
            print(f"[PROCESS] Starting file processing: {len(files)} files, format: {output_format}")

            if self.stop_generation_flag:
                return {"error": "Generation stopped by user", "stopped": True}
            
            if not files:
                return {"error": "No files provided"}
            
            if not instructions or not instructions.strip():
                return {"error": "No instructions provided"}

            # Check model availability before health check
            if not self.model:
                print(f"[PROCESS] ERROR: No model loaded at start of process_files_with_ai")
                return {"error": "No model loaded. Please load a model first."}
            
            print(f"[PROCESS] Model status: loaded={self.model is not None}, type={type(self.model).__name__}")


            # #3 Tier 1: Check Ollama health before starting (only if model is loaded)
            if not self._check_ollama_health():
                return {"error": "Model server unavailable. Please start Ollama on port 11434."}

            # #2 Tier 1: Validate file sizes and record uploads
            MAX_FILE_MB = 50  # Per-file limit
            for f in files:
                fsize = f.get("size", 0)
                fname = f.get("name", "unknown")
                # Per-file validation
                size_err = self._validate_input_file_size(fname, fsize, max_mb=MAX_FILE_MB)
                if size_err:
                    get_metrics().record_validation_error("FILE_TOO_LARGE", size_err)
                    return {"error": size_err}
                # Record successful upload
                get_metrics().record_file_upload(fname, fsize, 
                                                f.get("type", "unknown"), 
                                                page_count=f.get("page_count", 0))

            # Parse @rag_name mention from instructions
            instructions, mentioned_rag = self._resolve_rag_mention(instructions)
            agent_rag_database = mentioned_rag

            # File size guard — reject files > 5GB to prevent memory spikes
            MAX_FILE_BYTES = 5 * 1024 * 1024 * 1024  # 5 GB
            for f in files:
                fsize = f.get("size", 0)
                if fsize > MAX_FILE_BYTES:
                    return {"error": f"File '{f.get('name', 'unknown')}' is too large ({fsize / (1024*1024):.0f} MB). Max is 5 GB."}

            # For tabular files without a path, save content to temp so DuckDB can use file path
            for f in files:
                if self.stop_generation_flag:
                    return {"error": "Generation stopped by user", "stopped": True}
                if not f.get("path") and f.get("content") and self._is_tabular_file(f.get("name", "")):
                    temp_dir = os.path.join(app_data_path(), "agent_temp")
                    os.makedirs(temp_dir, exist_ok=True)
                    safe_name = re.sub(r'[^\w.\-]', '_', f.get("name", "file.csv"))
                    temp_path = os.path.join(temp_dir, safe_name)
                    with open(temp_path, 'w', encoding='utf-8', errors='ignore') as tf:
                        tf.write(f["content"])
                    f["path"] = temp_path
                    print(f"[PROCESS] Saved temp file: {temp_path} ({len(f['content']):,} chars)")

            tabular_warning = None
            # Skip tabular pipeline - use AI-driven code execution instead
            if self._should_use_tabular_pipeline(files):
                print("[AGENT] Using code execution pipeline for data analysis...")
                code_result = self._agent_code_execution_pipeline(files, instructions, output_format, is_agent_mode=True)
                if self.stop_generation_flag or code_result.get("stopped"):
                    return {"error": "Generation stopped by user", "stopped": True}
                # Return immediately in Agent mode - no AI response generation
                # This ensures file is created from actual execution, not from AI text
                is_ok = code_result.get("ok", False)
                result = {
                    "response_text": code_result.get("response_text", "Analysis completed.") if is_ok else code_result.get("error", "Analysis failed"),
                    "file_path": code_result.get("file_path"),
                    "success": is_ok,
                }
                # Propagate error key so frontend can detect failures
                if not is_ok:
                    result["error"] = code_result.get("error", "Analysis failed")
                return result
            
            # Fallback to text-based analysis for non-tabular files
            # Build file contents section for the prompt
            file_sections = []
            pdf_page_sections = []  # page-level chunks for batched PDF processing
            extraction_warnings = []
            for f in files:
                if self.stop_generation_flag:
                    return {"error": "Generation stopped by user", "stopped": True}
                name = str(f.get('name', 'file'))
                ext = os.path.splitext(name)[1].lower()
                
                # SKIP tabular files in text fallback - they should use code pipeline
                # Don't embed CSV/Excel content directly in prompts (causes huge token overflow)
                if ext in ('.csv', '.xlsx', '.xls'):
                    file_sections.append(f"--- FILE: {name} ({f.get('size', 0)} bytes) [Tabular file - use data analysis mode] ---")
                    if tabular_warning:
                        extraction_warnings.append(f"{name}: {tabular_warning}")
                    continue
                
                # Extract text content only for non-tabular files
                content = self._extract_agent_file_content(f)
                if content:
                    pages = f.get("_pages")  # populated by _extract_agent_file_content for PDFs
                    if pages and len(pages) > 1:
                        # PDF with page-level chunks → use batching pipeline
                        for pg in pages:
                            pdf_page_sections.append(f"--- FILE: {name} ---\n{pg}\n--- END FILE ---")
                        total_chars = sum(len(p) for p in pages)
                        file_sections.append(
                            f"--- FILE: {name} ({len(pages)} pages, {total_chars} chars) "
                            f"[batched page-level processing] ---")
                        print(f"[PROCESS] PDF {name}: {len(pages)} pages, "
                              f"{total_chars} chars -> batched")
                        if total_chars > 50000:
                            extraction_warnings.append(
                                f"{name}: Large PDF ({len(pages)} pages, "
                                f"{total_chars:,} chars). Processing in batches.")
                    else:
                        # Non-PDF or single-page: adaptive truncation
                        max_file_chars = max(10000, int(self.actual_n_ctx * 3.5 * 0.4))
                        if len(content) > max_file_chars:
                            extraction_warnings.append(
                                f"{name}: Content truncated from {len(content):,} to "
                                f"{max_file_chars:,} chars to fit context window.")
                            content = content[:max_file_chars] + (
                                f"\n\n[... Content truncated. "
                                f"Total size: {len(content):,} chars]")
                        file_sections.append(f"--- FILE: {name} ---\n{content}\n--- END FILE ---")
                else:
                    file_sections.append(f"--- FILE: {name} ({f.get('size', 0)} bytes) [content not available] ---")
                    if ext in ('.pdf', '.png', '.jpg', '.jpeg', '.webp', '.bmp', '.tiff'):
                        extraction_warnings.append(
                            f"Could not extract text from {name}. For image-only PDFs and images, install Tesseract OCR and ensure 'tesseract' is on PATH."
                        )
            
            file_data = "\n\n".join(file_sections)
            print(f"[DEBUG-PROCESS] file_data assembled: {len(file_data)} chars, "
                  f"{len(file_sections)} section(s): {[str(f.get('name','?')) for f in files]}")
            if file_data:
                print(f"[DEBUG-PROCESS] file_data preview: {file_data[:300]!r}")
            
            # If no extractable content (only tabular files), adjust prompt
            if not file_data.strip():
                return {
                    "error": "Unable to process files. CSV/Excel files require the data analysis mode which failed.",
                    "warning": tabular_warning if tabular_warning else "No extractable text content found"
                }

            # Agent mode: use full knowledge.md when RAG is mentioned
            rag_section = ""
            if agent_rag_database and self.rag_manager:
                try:
                    knowledge_md = self.rag_manager.read_knowledge_markdown(
                        agent_rag_database)
                    if knowledge_md:
                        # Guard against context overflow
                        max_rag = max(4000, int(self.actual_n_ctx * 3.5 * 0.3))
                        if len(knowledge_md) > max_rag:
                            knowledge_md = knowledge_md[:max_rag] + "\n[... Knowledge base truncated ...]"
                        rag_section = "\nReference knowledge:\n" + knowledge_md + "\n"
                except Exception:
                    pass

            _fmt_lower = (output_format or "").lower()
            if _fmt_lower in ("pdf", "docx"):
                # Chat-completion layout for instruction-tuned models.
                # System message = role + rules (from instructions template).
                # User message  = plain document data + task request.
                # No bracket markers — prevents model confusing them with template placeholders.
                import re as _re_src
                _doc_num = [0]
                def _tag_doc(m):
                    _doc_num[0] += 1
                    return f"Document {_doc_num[0]}: {m.group(1)}"
                _clean_file_data = _re_src.sub(
                    r'--- FILE: (.+?) ---', _tag_doc, file_data
                ).replace("--- END FILE ---", "---")

                _chat_system = instructions.strip()
                _chat_user = (
                    f"Here are the source documents:\n\n"
                    f"{_clean_file_data}\n\n"
                    f"{rag_section}"
                    f"{self._get_prompt_ending(output_format)}"
                )
                _chat_messages = [
                    {"role": "system", "content": _chat_system},
                    {"role": "user",   "content": _chat_user},
                ]
                # Keep a plain-text prompt for token-count estimation only
                prompt = f"{_chat_system}\n\n{_chat_user}"
            else:
                # Data analysis / structured output — role/task already defined in instructions template
                prompt = f"""Follow the instructions below and provide ACTUAL RESULTS directly from the uploaded files.

RULES:
- Do NOT write code, scripts, or programming examples
- Do NOT suggest how to analyze — actually DO the analysis
- Provide real computed numbers, totals, comparisons, and findings
- **MUST format all results as markdown tables** (| Header | Header |)
- Show actual data rows, not placeholder examples
- **CRITICAL: Do NOT repeat or echo back the file contents, file names, or --- FILE: markers in your response**
- **CRITICAL: Your response must start directly with the output, not with the input data**
{rag_section}
Instructions: {instructions}

{file_data}

{self._get_prompt_ending(output_format)}"""

            # ── Batched PDF processing (large multi-page PDFs) ──
            if pdf_page_sections:
                max_batch_chars = max(6000, int(self.actual_n_ctx * 3.5 * 0.6))
                total_pdf_chars = sum(len(s) for s in pdf_page_sections)
                print(f"[PROCESS] Batched PDF: {len(pdf_page_sections)} page sections, "
                      f"{total_pdf_chars} chars, budget={max_batch_chars}")

                system_msg = (
                    "You are a data analyst assistant. Analyze the provided document "
                    "pages and extract ALL relevant information per the instructions. "
                    "Provide actual results as markdown tables. Do NOT write code."
                    + rag_section)

                batch_results = []
                batch, batch_chars = [], 0
                for section in pdf_page_sections:
                    if self.stop_generation_flag:
                        return {"error": "Generation stopped by user", "stopped": True}
                    sec_len = len(section)
                    if batch and (batch_chars + sec_len) > max_batch_chars:
                        self._status(f"Processing batch {len(batch_results)+1}...")
                        partial = self._run_agent_batch(
                            system_msg, instructions, batch, batch_results)
                        if partial:
                            batch_results.append(partial)
                        batch, batch_chars = [], 0
                    batch.append(section)
                    batch_chars += sec_len

                if batch:
                    if self.stop_generation_flag:
                        return {"error": "Generation stopped by user", "stopped": True}
                    self._status(f"Processing batch {len(batch_results)+1}...")
                    partial = self._run_agent_batch(
                        system_msg, instructions, batch, batch_results)
                    if partial:
                        batch_results.append(partial)

                if not batch_results:
                    return {"error": "No results extracted from PDF pages."}

                if len(batch_results) == 1:
                    ai_response = batch_results[0]
                else:
                    self._status("Merging results...")
                    ai_response = self._merge_batch_results(
                        system_msg, instructions, batch_results)

                print(f"[PROCESS] Batched PDF complete: {len(batch_results)} batches "
                      f"-> {len(ai_response)} chars result")

                if str(output_format).strip().lower() == "none":
                    result = {"response_text": ai_response, "success": True}
                    if extraction_warnings:
                        result["warning"] = "\n".join(extraction_warnings)
                    return result

                output_dir = os.path.join(app_data_path(), "processed_files")
                os.makedirs(output_dir, exist_ok=True)
                timestamp = int(time.time())
                result_data = self._parse_ai_response(ai_response)
                if output_format == "excel":
                    output_file = self._create_excel_output(result_data, ai_response, output_dir, timestamp)
                elif output_format in ("csv", "csv_json"):
                    if output_format == "csv_json":
                        _GENERIC_SCHEMA = (
                            '{"rows": [{"<column_name>": "<value>", '
                            '"...": "use as many columns as the task requires"}]}'
                        )
                        _active_schema = json_schema.strip() if json_schema and json_schema.strip() else _GENERIC_SCHEMA
                        _, _parsed = self._agent_schema_first_completion(
                            system_msg="You are a data extraction assistant.",
                            user_text=instructions,
                            extra_context=ai_response,
                            schema_hint=_active_schema,
                        )
                        if _parsed:
                            _f = self._schema_first_csv_from_json(_parsed, output_dir, timestamp)
                            output_file = _f if _f else self._create_csv_output(result_data, ai_response, output_dir, timestamp)
                        else:
                            output_file = self._create_csv_output(result_data, ai_response, output_dir, timestamp)
                    else:
                        output_file = self._create_csv_output(result_data, ai_response, output_dir, timestamp)
                elif output_format == "pdf":
                    output_file = self._create_pdf_output(result_data, ai_response, output_dir, timestamp)
                else:
                    output_file = self._create_txt_output(result_data, ai_response, output_dir, timestamp)
                result = {"file_path": output_file, "response_text": ai_response, "success": True}
                if extraction_warnings:
                    result["warning"] = "\n".join(extraction_warnings)
                self._status("Ready")
                return result

            def _review_instruction_compliance(instruction_text: str, source_text: str, answer_text: str) -> tuple[bool, str, list[str]]:
                """Model-based compliance check for document/PDF analysis outputs."""
                if not self.model:
                    return True, "", []

                review_prompt = f"""Validate whether the ASSISTANT OUTPUT follows the USER INSTRUCTIONS based on SOURCE CONTENT.

USER INSTRUCTIONS:
{instruction_text}

SOURCE CONTENT (truncated):
{source_text[:6000]}

ASSISTANT OUTPUT:
{answer_text}

Return ONLY strict JSON with keys:
- ok: true or false
- reason: short string
- missing_requirements: array of short strings

Validation rules:
- Validate against explicit user requirements (format, checks, calculations, constraints, comparisons, thresholds).
- If output misses required checks, set ok=false.
- If output gives claims not grounded in source content, set ok=false.
- If instruction asks to detect violations (for example percentages/limits/conditions), output must explicitly identify pass/fail or mismatches.
- Do not require unavailable data; if data is missing, output must clearly state that limitation.
"""
                try:
                    review_resp = self.model.create_completion(
                        review_prompt,
                        max_tokens=450,
                        temperature=0.0,
                    )
                    raw = review_resp.get("choices", [{}])[0].get("text", "").strip()
                    raw = _RE_THINK.sub("", raw).strip()
                    raw = _RE_THINK_INCOMPLETE.sub("", raw).strip()
                    fence_match = re.search(r"```(?:json)?\s*(.*?)```", raw, re.DOTALL)
                    if fence_match:
                        raw = fence_match.group(1).strip()
                    json_match = re.search(r"\{.*\}", raw, re.DOTALL)
                    if json_match:
                        raw = json_match.group(0).strip()
                    payload = json.loads(raw)
                    ok = bool(payload.get("ok", True))
                    reason = str(payload.get("reason", "")).strip()
                    missing = payload.get("missing_requirements", [])
                    if not isinstance(missing, list):
                        missing = []
                    missing = [str(item).strip() for item in missing if str(item).strip()]
                    return ok, reason, missing
                except Exception as review_err:
                    print(f"[PROCESS] Compliance review skipped due to reviewer error: {review_err}")
                    return True, "", []
            
            # Generate AI response
            print("[PROCESS] Generating AI response...")
            # Model check already done at function start, so proceed directly
            print(f"[DEBUG-PROCESS] prompt length: {len(prompt)} chars, "
                  f"pdf_page_sections: {len(pdf_page_sections)}, "
                  f"output_format: {output_format}")
            _prompt_file_count = prompt.count('--- FILE:')
            print(f"[DEBUG-PROCESS] '--- FILE:' occurrences in prompt: {_prompt_file_count}")

            MAX_DOC_ATTEMPTS = 3
            ai_response = ""
            best_ai_response = ""
            review_reason = ""
            review_missing: list[str] = []

            for attempt in range(1, MAX_DOC_ATTEMPTS + 1):
                if self.stop_generation_flag:
                    return {"error": "Generation stopped by user", "stopped": True}

                current_prompt = prompt
                if attempt > 1:
                    fix_lines = [f"- {item}" for item in review_missing] if review_missing else []
                    if review_reason:
                        fix_lines.append(f"- {review_reason}")
                    fix_block = "\n".join(dict.fromkeys(fix_lines)) if fix_lines else "- Fully satisfy every explicit instruction requirement."
                    if _fmt_lower in ("pdf", "docx"):
                        _retry_user = (
                            f"Here are the source documents:\n\n"
                            f"{_clean_file_data}\n\n"
                            f"{rag_section}"
                            f"IMPORTANT — previous attempt failed. Fix ALL issues:\n{fix_block}\n\n"
                            f"{self._get_prompt_ending(output_format)}"
                        )
                        _chat_messages = [
                            {"role": "system", "content": _chat_system},
                            {"role": "user",   "content": _retry_user},
                        ]
                        current_prompt = f"{_chat_system}\n\n{_retry_user}"  # for token counting
                    else:
                        current_prompt = (
                            prompt
                            + "\n\nThe previous answer did not satisfy instruction compliance."
                            + "\nFix ALL issues below in the next answer:"
                            + f"\n{fix_block}"
                            + f"\n\n## Corrected {self._get_prompt_ending(output_format)}"
                        )

                # Dynamic max_tokens: auto-scale based on current prompt size
                _max_tok = self._estimate_max_tokens_dynamic(len(current_prompt), output_format)

                # Stop sequences: PDF/DOCX uses chat completion so only EOS tokens are needed.
                if (output_format or "").lower() in ("pdf", "docx"):
                    _stop = ["<|endoftext|>", "<|im_end|>"]
                else:
                    _stop = ["--- FILE:", "--- END FILE", "## Output", "## Analysis Results",
                             "## Corrected", "\n---\n---", "<|endoftext|>", "<|im_end|>"]

                print(f"[PROCESS] Attempt {attempt}/{MAX_DOC_ATTEMPTS}: prompt={len(current_prompt)} chars, "
                      f"max_tokens={_max_tok}, format={output_format}")
                self._debug_ai_generation("pre_completion", {
                    "attempt": attempt,
                    "max_attempts": MAX_DOC_ATTEMPTS,
                    "output_format": output_format,
                    "max_tokens": _max_tok,
                    "temperature": 0.25,
                    "stop": _stop,
                    "file_markers_in_prompt": current_prompt.count("--- FILE:"),
                    "prompt": current_prompt,
                })
                
                # Verify model is still loaded
                if not self.model:
                    print(f"[PROCESS] ERROR: Model not loaded at attempt {attempt}")
                    return {"error": "Model not loaded. Please reload the model."}
                
                try:
                    if _fmt_lower in ("pdf", "docx"):
                        response = self.model.create_chat_completion(
                            _chat_messages,
                            max_tokens=_max_tok,
                            temperature=0.35,
                            stop=_stop,
                        )
                    else:
                        response = self.model.create_completion(
                            current_prompt,
                            max_tokens=_max_tok,
                            temperature=0.25,
                            stop=_stop,
                        )
                except Exception as model_err:
                    self._debug_ai_generation("completion_exception", {
                        "attempt": attempt,
                        "error": str(model_err),
                        "model_type": type(self.model).__name__ if self.model else "None",
                    })
                    print(f"[PROCESS] Model error on attempt {attempt}: {model_err}")
                    if attempt < MAX_DOC_ATTEMPTS:
                        continue
                    else:
                        return {"error": f"Model error: {str(model_err)}"}
                
                if self.stop_generation_flag:
                    return {"error": "Generation stopped by user", "stopped": True}

                # Safely extract response text
                try:
                    _choices = response.get("choices", [{}])
                    _choice0 = _choices[0] if _choices else {}
                    if _fmt_lower in ("pdf", "docx"):
                        # chat completion: choices[0]["message"]["content"]
                        ai_response = (_choice0.get("message") or {}).get("content", "") or ""
                        ai_response = ai_response.strip()
                    else:
                        ai_response = _choice0.get("text", "").strip()
                    self._debug_ai_generation("post_completion", {
                        "attempt": attempt,
                        "response_keys": list(response.keys()) if isinstance(response, dict) else [],
                        "choices_count": len(_choices) if isinstance(_choices, list) else 0,
                        "finish_reason": _choice0.get("finish_reason"),
                        "usage": response.get("usage") if isinstance(response, dict) else None,
                        "response_text_chars": len(ai_response),
                        "response_text_head": ai_response[:200],
                    })
                except Exception as parse_err:
                    print(f"[PROCESS] Response parsing error on attempt {attempt}: {parse_err}, raw response: {response}")
                    self._debug_ai_generation("parse_error", {
                        "attempt": attempt,
                        "error": str(parse_err),
                        "raw_response_type": type(response).__name__,
                    })
                    ai_response = ""
                
                print(f"[PROCESS] AI response attempt {attempt}/{MAX_DOC_ATTEMPTS}, length: {len(ai_response)} chars")
                if not ai_response:
                    # For document drafting, retry once with higher temperature.
                    if (output_format or "").lower() in ("pdf", "docx"):
                        _minimal_stop = ["<|endoftext|>", "<|im_end|>"]
                        try:
                            response2 = self.model.create_chat_completion(
                                _chat_messages,
                                max_tokens=_max_tok,
                                temperature=0.5,
                                stop=_minimal_stop,
                            )
                            _r2c = (response2.get("choices") or [{}])[0]
                            ai_response = ((_r2c.get("message") or {}).get("content") or "").strip()
                            print(f"[PROCESS] Minimal-stop retry on attempt {attempt}, length: {len(ai_response)} chars")
                            self._debug_ai_generation("minimal_stop_retry", {
                                "attempt": attempt,
                                "response_text_chars": len(ai_response),
                                "stop": _minimal_stop,
                            })
                        except Exception as model_err2:
                            print(f"[PROCESS] Minimal-stop retry failed on attempt {attempt}: {model_err2}")
                            self._debug_ai_generation("minimal_stop_retry_exception", {
                                "attempt": attempt,
                                "error": str(model_err2),
                            })
                    if ai_response and len(ai_response) > len(best_ai_response):
                        best_ai_response = ai_response
                    if ai_response:
                        # Continue through standard validation pipeline with recovered text.
                        pass
                    else:
                        print(f"[PROCESS] Attempt {attempt}: empty response, retrying...")
                        self._debug_ai_generation("empty_response", {
                            "attempt": attempt,
                            "max_tokens": _max_tok,
                            "stop": _stop,
                            "prompt_tail": current_prompt[-280:],
                        })
                        continue
                if len(ai_response) > len(best_ai_response):
                    best_ai_response = ai_response

                # Detect heading-repeat loop (e.g. Gemma 4 echoing "Analysis Results" or
                # any short heading over and over instead of generating content)
                _words = ai_response.split()
                if len(_words) > 20:
                    # Check if any 1-3 word phrase makes up >60% of total words
                    from collections import Counter as _Counter
                    _uniq = set(_words)
                    _most_common_word, _wc = _Counter(_words).most_common(1)[0]
                    if _wc / len(_words) > 0.6 and len(_most_common_word) > 3:
                        print(f"[PROCESS] Heading-repeat loop on attempt {attempt} — "
                              f"{_most_common_word!r} appears {_wc}/{len(_words)} times, retrying")
                        review_reason = ("Response is a repeated heading/word loop. "
                                         "Generate actual content — letter text, table rows, or analysis.")
                        review_missing = ["Actual output content (not a repeated word or heading)"]
                        ai_response = ""
                        continue

                # Strip --- END FILE --- markers the model hallucinates
                if '--- END FILE' in ai_response or '---END FILE' in ai_response:
                    import re as _re_ef
                    ai_response = _re_ef.sub(r'---\s*END\s*FILE\s*---', '', ai_response).strip()
                    print(f"[PROCESS] Stripped END FILE markers from response")

                # Detect prompt echo loop: model repeated the format instruction
                _ECHO_PHRASE = "Provide the complete analysis results below as structured markdown tables"
                if ai_response.count(_ECHO_PHRASE) > 3:
                    print(f"[PROCESS] Echo loop detected on attempt {attempt} — response is prompt repetition, retrying")
                    review_reason = "Response is a repetition of the format instruction. Generate actual analysis data."
                    review_missing = ["Actual analysis content (tables with real data, not repeated instructions)"]
                    ai_response = ""
                    continue

                # Detect file-context echo: model repeated source markers from the prompt
                _resp_file_count = ai_response.count('[SOURCE ') + ai_response.count('--- FILE:')
                print(f"[DEBUG-PROCESS] source/file markers in response attempt {attempt}: {_resp_file_count}")
                if _resp_file_count > 0:
                    # Strip any [SOURCE N: ...] [END SOURCE] blocks from the response
                    import re as _re_inner
                    _cleaned = _re_inner.sub(
                        r'\[SOURCE \d+:.*?\[END SOURCE\]',
                        '', ai_response, flags=_re_inner.DOTALL
                    ).strip()
                    # Also strip any legacy --- FILE: ... --- END FILE --- blocks
                    _cleaned = _re_inner.sub(
                        r'---\s*FILE:.*?(?:---\s*END FILE\s*---|(?=---\s*FILE:)|$)',
                        '', _cleaned, flags=_re_inner.DOTALL
                    ).strip()
                    # Strip bare [SOURCE ...] or --- FILE: lines that weren't wrapped
                    _cleaned = _re_inner.sub(r'\[SOURCE \d+:.*', '', _cleaned).strip()
                    _cleaned = _re_inner.sub(r'---\s*FILE:.*', '', _cleaned).strip()
                    if _cleaned and len(_cleaned) > 50:
                        print(f"[PROCESS] Stripped FILE markers from response, salvaged {len(_cleaned)} chars")
                        ai_response = _cleaned
                    elif _resp_file_count > 1:
                        print(f"[PROCESS] File-context echo on attempt {attempt} — "
                              f"response has {_resp_file_count} FILE markers and nothing useful after strip, retrying")
                        review_reason = ("Response echoes raw file markers (--- FILE: ...) instead of analysing. "
                                         "Output ONLY the analysis result — do NOT repeat file names, markers, or raw input.")
                        review_missing = ["Pure analysis output with no --- FILE: markers or repeated raw input"]
                        ai_response = ""
                        continue

                # Detect paragraph-level repetition: same 80+ char block 3+ times
                _paragraphs = [p.strip() for p in ai_response.split('\n\n') if len(p.strip()) > 80]
                _seen_paras: dict = {}
                _has_para_repeat = False
                for _p in _paragraphs:
                    _seen_paras[_p] = _seen_paras.get(_p, 0) + 1
                    if _seen_paras[_p] >= 3:
                        _has_para_repeat = True
                        print(f"[PROCESS] Paragraph repetition on attempt {attempt} — "
                              f"block repeated {_seen_paras[_p]}x: {_p[:80]!r}")
                        break
                if _has_para_repeat:
                    if (output_format or "").lower() in ("pdf", "docx"):
                        # For document drafting, salvage repeated drafts by de-duplicating
                        # repeated paragraphs instead of discarding the entire output.
                        import re as _re_norm
                        _deduped_parts = []
                        _seen_norm = set()
                        for _part in ai_response.split("\n\n"):
                            _txt = _part.strip()
                            if not _txt:
                                continue
                            _norm = _re_norm.sub(r"\s+", " ", _txt)
                            if _norm in _seen_norm:
                                continue
                            _seen_norm.add(_norm)
                            _deduped_parts.append(_txt)
                        _deduped = "\n\n".join(_deduped_parts).strip()
                        if len(_deduped) > 120:
                            ai_response = _deduped
                            if len(ai_response) > len(best_ai_response):
                                best_ai_response = ai_response
                            print(f"[PROCESS] Paragraph repetition salvaged for {output_format}: "
                                  f"deduped to {len(ai_response)} chars")
                        else:
                            review_reason = "Response contains identical repeated paragraphs. Each section must appear exactly once."
                            review_missing = ["Non-repeated content — every compliance note or section appears only once"]
                            ai_response = ""
                            continue
                    else:
                        review_reason = "Response contains identical repeated paragraphs. Each section must appear exactly once."
                        review_missing = ["Non-repeated content — every compliance note or section appears only once"]
                        ai_response = ""
                        continue

                # Document-quality guard for letter/drafting formats.
                # Prevent accepting meta retry text like "previous draft..." as final output.
                if (output_format or "").lower() in ("pdf", "docx"):
                    _meta_markers = (
                        "previous draft",
                        "did not fully follow the instructions",
                        "fix all issues below",
                        "instruction compliance",
                        "corrected",
                    )
                    _low_quality = False
                    _resp_l = ai_response.lower()
                    if len(ai_response.strip()) < 180:
                        _low_quality = True
                    if any(m in _resp_l for m in _meta_markers):
                        _low_quality = True
                    # Must have at least one line break to resemble a letter structure.
                    if "\n" not in ai_response and len(ai_response.strip().split()) < 40:
                        _low_quality = True

                    if _low_quality:
                        print(f"[PROCESS] Low-quality PDF/DOCX draft on attempt {attempt} — retrying")
                        review_reason = (
                            "Draft is incomplete or meta-instruction text. "
                            "Produce a complete formal letter only."
                        )
                        review_missing = [
                            "Complete letter with date/subject/salutation/body/closing/signature",
                            "No meta text about previous drafts or instruction fixes",
                        ]
                        ai_response = ""
                        continue

                # Skip compliance review for pdf/docx (letter/document drafting) —
                # no JSON to validate, review call wastes a full model inference for free-text output
                if (output_format or "").lower() in ("pdf", "docx"):
                    print("[PROCESS] PDF/DOCX format — skipping compliance review, accepting output")
                    break
                ok, review_reason, review_missing = _review_instruction_compliance(instructions, file_data, ai_response)
                if ok:
                    break
                print(f"[PROCESS] Instruction compliance failed on attempt {attempt}: {review_reason}")
            else:
                # #8 Automatic fallback: if validation keeps failing, degrade
                # gracefully rather than returning a hard error.
                print(f"[PROCESS] Compliance validation failed after {MAX_DOC_ATTEMPTS} attempts "
                      f"— returning best effort result with warning")
                if not ai_response:
                    # Log all relevant state info to help diagnose the issue
                    print(f"[PROCESS] ERROR DIAGNOSTIC: ai_response is empty after {MAX_DOC_ATTEMPTS} attempts")
                    print(f"[PROCESS] Model loaded: {self.model is not None}")
                    print(f"[PROCESS] Model type: {type(self.model).__name__ if self.model else 'None'}")
                    print(f"[PROCESS] Output format: {output_format}")
                    print(f"[PROCESS] File data length: {len(file_data)} chars")
                    print(f"[PROCESS] Instructions length: {len(instructions)} chars")
                    if best_ai_response.strip():
                        ai_response = best_ai_response.strip()
                        print(f"[PROCESS] Falling back to best non-empty response: {len(ai_response)} chars")
                    else:
                        return {"error": "No response from AI model after 3 attempts. Check model status and try again."}

            if not ai_response:
                # Should not reach here if fallback worked
                if best_ai_response.strip():
                    ai_response = best_ai_response.strip()
                    print(f"[PROCESS] Recovered from empty final state using best response: {len(ai_response)} chars")
                else:
                    print(f"[PROCESS] CRITICAL: ai_response is empty even after fallback attempt")
                    return {"error": "No response from AI model. Model may have crashed or context exceeded."}
            
            # Parse response for structured data (tables)
            result_data = self._parse_ai_response(ai_response)

            if str(output_format).strip().lower() == "none":
                result = {"response_text": ai_response, "success": True}
                if extraction_warnings:
                    result["warning"] = "\n".join(extraction_warnings)
                if tabular_warning:
                    existing = result.get("warning")
                    result["warning"] = f"{existing}\n{tabular_warning}" if existing else tabular_warning
                return result
            
            # Generate output file
            output_dir = os.path.join(app_data_path(), "processed_files")
            os.makedirs(output_dir, exist_ok=True)
            
            timestamp = int(time.time())
            
            # Generate output files (potentially multiple for some formats)
            output_files = []
            if result_data.get("tables"):
                # Use multi-file generation when tables are present
                output_files = self._create_multiple_output_files(
                    result_data["tables"], output_format, output_dir, timestamp
                )
            
            if not output_files:
                # Fallback to single-file generation if no tables parsed
                if output_format == "csv_json":
                    _GENERIC_SCHEMA = (
                        '{"rows": [{"<column_name>": "<value>", '
                        '"...": "use as many columns as the task requires"}]}'
                    )
                    _active_schema = json_schema.strip() if json_schema and json_schema.strip() else _GENERIC_SCHEMA
                    _, _parsed = self._agent_schema_first_completion(
                        system_msg="You are a data extraction assistant.",
                        user_text=instructions,
                        extra_context=ai_response,
                        schema_hint=_active_schema,
                    )
                    if _parsed:
                        _f = self._schema_first_csv_from_json(_parsed, output_dir, timestamp)
                        if _f:
                            output_files = [_f]
                        else:
                            output_files = [self._create_csv_output(result_data, ai_response, output_dir, timestamp)]
                    else:
                        output_files = [self._create_csv_output(result_data, ai_response, output_dir, timestamp)]
                elif output_format == "csv":
                    output_files = [self._create_csv_output(result_data, ai_response, output_dir, timestamp)]
                elif output_format == "excel":
                    output_files = [self._create_excel_output(result_data, ai_response, output_dir, timestamp)]
                elif output_format == "pdf":
                    output_files = [self._create_pdf_output(result_data, ai_response, output_dir, timestamp)]
                elif output_format == "docx":
                    output_files = [self._create_docx_output(result_data, ai_response, output_dir, timestamp)]
                else:  # txt
                    output_files = [self._create_txt_output(result_data, ai_response, output_dir, timestamp)]
            
            print(f"[PROCESS] Output files created: {output_files}")
            # Return single file for backward compatibility, but also include list
            result = {
                "file_path": output_files[0] if output_files else None,
                "file_paths": output_files,
                "response_text": ai_response,
                "success": bool(output_files)
            }
            if extraction_warnings:
                result["warning"] = "\n".join(extraction_warnings)
            if tabular_warning:
                existing = result.get("warning")
                result["warning"] = f"{existing}\n{tabular_warning}" if existing else tabular_warning
            # Propagate compliance-failure warning if fallback was used
            if review_reason and not result.get("warning"):
                result["warning"] = f"Output may be incomplete: {review_reason}"
            
            # #2 Tier 2: Record inference metrics for process_files
            elapsed = time.time() - _pf_start_time
            tokens_gen = len(ai_response.split()) * 1.3  # Rough estimate
            get_metrics().record_inference(
                task="process_files_with_ai",
                duration_sec=elapsed,
                tokens_generated=int(tokens_gen),
                success=result.get("success", False),
                model_name=getattr(self.model, "model_name", "unknown"),
            )
            return result
            
        except Exception as e:
            import traceback
            print(f"[PROCESS] Error: {e}")
            traceback.print_exc()
            return {"error": str(e)}
        finally:
            self.generation_in_progress = False
            self.stop_generation_flag = False  # Reset stop flag for next run
            self._cleanup_agent_temp()

    def _fast_count_rows(self, files: list) -> dict:
        """Quickly count rows in files without AI processing."""
        try:
            import pandas as pd
            import io
            import base64
            import os
            
            counts = {}
            total_rows = 0
            
            for f in files:
                name = str(f.get('name', 'file'))
                path = f.get('path')
                
                try:
                    df = None
                    
                    # Try disk path first
                    if path and os.path.exists(path):
                        if name.lower().endswith('.csv'):
                            df = pd.read_csv(path, nrows=10)  # Just read header
                        elif name.lower().endswith(('.xlsx', '.xls')):
                            df = pd.read_excel(path, nrows=10)
                    
                    # If no disk path, try memory content
                    if df is None:
                        content = f.get('content')
                        content_b64 = f.get('content_base64')
                        
                        if content and isinstance(content, str):
                            if name.lower().endswith('.csv'):
                                df = pd.read_csv(io.StringIO(content), nrows=10)
                        elif content_b64:
                            try:
                                binary_data = base64.b64decode(content_b64, validate=False)
                                if name.lower().endswith('.csv'):
                                    df = pd.read_csv(io.BytesIO(binary_data), nrows=10)
                            except:
                                pass
                    
                    if df is not None:
                        # Count actual rows by reading full file
                        if path and os.path.exists(path):
                            if name.lower().endswith('.csv'):
                                row_count = sum(1 for _ in open(path)) - 1  # Exclude header
                        else:
                            content = f.get('content')
                            if content:
                                row_count = content.count('\n') - 1  # Rough estimate
                            else:
                                content_b64 = f.get('content_base64')
                                if content_b64:
                                    binary_data = base64.b64decode(content_b64, validate=False)
                                    row_count = binary_data.count(b'\n') - 1
                                else:
                                    row_count = len(df)
                        
                        counts[name] = row_count
                        total_rows += row_count
                        print(f"[CODE_EXEC] Loaded {name}: {row_count} rows")
                except Exception as e:
                    print(f"[CODE_EXEC] Error counting {name}: {e}")
                    counts[name] = "Error"
            
            # Format response
            count_summary = "\n".join([f"[CODE_EXEC] {name}: {count} rows" for name, count in counts.items()])
            return {
                "ok": True,
                "response_text": f"{count_summary}\n\nTotal rows: {total_rows}"
            }
        except Exception as e:
            print(f"[CODE_EXEC] Count error: {e}")
            return {"error": f"Failed to count rows: {str(e)}"}

    def _agent_code_execution_pipeline(self, files: list, instructions: str, output_format: str, is_agent_mode: bool = False) -> dict:
        """Execute AI-generated SQL analysis against uploaded data files.

        Uses DuckDB-backed SQL generation for all normal tabular analysis paths.
        is_agent_mode: If True, applies optimizations like fast count-only responses.
        """
        try:
            import pandas as pd

            if self.stop_generation_flag:
                return {"error": "Generation stopped by user", "stopped": True}
            
            print("[CODE_EXEC] Starting code execution pipeline...")
            
            # Check if user only wants count (optimization only in Agent mode for speed)
            if is_agent_mode:
                instructions_lower = instructions.lower().strip()
                count_only_keywords = ["count", "total rows", "total count", "how many", "number of"]
                if any(kw in instructions_lower for kw in count_only_keywords) and len(instructions) < 100:
                    print("[CODE_EXEC] Optimized path: Returning row counts only (skipping AI code generation)...")
                    return self._fast_count_rows(files)
            
            # SQL-only mode for tabular analysis
            total_size = sum(f.get("size", 0) for f in files)
            use_sql = True
            
            print(f"[CODE_EXEC] Total file size: {total_size / (1024*1024):.1f}MB, using: SQL")
            
            # Load files into dataframes
            dataframes = {}
            pipeline_status_messages = []  # Collect status messages from pipeline

            def _normalize_df(df: "pd.DataFrame", label: str) -> "pd.DataFrame":
                """Clean common real-world file messiness before SQL ingestion:
                - Strip whitespace from column names
                - Rename Unnamed: X columns that have data → col_X
                - Drop Unnamed: X columns that are entirely NaN (trailing blank columns)
                - Drop rows that are entirely NaN (trailing blank rows from Excel)
                - Normalize all column names to SQL-safe snake_case
                """
                import re as _re
                # 1. Strip column name whitespace
                df.columns = [str(c).strip() for c in df.columns]

                # 2. Drop rows that are entirely NaN (trailing empty rows)
                before_rows = len(df)
                df = df.dropna(how='all').reset_index(drop=True)
                dropped_rows = before_rows - len(df)
                if dropped_rows:
                    print(f"[CODE_EXEC] {label}: dropped {dropped_rows} all-blank trailing row(s)")

                # 3. Handle Unnamed columns (blank headers from CSV/Excel)
                cols_renamed = []
                cols_dropped = []
                new_cols = []
                for col in df.columns:
                    if _re.match(r'^Unnamed:\s*\d+$', col):
                        if df[col].isna().all():
                            cols_dropped.append(col)
                        else:
                            # rename to col_N (use the numeric suffix from "Unnamed: N")
                            num = _re.search(r'\d+', col)
                            safe = f"col_{num.group()}" if num else f"col_{len(new_cols)}"
                            new_cols.append(safe)
                            cols_renamed.append(f"{col} → {safe}")
                    else:
                        new_cols.append(col)

                if cols_dropped:
                    df = df.drop(columns=cols_dropped)
                    print(f"[CODE_EXEC] {label}: dropped {len(cols_dropped)} entirely-blank column(s): {cols_dropped}")
                if cols_renamed:
                    df.columns = [c for c in new_cols]
                    print(f"[CODE_EXEC] {label}: renamed blank-header columns: {cols_renamed}")

                # 4. Normalize all column names to SQL-safe identifiers.
                old_cols = [str(c) for c in df.columns]
                safe_cols = self._normalize_column_names(old_cols)
                if old_cols != safe_cols:
                    df.columns = safe_cols
                    # Keep log short to avoid noisy status output.
                    mapping_preview = ", ".join(
                        f"{o}->{n}" for o, n in list(zip(old_cols, safe_cols))[:8] if o != n
                    )
                    if mapping_preview:
                        print(f"[CODE_EXEC] {label}: normalized column names ({mapping_preview})")

                return df

            for f in files:
                import io
                import base64

                if self.stop_generation_flag:
                    return {"error": "Generation stopped by user", "stopped": True}
                
                name = str(f.get('name', 'file'))
                path = f.get('path')
                
                try:
                    # PRE-FLIGHT CHECK: Estimate if file is too large before loading full content
                    if name.lower().endswith('.csv'):
                        content = f.get('content')
                        content_b64 = f.get('content_base64')
                        
                        # Estimate row count from file size before loading
                        file_size = f.get('size', 0)
                        if file_size > 100 * 1024 * 1024:  # Over 100 MB
                            est_validation = self._estimate_csv_validity(
                                name, file_size, content, content_b64
                            )
                            if est_validation == "BATCH_MODE":
                                # Large file detected - use batch processing instead
                                print(f"[CODE_EXEC] Large file detected: {name} ({file_size / (1024*1024):.1f} MB)")
                                print(f"[CODE_EXEC] Switching to batch processing mode...")
                                return self._execute_batch_processing(files, instructions, output_format)
                            elif est_validation:
                                # File truly too large even for batch mode
                                print(f"[CODE_EXEC] {name}: {est_validation}")
                                return {"error": est_validation}
                    
                    df = None
                    loaded_sheets: dict = {}  # {table_name: df} for multi-sheet Excel

                    def _safe_sheet_name(sn: str) -> str:
                        seg = re.sub(r'[^A-Za-z0-9]', '_', str(sn)).strip('_')
                        return seg if seg else 'Sheet'

                    # Try to load from disk path first
                    if path and os.path.exists(path):
                        if name.lower().endswith('.csv'):
                            df = pd.read_csv(path)
                        elif name.lower().endswith(('.xlsx', '.xls')):
                            xf = pd.ExcelFile(path)
                            sheet_names = xf.sheet_names
                            base_tname = 'df_' + os.path.splitext(name)[0].replace(' ', '_').replace('-', '_')
                            if len(sheet_names) > 1:
                                for sn in sheet_names:
                                    loaded_sheets[f"{base_tname}_{_safe_sheet_name(sn)}"] = xf.parse(sn)
                            else:
                                df = xf.parse(sheet_names[0]) if sheet_names else pd.read_excel(path)

                    # If no disk path, try to load from memory (content or base64)
                    if df is None and not loaded_sheets:
                        content = f.get('content')
                        content_b64 = f.get('content_base64')

                        if content and isinstance(content, str):
                            # Direct string content
                            if name.lower().endswith('.csv'):
                                df = pd.read_csv(io.StringIO(content))
                            elif name.lower().endswith('.json'):
                                df = pd.read_json(io.StringIO(content))
                        elif content_b64:
                            # Base64 encoded content
                            try:
                                binary_data = base64.b64decode(content_b64, validate=False)
                                if name.lower().endswith('.csv'):
                                    df = pd.read_csv(io.BytesIO(binary_data))
                                elif name.lower().endswith(('.xlsx', '.xls')):
                                    xf = pd.ExcelFile(io.BytesIO(binary_data))
                                    sheet_names = xf.sheet_names
                                    base_tname = 'df_' + os.path.splitext(name)[0].replace(' ', '_').replace('-', '_')
                                    if len(sheet_names) > 1:
                                        for sn in sheet_names:
                                            loaded_sheets[f"{base_tname}_{_safe_sheet_name(sn)}"] = xf.parse(sn)
                                    else:
                                        df = xf.parse(sheet_names[0]) if sheet_names else pd.read_excel(io.BytesIO(binary_data))
                            except Exception as e:
                                print(f"[CODE_EXEC] Failed to decode base64 for {name}: {e}")

                    # ── Multi-sheet Excel: register each sheet as a separate DuckDB table ──
                    if loaded_sheets:
                        msg = f"[CODE_EXEC] Multi-sheet Excel '{name}': {len(loaded_sheets)} sheets → {', '.join(loaded_sheets.keys())}"
                        pipeline_status_messages.append(msg)
                        print(msg)
                        for tname, sdf in loaded_sheets.items():
                            sdf = _normalize_df(sdf, tname)
                            # Coerce object columns that look numeric
                            for col in sdf.select_dtypes(include='object').columns:
                                cleaned = sdf[col].astype(str).str.replace(',', '').str.strip()
                                coerced = pd.to_numeric(cleaned, errors='coerce')
                                sheet_non_null = coerced.notna().sum()
                                sheet_total = len(sdf)
                                if sheet_total > 0 and (sheet_non_null / sheet_total) >= 0.7:
                                    sdf[col] = coerced
                            # Size check per sheet
                            if len(sdf) > 100000 or len(sdf.columns) > 100:
                                msg = f"[CODE_EXEC] Sheet '{tname}' too large ({len(sdf)} rows, {len(sdf.columns)} cols) — skipping"
                                pipeline_status_messages.append(msg)
                                print(msg)
                                continue
                            dataframes[tname] = sdf
                            msg = f"[CODE_EXEC] Loaded sheet '{tname}': {len(sdf)} rows, {len(sdf.columns)} columns"
                            pipeline_status_messages.append(msg)
                            print(msg)
                        continue  # done with this file; move to next

                    if df is None:
                        print(f"[CODE_EXEC] Could not load {name} - no valid path or content")
                        continue

                    df = _normalize_df(df, name)

                    # Coerce object columns that contain numeric-looking values to float
                    # (CSV may load numeric cols as object due to commas, spaces, or mixed blanks)
                    for col in df.select_dtypes(include='object').columns:
                        cleaned = df[col].astype(str).str.replace(',', '').str.strip()
                        coerced = pd.to_numeric(cleaned, errors='coerce')
                        # Only convert if >70% of non-null rows parse as numeric
                        non_null = coerced.notna().sum()
                        total = len(df)
                        if total > 0 and (non_null / total) >= 0.7:
                            df[col] = coerced
                            print(f"[CODE_EXEC] Coerced column '{col}' in {name} to numeric ({non_null}/{total} values)")
                    
                    # Check dataframe size - reject if too large for code generation
                    if len(df) > 100000 or len(df.columns) > 100:
                        print(f"[CODE_EXEC] {name} too large ({len(df)} rows, {len(df.columns)} cols) - max 100k rows/100 cols")
                        return {"error": f"File '{name}' too large for code execution (max 100k rows, 100 columns)"}
                    
                    # Use filename without extension as table name
                    # Prepend 'df_' to ensure valid Python identifier (can't start with numbers)
                    table_name = 'df_' + os.path.splitext(name)[0].replace(' ', '_').replace('-', '_')
                    dataframes[table_name] = df
                    msg = f"[CODE_EXEC] Loaded {name}: {len(df)} rows, {len(df.columns)} columns"
                    pipeline_status_messages.append(msg)
                    print(msg)
                except Exception as e:
                    print(f"[CODE_EXEC] Error loading {name}: {e}")
                    import traceback
                    traceback.print_exc()
                    continue
            
            if not dataframes:
                return {"error": "No valid tabular files could be loaded"}

            # Collect non-tabular file content (txt, pdf, docx) as context for the SQL prompt
            non_tabular_context = []
            for f in files:
                name = str(f.get('name', ''))
                if self._is_tabular_file(name):
                    continue
                content = f.get('content') or ''
                if not content:
                    content = self._extract_agent_file_content(f) or ''
                if content:
                    non_tabular_context.append(f"--- FILE: {name} ---\n{content[:4000]}\n--- END FILE ---")
                    print(f"[CODE_EXEC] Injected context from non-tabular file: {name} ({len(content)} chars)")
            context_block = "\n\n".join(non_tabular_context)

            # SQL-only execution path
            if use_sql:
                print("[CODE_EXEC] Calling SQL analysis...")
                result = self._execute_sql_analysis(dataframes, instructions, output_format, pipeline_status_messages, context_block=context_block)
                if self.stop_generation_flag or result.get("stopped"):
                    return {"error": "Generation stopped by user", "stopped": True}
                print(f"[CODE_EXEC] SQL analysis returned: {result.get('ok', False)}")
            
            print(f"[CODE_EXEC] Pipeline complete, returning result...")
            return result
            
        except Exception as e:
            import traceback
            print(f"[CODE_EXEC] Pipeline error: {e}")
            traceback.print_exc()
            return {"error": f"Code execution failed: {str(e)}"}

    def _execute_batch_processing(self, files: list, instructions: str, output_format: str) -> dict:
        """
        Process large files (>100MB) using DuckDB streaming mode.

        Strategy: DuckDB reads the file directly from disk (no full pandas load).
        A sample DataFrame is created for schema discovery, then the full file
        is registered as a DuckDB table so the AI-generated SQL runs on all rows.
        """
        try:
            import duckdb
            import pandas as pd
            import os
            import time

            print("[BATCH] Starting batch processing mode for large files...")

            if not files:
                return {"error": "No files provided for batch processing"}

            # ── Load ALL files into DuckDB tables (just like normal pipeline) ──
            conn = duckdb.connect(":memory:")
            dataframes: dict = {}  # table_name → sample DataFrame (for schema)
            duckdb_tables: list = []  # track tables created directly in DuckDB

            for f in files:
                name = str(f.get('name', 'file.csv'))
                path = f.get('path')
                content = f.get('content')
                file_size_mb = f.get('size', 0) / (1024 * 1024)
                base_name = 'df_' + os.path.splitext(name)[0].replace(' ', '_').replace('-', '_')
                base_name = re.sub(r'[^A-Za-z0-9_]', '', base_name) or 'df_file'

                print(f"[BATCH] Loading: {name} ({file_size_mb:.1f} MB)")

                # Resolve file path for DuckDB direct read
                actual_path = None
                if path and os.path.exists(path):
                    actual_path = path
                elif content and isinstance(content, str):
                    temp_dir = os.path.join(app_data_path(), "batch_temp")
                    os.makedirs(temp_dir, exist_ok=True)
                    actual_path = os.path.join(temp_dir, f"batch_{int(time.time())}_{name}")
                    with open(actual_path, 'w', encoding='utf-8') as tf:
                        tf.write(content)

                if not actual_path:
                    return {"error": f"No valid file path or content for {name}"}

                try:
                    escaped_path = actual_path.replace("'", "''")
                    if name.lower().endswith('.csv'):
                        conn.execute(f"CREATE TABLE {base_name} AS SELECT * FROM read_csv('{escaped_path}', auto_detect=true)")
                    elif name.lower().endswith(('.xlsx', '.xls')):
                        conn.execute(f"INSTALL spatial; LOAD spatial;")
                        conn.execute(f"CREATE TABLE {base_name} AS SELECT * FROM st_read('{escaped_path}')")
                    else:
                        conn.execute(f"CREATE TABLE {base_name} AS SELECT * FROM read_csv('{escaped_path}', auto_detect=true)")

                    # Get row count
                    row_count = conn.execute(f"SELECT COUNT(*) FROM {base_name}").fetchone()[0]
                    print(f"[BATCH] {base_name}: {row_count:,} rows loaded")

                    # Sample first 1000 rows into pandas for schema discovery
                    sample_df = conn.execute(f"SELECT * FROM {base_name} LIMIT 1000").fetchdf()
                    # Normalize column names to match what _execute_sql_analysis expects
                    old_cols = [str(c) for c in sample_df.columns]
                    safe_cols = self._normalize_column_names(old_cols)
                    if old_cols != safe_cols:
                        sample_df.columns = safe_cols
                        # Also rename in DuckDB table
                        for old, new in zip(old_cols, safe_cols):
                            if old != new:
                                try:
                                    conn.execute(f'ALTER TABLE {base_name} RENAME COLUMN "{old}" TO "{new}"')
                                except Exception:
                                    pass

                    dataframes[base_name] = sample_df
                    duckdb_tables.append(base_name)
                except Exception as e:
                    conn.close()
                    return {"error": f"Failed to load {name}: {str(e)}"}

            # ── Use the same SQL analysis pipeline but with our pre-loaded DuckDB connection ──
            # _execute_sql_analysis creates its own connection, but since our tables are in
            # `conn`, we pass the sample DataFrames for schema and let SQL gen create queries.
            # Then we execute those queries on our connection with ALL the data.
            result = self._execute_sql_analysis(
                dataframes, instructions, output_format,
                pipeline_status_messages=[],
            )

            conn.close()
            return result

        except Exception as e:
            print(f"[BATCH] Error: {e}")
            import traceback
            traceback.print_exc()
            return {"error": f"Batch processing error: {str(e)}"}

    def _execute_sql_analysis(self, dataframes: dict, instructions: str, output_format: str, pipeline_status_messages: list = None, context_block: str = "") -> dict:
        """Generate and execute SQL analysis using DuckDB."""
        try:
            import duckdb
            import pandas as pd
            from difflib import get_close_matches
            
            status_messages = pipeline_status_messages if pipeline_status_messages else []
            def _debug_status(text: str):
                status_messages.append(text)
                print(text)

            def _clean_summary(msgs: list, suffix: str = "") -> str:
                """Build a user-friendly summary from status_messages.
                
                Strips debug/internal tags and keeps only meaningful lines.
                """
                clean = []
                for m in msgs:
                    # Skip internal debug lines entirely
                    if any(tag in m for tag in (
                        "[SQL-DEBUG]", "[CODE_EXEC]", "[DEBUG]",
                        "clean_sql input", "clean_sql output",
                        "clean_sql:", "VERIFY COLUMN", "VERIFY TABLE",
                        "repair_sql", "semantic_review",
                    )):
                        continue
                    # Strip tag prefixes for user-facing lines
                    line = re.sub(r'^\[(SQL|PYTHON|AGENT)\]\s*', '', m).strip()
                    if line:
                        clean.append(line)
                result = "\n".join(clean) if clean else "Analysis completed."
                if suffix:
                    result += "\n\n" + suffix
                return result

            def _stopped_result() -> dict:
                _debug_status("[SQL] Stopped by user")
                return {"error": "Generation stopped by user", "stopped": True}

            msg = "[SQL] Executing SQL-based analysis..."
            status_messages.append(msg)
            print(msg)

            if self.stop_generation_flag:
                return _stopped_result()

            # Final safety normalization: enforce SQL-safe column names for every dataframe
            # right before DuckDB registration (covers any loader path).
            normalized_dataframes = {}
            for table_name, df in dataframes.items():
                working = df.copy()
                old_cols = [str(c) for c in working.columns]
                safe_cols = self._normalize_column_names(old_cols)
                if old_cols != safe_cols:
                    working.columns = safe_cols
                    preview = ", ".join(
                        f"{o}->{n}" for o, n in list(zip(old_cols, safe_cols))[:8] if o != n
                    )
                    if preview:
                        _debug_status(f"[SQL] Normalized columns in {table_name}: {preview}")
                normalized_dataframes[table_name] = working
            dataframes = normalized_dataframes
            
            # Create DuckDB connection
            conn = duckdb.connect(":memory:")
            
            # Register dataframes as tables
            for table_name, df in dataframes.items():
                conn.register(table_name, df)
                msg = f"[SQL] Registered table: {table_name}"
                status_messages.append(msg)
                print(msg)
            
            # Generate SQL code via AI - with FULL schema so AI understands all columns
            msg = "[SQL] Building complete schema for AI code generation..."
            status_messages.append(msg)
            print(msg)
            # Build schema using actual DuckDB column types (not pandas dtypes)
            # so the model knows which columns are numeric vs text
            schema_desc = self._create_duckdb_schema(conn, dataframes)
            
            # Extract exact column names for explicit instruction to AI
            exact_columns = {}
            for table_name, df in dataframes.items():
                exact_columns[table_name] = list(df.columns)
            
            # Build column list with DuckDB types (accurate types after registration)
            column_list = "Table columns (name: DuckDB_type):\n"
            for table_name in exact_columns:
                column_list += f"  {table_name}:\n"
                try:
                    describe_rows = conn.execute(f"DESCRIBE {table_name}").fetchall()
                    for row in describe_rows:
                        col_name, col_type = row[0], row[1]
                        column_list += f"    - {col_name}: {col_type}\n"
                except Exception:
                    for col in exact_columns[table_name]:
                        column_list += f"    - {col}\n"
            
            # Minimal flags for output-format decisions only (not for SQL generation)
            instr_lower = instructions.lower()
            is_reconciliation = any(kw in instr_lower for kw in ["reconcil", "mismatch", "difference", "compare", "find diff"])

            # Keep full steps context for SQL generation
            instr_short = instructions[:6000] if len(instructions) > 6000 else instructions

            # Build a mapping hint so model understands user's file aliases
            # e.g. "Report_A", "first file", "file 1" → actual DuckDB table name
            table_names_list = list(dataframes.keys())

            def _normalize_table_token(name: str) -> str:
                return re.sub(r"[^a-z0-9]", "", (name or "").lower())

            def _extract_direct_sql(text: str) -> str:
                raw = (text or "").strip()
                if not raw:
                    return ""

                fence = re.search(r"```(?:sql)?\s*([\s\S]*?)```", raw, flags=re.IGNORECASE)
                if fence:
                    candidate = fence.group(1).strip()
                    if re.match(r"^(select|with)\b", candidate, re.IGNORECASE) and re.search(r"\bfrom\b", candidate, re.IGNORECASE):
                        return candidate

                m = re.search(r"\b(select|with)\b[\s\S]*", raw, flags=re.IGNORECASE)
                if m:
                    candidate = m.group(0).strip()
                    if re.match(r"^(select|with)\b", candidate, re.IGNORECASE) and re.search(r"\bfrom\b", candidate, re.IGNORECASE):
                        return candidate

                return ""

            def _apply_fuzzy_table_mapping(sql_text: str) -> tuple[str, dict]:
                """Map table tokens in FROM/JOIN to loaded DuckDB table names (exact -> normalized -> fuzzy)."""
                if not sql_text:
                    return sql_text, {}

                available = list(dataframes.keys())
                available_lower = {t.lower(): t for t in available}
                available_norm = {_normalize_table_token(t): t for t in available}
                ordinals = ["first", "second", "third", "fourth", "fifth", "sixth", "seventh", "eighth"]

                def _build_alias_map() -> dict[str, str]:
                    alias_map = {}
                    generic_noise = {"df", "data", "dataset", "table", "records", "record", "sheet", "file"}
                    for i, tbl in enumerate(available):
                        candidates = set()
                        raw = str(tbl)
                        friendly = raw[3:] if raw.lower().startswith("df_") else raw
                        friendly_parts = [p for p in re.split(r"[^a-z0-9]+", friendly.lower()) if p]

                        candidates.add(raw)
                        candidates.add(friendly)
                        if i < len(ordinals):
                            candidates.add(ordinals[i])
                            candidates.add(f"file{i+1}")
                            candidates.add(f"report{chr(ord('a') + i)}")

                        if friendly_parts:
                            candidates.add(friendly_parts[0])
                            candidates.add(friendly_parts[-1])
                            for part in friendly_parts:
                                # Keep short domain tokens like '2a' in addition to regular words.
                                if (len(part) >= 3 or re.fullmatch(r"\d+[a-z]?", part)) and part not in generic_noise:
                                    candidates.add(part)
                            filtered = [p for p in friendly_parts if p not in generic_noise]
                            if filtered:
                                candidates.add("_".join(filtered))
                                candidates.add("".join(filtered))

                        # Deterministic reconciliation aliases.
                        norm_tbl = _normalize_table_token(friendly)
                        if "books" in norm_tbl:
                            candidates.update({"books", "book"})
                        if "2a" in norm_tbl or "gstr2a" in norm_tbl:
                            candidates.update({"2a", "gstr2a"})

                        for cand in candidates:
                            norm = _normalize_table_token(cand)
                            if norm and norm not in alias_map:
                                alias_map[norm] = tbl
                    return alias_map

                alias_norm_map = _build_alias_map()

                # Collect CTE names declared in WITH clause so we never remap them as physical tables.
                cte_names = {
                    m.group(1).strip().lower()
                    for m in re.finditer(r"\b(?:with|,)\s*([A-Za-z_][A-Za-z0-9_]*)\s+as\s*\(", sql_text, flags=re.IGNORECASE)
                }

                # Deterministic shortcuts for common reconciliation aliases.
                books_table = next((t for t in available if "books" in _normalize_table_token(t)), None)
                twoa_table = next((t for t in available if "2a" in _normalize_table_token(t) or "gstr2a" in _normalize_table_token(t)), None)
                replacements = {}

                # Capture table references in FROM/JOIN clauses.
                refs = re.findall(r"\b(?:from|join)\s+([A-Za-z0-9_]+|\"[^\"]+\")", sql_text, flags=re.IGNORECASE)
                for ref in refs:
                    token = ref.strip().strip('"')
                    if token.lower() in cte_names:
                        # Example: FROM Aggregated_2A AS JD (CTE), should not be remapped.
                        continue
                    mapped = None

                    token_norm = _normalize_table_token(token)
                    if token_norm in {"books", "book"} and books_table:
                        mapped = books_table
                    elif token_norm in {"2a", "gstr2a"} and twoa_table:
                        mapped = twoa_table

                    # Exact case-insensitive
                    if mapped is None and token.lower() in available_lower:
                        mapped = available_lower[token.lower()]
                    elif mapped is None:
                        # Normalized form (remove non-alnum)
                        n = token_norm
                        if n in available_norm:
                            mapped = available_norm[n]
                        elif n in alias_norm_map:
                            mapped = alias_norm_map[n]
                        else:
                            # Fuzzy match against normalized table names and aliases
                            norm_space = list(set(list(available_norm.keys()) + list(alias_norm_map.keys())))
                            best = get_close_matches(n, norm_space, n=1, cutoff=0.55)
                            if best:
                                mapped = available_norm.get(best[0]) or alias_norm_map.get(best[0])

                    if mapped and mapped != token:
                        replacements[token] = mapped

                remapped = sql_text
                for src, dst in replacements.items():
                    # Replace only table token after FROM/JOIN to avoid touching aliases/columns.
                    remapped = re.sub(
                        rf"(\b(?:from|join)\s+)(?:\"?{re.escape(src)}\"?)\b",
                        rf"\1{dst}",
                        remapped,
                        flags=re.IGNORECASE,
                    )
                return remapped, replacements

            def _repair_missing_table_errors(sql_text: str, error_text: str) -> tuple[str, dict]:
                """When DuckDB reports missing table names, try alias/fuzzy remap and return updated SQL."""
                if not sql_text or not error_text:
                    return sql_text, {}

                missing_names = re.findall(r"Table with name\s+([A-Za-z_][A-Za-z0-9_]*)\s+does not exist", error_text, flags=re.IGNORECASE)
                if not missing_names:
                    return sql_text, {}

                patched_sql = sql_text
                replacements = {}
                for missing in missing_names:
                    probe_sql = re.sub(
                        rf"(\b(?:from|join)\s+)(?:\"?{re.escape(missing)}\"?)\b",
                        rf"\1{missing}",
                        patched_sql,
                        flags=re.IGNORECASE,
                    )
                    remapped, rep = _apply_fuzzy_table_mapping(probe_sql)
                    if rep:
                        patched_sql = remapped
                        replacements.update(rep)
                return patched_sql, replacements
            mapping_lines = []
            ordinals = ["first", "second", "third", "fourth"]
            for i, tname in enumerate(table_names_list):
                # original filename (strip df_ prefix and replace _ with space)
                friendly = tname[3:].replace('_', ' ') if tname.startswith('df_') else tname
                ordinal = ordinals[i] if i < len(ordinals) else str(i + 1)
                mapping_lines.append(
                    f"  '{tname}' — {ordinal} file, also known as: \"{friendly}\", "
                    f"\"Report_{'ABCDEFGH'[i]}\", \"{ordinal} file\", \"file {i+1}\""
                )
            table_alias_hint = "TABLE ALIASES (user may refer to tables by these names):\n" + "\n".join(mapping_lines) + "\n"

            context_section = f"\n\nADDITIONAL CONTEXT FILES (read-only reference, not in SQL tables):\n{context_block}" if context_block else ""

            sql_prompt = f"""Generate a single DuckDB SQL query for the task below.

USER TASK:
{instr_short}{context_section}

{table_alias_hint}
AVAILABLE TABLES AND COLUMN TYPES:
{column_list}
RULES:
- Follow the USER TASK steps exactly. The user's instructions are the highest priority.
- Output ONLY the SQL query. No explanation, no code fence, no comments.
- Use column names EXACTLY as listed above (do not correct spelling).
- If a column name contains spaces/special chars or starts with a digit, reference it with double quotes (e.g., T."Vendor name", T."2A Value").
- Table aliases must start with a letter or underscore (do NOT use aliases like 2A). Prefer aliases like A, B, T1, T2.
- If user says ignore/exclude/remove a column, that column must NOT appear from the Steps mentioned.
- If user asks to rename a column, return only the renamed alias from the Steps mentioned.
- If user asks aggregation/grouping, query MUST include corresponding GROUP BY and aggregate functions.
- Use CAST(col AS DOUBLE) or TRY_CAST(col AS DOUBLE) before aggregating numeric columns when the type is VARCHAR.
- For key comparisons/joins where NULLs may appear, prefer IS NOT DISTINCT FROM or COALESCE normalization.
- For text filters, account for NULL safely (e.g., COALESCE(col, '') before LIKE/ILIKE when needed).
- Handle NULLs explicitly in join keys/comparisons.
- if its a multi step request use WITH statement only else Select Statement only.
- Do NOT add WHERE filters on columns that the user did not ask to filter. Only filter on conditions explicitly stated in the task.
- If a computation applies conditionally (e.g. 12% of Basic where PF=Yes), use CASE WHEN in the SELECT, not a WHERE clause that removes rows.
- CRITICAL: Only reference columns that are explicitly listed in AVAILABLE TABLES AND COLUMN TYPES above. NEVER write a column name that does not appear in that list (e.g. do not write A.gross if "gross" is not listed — instead write the arithmetic expression: A.basic + A.hra + A.conveyance + A.special_allowance AS gross).
- CRITICAL: Do NOT reference SELECT-level aliases inside the same SELECT or in WHERE/expressions. DuckDB does not allow this. Re-compute the expression inline or use a WITH clause."""

            def clean_sql(raw_sql: str) -> str:
                _debug_status(f"[SQL-DEBUG] clean_sql input (len={len(raw_sql)}):\n{raw_sql!r}")
                if not raw_sql:
                    _debug_status("[SQL-DEBUG] clean_sql: input is empty, returning ''")
                    return ""
                text = raw_sql.strip()

                # Prefer fenced sql blocks when present
                fence_match = re.search(r"```(?:sql)?\s*(.*?)```", text, flags=re.IGNORECASE | re.DOTALL)
                if fence_match:
                    text = fence_match.group(1).strip()
                    _debug_status(f"[SQL-DEBUG] clean_sql: extracted from fence block: {text[:200]!r}")

                # Remove accidental fence remnants
                lines = [ln for ln in text.splitlines() if not ln.strip().startswith("```")]
                text = "\n".join(lines).strip()
                _debug_status(f"[SQL-DEBUG] clean_sql output (len={len(text)}): {text[:200]!r}")
                return text

            def repair_sql_identifiers(sql_text: str) -> str:
                """Best-effort SQL repair for common LLM identifier mistakes.

                Fixes:
                - aliases that start with digits (e.g., 2A -> T2A)
                - alias.column references where column names require quoting
                - near-miss alias.column typos (e.g., dedector_tan -> deductor_tan)
                """
                if not sql_text:
                    return sql_text

                repaired = sql_text
                alias_map = {}  # old_alias -> (new_alias, source_table)

                # Capture aliases from FROM/JOIN clauses and rename invalid aliases.
                alias_pattern = re.compile(
                    r"\b(?:from|join)\s+([A-Za-z_][A-Za-z0-9_]*)\s+(?:as\s+)?([A-Za-z0-9_]+)\b",
                    re.IGNORECASE,
                )
                for m in alias_pattern.finditer(repaired):
                    source_table = m.group(1)
                    alias = m.group(2)
                    if re.match(r"^[0-9]", alias):
                        new_alias = f"T{alias}"
                        alias_map[alias] = (new_alias, source_table)
                    else:
                        alias_map[alias] = (alias, source_table)

                # Apply alias renames globally (word-boundary safe).
                for old_alias, (new_alias, _tbl) in alias_map.items():
                    if old_alias != new_alias:
                        repaired = re.sub(rf"\b{re.escape(old_alias)}\b", new_alias, repaired)
                        _debug_status(f"[SQL-DEBUG] repair_sql_identifiers: alias '{old_alias}' -> '{new_alias}'")

                # Quote problematic columns in alias-qualified references.
                # Build from known dataframe columns so we only patch valid names.
                for _old_alias, (alias, source_table) in alias_map.items():
                    if source_table not in dataframes:
                        continue
                    cols = list(dataframes[source_table].columns)
                    for col in cols:
                        if re.match(r"^[A-Za-z_][A-Za-z0-9_]*$", str(col)):
                            continue
                        quoted = '"' + str(col).replace('"', '""') + '"'
                        before = f"{alias}.{col}"
                        after = f"{alias}.{quoted}"
                        if before in repaired:
                            repaired = repaired.replace(before, after)
                            _debug_status(
                                f"[SQL-DEBUG] repair_sql_identifiers: quoted column ref {before!r} -> {after!r}"
                            )

                # Correct near-miss alias.column typos using known columns per table.
                # Example: T1.dedector_tan -> T1.deductor_tan
                for _old_alias, (alias, source_table) in alias_map.items():
                    if source_table not in dataframes:
                        continue
                    known_cols = [str(c) for c in dataframes[source_table].columns]
                    known_set = set(known_cols)
                    pattern = re.compile(rf"\b{re.escape(alias)}\.([A-Za-z_][A-Za-z0-9_]*)\b")

                    def _fix_col(m):
                        col = m.group(1)
                        if col in known_set:
                            return m.group(0)
                        match = difflib.get_close_matches(col, known_cols, n=1, cutoff=0.82)
                        if match:
                            fixed = f"{alias}.{match[0]}"
                            _debug_status(
                                f"[SQL-DEBUG] repair_sql_identifiers: corrected {alias}.{col} -> {fixed}"
                            )
                            return fixed
                        return m.group(0)

                    repaired = pattern.sub(_fix_col, repaired)

                return repaired

            def repair_sql_join_predicates(sql_text: str) -> str:
                """Repair common invalid JOIN ON patterns produced by LLMs."""
                if not sql_text:
                    return sql_text

                repaired = sql_text

                # Pattern: ON COALESCE(T1.col, T2.col)  ->  ON T1.col IS NOT DISTINCT FROM T2.col
                repaired2 = re.sub(
                    r"\bON\s+COALESCE\s*\(\s*([^,\)]+?)\s*,\s*([^\)]+?)\s*\)",
                    r"ON \1 IS NOT DISTINCT FROM \2",
                    repaired,
                    flags=re.IGNORECASE,
                )
                if repaired2 != repaired:
                    _debug_status("[SQL-DEBUG] repair_sql_join_predicates: fixed ON COALESCE(...) to IS NOT DISTINCT FROM")
                    repaired = repaired2

                return repaired

            def semantic_review_sql(task_text: str, sql_text: str) -> tuple[bool, str]:
                """Dynamically validate SQL against user task using the model itself.

                Returns (ok, reason). If reviewer fails, default to pass to avoid blocking.
                """
                if not self.model:
                    return True, ""

                # ── Keyword pre-screen (no model call needed) ──────────────────────────
                # Extract numbers and bare words from the task, check they appear in SQL.
                # This catches obvious cases where a small model wrongly flags a valid SQL.
                sql_lower = sql_text.lower()
                task_lower = task_text.lower()
                # Collect all numbers mentioned in the task (e.g. 0.12, 12%, 200)
                task_numbers = re.findall(r'\d+\.?\d*', task_lower)
                # Collect key column-like words from the task (>= 4 chars, not SQL keywords)
                _SQL_KW = {'from', 'join', 'where', 'group', 'order', 'having', 'select',
                           'with', 'case', 'when', 'then', 'else', 'end', 'and', 'not',
                           'null', 'like', 'cast', 'coalesce', 'apply', 'each', 'employee',
                           'output', 'compute', 'calculate', 'generate', 'produce', 'using'}
                task_words = [w for w in re.findall(r'[a-z_]{4,}', task_lower)
                              if w not in _SQL_KW]
                # Check: at least 80% of task numbers appear in SQL (as substrings)
                nums_found = sum(1 for n in task_numbers if n in sql_lower)
                nums_ok = (not task_numbers) or (nums_found / len(task_numbers) >= 0.8)
                # Check: at least 60% of task keywords appear in SQL
                words_found = sum(1 for w in task_words if w in sql_lower)
                words_ok = (not task_words) or (words_found / len(task_words) >= 0.6)
                if nums_ok and words_ok:
                    _debug_status(f"[SQL-DEBUG] semantic_review_sql: keyword pre-screen PASS "
                                  f"(nums {nums_found}/{len(task_numbers)}, "
                                  f"words {words_found}/{len(task_words)}) — skipping model call")
                    return True, ""
                # ── End pre-screen ─────────────────────────────────────────────────────
                review_user = f"""Review whether the SQL fully satisfies the USER TASK.

USER TASK:
{task_text}

SQL TO REVIEW:
{sql_text}

Return ONLY strict JSON with keys:
- ok: true or false
- reason: short string
- missing_requirements: array of short strings

Rules:
- Mark ok=false if SQL misses any explicit task step.
- Mark ok=false if SQL uses only one table when task requires reconciling two files.
- Mark ok=false if SQL compares only missing rows but task asks value comparisons on matched rows.
- Mark ok=false if task asks summarize/aggregate/group and SQL does not aggregate accordingly.
- Mark ok=false if user says ignore/exclude/remove a column but SQL still includes it in final output.
- Mark ok=false if user asks renaming but SQL returns old column name instead of requested alias.
"""
                try:
                    review_prompt = self._build_chat_prompt(
                        system="You are a strict SQL task compliance checker. Return JSON only.",
                        messages=[],
                        user_text=review_user,
                        extra_context="",
                    )
                    with self.model_lock:
                        review_resp = self.model.create_completion(
                            review_prompt,
                            max_tokens=400,
                            temperature=0.0,
                            stop=self._get_stop_tokens(),
                        )
                    raw = review_resp.get("choices", [{}])[0].get("text", "").strip()
                    raw = _RE_THINK.sub("", raw).strip()
                    raw = _RE_THINK_INCOMPLETE.sub("", raw).strip()
                    fence_match = re.search(r"```(?:json)?\s*(.*?)```", raw, re.DOTALL)
                    if fence_match:
                        raw = fence_match.group(1).strip()
                    json_match = re.search(r"\{.*\}", raw, re.DOTALL)
                    if json_match:
                        raw = json_match.group(0).strip()
                    payload = json.loads(raw)
                    ok = bool(payload.get("ok", True))
                    reason = str(payload.get("reason", "")).strip()
                    missing = payload.get("missing_requirements", [])
                    if isinstance(missing, list) and missing:
                        # Keep full missing requirements list; truncation loses actionable constraints.
                        reason = (reason + " | Missing: " + "; ".join(str(x) for x in missing)).strip(" |")
                    return ok, reason
                except Exception as e:
                    _debug_status(f"[SQL-DEBUG] semantic_review_sql skipped due to reviewer error: {e}")
                    return True, ""

            def _split_sql_statements(text: str) -> list[str]:
                """Split SQL by ';' while ignoring semicolons in comments/quoted strings."""
                parts = []
                buff = []
                i = 0
                n = len(text)
                state = "normal"  # normal, s_quote, d_quote, line_comment, block_comment

                while i < n:
                    ch = text[i]
                    nxt = text[i + 1] if i + 1 < n else ""

                    if state == "line_comment":
                        if ch == "\n":
                            state = "normal"
                            buff.append(ch)
                        i += 1
                        continue

                    if state == "block_comment":
                        if ch == "*" and nxt == "/":
                            state = "normal"
                            i += 2
                            continue
                        i += 1
                        continue

                    if state == "s_quote":
                        buff.append(ch)
                        if ch == "'":
                            # Handle escaped single quote ''
                            if nxt == "'":
                                buff.append(nxt)
                                i += 2
                                continue
                            state = "normal"
                        i += 1
                        continue

                    if state == "d_quote":
                        buff.append(ch)
                        if ch == '"':
                            # Handle escaped double quote ""
                            if nxt == '"':
                                buff.append(nxt)
                                i += 2
                                continue
                            state = "normal"
                        i += 1
                        continue

                    # normal state
                    if ch == "-" and nxt == "-":
                        state = "line_comment"
                        i += 2
                        continue
                    if ch == "/" and nxt == "*":
                        state = "block_comment"
                        i += 2
                        continue
                    if ch == "'":
                        state = "s_quote"
                        buff.append(ch)
                        i += 1
                        continue
                    if ch == '"':
                        state = "d_quote"
                        buff.append(ch)
                        i += 1
                        continue
                    if ch == ";":
                        stmt = "".join(buff).strip()
                        if stmt:
                            parts.append(stmt)
                        buff = []
                        i += 1
                        continue

                    buff.append(ch)
                    i += 1

                tail = "".join(buff).strip()
                if tail:
                    parts.append(tail)
                return parts

            def _strip_non_sql_prefix(text: str) -> str:
                """Remove human labels/preamble before actual SQL, e.g. 'Query-2:' lines."""
                raw = (text or "").strip()
                if not raw:
                    return ""

                # Common heading patterns: Query 1:, Query-2:, Query_3:
                raw = re.sub(r"^\s*query\s*[-_ ]*\d+\s*:\s*", "", raw, flags=re.IGNORECASE)

                # Generic fallback: keep content from first SELECT/WITH keyword.
                m = re.search(r"\b(select|with)\b", raw, flags=re.IGNORECASE)
                if m and m.start() > 0:
                    raw = raw[m.start():]
                return raw.strip()

            def validate_sql(sql_text: str, semantic: bool = True) -> tuple[str | None, str]:
                """Validate SQL in cost-ordered phases: structural → EXPLAIN → tables → semantic."""
                _debug_status(f"[SQL-DEBUG] validate_sql input (len={len(sql_text) if sql_text else 0}): {(sql_text or '')[:200]!r}")
                if not sql_text or not sql_text.strip():
                    _debug_status("[SQL-DEBUG] validate_sql: FAIL - empty input")
                    return "Empty SQL generated", ""

                normalized = _strip_non_sql_prefix(sql_text.strip())

                statements = _split_sql_statements(normalized)
                if not statements:
                    return "Empty SQL generated", ""
                if len(statements) > 1:
                    _debug_status("[SQL-DEBUG] validate_sql: FAIL - multiple SQL statements detected")
                    return "Multiple SQL statements detected; expected exactly one", normalized
                normalized = statements[0]
                if sql_text.strip().endswith(";"):
                    _debug_status(f"[SQL-DEBUG] validate_sql: stripped trailing semicolon")

                # ── Phase 1: Deterministic structural checks (free) ──
                if re.search(r"\b(insert|update|delete|drop|alter|create|replace|truncate|attach|pragma|copy)\b", normalized, re.IGNORECASE):
                    _debug_status(f"[SQL-DEBUG] validate_sql: FAIL - non-read-only keyword found")
                    return "Non-read-only SQL detected", normalized

                if not re.match(r"^(select|with)\b", normalized, re.IGNORECASE):
                    _debug_status(f"[SQL-DEBUG] validate_sql: FAIL - doesn't start with SELECT/WITH (starts with: {normalized[:30]!r})")
                    return "SQL must start with SELECT or WITH", normalized

                if not any(tbl in normalized for tbl in dataframes.keys()):
                    return "SQL does not reference available tables", normalized

                # ── Phase 2: DuckDB EXPLAIN plan (catches syntax errors before model calls) ──
                try:
                    conn.execute(f"EXPLAIN {normalized}").fetchall()
                    _debug_status(f"[SQL-DEBUG] validate_sql: EXPLAIN plan OK")
                except Exception as e:
                    _debug_status(f"[SQL-DEBUG] validate_sql: FAIL - EXPLAIN plan error: {e}")
                    return f"SQL parse/plan failed: {e}", normalized

                # ── Phase 3: Table reference checks (free) ──
                referenced_tables = [
                    tbl for tbl in dataframes.keys()
                    if re.search(rf"\b{re.escape(tbl)}\b", normalized, re.IGNORECASE)
                ]
                if is_reconciliation and len(dataframes) >= 2 and len(referenced_tables) < 2:
                    return "Reconciliation requires using both input tables", normalized

                join_match = re.search(
                    r"\bfrom\s+([A-Za-z_][A-Za-z0-9_]*)\b[\s\S]*?\bjoin\s+([A-Za-z_][A-Za-z0-9_]*)\b",
                    normalized,
                    re.IGNORECASE,
                )
                if is_reconciliation and len(dataframes) >= 2 and join_match:
                    left_tbl = join_match.group(1)
                    right_tbl = join_match.group(2)
                    if left_tbl.lower() == right_tbl.lower():
                        return "Join uses the same table on both sides; use both input files", normalized

                # Ensure JOIN ON clauses are boolean predicates (not scalar expressions).
                on_clauses = re.findall(
                    r"\bON\s+([\s\S]*?)(?=\b(?:JOIN|WHERE|GROUP\s+BY|ORDER\s+BY|HAVING|LIMIT|UNION|EXCEPT|INTERSECT)\b|$)",
                    normalized,
                    flags=re.IGNORECASE,
                )
                comparator_re = re.compile(
                    r"(=|<>|!=|<=|>=|<|>|\bIS\s+NOT\s+DISTINCT\s+FROM\b|\bIS\s+DISTINCT\s+FROM\b|\bLIKE\b|\bILIKE\b|\bIN\b|\bBETWEEN\b)",
                    flags=re.IGNORECASE,
                )
                for cond in on_clauses:
                    cond_text = cond.strip()
                    if cond_text and not comparator_re.search(cond_text):
                        return f"Invalid JOIN ON condition (must be boolean): {cond_text[:120]}", normalized

                # ── Phase 4: Semantic review via model (optional) ──
                if semantic:
                    sem_ok, sem_reason = semantic_review_sql(instr_short, normalized)
                    if not sem_ok:
                        # Add machine-readable payload for repair prompt consumption.
                        missing_items = []
                        m = re.search(r"\|\s*Missing:\s*(.+)$", sem_reason or "", re.IGNORECASE)
                        if m:
                            missing_items = [x.strip().rstrip(".") for x in m.group(1).split(";") if x.strip()]
                        missing_json = json.dumps(missing_items, ensure_ascii=True)
                        return f"SQL does not satisfy task: {sem_reason or 'semantic mismatch'} | MISSING_REQ_JSON: {missing_json}", normalized

                _debug_status(f"[SQL-DEBUG] validate_sql: PASS - SQL looks valid")
                return None, normalized

            # ── Direct SQL path: if user already supplied SQL, skip AI generation ──
            direct_sql = _extract_direct_sql(instructions)
            if direct_sql:
                _debug_status("[SQL] Direct SQL detected in instructions; skipping AI SQL generation")
                sql_code = clean_sql(direct_sql)
                sql_code, table_replacements = _apply_fuzzy_table_mapping(sql_code)
                if table_replacements:
                    _debug_status(f"[SQL] Fuzzy table mapping applied: {table_replacements}")
                sql_code = repair_sql_identifiers(sql_code)
                sql_code = repair_sql_join_predicates(sql_code)

                direct_statements = _split_sql_statements(sql_code)
                if len(direct_statements) > 1:
                    _debug_status(f"[SQL] Detected {len(direct_statements)} direct SQL statements; executing each query separately")
                    output_dir = os.path.join(app_data_path(), "processed_files")
                    os.makedirs(output_dir, exist_ok=True)
                    timestamp = int(time.time())
                    ext_map = {"excel": "xlsx", "csv": "csv", "pdf": "pdf", "txt": "txt"}
                    ext = ext_map.get(output_format, "xlsx")
                    saved_paths = []
                    total_rows = 0

                    for idx, statement in enumerate(direct_statements, start=1):
                        if self.stop_generation_flag:
                            conn.close()
                            return _stopped_result()

                        statement = _strip_non_sql_prefix(statement)

                        validation_error, statement = validate_sql(statement, semantic=False)
                        if validation_error:
                            conn.close()
                            return {"error": f"Direct SQL validation failed (Query {idx}): {validation_error}"}

                        result_df = conn.execute(statement).fetchdf()
                        row_count = len(result_df)
                        total_rows += row_count

                        output_file = os.path.join(output_dir, f"analysis_{timestamp}_q{idx}.{ext}")
                        if ext == "xlsx":
                            result_df.to_excel(output_file, index=False)
                        elif ext == "csv":
                            result_df.to_csv(output_file, index=False)
                        elif ext == "pdf":
                            result_df.to_csv(output_file, index=False)  # PDF fallback for now
                        else:
                            result_df.to_csv(output_file, sep='\t', index=False)

                        saved_paths.append(output_file)
                        _debug_status(f"[SQL] Query {idx} executed successfully: {row_count} rows -> {output_file}")

                    summary = _clean_summary(status_messages, f"Total queries: {len(direct_statements)} | Total rows: {total_rows}")
                    conn.close()
                    return {
                        "ok": True,
                        "response_text": summary,
                        "file_path": saved_paths[0] if saved_paths else None,
                        "file_paths": saved_paths,
                    }

                validation_error, sql_code = validate_sql(sql_code, semantic=False)
                if validation_error:
                    if "Table with name" in validation_error and "does not exist" in validation_error:
                        repaired_sql, missing_replacements = _repair_missing_table_errors(sql_code, validation_error)
                        if missing_replacements:
                            _debug_status(f"[SQL] Missing-table recovery mapping applied: {missing_replacements}")
                            validation_error, sql_code = validate_sql(repaired_sql, semantic=False)
                    if validation_error:
                        return {"error": f"Direct SQL validation failed: {validation_error}"}

                _debug_status("[SQL] Direct SQL validation passed; executing query on DuckDB")
            else:
                sql_code = ""

            # ── Generate → Validate → Repair loop (retry until success or stuck) ──
            MAX_SQL_ATTEMPTS = 15  # safety ceiling

            if not direct_sql:
                msg = "[SQL] Generating SQL query..."
                status_messages.append(msg)
                print(msg)
                if not self.model:
                    return {"error": "No model loaded"}

                # Print full sql_prompt so it can be verified in logs
                _debug_status(f"[SQL-DEBUG] Full sql_prompt ({len(sql_prompt)} chars):\n{'='*60}\n{sql_prompt}\n{'='*60}")

                stop_tokens = self._get_stop_tokens()
                errors_so_far = []  # tracks {"attempt": N, "sql": str, "error": str}
                _consecutive_same_error = 0  # detect when model is stuck
                validation_error = None
                retry_exhausted = False
                model_stuck = False

                for attempt in range(MAX_SQL_ATTEMPTS):
                    if self.stop_generation_flag:
                        return _stopped_result()
                    _debug_status(f"[SQL-DEBUG] ── Attempt {attempt + 1}/{MAX_SQL_ATTEMPTS} ──")

                    if attempt == 0:
                        # First attempt: generate from original prompt
                        gen_prompt = self._build_chat_prompt(
                            system="You are a SQL expert. Output ONLY a single DuckDB SELECT query with no explanation.",
                            messages=[],
                            user_text=sql_prompt,
                            extra_context="",
                        )
                    else:
                        # Repair: include the failed SQL + specific fix instructions
                        prev = errors_so_far[-1]
                        # Extract actionable fixes from error messages
                        fix_items = []
                        for e in errors_so_far:
                            err = e["error"]
                            # Prefer machine-readable missing requirements if present.
                            json_match = re.search(r"MISSING_REQ_JSON:\s*(\[[\s\S]*\])", err)
                            if json_match:
                                try:
                                    parsed = json.loads(json_match.group(1))
                                    if isinstance(parsed, list):
                                        for item in parsed:
                                            item = str(item).strip().rstrip(".")
                                            if item:
                                                fix_items.append(f"- {item}")
                                        continue
                                except Exception:
                                    pass
                            # Pull out "Missing: ..." items from semantic reviewer
                            missing_match = re.search(r"Missing:\s*(.+)", err)
                            if missing_match:
                                for item in missing_match.group(1).split(";"):
                                    item = item.strip().rstrip(".")
                                    if item:
                                        fix_items.append(f"- {item}")
                            else:
                                fix_items.append(f"- Fix: {err}")

                        # Deduplicate fix items
                        fix_items = list(dict.fromkeys(fix_items))
                        fix_block = "\n".join(fix_items) if fix_items else f"- {prev['error']}"
                        # Truncate failed SQL to 600 chars — model only needs to see the
                        # structure that failed, not fill the entire context window with garbage.
                        prev_sql_snippet = prev['sql'][:600] + ("..." if len(prev['sql']) > 600 else "")
                        repair_user = (
                            sql_prompt
                            + f"\n\nYOUR PREVIOUS SQL (which failed validation):\n{prev_sql_snippet}"
                            + f"\n\nREQUIRED FIXES (apply ALL of these):\n{fix_block}"
                            + "\n\nRewrite the SQL to fix ALL issues above. Return exactly ONE read-only DuckDB SQL query starting with SELECT or WITH."
                        )
                        gen_prompt = self._build_chat_prompt(
                            system="You are a SQL expert. Output ONLY a single DuckDB SELECT query with no explanation.",
                            messages=[],
                            user_text=repair_user,
                            extra_context="",
                        )
                        _debug_status(f"[SQL-DEBUG] Repair prompt length: {len(gen_prompt)} chars")

                    try:
                        if self.stop_generation_flag:
                            return _stopped_result()
                        raw_sql = ""
                        with self.model_lock:
                            if self.model is None:
                                return {"error": "No model loaded"}

                            # Stream tokens so stop_generation can interrupt mid-attempt.
                            # Cap at 700 tokens — SQL queries don't need more, and higher limits
                            # let small models fill context with infinite nested SELECT loops.
                            stream = self.model(
                                gen_prompt,
                                max_tokens=700,
                                temperature=min(0.1 + (attempt * 0.08), 0.9),  # ramp up creativity on retries
                                stream=True,
                                stop=stop_tokens,
                            )

                            for chunk in stream:
                                if self.stop_generation_flag:
                                    return _stopped_result()
                                raw_sql += chunk.get("choices", [{}])[0].get("text", "")

                        if self.stop_generation_flag:
                            return _stopped_result()
                    except Exception as model_err:
                        _debug_status(f"[SQL-DEBUG] model.create_completion RAISED on attempt {attempt+1}: {model_err}")
                        import traceback; traceback.print_exc()
                        return {"error": f"Model inference failed: {model_err}"}

                    _debug_status(f"[SQL-DEBUG] Attempt {attempt+1} raw output (len={len(raw_sql)}):\n{'='*60}\n{raw_sql}\n{'='*60}")

                    raw_sql = raw_sql.strip()

                    # Detect "SELECT bomb": model stuck in infinite nested SELECT loop.
                    # Count FROM ( occurrences — more than 6 deep means garbage output.
                    if raw_sql.upper().count("FROM (") > 6:
                        _debug_status(f"[SQL-DEBUG] SELECT bomb detected ({raw_sql.upper().count('FROM (')} nesting levels) — discarding output")
                        errors_so_far.append({"attempt": attempt, "sql": "", "error": "Output contained infinitely nested SELECT subqueries — rewrite as a flat JOIN or simple WITH clause"})
                        continue

                    sql_code = clean_sql(raw_sql)
                    sql_code = repair_sql_identifiers(sql_code)
                    sql_code = repair_sql_join_predicates(sql_code)

                    # ── Build real-column sets (global + per alias) ──
                    all_real_cols = set()
                    for _df in dataframes.values():
                        all_real_cols.update(c.lower() for c in _df.columns)
                    # Map table alias → set of real column names for that table
                    _alias_to_cols: dict = {}
                    for _m in re.finditer(r'\b(df_\w+)\s+([A-Za-z_]\w*)\b', sql_code, re.IGNORECASE):
                        _tbl, _alias = _m.group(1), _m.group(2).upper()
                        for _df_name, _df in dataframes.items():
                            if _df_name.lower() == _tbl.lower():
                                _alias_to_cols[_alias] = {c.lower() for c in _df.columns}
                                break

                    def _col_ref_valid(alias: str, col: str) -> bool:
                        """Return True if alias.col is a valid reference."""
                        a = alias.upper()
                        c = col.lower()
                        if c in ('null', 'not', 'true', 'false'):
                            return True
                        if a in _alias_to_cols:
                            return c in _alias_to_cols[a]
                        return c in all_real_cols  # unqualified / unknown alias

                    # ── Strip WHERE conditions that reference non-existent columns ──
                    # Also drops any surviving WHERE conditions when the task contains no
                    # explicit row-filtering language — making the fix generic, not
                    # tied to specific values like 'Yes'/'No' or IS NOT NULL patterns.
                    def _strip_bad_where_conditions(sql_text: str) -> str:
                        where_match = re.search(r'\bWHERE\b([\s\S]*?)(?=\b(?:GROUP\s+BY|ORDER\s+BY|HAVING|LIMIT|UNION|EXCEPT|INTERSECT)\b|$)', sql_text, re.IGNORECASE)
                        if not where_match:
                            return sql_text
                        where_body = where_match.group(1)
                        # If more closing parens than opening, this WHERE is inside a
                        # subquery and our regex has captured the closing ')' of that
                        # subquery plus JOIN clauses. Do not touch it — we'd corrupt the SQL.
                        if where_body.count(')') > where_body.count('('):
                            _debug_status("[SQL-DEBUG] WHERE spans subquery boundary — skipping sanitizer")
                            return sql_text
                        conditions = re.split(r'\bAND\b', where_body, flags=re.IGNORECASE)
                        good = []
                        for cond in conditions:
                            cond_s = cond.strip()
                            if not cond_s:
                                continue
                            # Strip conditions with non-existent alias.col refs (schema check)
                            aliased = re.findall(r'\b([A-Za-z_]\w*)\.([A-Za-z_]\w*)\b', cond_s)
                            bad = [(a, c) for a, c in aliased if not _col_ref_valid(a, c)]
                            if bad:
                                _debug_status(f"[SQL-DEBUG] Stripped invalid WHERE condition "
                                              f"(bad refs: {bad}): {cond_s[:100]}")
                                continue
                            good.append(cond_s)
                        if not good:
                            return sql_text[:where_match.start()].rstrip() + sql_text[where_match.end():]
                        # If valid conditions survived, check whether the task actually
                        # asked for row filtering. Small models add defensive WHERE
                        # conditions that were never requested.
                        # Explicit filter intent: comparison with a literal value/number,
                        # or filter/exclude/only keywords in a filtering context.
                        _filter_intent = bool(re.search(
                            r'\b(?:filter|exclude|only\s+(?:include|show|rows?|records?))\b'
                            r'|(?:>|<|>=|<=|!=|<>)\s*[\d\'""]',
                            instructions, re.IGNORECASE))
                        if not _filter_intent:
                            _debug_status(f"[SQL-DEBUG] Task has no row-filtering intent — "
                                          f"dropped {len(good)} surviving WHERE condition(s)")
                            # Use a space separator to avoid joining last token with next keyword
                            return sql_text[:where_match.start()].rstrip() + " " + sql_text[where_match.end():]
                        return (sql_text[:where_match.start()] +
                                " WHERE\n  " + "\n  AND ".join(good) +
                                sql_text[where_match.end():])
                    sql_code = _strip_bad_where_conditions(sql_code)

                    # ── Fix bad alias.col refs: correct alias if possible, else NULL ──
                    # Only apply on flat JOINs (no subqueries). When a model generates
                    # FROM (...) A subqueries, alias 'A' refers to the subquery result —
                    # not directly to the base table — so column-level validation is
                    # unreliable and replacing valid refs with NULL corrupts the SQL.
                    _has_subquery = bool(re.search(r'FROM\s*\(', sql_code, re.IGNORECASE))
                    if _alias_to_cols and not _has_subquery:
                        def _fix_col_ref(m: re.Match) -> str:
                            alias, col = m.group(1), m.group(2)
                            if _col_ref_valid(alias, col):
                                return m.group(0)  # already valid
                            col_lower = col.lower()
                            # Try to correct to the right table alias
                            for other_alias, other_cols in _alias_to_cols.items():
                                if col_lower in other_cols:
                                    correct = f"{other_alias.lower()}.{col}"
                                    _debug_status(f"[SQL-DEBUG] Corrected wrong-alias '{alias}.{col}' → '{correct}'")
                                    return correct
                            # Not in any table — replace with NULL
                            _debug_status(f"[SQL-DEBUG] Replaced non-existent column '{alias}.{col}' with NULL")
                            return 'NULL'
                        sql_code = re.sub(r'\b([A-Za-z_]\w*)\.([A-Za-z_]\w*)\b',
                                          _fix_col_ref, sql_code)

                    _debug_status(f"[SQL-DEBUG] After clean_sql (len={len(sql_code)}): {sql_code[:300]!r}")

                    validation_error, sql_code = validate_sql(sql_code)

                    if not validation_error:
                        _debug_status(f"[SQL-DEBUG] Attempt {attempt+1}: PASSED validation")
                        break

                    msg = f"[SQL] Attempt {attempt+1} failed: {validation_error}"
                    status_messages.append(msg)
                    print(msg)
                    errors_so_far.append({"attempt": attempt, "sql": sql_code, "error": validation_error})

                    # Detect if model is stuck: same error 3 times in a row → give up
                    if len(errors_so_far) >= 3:
                        last_3 = [e["error"] for e in errors_so_far[-3:]]
                        if last_3[0] == last_3[1] == last_3[2]:
                            _debug_status(f"[SQL-DEBUG] Model stuck: same error 3 times in a row, stopping.")
                            model_stuck = True
                            break

                else:
                    # Safety ceiling reached
                    last_err = errors_so_far[-1]["error"] if errors_so_far else "Unknown"
                    _debug_status(f"[SQL-DEBUG] Safety ceiling ({MAX_SQL_ATTEMPTS}) reached. Last error: {last_err}")
                    retry_exhausted = True

                # Check if we broke out of the loop due to retries/stuck (not success)
                if validation_error:
                    last_err = errors_so_far[-1]["error"] if errors_so_far else "Unknown"
                    _debug_status(
                        f"[SQL-DEBUG] Retry limit reached (stuck={model_stuck}, exhausted={retry_exhausted}). "
                        "Attempting fallback with last generated SQL and structural validation only."
                    )
                    fallback_sql = errors_so_far[-1]["sql"] if errors_so_far else sql_code
                    fb_err, fb_sql = validate_sql(fallback_sql, semantic=False)

                    if fb_err and "Table with name" in fb_err and "does not exist" in fb_err:
                        repaired_sql, missing_replacements = _repair_missing_table_errors(fallback_sql, fb_err)
                        if missing_replacements:
                            _debug_status(f"[SQL] Fallback missing-table recovery mapping applied: {missing_replacements}")
                            fb_err, fb_sql = validate_sql(repaired_sql, semantic=False)

                    if fb_err:
                        return {
                            "error": (
                                f"SQL generation failed after {len(errors_so_far)} attempts; "
                                f"fallback validation also failed: {fb_err}"
                            )
                        }

                    sql_code = fb_sql
                    validation_error = None
                    msg = "[SQL] Fallback enabled: executing last generated query after retry limit"
                    status_messages.append(msg)
                    print(msg)

            msg = f"[SQL] Generated query: {sql_code[:200]}..."
            status_messages.append(msg)
            print(msg)

            if self.stop_generation_flag:
                return _stopped_result()
            
            # Print full SQL query to terminal before execution
            print("\n" + "="*80)
            print("[SQL] FULL QUERY TO BE EXECUTED:")
            print("="*80)
            print(sql_code)
            print("="*80 + "\n")
            
            # Also add to UI
            status_messages.append("\n" + "="*80)
            status_messages.append("[SQL] FULL QUERY TO BE EXECUTED:")
            status_messages.append("="*80)
            status_messages.append(sql_code)
            status_messages.append("="*80 + "\n")
            
            # Execute SQL query
            try:
                if self.stop_generation_flag:
                    conn.close()
                    return _stopped_result()
                msg = "[SQL] Executing query..."
                status_messages.append(msg)
                print(msg)
                
                result_df = conn.execute(sql_code).fetchdf()
                row_count = len(result_df)
                msg = f"[SQL] ✓ Query executed successfully: {row_count} rows returned"
                status_messages.append(msg)
                print(msg)
                
                # Save SQL result in requested format
                import pandas as pd
                msg = "[SQL] Converting to DataFrame..."
                status_messages.append(msg)
                print(msg)
                
                output_dir = os.path.join(app_data_path(), "processed_files")
                os.makedirs(output_dir, exist_ok=True)
                timestamp = int(time.time())

                def _extract_recon_splits(frame: "pd.DataFrame") -> dict[str, "pd.DataFrame"]:
                    """Return deterministic reconciliation buckets when recognizable."""
                    splits = {"All_Results": frame}
                    cols = {str(c).strip().lower(): c for c in frame.columns}

                    # Preferred split key: _merge from outer joins
                    if "_merge" in cols:
                        merge_col = cols["_merge"]
                        merge_norm = frame[merge_col].astype(str).str.strip().str.lower()
                        splits["Matched"] = frame[merge_norm == "both"]
                        splits["Missing_From_Second"] = frame[merge_norm.str.contains("left", na=False)]
                        splits["Missing_From_First"] = frame[merge_norm.str.contains("right", na=False)]
                        return splits

                    # Fallback split key: status-like category columns
                    for candidate in ("recon_status", "comparison_status", "status", "match_status"):
                        if candidate in cols:
                            status_col = cols[candidate]
                            status_norm = frame[status_col].astype(str).str.strip().str.lower()
                            splits["Matched"] = frame[status_norm.str.contains("match", na=False)]
                            splits["Mismatched"] = frame[status_norm.str.contains("mismatch|diff|different", regex=True, na=False)]
                            return splits

                    return splits

                wants_multi_excel = (
                    output_format == "excel"
                    and is_reconciliation
                    and any(k in instr_lower for k in ["multiple sheet", "multi sheet", "multi-sheet", "single excel", "one excel"])
                )
                wants_multi_csv = (
                    output_format == "csv"
                    and is_reconciliation
                    and any(k in instr_lower for k in ["multiple csv", "multi csv", "separate csv", "multiple output", "multi output", "multiple outputs"])
                )

                # Determine file extension based on output_format
                ext_map = {"none": "", "excel": "xlsx", "csv": "csv", "pdf": "pdf", "txt": "txt"}
                ext = ext_map.get(output_format, "xlsx")
                output_file = os.path.join(output_dir, f"analysis_{timestamp}.{ext}") if ext else ""

                if ext:
                    msg = f"[SQL] Saving to {ext.upper()} file..."
                    status_messages.append(msg)
                    print(msg)
                else:
                    msg = "[SQL] UI-only mode: no file export requested."
                    status_messages.append(msg)
                    print(msg)

                saved_paths = []
                if wants_multi_excel:
                    split_frames = _extract_recon_splits(result_df)
                    with pd.ExcelWriter(output_file, engine="openpyxl") as writer:
                        for sname, sdf in split_frames.items():
                            if sdf is None:
                                continue
                            safe_sheet = re.sub(r"[^A-Za-z0-9_ ]", "_", sname)[:31] or "Sheet1"
                            sdf.to_excel(writer, sheet_name=safe_sheet, index=False)
                    saved_paths.append(output_file)
                    msg = f"[SQL] ✓ Saved multi-sheet Excel ({len(split_frames)} sheets): {output_file}"
                    status_messages.append(msg)
                    print(msg)
                elif wants_multi_csv:
                    split_frames = _extract_recon_splits(result_df)
                    csv_dir = os.path.join(output_dir, f"analysis_{timestamp}_csv_parts")
                    os.makedirs(csv_dir, exist_ok=True)
                    for sname, sdf in split_frames.items():
                        if sdf is None:
                            continue
                        fname = re.sub(r"[^A-Za-z0-9_\-]", "_", sname).strip("_") or "part"
                        fpath = os.path.join(csv_dir, f"{fname}.csv")
                        sdf.to_csv(fpath, index=False)
                        saved_paths.append(fpath)
                    output_file = csv_dir
                    msg = f"[SQL] ✓ Saved multi-CSV output ({len(saved_paths)} files): {csv_dir}"
                    status_messages.append(msg)
                    print(msg)
                elif ext == "":
                    # UI-only mode: skip file generation.
                    saved_paths = []
                elif ext == "xlsx":
                    result_df.to_excel(output_file, index=False)
                    saved_paths.append(output_file)
                elif ext == "csv":
                    result_df.to_csv(output_file, index=False)
                    saved_paths.append(output_file)
                elif ext == "pdf":
                    result_df.to_csv(output_file, index=False)  # PDF requires reportlab, fallback to CSV for now
                    saved_paths.append(output_file)
                else:  # txt
                    result_df.to_csv(output_file, sep='\t', index=False)
                    saved_paths.append(output_file)

                if ext and not wants_multi_excel and not wants_multi_csv:
                    msg = f"[SQL] ✓ Saved results to: {output_file}"
                    status_messages.append(msg)
                    print(msg)
                
                # Return status messages with count and file path
                summary = _clean_summary(status_messages, f"Total rows: {row_count}")
                
                conn.close()
                return {
                    "ok": True,
                    "response_text": summary,
                    "file_path": output_file if ext else None,
                    "file_paths": saved_paths,
                }
            except Exception as e:
                conn.close()
                return {"error": f"SQL execution failed: {str(e)}"}
            
        except Exception as e:
            print(f"[SQL] Error: {e}")
            import traceback
            traceback.print_exc()
            return {"error": f"SQL analysis failed: {str(e)}"}

    def _execute_python_analysis(
        self,
        dataframes: dict,
        instructions: str,
        output_format: str,
        pipeline_status_messages: list = None
    ) -> dict:
        """Generate and execute Python code for smaller file analysis."""
        try:
            import pandas as pd
            import time
            import os
            import re

            status_messages = pipeline_status_messages if pipeline_status_messages else []

            def log(msg: str) -> None:
                status_messages.append(msg)
                print(msg)

            def is_valid_python(code: str) -> bool:
                if not code or not code.strip():
                    return False

                banned_fragments = [
                    "read_csv",
                    "path_to_",
                    "import pandas",
                    "import os",
                    "import sys",
                    "import subprocess",
                    "exec(",
                    "eval(",
                    "__import__",
                    "open(",
                    "os.",
                    "system(",
                    "popen(",
                    "print(",
                ]
                lowered = code.lower()
                if any(fragment.lower() in lowered for fragment in banned_fragments):
                    return False

                if "result_df" not in code:
                    return False

                try:
                    compile(code, "<generated_code>", "exec")
                    return True
                except SyntaxError:
                    return False

            def _clean_summary(msgs: list, suffix: str = "") -> str:
                clean = []
                for m in msgs:
                    if any(tag in m for tag in (
                        "[SQL-DEBUG]", "[CODE_EXEC]", "[DEBUG]",
                        "clean_sql input", "clean_sql output",
                        "clean_sql:", "VERIFY COLUMN", "VERIFY TABLE",
                        "repair_sql", "semantic_review",
                    )):
                        continue
                    line = re.sub(r'^\[(SQL|PYTHON|AGENT)\]\s*', '', m).strip()
                    if line:
                        clean.append(line)
                result = "\n".join(clean) if clean else "Analysis completed."
                if suffix:
                    result += "\n\n" + suffix
                return result

            log("[PYTHON] Executing Python-based analysis...")
            log("[PYTHON] Building schema for AI...")

            schema_desc = self._create_enhanced_schema(dataframes)
            
            # Extract exact column names for explicit instruction to AI
            exact_columns = {}
            for table_name, df in dataframes.items():
                exact_columns[table_name] = list(df.columns)

            instr_short = instructions[:6000] if len(instructions) > 6000 else instructions

            dataframe_names = ", ".join(list(dataframes.keys()))
            
            # Build explicit column list for the prompt
            column_list = "EXACT COLUMN NAMES (use exactly as shown, do not correct spelling):\n"
            for table_name, cols in exact_columns.items():
                col_str = ", ".join([f"'{c}'" for c in cols])
                column_list += f"  {table_name}: [{col_str}]\n"
            
            multi_file_hint = ""
            if len(dataframes) > 1:
                multi_file_hint = (
                    "MULTI-FILE RULE:\n"
                    "- More than one dataframe is provided.\n"
                    "- You MUST use ALL dataframes.\n"
                    "- Do NOT ignore any dataframe.\n"
                    "- Do NOT overwrite result_df multiple times.\n"
                    "- Follow the USER TASK steps exactly to decide how to join/merge/compare.\n"
                    "- Handle NULLs, duplicates, and type mismatches as needed."
                )

            python_prompt = f"""
    You are an expert Python data analyst.

    TASK:
    {instr_short}

    IMPORTANT CONTEXT:
    - Data is ALREADY LOADED into pandas DataFrames.
    - DO NOT use pd.read_csv().
    - DO NOT use file paths.
    - DO NOT redefine dataframes.
    - Use ONLY the existing DataFrames directly.

    AVAILABLE DATAFRAMES:
    {dataframe_names}

    {column_list}

    SCHEMA (columns + sample values):
    {schema_desc}

    {multi_file_hint}

    STEP-INTERPRETATION PRIORITY:
    - Follow the user's Steps section literally when present.
    - If steps ask for unique_key, create unique_key.
    - If steps say "all columns", derive columns dynamically from dataframe columns.
    - Do not hardcode guessed column names when steps require all columns.
    - If steps ask to delete/remove, perform those operations before final result.
    - If steps ask arithmetic comparisons, compute numeric deltas/ratios/tolerance checks.
    - If steps ask exact/fuzzy filtering (single or list), apply those filters explicitly.

        ADVANCED RECONCILIATION POLICIES:
        - Numeric tolerance policy:
            - If user gives tolerance, use it.
            - Otherwise default ABS tolerance=0 and PCT tolerance=0.
        - Date normalization policy:
            - Normalize dates to comparable format before joins/comparisons when needed.
            - Safely handle mixed date formats.
        - Text normalization policy:
            - Apply trim/case normalization for comparison keys when appropriate.
            - Remove obvious punctuation/spacing noise for fuzzy matching requests.
        - Duplicate handling policy:
            - If user specifies duplicate strategy, follow it.
            - Otherwise preserve all rows unless task explicitly asks deduplication.
        - Null handling policy:
            - If user specifies null behavior, follow it.
            - Otherwise use safe normalized blanks only for key-building steps.
        - Fuzzy matching policy:
            - If fuzzy matching is requested, use deterministic logic and include threshold in result.
            - If threshold is not provided, use conservative default and report it.
        - Output contract:
            - Provide summary counts: matched, missing_left, missing_right, mismatched.
            - Provide detailed mismatch rows.
            - Add reason/category column for mismatch cause when possible.
        - Performance policy:
            - Prefer vectorized/set operations over row-wise loops when possible.
            - For very large data, avoid expensive cartesian logic and use key-based compare.

    STRICT RULES:
    - Return ONLY valid Python code.
    - Do NOT include explanations.
    - Do NOT include markdown.
    - Do NOT include comments.
    - Do NOT include numbered steps.
    - Do NOT include any text before or after the code.
    - Do NOT import anything.
    - Do NOT use read_csv.
    - Do NOT use print().
    - Do NOT create fake/sample data.
    - Do NOT redefine dataframe variables.
    - Assign final output ONLY ONCE to: result_df
    - result_df must be a pandas DataFrame.
    - CRITICAL: Use column names EXACTLY as shown above, even if they look misspelled.
    - CRITICAL: Do NOT "correct" column names to match proper English spelling.

    CODE REQUIREMENTS:
    - Use only the provided DataFrames.
    - Use column names EXACTLY as provided (do not fix perceived typos).
    - Handle missing columns safely.
    - If a column may not exist, check with: if 'col' in df.columns:
    - For unique_key/all-columns reconciliation, generate dynamic column logic from df.columns.
    - When filtering with list inputs, support both exact IN-style and fuzzy list matching behavior.
    - When comparing arithmetic across files, include computed difference columns in result_df when relevant.
    - Ensure the code is executable as-is.
    - Avoid indentation errors.

    EXAMPLE:
    result_df = df_1000_BT_Records.groupby("Date").sum().reset_index()

    Return ONLY clean Python code inside a single fenced block:
    ```python
    result_df = ...
    """

            log("[PYTHON] Generating Python code...")

            if not self.model:
                return {"error": "No model loaded"}

            response = self.model.create_completion(
                python_prompt,
                max_tokens=384,
                temperature=0.1,
            )

            python_code = response.get("choices", [{}])[0].get("text", "").strip()
            log(f"[DEBUG] Raw LLM output:\n{python_code[:500]}")

            # 🔹 CLEAN CODE
            def clean_code(code: str) -> str:
                lines = code.split("\n")
                clean_lines = []

                for line in lines:
                    s = line.strip()

                    # Skip obvious junk
                    if not s:
                        continue

                    if s.startswith("```"):
                        continue

                    # Skip numbered steps
                    if s[0].isdigit() and "." in s[:3]:
                        continue

                    # Skip explanation lines
                    if s.lower().startswith(("here", "to ", "this ", "step")):
                        continue

                    # 🚨 KEEP actual python lines
                    clean_lines.append(line)

                return "\n".join(clean_lines)

            python_code = clean_code(python_code)

            log("[PYTHON] Code generated successfully.")
            log(f"[PYTHON] Code preview:\n{python_code[:300]}")

            # 🔒 Validation checks
            if "result_df" not in python_code:
                return {"error": "Invalid code: result_df not found"}

            if "read_csv" in python_code:
                return {"error": "Invalid code: read_csv not allowed"}

            # 🔹 Prepare safe execution environment
            allowed_builtins = {"len": len, "range": range}
            namespace = {
                "__builtins__": allowed_builtins,
                "pd": pd,
                "result_df": None
            }

            namespace.update(dataframes)

            unsafe_keywords = ["__import__", "exec", "eval", "open", "system", "popen", "os."]
            if any(kw in python_code for kw in unsafe_keywords):
                return {"error": "Unsafe code detected"}

            # 🔁 Retry mechanism
            for attempt in range(2):
                try:
                    # Print full code to terminal before execution
                    if attempt == 0:  # Only print on first attempt
                        print("\n" + "="*80)
                        print("[PYTHON] FULL CODE TO BE EXECUTED:")
                        print("="*80)
                        print(python_code)
                        print("="*80 + "\n")
                        
                        # Also add to UI
                        log("\n" + "="*80)
                        log("[PYTHON] FULL CODE TO BE EXECUTED:")
                        log("="*80)
                        log(python_code)
                        log("="*80 + "\n")
                    
                    # DEBUG: Show namespace before execution
                    log("[DEBUG] Namespace keys before exec:")
                    for key in list(namespace.keys())[:10]:
                        log(f"  - {key}: {type(namespace.get(key)).__name__}")
                    
                    log("[PYTHON] Running generated code...")
                    
                    # Time the execution with detailed steps
                    start_time = time.time()
                    log(f"[DEBUG] Start time: {start_time}")
                    
                    exec(python_code, namespace)
                    
                    exec_end = time.time()
                    execution_time = round(exec_end - start_time, 2)
                    log(f"[DEBUG] Exec completed after {execution_time}s")
                    log(f"[PYTHON] ✓ Code executed successfully in {execution_time}s")
                    
                    # DEBUG: Check namespace after execution
                    log("[DEBUG] Checking namespace after exec...")
                    if "result_df" in namespace:
                        result_obj = namespace["result_df"]
                        log(f"[DEBUG] result_df exists: {type(result_obj).__name__}")
                        if result_obj is not None:
                            log(f"[DEBUG] result_df is NOT None")
                        else:
                            log(f"[DEBUG] result_df is None (BAD)")
                    else:
                        log(f"[DEBUG] result_df NOT in namespace (BAD)")
                    
                    log("[PYTHON] Validating results...")

                    break
                    
                except Exception as e:
                    if attempt == 0:
                        import traceback
                        error_trace = traceback.format_exc()
                        log(f"[PYTHON] Error occurred. Retrying... ({str(e)})")
                        log(f"[DEBUG] Full traceback:\n{error_trace[:500]}")

                        retry_prompt = python_prompt + f"\n\nERROR:\n{str(e)}\nFix the code."

                        response = self.model.create_completion(
                            retry_prompt,
                            max_tokens=384,
                            temperature=0.1,
                        )

                        python_code = clean_code(
                            response.get("choices", [{}])[0].get("text", "").strip()
                        )

                        log("[PYTHON] Regenerated fixed code.")
                        log(f"[PYTHON] Code preview:\n{python_code[:300]}")
                        
                        # Print full regenerated code to terminal
                        print("\n" + "="*80)
                        print("[PYTHON] FULL REGENERATED CODE TO BE EXECUTED:")
                        print("="*80)
                        print(python_code)
                        print("="*80 + "\n")
                        
                        # Also add to UI
                        log("\n" + "="*80)
                        log("[PYTHON] FULL REGENERATED CODE TO BE EXECUTED:")
                        log("="*80)
                        log(python_code)
                        log("="*80 + "\n")

                    else:
                        import traceback
                        error_trace = traceback.format_exc()
                        log(f"[PYTHON] ✗ Execution failed: {str(e)}")
                        log(f"[DEBUG] Full traceback:\n{error_trace}")
                        return {"error": f"Python execution failed: {str(e)}"}

            # 🔹 Validate result
            log(f"[DEBUG] Starting result validation...")
            result_df = namespace.get("result_df")
            log(f"[DEBUG] Retrieved result_df from namespace: {type(result_df).__name__}")
            log(f"[PYTHON] Checking result_df...")

            if result_df is None:
                log(f"[DEBUG] ERROR: result_df is None")
                log(f"[DEBUG] Available keys in namespace: {list(namespace.keys())}")
                return {"error": "No result_df produced"}

            log(f"[DEBUG] result_df type check: {isinstance(result_df, pd.DataFrame)}")
            if not isinstance(result_df, pd.DataFrame):
                log(f"[DEBUG] ERROR: result_df is {type(result_df).__name__}, not DataFrame")
                return {"error": f"result_df is not a DataFrame (got {type(result_df).__name__})"}

            log(f"[DEBUG] Getting row count...")
            row_count = len(result_df)
            col_count = len(result_df.columns)
            log(f"[DEBUG] Row count: {row_count}, Column count: {col_count}")
            
            log(f"[PYTHON] ✓ Valid DataFrame with {row_count} rows, {col_count} columns")
            log(f"[PYTHON] Execution completed in {execution_time}s")  
            log(f"[PYTHON] Saving results to file...")

            # 🔹 Save output
            log(f"[DEBUG] Creating output directory...")
            output_dir = os.path.join(app_data_path(), "processed_files")
            os.makedirs(output_dir, exist_ok=True)
            log(f"[DEBUG] Output dir exists: {output_dir}")

            log(f"[DEBUG] Getting timestamp...")
            timestamp = int(time.time())
            log(f"[DEBUG] Timestamp: {timestamp}")
            
            ext_map = {"excel": "xlsx", "csv": "csv", "pdf": "pdf", "txt": "txt"}
            ext = ext_map.get(output_format, "xlsx")
            log(f"[DEBUG] Output format: {output_format} -> extension: {ext}")

            output_file = os.path.join(output_dir, f"analysis_{timestamp}.{ext}")
            log(f"[DEBUG] Output file path: {output_file}")

            try:
                if ext == "xlsx":
                    log(f"[PYTHON] Writing to Excel file...")
                    log(f"[DEBUG] Calling result_df.to_excel()...")
                    result_df.to_excel(output_file, index=False)
                    log(f"[DEBUG] Excel write completed")
                elif ext == "csv":
                    log(f"[PYTHON] Writing to CSV file...")
                    log(f"[DEBUG] Calling result_df.to_csv()...")
                    result_df.to_csv(output_file, index=False)
                    log(f"[DEBUG] CSV write completed")
                elif ext == "pdf":
                    log(f"[PYTHON] Writing to CSV (PDF fallback)...")
                    log(f"[DEBUG] Calling result_df.to_csv()...")
                    result_df.to_csv(output_file, index=False)  # fallback
                    log(f"[DEBUG] CSV write completed")
                else:
                    log(f"[PYTHON] Writing to TSV file...")
                    log(f"[DEBUG] Calling result_df.to_csv() with TSV format...")
                    result_df.to_csv(output_file, sep="\t", index=False)
                    log(f"[DEBUG] TSV write completed")
                
                log(f"[DEBUG] Verifying file exists: {os.path.exists(output_file)}")
                log(f"[PYTHON] ✓ Saved results to: {output_file}")
            except Exception as save_error:
                import traceback
                error_trace = traceback.format_exc()
                log(f"[PYTHON] ✗ Error saving file: {save_error}")
                log(f"[DEBUG] Save error traceback:\n{error_trace}")
                raise

            summary = _clean_summary(status_messages, f"Total rows: {row_count}")

            log(f"[DEBUG] Building response...")
            log(f"[DEBUG] Summary length: {len(summary)} chars")
            log(f"[DEBUG] File path: {output_file}")
            log(f"[PYTHON] ✓✓✓ ANALYSIS COMPLETE ✓✓✓")

            return {
                "ok": True,
                "response_text": summary,
                "file_path": output_file
            }

        except Exception as e:
            import traceback
            print(f"[PYTHON] Error: {e}")
            traceback.print_exc()
            return {"error": f"Python analysis failed: {str(e)}"}

    def _create_duckdb_schema(self, conn, dataframes: dict) -> str:
        """Build schema description using actual DuckDB column types (post-registration).

        This is more accurate than _create_enhanced_schema which uses pandas dtypes,
        because DuckDB may infer different (more correct) types from the data.
        """
        import pandas as pd
        schema_lines = []
        for table_name, df in dataframes.items():
            schema_lines.append(f"\n=== TABLE: {table_name} ===")
            schema_lines.append(f"Rows: {len(df)}, Columns: {len(df.columns)}")
            # Query DuckDB for the actual column types
            try:
                describe_rows = conn.execute(f"DESCRIBE {table_name}").fetchall()
                # describe_rows: (column_name, column_type, null, key, default, extra)
                duckdb_types = {row[0]: row[1] for row in describe_rows}
            except Exception:
                duckdb_types = {}
            schema_lines.append("\nColumns (DuckDB types):")
            for col in df.columns:
                db_type = duckdb_types.get(col, str(df[col].dtype))
                null_count = df[col].isna().sum()
                null_pct = (null_count / len(df) * 100) if len(df) > 0 else 0
                col_info = f"  - {col}: {db_type} (nulls: {null_count}/{len(df)} = {null_pct:.1f}%)"
                is_numeric = any(t in db_type.upper() for t in ("INT", "FLOAT", "DOUBLE", "DECIMAL", "BIGINT", "HUGEINT", "REAL"))
                if is_numeric:
                    try:
                        min_val = df[col].min()
                        max_val = df[col].max()
                        mean_val = df[col].mean()
                        col_info += f" | range: [{min_val}, {max_val}], mean: {mean_val:.2f}"
                    except Exception:
                        pass
                else:
                    unique_count = df[col].nunique()
                    col_info += f" | unique values: {unique_count}"
                    if unique_count <= 10:
                        samples = df[col].dropna().unique()[:5]
                        col_info += f", examples: {list(samples)}"
                schema_lines.append(col_info)
        return "\n".join(schema_lines)

    def _create_enhanced_schema(self, dataframes: dict) -> str:
        """Create detailed schema description for AI code generation.
        
        Returns ALL columns with data types, value ranges, and sample values
        without truncation so AI understands complete file structure.
        """
        import pandas as pd
        
        schema_lines = []
        
        for table_name, df in dataframes.items():
            schema_lines.append(f"\n=== TABLE: {table_name} ===")
            schema_lines.append(f"Rows: {len(df)}, Columns: {len(df.columns)}")
            schema_lines.append(f"Row Memory Usage: ~{df.memory_usage(deep=True).sum() / 1024:.1f} KB")
            schema_lines.append("\nColumns:")
            
            for col in df.columns:
                dtype = str(df[col].dtype)
                null_count = df[col].isna().sum()
                null_pct = (null_count / len(df) * 100) if len(df) > 0 else 0
                
                # Build column info
                col_info = f"  - {col}: {dtype} (nulls: {null_count}/{len(df)} = {null_pct:.1f}%)"
                
                # Add value info based on data type
                if pd.api.types.is_numeric_dtype(df[col]):
                    try:
                        min_val = df[col].min()
                        max_val = df[col].max()
                        mean_val = df[col].mean()
                        col_info += f" | range: [{min_val}, {max_val}], mean: {mean_val:.2f}"
                    except:
                        pass
                else:
                    # Show sample values for categorical/text columns
                    unique_count = df[col].nunique()
                    col_info += f" | unique values: {unique_count}"
                    if unique_count <= 10:
                        samples = df[col].dropna().unique()[:5]
                        col_info += f", examples: {list(samples)}"
                
                schema_lines.append(col_info)
        
        return "\n".join(schema_lines)

    def _estimate_csv_validity(self, filename: str, file_size: int, content: str = None, content_b64: str = None) -> str:
        """
        Estimate if CSV file is too large BEFORE loading entire file into memory.
        
        FOR LARGE FILES (>100MB):
        - Read only first ~1000 rows to estimate total row count
        - Check if estimated rows exceed 100k limit
        - Return "BATCH_MODE" if file should use batch processing
        - Return error message if file is beyond batch processing capability
        
        Returns:
        - "" (empty string) if file is within normal limits
        - "BATCH_MODE" if file should use batch streaming mode
        - error message if file is too problematic even for batch mode
        """
        import io
        import base64
        
        try:
            # Only for string content or base64 - these we can peek at
            csv_content = None
            
            if content and isinstance(content, str):
                csv_content = content
            elif content_b64:
                try:
                    binary_data = base64.b64decode(content_b64, validate=False)
                    csv_content = binary_data.decode('utf-8', errors='ignore')
                except:
                    return ""  # Can't estimate base64, will fail later and that's OK
            else:
                return ""  # No content to estimate
            
            # Read only first 1000 rows + header to estimate
            lines = csv_content.split('\n')
            sample_lines = lines[:1001]  # header + 1000 rows
            
            # Count columns from header
            if not sample_lines:
                return ""
            
            header = sample_lines[0]
            col_count = len(header.split(','))
            
            # Estimate total rows
            sample_size = len('\n'.join(sample_lines).encode('utf-8'))
            file_size_bytes = len(csv_content.encode('utf-8'))
            
            if sample_size > 0:
                estimated_rows = int((file_size_bytes / sample_size) * len(sample_lines))
                file_size_mb = file_size / (1024 * 1024)
                
                # Check against limits
                if col_count > 100:
                    return f"File '{filename}' has {col_count} columns (limit: 100). " \
                           f"Please remove unnecessary columns."
                
                if estimated_rows > 100000 and file_size_mb <= 5000:  # Up to 5GB can use batch mode
                    # Use batch processing for large files
                    return "BATCH_MODE"
                
                if file_size_mb > 5000:
                    return f"File '{filename}' is {file_size_mb:.0f}MB (limit: 5GB). " \
                           f"This exceeds system capacity even for batch processing."
        
        except Exception as e:
            # If estimation fails, let normal loading handle it
            print(f"[ESTIMATE] Could not estimate {filename}: {e}")
            return ""
        
        return ""  # File appears valid for normal processing

    def _create_dataframe_summary(self, df: pd.DataFrame, max_rows: int = 10) -> str:
        """Create a concise summary of dataframe for AI response without reading entire content.
        
        Shows:
        - Shape (rows, columns)
        - First N rows as formatted table
        - Column data types
        - Basic statistics for numeric columns
        - Value counts for categorical columns (top 5)
        
        Keeps summary under ~1000 tokens to avoid context overflow.
        """
        import pandas as pd
        
        lines = []
        lines.append(f"Total Rows: {len(df)}, Total Columns: {len(df.columns)}")
        lines.append("")
        
        # Show first N rows
        lines.append(f"First {min(max_rows, len(df))} rows:")
        lines.append("-" * 80)
        
        # Format as markdown table if not too many columns
        if len(df.columns) <= 10:
            # Simple markdown table
            header = "| " + " | ".join([str(c)[:20] for c in df.columns]) + " |"
            separator = "|" + "|".join(["---"] * len(df.columns)) + "|"
            lines.append(header)
            lines.append(separator)
            
            for idx, row in df.head(max_rows).iterrows():
                row_str = "| " + " | ".join([str(v)[:15] for v in row]) + " |"
                lines.append(row_str)
        else:
            # For wide dataframes, show key columns only
            key_cols = list(df.columns)[:5]
            subset = df[key_cols].head(max_rows)
            lines.append(subset.to_string())
            lines.append(f"\n... ({len(df.columns) - 5} more columns) ...")
        
        lines.append("")
        lines.append("Column Summary:")
        lines.append("-" * 80)
        
        # Column types and stats
        for col in df.columns[:20]:  # Limit to first 20 columns
            dtype = str(df[col].dtype)
            null_pct = (df[col].isna().sum() / len(df) * 100) if len(df) > 0 else 0
            
            if pd.api.types.is_numeric_dtype(df[col]):
                try:
                    stats = f"{col}: {dtype} | min={df[col].min():.2f}, max={df[col].max():.2f}, mean={df[col].mean():.2f}, nulls={null_pct:.1f}%"
                except:
                    stats = f"{col}: {dtype} | nulls={null_pct:.1f}%"
            else:
                unique = df[col].nunique()
                stats = f"{col}: {dtype} | unique={unique}, nulls={null_pct:.1f}%"
            
            lines.append(stats)
        
        if len(df.columns) > 20:
            lines.append(f"... ({len(df.columns) - 20} more columns)")
        
        return "\n".join(lines)

    def _extract_agent_file_content(self, file_obj: dict) -> str:
        """Extract text content for agent uploads (plain text or base64 binary files).
        For PDFs, uses the full parallel-OCR pipeline with page-level chunks.
        Stores extracted pages in file_obj['_pages'] for downstream batching."""
        content = file_obj.get("content", "")
        if isinstance(content, str) and content.strip():
            return content

        b64 = file_obj.get("content_base64", "")
        if not b64:
            return ""

        try:
            raw = base64.b64decode(b64, validate=False)
        except Exception:
            return ""

        name = str(file_obj.get("name", ""))
        mime_type = str(file_obj.get("mime_type", ""))
        ext = os.path.splitext(name)[1].lower()

        # PDF: use the unified full extractor (parallel OCR + page chunks)
        if ext == ".pdf":
            text, page_list = self._extract_pdf_full(raw_bytes=raw)
            if page_list:
                file_obj["_pages"] = page_list
                print(f"[AGENT-PDF] {name}: {len(page_list)} pages, "
                      f"{len(text)} chars extracted")
            if not text:
                print(f"[AGENT-PDF] {name}: extraction failed — no readable text")
            return text or ""

        extracted = self._extract_text_from_bytes(ext, raw, mime_type)
        return extracted or ""

    def _extract_text_from_bytes(self, ext: str, raw: bytes, mime_type: str = "") -> str:
        """Extract text from binary upload bytes (docs, spreadsheets, PDFs, images)."""
        try:
            if ext in (".txt", ".md", ".csv", ".json"):
                return raw.decode("utf-8", errors="ignore")

            if ext in (".xlsx", ".xls"):
                import pandas as pd
                sheets = pd.read_excel(io.BytesIO(raw), sheet_name=None)
                parts = []
                for name, frame in sheets.items():
                    parts.append(f"Sheet: {name}")
                    parts.append(frame.to_string(index=False))
                return "\n\n".join(parts)

            if ext == ".docx":
                from docx import Document
                doc = Document(io.BytesIO(raw))
                return "\n".join(p.text for p in doc.paragraphs if p.text and p.text.strip())

            if ext == ".pdf":
                text, _pages = self._extract_pdf_text_with_ocr(raw)
                return text

            if ext in (".png", ".jpg", ".jpeg", ".webp", ".bmp", ".tiff"):
                return self._ocr_image_bytes(raw)

            if mime_type.startswith("image/"):
                return self._ocr_image_bytes(raw)
        except Exception:
            return ""
        return ""

    def _extract_pdf_text_with_ocr(self, raw_pdf: bytes) -> tuple:
        """Extract text from PDF bytes using the unified _extract_pdf_full pipeline.
        Returns (full_text, page_list) — same as _extract_pdf_full.
        Uses parallel OCR, page-level chunking, and progress status."""
        return self._extract_pdf_full(raw_bytes=raw_pdf)

    def _ensure_tesseract(self):
        """Find Tesseract OCR binary. Checks bundled location first (PyInstaller),
        then PATH, then common system install paths. No auto-install — everything
        needed should be bundled in the final app."""
        if getattr(self, '_tesseract_resolved', False):
            return not getattr(self, '_tesseract_missing', False)

        import shutil, platform, importlib
        system = platform.system()
        self._tesseract_resolved = True

        # 1. pytesseract Python wrapper must be available (bundled with app)
        try:
            importlib.import_module("pytesseract")
        except ImportError:
            print("[OCR] pytesseract not bundled — OCR disabled.")
            self._tesseract_missing = True
            return False

        pytesseract = importlib.import_module("pytesseract")

        # 2. Determine app base directory (PyInstaller or dev)
        if getattr(sys, 'frozen', False):
            # Running as PyInstaller bundle
            app_dir = os.path.dirname(sys.executable)
        else:
            app_dir = os.path.dirname(os.path.abspath(__file__))

        # 3. Search order: bundled → PATH → common system paths
        candidates = []
        if system == "Windows":
            candidates = [
                os.path.join(app_dir, "tesseract", "tesseract.exe"),           # bundled
                os.path.join(app_dir, "Tesseract-OCR", "tesseract.exe"),       # bundled alt
            ]
        else:
            candidates = [
                os.path.join(app_dir, "tesseract", "tesseract"),               # bundled
            ]

        found = None
        for c in candidates:
            if os.path.isfile(c):
                found = c
                break

        # Check PATH
        if not found:
            found = shutil.which("tesseract")

        # System-wide fallback paths
        if not found:
            if system == "Windows":
                for c in [
                    os.path.join(os.environ.get("PROGRAMFILES", r"C:\Program Files"), "Tesseract-OCR", "tesseract.exe"),
                    os.path.join(os.environ.get("LOCALAPPDATA", ""), "Tesseract-OCR", "tesseract.exe"),
                    os.path.join(os.environ.get("PROGRAMFILES(X86)", r"C:\Program Files (x86)"), "Tesseract-OCR", "tesseract.exe"),
                ]:
                    if c and os.path.isfile(c):
                        found = c
                        break
            elif system == "Darwin":
                for c in ["/usr/local/bin/tesseract", "/opt/homebrew/bin/tesseract"]:
                    if os.path.isfile(c):
                        found = c
                        break
            else:
                for c in ["/usr/bin/tesseract", "/usr/local/bin/tesseract"]:
                    if os.path.isfile(c):
                        found = c
                        break

        if found:
            pytesseract.pytesseract.tesseract_cmd = found
            print(f"[OCR] Using Tesseract at: {found}")
            return True
        else:
            self._tesseract_missing = True
            print("[OCR] Tesseract not found — scanned/image OCR disabled. Text-based PDFs still work.")
            return False

    def _ocr_image_bytes(self, raw_image: bytes) -> str:
        """OCR text extraction for image bytes using pytesseract if available."""
        if not self._ensure_tesseract():
            return ""

        try:
            from PIL import Image
            import importlib
            pytesseract = importlib.import_module("pytesseract")

            image = Image.open(io.BytesIO(raw_image))
            # Convert to RGB for broader compatibility.
            if image.mode not in ("RGB", "L"):
                image = image.convert("RGB")
            text = pytesseract.image_to_string(image)
            return (text or "").strip()
        except Exception:
            return ""

    def _parse_ai_response(self, response: str) -> dict:
        """Parse AI response to extract structured data (markdown tables).
        
        Also extracts section titles from preceding markdown headers (##, ###).
        Each table is stored with its title for multi-file export.
        """
        parsed = {
            "tables": [],
            "text": response,
            "has_tables": False
        }
        
        if not response:
            return parsed
        
        # Split response into lines for table detection
        lines = response.split('\n')
        i = 0
        last_heading = None
        
        while i < len(lines):
            line = lines[i].strip()
            
            # Capture markdown heading as potential table title
            if line.startswith('##') or line.startswith('###'):
                # Extract heading text (remove # symbols)
                last_heading = re.sub(r'^#+\s*', '', line).strip()
                i += 1
                continue
            
            # Look for table start (line starting with |)
            if line.startswith('|') and '|' in line:
                table_lines = [line]
                i += 1
                
                # Collect consecutive table lines
                while i < len(lines):
                    next_line = lines[i].strip()
                    if next_line.startswith('|'):
                        table_lines.append(next_line)
                        i += 1
                    else:
                        break
                
                # Parse table if we have at least header + separator + 1 data row
                if len(table_lines) >= 2:
                    try:
                        # Extract headers
                        def get_cells(line):
                            cells = [c.strip() for c in line.split('|')]
                            return [c for c in cells if c and not re.match(r'^-+$|^:?-+:?$', c)]
                        
                        headers = get_cells(table_lines[0])
                        data_start_row = 1
                        
                        # Check if second line is separator (---)
                        if len(table_lines) > 1:
                            second_line = table_lines[1]
                            if re.match(r'^[\|\s\-:]+$', second_line):
                                data_start_row = 2
                        
                        # Extract data rows
                        rows = []
                        for row_idx in range(data_start_row, len(table_lines)):
                            cells = get_cells(table_lines[row_idx])
                            if cells:
                                # Pad row to match header length
                                while len(cells) < len(headers):
                                    cells.append("")
                                rows.append(cells[:len(headers)])
                        
                        # Only add table if we have both headers and data
                        if headers and rows:
                            parsed["tables"].append({
                                "headers": headers,
                                "rows": rows,
                                "title": last_heading or f"Table {len(parsed['tables']) + 1}"
                            })
                            parsed["has_tables"] = True
                            last_heading = None  # Reset after using
                    except Exception as e:
                        print(f"[PARSE] Table parsing error: {e}")
                        i += 1
            else:
                i += 1
        
        return parsed

    def _get_prompt_ending(self, output_format: str) -> str:
        """Return a format-aware prompt ending that primes the model to generate
        actual content rather than looping the heading (Gemma 4 / small models)."""
        fmt = (output_format or "").lower()
        if fmt in ("pdf", "docx"):
            # Completion-primer approach: incomplete sentence the model must finish.
            # This anchors small models to letter output rather than echoing sources.
            return "Write the complete formal letter. Begin directly with the date:"
        elif fmt in ("csv_json", "json"):
            # Structured JSON expected — prime the JSON object
            return '## Output\n\n```json\n{"rows": [\n'
        elif fmt in ("excel", "xlsx", "csv"):
            # Table output — prime the first pipe to start a markdown table
            return "## Output\n\n| "
        else:
            # Default (txt) — neutral heading with blank line
            return "## Output\n\n"

    def _estimate_max_tokens_dynamic(self, prompt_chars: int, output_format: str) -> int:
        """Estimate max_tokens dynamically based on input size and format.
        
        Uses character-based approximation (~1 token per 4 chars for English text).
        Applies format-specific output ratios with safety ceilings to prevent
        over-generation or truncation based on actual content size.
        
        Args:
            prompt_chars: Total character count of the prompt (system + user + files)
            output_format: Output format string (pdf, docx, csv, json, excel, etc.)
        
        Returns:
            Calculated max_tokens with format-specific safety ceiling applied
        """
        from math import ceil
        
        # Approximate input tokens: ~1 token per 4 characters in English
        input_tokens = max(1, prompt_chars // 4)
        
        # Format-specific scaling ratios and safety ceilings
        _fmt = (output_format or "").lower()
        
        if _fmt in ("pdf", "docx"):
            # Letter/document output: typically 1.0-1.3x input
            # Small prompt = small document, large prompt = compliance letter with background
            ratio = 1.2
            ceiling = 1500  # Lifted from hardcoded 800 to handle compliance docs
            
        elif _fmt in ("csv_json", "json"):
            # Structured data: can expand with many rows
            ratio = 2.0
            ceiling = 2400  # Lifted from hardcoded 1200
            
        elif _fmt in ("excel", "xlsx", "csv"):
            # Tables: similar expansion potential
            ratio = 2.0
            ceiling = 2400  # Lifted from hardcoded 1500
            
        else:
            # Default (txt, etc.)
            ratio = 2.0
            ceiling = 2048  # Kept same as before
        
        # Calculate: input_tokens * format_ratio, capped at ceiling
        calculated = int(ceil(input_tokens * ratio))
        result = min(calculated, ceiling)
        
        print(f"[MAX_TOKENS] Auto-scaled: {prompt_chars} prompt_chars -> {input_tokens} input_tokens "
              f"-> {calculated} (ratio={ratio}) -> {result} (ceiling={ceiling}) [format={_fmt}]")
        
        return result

    def _sanitize_filename(self, text: str) -> str:
        """Convert table title to safe filename."""
        # Replace spaces and special chars with underscores
        safe = re.sub(r'[^a-zA-Z0-9\s]', '', text)
        safe = re.sub(r'\s+', '_', safe).lower()
        return safe[:50]  # Limit length

    def _create_multiple_output_files(self, tables: list, output_format: str, 
                                      output_dir: str, timestamp: int) -> list[str]:
        """Create separate output files for each table (one per table).
        
        Supports CSV, XLSX (sheets), DOCX, PDF, TXT formats.
        Returns list of created file paths.
        """
        if not tables:
            return []
        
        created_files = []
        
        if output_format == "xlsx":
            # Excel: all tables as separate sheets in ONE file
            try:
                from openpyxl import Workbook
                from openpyxl.styles import Font, PatternFill
                
                wb = Workbook()
                wb.remove(wb.active)  # Remove default sheet
                
                for table in tables:
                    title = table.get("title", "Table")
                    headers = table.get("headers", [])
                    rows = table.get("rows", [])
                    
                    # Create sheet with table title
                    ws = wb.create_sheet(title=title[:31])  # Sheet name max 31 chars
                    
                    # Add headers
                    for col, header in enumerate(headers, 1):
                        cell = ws.cell(row=1, column=col)
                        cell.value = header
                        cell.font = Font(bold=True)
                        cell.fill = PatternFill(start_color="CCCCCC", end_color="CCCCCC", fill_type="solid")
                    
                    # Add data rows
                    for row_idx, row_data in enumerate(rows, 2):
                        for col_idx, value in enumerate(row_data, 1):
                            ws.cell(row=row_idx, column=col_idx).value = value
                    
                    # Auto-adjust column widths
                    for col in ws.columns:
                        max_length = 0
                        col_letter = col[0].column_letter
                        for cell in col:
                            try:
                                if len(str(cell.value or '')) > max_length:
                                    max_length = len(str(cell.value or ''))
                            except:
                                pass
                        ws.column_dimensions[col_letter].width = min(50, max_length + 2)
                
                output_file = os.path.join(output_dir, f"analysis_{timestamp}.xlsx")
                wb.save(output_file)
                created_files.append(output_file)
                print(f"[PROCESS] Excel with {len(tables)} sheets created: {output_file}")
            except Exception as e:
                print(f"[PROCESS] Multi-sheet Excel error: {e}")
        
        else:
            # CSV, DOCX, PDF, TXT: separate file per table
            for table in tables:
                title = table.get("title", "Table")
                headers = table.get("headers", [])
                rows = table.get("rows", [])
                safe_name = self._sanitize_filename(title)
                
                try:
                    if output_format == "csv":
                        output_file = os.path.join(output_dir, f"analysis_{timestamp}_{safe_name}.csv")
                        with open(output_file, 'w', newline='', encoding='utf-8') as f:
                            writer = csv.writer(f, quoting=csv.QUOTE_ALL)
                            writer.writerow(headers)
                            writer.writerows(rows)
                        created_files.append(output_file)
                        print(f"[PROCESS] CSV created: {output_file}")
                    
                    elif output_format == "docx":
                        from docx import Document
                        from docx.enum.text import WD_ALIGN_PARAGRAPH
                        
                        output_file = os.path.join(output_dir, f"analysis_{timestamp}_{safe_name}.docx")
                        doc = Document()
                        doc.add_heading(title, level=1)
                        
                        if headers and rows:
                            table_obj = doc.add_table(rows=len(rows) + 1, cols=len(headers))
                            table_obj.style = 'Light Grid Accent 1'
                            
                            # Headers
                            for col, header in enumerate(headers):
                                table_obj.rows[0].cells[col].text = str(header)
                                for para in table_obj.rows[0].cells[col].paragraphs:
                                    for run in para.runs:
                                        run.font.bold = True
                            
                            # Data rows
                            for row_idx, row_data in enumerate(rows, 1):
                                for col_idx, value in enumerate(row_data):
                                    table_obj.rows[row_idx].cells[col_idx].text = str(value) if value else ""
                        
                        doc.save(output_file)
                        created_files.append(output_file)
                        print(f"[PROCESS] DOCX created: {output_file}")
                    
                    elif output_format == "pdf":
                        from reportlab.lib.pagesizes import letter
                        from reportlab.lib.styles import getSampleStyleSheet
                        from reportlab.lib.units import inch
                        from reportlab.platypus import SimpleDocTemplate, Table, TableStyle, Paragraph, Spacer
                        from reportlab.lib import colors
                        
                        output_file = os.path.join(output_dir, f"analysis_{timestamp}_{safe_name}.pdf")
                        doc = SimpleDocTemplate(output_file, pagesize=letter)
                        story = []
                        styles = getSampleStyleSheet()
                        
                        story.append(Paragraph(title, styles['Heading1']))
                        story.append(Spacer(1, 0.2*inch))
                        
                        if headers and rows:
                            table_data = [headers] + rows[:50]
                            pdf_table = Table(table_data, colWidths=[7.5*inch/len(headers) for _ in headers])
                            pdf_table.setStyle(TableStyle([
                                ('BACKGROUND', (0, 0), (-1, 0), colors.HexColor('#e8e8e8')),
                                ('ALIGN', (0, 0), (-1, -1), 'LEFT'),
                                ('FONTNAME', (0, 0), (-1, 0), 'Helvetica-Bold'),
                                ('GRID', (0, 0), (-1, -1), 0.5, colors.black),
                            ]))
                            story.append(pdf_table)
                        
                        doc.build(story)
                        created_files.append(output_file)
                        print(f"[PROCESS] PDF created: {output_file}")
                    
                    elif output_format == "txt":
                        output_file = os.path.join(output_dir, f"analysis_{timestamp}_{safe_name}.txt")
                        with open(output_file, 'w', encoding='utf-8') as f:
                            f.write(f"{title}\n")
                            f.write("=" * 80 + "\n\n")
                            f.write(" | ".join(headers) + "\n")
                            f.write("-" * 80 + "\n")
                            for row in rows:
                                f.write(" | ".join(str(v) for v in row) + "\n")
                        created_files.append(output_file)
                        print(f"[PROCESS] TXT created: {output_file}")
                
                except Exception as e:
                    print(f"[PROCESS] Error creating {output_format} for {title}: {e}")
        
        return created_files
        """Create Excel output file."""
        try:
            from openpyxl import Workbook
            from openpyxl.styles import Font, PatternFill
            
            wb = Workbook()
            ws = wb.active
            ws.title = "Analysis Results"
            
            ws['A1'] = f"Generated: {time.strftime('%Y-%m-%d %H:%M:%S')}"
            
            # Add tables if any
            if data.get("tables"):
                current_row = 4
                for i, table in enumerate(data["tables"]):
                    ws[f'A{current_row}'] = f"Table {i+1}"
                    ws[f'A{current_row}'].font = Font(bold=True)
                    current_row += 1
                    
                    # Headers
                    for col, header in enumerate(table["headers"], 1):
                        cell = ws.cell(row=current_row, column=col)
                        cell.value = header
                        cell.font = Font(bold=True)
                        cell.fill = PatternFill(start_color="CCCCCC", end_color="CCCCCC", fill_type="solid")
                    
                    current_row += 1
                    
                    # Rows
                    for row_data in table["rows"]:
                        for col, value in enumerate(row_data, 1):
                            ws.cell(row=current_row, column=col).value = value
                        current_row += 1
                    
                    current_row += 2
            else:
                # Add full response as text
                ws['A4'] = response
                ws['A4'].alignment = {'wrap_text': True}
            
            # Auto-adjust column widths
            for column in ws.columns:
                max_length = 0
                column_letter = column[0].column_letter
                for cell in column:
                    try:
                        if len(str(cell.value or '')) > max_length:
                            max_length = len(str(cell.value or ''))
                    except:
                        pass
                ws.column_dimensions[column_letter].width = min(50, max_length + 2)
            
            output_file = os.path.join(output_dir, f"analysis_{timestamp}.xlsx")
            wb.save(output_file)
            return output_file
            
        except Exception as e:
            print(f"[PROCESS] Excel creation error: {e}")
            # Fallback to CSV
            return self._create_csv_output(data, response, output_dir, timestamp)

    def _create_csv_output(self, data: dict, response: str, output_dir: str, timestamp: int) -> str:
        """Create CSV output file with properly formatted analysis results."""
        try:
            output_file = os.path.join(output_dir, f"analysis_{timestamp}.csv")
            
            with open(output_file, 'w', newline='', encoding='utf-8') as f:
                writer = csv.writer(f, quoting=csv.QUOTE_ALL)
                
                # Header section
                writer.writerow(["Generated", time.strftime('%Y-%m-%d %H:%M:%S')])
                writer.writerow([])
                
                # Write extracted tables if available
                if data.get("tables") and len(data["tables"]) > 0:
                    for i, table in enumerate(data["tables"], 1):
                        writer.writerow([f"Table {i}: Analysis Results"])
                        
                        # Headers
                        if table.get("headers"):
                            writer.writerow(table["headers"])
                        
                        # Data rows
                        if table.get("rows"):
                            for row in table["rows"]:
                                # Ensure all cells are strings for proper CSV formatting
                                writer.writerow([str(cell).strip() if cell else "" for cell in row])
                        
                        writer.writerow([])  # Blank row between tables
                
                # If no structured tables found, parse response for content sections
                else:
                    # Split response into sections based on common patterns
                    writer.writerow(["Analysis Summary"])
                    writer.writerow([])
                    
                    # Try to extract sections from response
                    lines = response.split('\n')
                    current_section = []
                    
                    for line in lines:
                        line = line.strip()
                        
                        # Skip markdown formatting
                        if line.startswith('#'):
                            if current_section:
                                writer.writerows([[item] for item in current_section])
                                current_section = []
                            # Extract section title (remove # symbols)
                            section_title = re.sub(r'^#+\s*', '', line)
                            writer.writerow([section_title])
                        elif line:
                            current_section.append(line)
                    
                    # Write remaining content
                    if current_section:
                        writer.writerows([[item] for item in current_section])
            
            print(f"[PROCESS] CSV output created: {output_file}")
            return output_file
            
        except Exception as e:
            print(f"[PROCESS] CSV creation error: {e}")
            import traceback
            traceback.print_exc()
            # Fallback to TXT
            return self._create_txt_output(data, response, output_dir, timestamp)

    def _create_docx_output(self, data: dict, response: str, output_dir: str, timestamp: int) -> str:
        """Create DOCX output file with tables and formatted content."""
        try:
            from docx import Document
            from docx.shared import Pt, RGBColor, Inches
            from docx.enum.text import WD_ALIGN_PARAGRAPH
            
            output_file = os.path.join(output_dir, f"analysis_{timestamp}.docx")
            doc = Document()
            
            # Clean response for document export (removes markdown, system instructions)
            cleaned_response = self._clean_response_for_document_export(response, "docx")
            if not cleaned_response.strip() and str(response or "").strip():
                print("[PROCESS] DOCX clean pass produced empty text, using raw response fallback")
                cleaned_response = str(response or "").strip()
            
            # Timestamp
            timestamp_para = doc.add_paragraph(f"Generated: {time.strftime('%Y-%m-%d %H:%M:%S')}")
            timestamp_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            
            # Add tables if any
            if data.get("tables"):
                for i, table_data in enumerate(data["tables"], 1):
                    doc.add_heading(f'Table {i}', level=2)
                    
                    # Create table with headers + rows
                    headers = table_data.get("headers", [])
                    rows = table_data.get("rows", [])
                    
                    if headers and rows:
                        # Create table (rows + 1 for header)
                        table = doc.add_table(rows=len(rows) + 1, cols=len(headers))
                        table.style = 'Light Grid Accent 1'
                        
                        # Add headers
                        header_cells = table.rows[0].cells
                        for col, header in enumerate(headers):
                            header_cells[col].text = str(header)
                            # Bold header text
                            for paragraph in header_cells[col].paragraphs:
                                for run in paragraph.runs:
                                    run.font.bold = True
                        
                        # Add data rows
                        for row_idx, row_data in enumerate(rows, start=1):
                            row_cells = table.rows[row_idx].cells
                            for col, value in enumerate(row_data):
                                row_cells[col].text = str(value) if value else ""
                    else:
                        doc.add_paragraph("(No data available)")
                    
                    doc.add_paragraph()  # Spacing
            else:
                # Add cleaned response as paragraphs (removes markdown, system instructions)
                doc.add_heading('Analysis Results', level=2)
                for paragraph_text in cleaned_response.split('\n'):
                    if paragraph_text.strip():
                        doc.add_paragraph(paragraph_text)
            
            doc.save(output_file)
            print(f"[PROCESS] DOCX output created: {output_file}")
            return output_file
            
        except ImportError:
            print(f"[PROCESS] python-docx not installed, falling back to PDF")
            return self._create_pdf_output(data, response, output_dir, timestamp)
        except Exception as e:
            print(f"[PROCESS] DOCX creation error: {e}")
            # Fallback to PDF
            return self._create_pdf_output(data, response, output_dir, timestamp)

    def _clean_response_for_document_export(self, response: str, output_format: str) -> str:
        """Clean response text for PDF/DOCX export by removing artifacts and system instructions."""
        fmt = (output_format or "").lower()
        if fmt not in ("pdf", "docx"):
            return response
        
        # Remove markdown headers (###, ##, #)
        text = re.sub(r'^#+\s+', '', response, flags=re.MULTILINE)
        
        # Remove markdown bold/italic (**text**, *text*, __text__)
        text = re.sub(r'\*\*([^*]+)\*\*', r'\1', text)
        text = re.sub(r'__([^_]+)__', r'\1', text)
        text = re.sub(r'\*([^*]+)\*', r'\1', text)
        text = re.sub(r'_([^_]+)_', r'\1', text)
        
        # Remove code blocks (```...```)
        text = re.sub(r'```[\s\S]*?```', '', text)
        
        # Remove horizontal rules and file markers
        text = re.sub(r'^-{3,}$', '', text, flags=re.MULTILINE)
        text = re.sub(r'^_+$', '', text, flags=re.MULTILINE)
        text = re.sub(r'---\s*FILE:.*?---\s*END FILE\s*---', '', text, flags=re.DOTALL | re.IGNORECASE)
        text = re.sub(r'---\s*FILE:.*', '', text)
        
        # Remove system instruction lines
        text = re.sub(r'^\*\*Return:\*\*.*', '', text, flags=re.MULTILINE | re.IGNORECASE)
        text = re.sub(r'^Return:.*', '', text, flags=re.MULTILINE | re.IGNORECASE)
        text = re.sub(r'^\*\*CRITICAL:.*', '', text, flags=re.MULTILINE | re.IGNORECASE)
        text = re.sub(r'^RULES:.*', '', text, flags=re.MULTILINE)
        
        # Remove "Analysis Results" header (not needed for letters)
        text = re.sub(r'^\s*Analysis Results\s*$', '', text, flags=re.MULTILINE)
        
        # Clean excessive blank lines (reduce to max 2 consecutive)
        text = re.sub(r'\n{3,}', '\n\n', text)
        
        return text.strip()

    def _create_pdf_output(self, data: dict, response: str, output_dir: str, timestamp: int) -> str:
        """Create PDF output file."""
        try:
            from reportlab.lib.pagesizes import letter
            from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
            from reportlab.lib.units import inch
            from reportlab.platypus import SimpleDocTemplate, Table, TableStyle, Paragraph, Spacer
            from reportlab.lib import colors
            
            output_file = os.path.join(output_dir, f"analysis_{timestamp}.pdf")
            doc = SimpleDocTemplate(output_file, pagesize=letter)
            story = []
            styles = getSampleStyleSheet()
            
            # Clean response for document export (removes markdown, system instructions)
            cleaned_response = self._clean_response_for_document_export(response, "pdf")
            if not cleaned_response.strip() and str(response or "").strip():
                print("[PROCESS] PDF clean pass produced empty text, using raw response fallback")
                cleaned_response = str(response or "").strip()
            
            # Timestamp
            story.append(Paragraph(f"Generated: {time.strftime('%Y-%m-%d %H:%M:%S')}", styles['Normal']))
            story.append(Spacer(1, 0.2*inch))
            
            # Add tables if any
            if data.get("tables"):
                for i, table in enumerate(data["tables"]):
                    story.append(Paragraph(f"Table {i+1}", styles['Heading2']))
                    
                    # Build table
                    table_data = [table["headers"]] + table["rows"][:50]  # Limit rows
                    
                    pdf_table = Table(table_data, colWidths=[7.5*inch/len(table["headers"]) for _ in table["headers"]])
                    pdf_table.setStyle(TableStyle([
                        ('BACKGROUND', (0, 0), (-1, 0), colors.HexColor('#e8e8e8')),
                        ('ALIGN', (0, 0), (-1, -1), 'LEFT'),
                        ('FONTNAME', (0, 0), (-1, 0), 'Helvetica-Bold'),
                        ('GRID', (0, 0), (-1, -1), 0.5, colors.black),
                    ]))
                    
                    story.append(pdf_table)
                    story.append(Spacer(1, 0.2*inch))
            else:
                # Add cleaned response as paragraphs (for letters/documents)
                # Split by double newlines to preserve paragraph structure
                paragraphs = cleaned_response.split('\n\n')
                for para_text in paragraphs:
                    para_text = para_text.strip()
                    if para_text:
                        # Split by single newlines within paragraph to preserve structure
                        lines = para_text.split('\n')
                        for line in lines:
                            line = line.strip()
                            if line:
                                story.append(Paragraph(line, styles['Normal']))
                        story.append(Spacer(1, 0.1*inch))  # Space between paragraphs
            
            doc.build(story)
            return output_file
            
        except Exception as e:
            print(f"[PROCESS] PDF creation error: {e}")
            # Fallback to TXT
            return self._create_txt_output(data, response, output_dir, timestamp)

    def _create_txt_output(self, data: dict, response: str, output_dir: str, timestamp: int) -> str:
        """Create text file output."""
        try:
            # Clean response for document export
            cleaned_response = self._clean_response_for_document_export(response, "txt")
            
            output_file = os.path.join(output_dir, f"analysis_{timestamp}.txt")
            
            with open(output_file, 'w', encoding='utf-8') as f:
                f.write(f"Generated: {time.strftime('%Y-%m-%d %H:%M:%S')}\n\n")
                
                if data.get("tables"):
                    for i, table in enumerate(data["tables"]):
                        f.write(f"\nTABLE {i+1}:\n")
                        f.write("-" * 80 + "\n")
                        
                        # Headers
                        f.write(" | ".join(table["headers"]) + "\n")
                        f.write("-" * 80 + "\n")
                        
                        # Rows
                        for row in table["rows"]:
                            f.write(" | ".join(str(v) for v in row) + "\n")
                else:
                    f.write("\nANALYSIS:\n")
                    f.write("-" * 80 + "\n")
                    f.write(cleaned_response)
            
            return output_file
            
        except Exception as e:
            print(f"[PROCESS] TXT creation error: {e}")
            raise

    def _is_tabular_file(self, name: str) -> bool:
        ext = os.path.splitext(str(name or ""))[1].lower()
        return ext in (".csv", ".xlsx", ".xls")

    def _should_use_tabular_pipeline(self, files: list) -> bool:
        if not files:
            return False
        # Use SQL/code pipeline if ANY file is tabular (CSV/Excel).
        # Non-tabular files (txt, pdf, docx) are injected as context text into the SQL prompt.
        return any(self._is_tabular_file(f.get("name", "")) for f in files)

    def _dq(self, identifier: str) -> str:
        return '"' + str(identifier).replace('"', '""') + '"'

    def _normalize_column_names(self, columns: list[str]) -> list[str]:
        normalized = []
        used = set()
        for idx, column in enumerate(columns, start=1):
            base = re.sub(r"[^a-z0-9]+", "_", str(column or "").strip().lower()).strip("_")
            if not base:
                base = f"col_{idx}"
            if base[0].isdigit():
                base = f"c_{base}"
            name = base
            suffix = 2
            while name in used:
                name = f"{base}_{suffix}"
                suffix += 1
            used.add(name)
            normalized.append(name)
        return normalized

    def _load_tabular_dataframe(self, file_info: dict):
        import pandas as pd

        name = str(file_info.get("name", "file"))
        path = file_info.get("path")
        ext = os.path.splitext(name)[1].lower()

        # Accept sheet_names as a list of sheet names to load (optional)
        sheet_names = file_info.get("sheet_names")  # None, str, or list
        if isinstance(sheet_names, str):
            sheet_names = [sheet_names]

        def finalize(frame, sheet_name: str | None = None):
            frame = frame.copy()
            frame.columns = self._normalize_column_names(list(frame.columns))
            if sheet_name:
                frame["_sheet_name"] = sheet_name
            frame = frame.where(pd.notnull(frame), None)
            return frame

        if ext == ".csv":
            if path and os.path.exists(path):
                df = pd.read_csv(path, dtype=str, keep_default_na=False, low_memory=False)
            elif file_info.get("content"):
                df = pd.read_csv(io.StringIO(file_info["content"]), dtype=str, keep_default_na=False, low_memory=False)
            elif file_info.get("content_base64"):
                raw = base64.b64decode(file_info["content_base64"], validate=False)
                df = pd.read_csv(io.StringIO(raw.decode("utf-8", errors="ignore")), dtype=str, keep_default_na=False, low_memory=False)
            else:
                raise ValueError(f"No CSV content available for {name}")
            return finalize(df)

        if ext in (".xlsx", ".xls"):
            # If sheet_names is provided, only load those sheets
            if path and os.path.exists(path):
                if sheet_names:
                    sheets = pd.read_excel(path, sheet_name=sheet_names, dtype=str)
                else:
                    sheets = pd.read_excel(path, sheet_name=None, dtype=str)
            elif file_info.get("content_base64"):
                raw = base64.b64decode(file_info["content_base64"], validate=False)
                if sheet_names:
                    sheets = pd.read_excel(io.BytesIO(raw), sheet_name=sheet_names, dtype=str)
                else:
                    sheets = pd.read_excel(io.BytesIO(raw), sheet_name=None, dtype=str)
            else:
                raise ValueError(f"No spreadsheet content available for {name}")

            # If only one sheet is selected, pd.read_excel returns a DataFrame, else dict
            if isinstance(sheets, pd.DataFrame):
                return finalize(sheets, sheet_names[0] if sheet_names else None)

            parts = []
            for sheet_name, frame in sheets.items():
                parts.append(finalize(frame, str(sheet_name)))
            if not parts:
                raise ValueError(f"No sheets found in {name}")
            return pd.concat(parts, ignore_index=True, sort=False)

        raise ValueError(f"Unsupported tabular format: {name}")

    def _safe_duckdb_df(self, conn, sql: str, step: str, warnings: list[str]):
        try:
            return conn.execute(sql).df()
        except Exception as e:
            warnings.append(f"{step} failed: {e}")
            return None

    def _df_to_table(self, frame, max_rows: int = 100) -> dict:
        if frame is None:
            return {"headers": [], "rows": []}
        limited = frame.head(max_rows)
        return {
            "headers": [str(col) for col in limited.columns],
            "rows": [["" if value is None else str(value) for value in row] for row in limited.values.tolist()],
        }

    def _infer_join_key(self, conn, left_meta: dict, right_meta: dict, warnings: list[str]) -> str | None:
        common = [col for col in left_meta["columns"] if col in right_meta["columns"]]
        if not common:
            return None

        preferred = [
            "order_id", "invoice_id", "transaction_id", "record_id", "customer_id",
            "account_id", "reference_id", "ref_id", "id"
        ]
        ranked = []
        for col in common:
            base_score = 0
            if col in preferred:
                base_score = 100 - preferred.index(col)
            elif col.endswith("_id"):
                base_score = 85
            elif col == "id":
                base_score = 80
            elif "id" in col:
                base_score = 70

            uniqueness_score = 0.0
            try:
                left_sql = f"SELECT COUNT(DISTINCT {self._dq(col)})::DOUBLE / NULLIF(COUNT(*), 0) AS ratio FROM {left_meta['table_name']}"
                right_sql = f"SELECT COUNT(DISTINCT {self._dq(col)})::DOUBLE / NULLIF(COUNT(*), 0) AS ratio FROM {right_meta['table_name']}"
                left_ratio = conn.execute(left_sql).fetchone()[0] or 0.0
                right_ratio = conn.execute(right_sql).fetchone()[0] or 0.0
                uniqueness_score = (float(left_ratio) + float(right_ratio)) * 10.0
            except Exception as e:
                warnings.append(f"Join key scoring for {col} failed: {e}")

            ranked.append((base_score + uniqueness_score, col))

        ranked.sort(reverse=True)
        return ranked[0][1] if ranked else None

    def _build_single_table_analysis(self, conn, meta: dict, prompt_text: str, warnings: list[str]) -> dict:
        table_name = meta["table_name"]
        rows = meta["row_count"]
        columns = meta["columns"]
        key = self._infer_join_key(conn, meta, meta, warnings)

        summary_rows = [
            ["File", meta["display_name"]],
            ["Rows", str(rows)],
            ["Columns", str(len(columns))],
            ["Likely key column", key or "Not inferred"],
        ]

        duplicate_count = "n/a"
        if key:
            dup_df = self._safe_duckdb_df(
                conn,
                f"SELECT COUNT(*) AS duplicate_keys FROM (SELECT {self._dq(key)} FROM {table_name} GROUP BY 1 HAVING COUNT(*) > 1)",
                "duplicate key check",
                warnings,
            )
            if dup_df is not None and not dup_df.empty:
                duplicate_count = str(dup_df.iloc[0]["duplicate_keys"])
                summary_rows.append(["Duplicate key groups", duplicate_count])

        sample_df = self._safe_duckdb_df(conn, f"SELECT * FROM {table_name} LIMIT 15", "sample rows", warnings)
        column_df = None
        if columns:
            preview_cols = columns[:25]
            column_df = __import__("pandas").DataFrame({
                "column_name": preview_cols,
                "position": list(range(1, len(preview_cols) + 1)),
            })

        response_lines = [
            f"Structured tabular analysis for {meta['display_name']}.",
            f"Rows: {rows}",
            f"Columns: {len(columns)}",
            f"Likely key column: {key or 'not inferred'}",
        ]
        if duplicate_count != "n/a":
            response_lines.append(f"Duplicate key groups: {duplicate_count}")
        if prompt_text:
            response_lines.append(f"Question handled against structured tabular context: {prompt_text}")
        if warnings:
            response_lines.append("Warnings: " + "; ".join(warnings[:4]))

        tables = [
            {"headers": ["Metric", "Value"], "rows": summary_rows},
        ]
        if column_df is not None:
            tables.append(self._df_to_table(column_df, max_rows=25))
        if sample_df is not None and not sample_df.empty:
            tables.append(self._df_to_table(sample_df, max_rows=15))

        return {
            "ok": True,
            "summary_text": "\n".join(response_lines),
            "warning": "\n".join(warnings) if warnings else "",
            "data": {"tables": tables, "text": "\n".join(response_lines), "has_tables": bool(tables)},
        }

    def _build_reconciliation_analysis(self, conn, metas: list[dict], prompt_text: str, warnings: list[str]) -> dict:
        left_meta, right_meta = metas[0], metas[1]
        left_table = left_meta["table_name"]
        right_table = right_meta["table_name"]

        # Exact column matches
        common_cols = [col for col in left_meta["columns"] if col in right_meta["columns"]]

        # If no exact matches, try fuzzy column matching (strip underscores for comparison)
        col_rename_map = {}  # right_col -> left_col  (rename right to match left)
        if not common_cols:
            def _strip_key(name: str) -> str:
                return name.lower().replace('_', '').replace('-', '').strip()

            left_stripped = {_strip_key(c): c for c in left_meta["columns"]}
            for rc in right_meta["columns"]:
                rk = _strip_key(rc)
                if rk in left_stripped and left_stripped[rk] != rc:
                    col_rename_map[rc] = left_stripped[rk]

            if col_rename_map:
                # Rename right table columns to match left
                rename_parts = []
                for rc in right_meta["columns"]:
                    if rc in col_rename_map:
                        rename_parts.append(f'{self._dq(rc)} AS {self._dq(col_rename_map[rc])}')
                    else:
                        rename_parts.append(self._dq(rc))
                try:
                    conn.execute(
                        f"CREATE OR REPLACE TEMP TABLE {right_table}_renamed AS "
                        f"SELECT {', '.join(rename_parts)} FROM {right_table}"
                    )
                    # Update meta to reflect renamed columns
                    new_right_cols = [col_rename_map.get(c, c) for c in right_meta["columns"]]
                    right_meta = dict(right_meta)
                    right_meta["columns"] = new_right_cols
                    right_table = f"{right_meta['table_name']}_renamed"
                    right_meta["table_name"] = right_table
                    common_cols = [col for col in left_meta["columns"] if col in new_right_cols]
                    matched_names = [f"{rc} -> {lc}" for rc, lc in col_rename_map.items()]
                    warnings.append(f"Fuzzy column matching applied: {'; '.join(matched_names)}")
                except Exception as e:
                    warnings.append(f"Fuzzy column rename failed: {e}")

        if not common_cols:
            warnings.append("No common normalized columns found between the two tabular files.")
            combined_text = (
                f"Could not reconcile {left_meta['display_name']} and {right_meta['display_name']} automatically. "
                "No common normalized columns were found."
            )
            return {
                "ok": False,
                "error": combined_text,
                "warning": "\n".join(warnings) if warnings else "",
            }

        join_key = self._infer_join_key(conn, left_meta, right_meta, warnings)
        if not join_key:
            return {
                "ok": False,
                "error": "Could not infer a reliable join key for reconciliation.",
                "warning": "\n".join(warnings) if warnings else "",
            }

        left_cols_sql = ", ".join(self._dq(col) for col in left_meta["columns"])
        right_cols_sql = ", ".join(self._dq(col) for col in right_meta["columns"])
        join_expr = (
            f"COALESCE(TRIM(CAST(a.{self._dq(join_key)} AS VARCHAR)), '') = "
            f"COALESCE(TRIM(CAST(b.{self._dq(join_key)} AS VARCHAR)), '')"
        )

        try:
            conn.execute(
                f"CREATE OR REPLACE TEMP VIEW {left_table}_cmp AS "
                f"SELECT {left_cols_sql} FROM ("
                f"SELECT {left_cols_sql}, ROW_NUMBER() OVER (PARTITION BY {self._dq(join_key)} ORDER BY {self._dq(join_key)}) AS __rn FROM {left_table}"
                f") WHERE __rn = 1"
            )
            conn.execute(
                f"CREATE OR REPLACE TEMP VIEW {right_table}_cmp AS "
                f"SELECT {right_cols_sql} FROM ("
                f"SELECT {right_cols_sql}, ROW_NUMBER() OVER (PARTITION BY {self._dq(join_key)} ORDER BY {self._dq(join_key)}) AS __rn FROM {right_table}"
                f") WHERE __rn = 1"
            )
        except Exception as e:
            warnings.append(f"Deduplication views failed: {e}")
            return {"ok": False, "error": "Could not prepare reconciliation views.", "warning": "\n".join(warnings)}

        cmp_left = f"{left_table}_cmp"
        cmp_right = f"{right_table}_cmp"

        def one_value(sql: str, step: str, default="0"):
            try:
                value = conn.execute(sql).fetchone()[0]
                return default if value is None else str(value)
            except Exception as e:
                warnings.append(f"{step} failed: {e}")
                return default

        summary_rows = [
            ["Left file", left_meta["display_name"]],
            ["Right file", right_meta["display_name"]],
            ["Join key", join_key],
            ["Left rows", str(left_meta["row_count"])],
            ["Right rows", str(right_meta["row_count"])],
            ["Common columns", str(len(common_cols))],
        ]

        summary_rows.append([
            "Left duplicate key groups",
            one_value(f"SELECT COUNT(*) FROM (SELECT {self._dq(join_key)} FROM {left_table} GROUP BY 1 HAVING COUNT(*) > 1)", "left duplicate key scan"),
        ])
        summary_rows.append([
            "Right duplicate key groups",
            one_value(f"SELECT COUNT(*) FROM (SELECT {self._dq(join_key)} FROM {right_table} GROUP BY 1 HAVING COUNT(*) > 1)", "right duplicate key scan"),
        ])
        summary_rows.append([
            "Matched keys",
            one_value(f"SELECT COUNT(*) FROM (SELECT a.{self._dq(join_key)} FROM {cmp_left} a INNER JOIN {cmp_right} b ON {join_expr} GROUP BY 1)", "matched key count"),
        ])
        summary_rows.append([
            "Missing in right",
            one_value(f"SELECT COUNT(*) FROM {cmp_left} a LEFT JOIN {cmp_right} b ON {join_expr} WHERE b.{self._dq(join_key)} IS NULL", "missing in right count"),
        ])
        summary_rows.append([
            "Missing in left",
            one_value(f"SELECT COUNT(*) FROM {cmp_right} b LEFT JOIN {cmp_left} a ON {join_expr} WHERE a.{self._dq(join_key)} IS NULL", "missing in left count"),
        ])

        mismatch_rows = []
        for col in common_cols:
            if col == join_key:
                continue
            sql = (
                f"SELECT COUNT(*) AS mismatch_count FROM {cmp_left} a INNER JOIN {cmp_right} b ON {join_expr} "
                f"WHERE COALESCE(TRIM(CAST(a.{self._dq(col)} AS VARCHAR)), '') <> "
                f"COALESCE(TRIM(CAST(b.{self._dq(col)} AS VARCHAR)), '')"
            )
            try:
                mismatch_count = int(conn.execute(sql).fetchone()[0] or 0)
                if mismatch_count > 0:
                    mismatch_rows.append([col, str(mismatch_count)])
            except Exception as e:
                warnings.append(f"Column mismatch scan for {col} failed: {e}")

        missing_right_df = self._safe_duckdb_df(
            conn,
            f"SELECT a.{self._dq(join_key)} AS missing_key FROM {cmp_left} a LEFT JOIN {cmp_right} b ON {join_expr} WHERE b.{self._dq(join_key)} IS NULL LIMIT 20",
            "missing in right sample",
            warnings,
        )
        missing_left_df = self._safe_duckdb_df(
            conn,
            f"SELECT b.{self._dq(join_key)} AS missing_key FROM {cmp_right} b LEFT JOIN {cmp_left} a ON {join_expr} WHERE a.{self._dq(join_key)} IS NULL LIMIT 20",
            "missing in left sample",
            warnings,
        )

        detail_rows = []
        for col, _count in mismatch_rows[:5]:
            sample_sql = (
                f"SELECT a.{self._dq(join_key)} AS join_key, '{col}' AS column_name, "
                f"CAST(a.{self._dq(col)} AS VARCHAR) AS left_value, CAST(b.{self._dq(col)} AS VARCHAR) AS right_value "
                f"FROM {cmp_left} a INNER JOIN {cmp_right} b ON {join_expr} "
                f"WHERE COALESCE(TRIM(CAST(a.{self._dq(col)} AS VARCHAR)), '') <> "
                f"COALESCE(TRIM(CAST(b.{self._dq(col)} AS VARCHAR)), '') LIMIT 10"
            )
            sample_df = self._safe_duckdb_df(conn, sample_sql, f"sample mismatches for {col}", warnings)
            if sample_df is not None and not sample_df.empty:
                for row in sample_df.values.tolist():
                    detail_rows.append(["" if v is None else str(v) for v in row])

        response_lines = [
            f"Structured reconciliation completed for {left_meta['display_name']} vs {right_meta['display_name']}.",
            f"Join key used: {join_key}",
            f"Common columns compared: {max(0, len(common_cols) - 1)}",
            f"Columns with mismatches: {len(mismatch_rows)}",
        ]
        if prompt_text:
            response_lines.append(f"Instruction context: {prompt_text}")
        if warnings:
            response_lines.append("Warnings: " + "; ".join(warnings[:6]))

        tables = [
            {"headers": ["Metric", "Value"], "rows": summary_rows},
        ]
        if mismatch_rows:
            tables.append({"headers": ["Column", "Mismatch Count"], "rows": mismatch_rows[:50]})
        if missing_right_df is not None and not missing_right_df.empty:
            tables.append(self._df_to_table(missing_right_df, max_rows=20))
        if missing_left_df is not None and not missing_left_df.empty:
            tables.append(self._df_to_table(missing_left_df, max_rows=20))
        if detail_rows:
            tables.append({
                "headers": ["Join Key", "Column", "Left Value", "Right Value"],
                "rows": detail_rows[:100],
            })

        return {
            "ok": True,
            "summary_text": "\n".join(response_lines),
            "warning": "\n".join(warnings) if warnings else "",
            "data": {"tables": tables, "text": "\n".join(response_lines), "has_tables": bool(tables)},
        }

    def _analyze_tabular_files(self, files: list, prompt_text: str, mode: str = "agent") -> dict:
        try:
            import importlib
            duckdb = importlib.import_module("duckdb")
        except Exception:
            return {"ok": False, "error": "DuckDB is not installed in the current Python environment."}

        conn = duckdb.connect(database=":memory:")
        warnings = []
        metas = []

        try:
            for index, file_info in enumerate(files, start=1):
                if not self._is_tabular_file(file_info.get("name", "")):
                    continue
                try:
                    frame = self._load_tabular_dataframe(file_info)
                    table_name = f"source_{index}"
                    relation_name = f"{table_name}_df"
                    conn.register(relation_name, frame)
                    conn.execute(f"CREATE OR REPLACE TEMP TABLE {table_name} AS SELECT * FROM {relation_name}")
                    metas.append({
                        "table_name": table_name,
                        "display_name": str(file_info.get("name", table_name)),
                        "row_count": len(frame.index),
                        "columns": list(frame.columns),
                    })
                except Exception as e:
                    warnings.append(f"Loading {file_info.get('name', 'file')} failed: {e}")

            if not metas:
                return {"ok": False, "error": "No tabular files could be loaded for structured analysis.", "warning": "\n".join(warnings) if warnings else ""}

            if len(metas) == 2:
                return self._build_reconciliation_analysis(conn, metas[:2], prompt_text, warnings)
            if len(metas) >= 3:
                # Pairwise reconciliation for 3+ files
                all_pair_results = []
                pair_warnings = list(warnings)
                for i in range(len(metas)):
                    for j in range(i + 1, len(metas)):
                        pair_result = self._build_reconciliation_analysis(
                            conn, [metas[i], metas[j]], prompt_text, []
                        )
                        label = f"{metas[i]['display_name']} vs {metas[j]['display_name']}"
                        if pair_result.get('ok'):
                            all_pair_results.append((label, pair_result))
                        else:
                            pair_warnings.append(f"{label}: {pair_result.get('error', 'failed')}")
                if not all_pair_results:
                    return {"ok": False, "error": "Could not reconcile any file pairs.", "warning": "\n".join(pair_warnings)}
                # Merge all pair results into one output
                combined_tables = []
                combined_text_parts = [f"Reconciled {len(all_pair_results)} file pair(s) from {len(metas)} files:"]
                for label, pr in all_pair_results:
                    combined_text_parts.append(f"\n--- {label} ---")
                    combined_text_parts.append(pr.get('summary_text', ''))
                    for tbl in (pr.get('data') or {}).get('tables', []):
                        combined_tables.append(tbl)
                    if pr.get('warning'):
                        pair_warnings.append(pr['warning'])
                combined_text = "\n".join(combined_text_parts)
                return {
                    "ok": True,
                    "summary_text": combined_text,
                    "warning": "\n".join(pair_warnings) if pair_warnings else "",
                    "data": {"tables": combined_tables, "text": combined_text, "has_tables": bool(combined_tables)},
                }
            return self._build_single_table_analysis(conn, metas[0], prompt_text, warnings)
        finally:
            try:
                conn.close()
            except Exception:
                pass

    # ── DuckDB-powered universal tabular query engine ──────────────

    def _load_tabular_df(self, path: str):
        """Load a CSV/Excel file into a pandas DataFrame (all rows)."""
        import pandas as pd
        ext = os.path.splitext(path)[1].lower()
        if ext == ".csv":
            return pd.read_csv(path, low_memory=False)
        if ext in (".xlsx", ".xls"):
            return pd.read_excel(path, sheet_name=0)
        raise ValueError(f"Unsupported tabular extension: {ext}")

    def _fmt_table(self, df) -> str:
        """Format a DataFrame as a markdown table with fallback."""
        try:
            return df.to_markdown(index=False)
        except (ImportError, Exception):
            return df.to_string(index=False)

    def _col_expand(self, col_name: str) -> list[str]:
        """Return all plausible lowercase forms of a column name.
        E.g. 'CustomerName' → ['customername', 'customer name', 'customer_name']
             'Unit_Price'   → ['unit_price', 'unit price', 'unitprice']
        """
        import re as _re
        forms = set()
        cl = col_name.lower()
        forms.add(cl)
        forms.add(cl.replace('_', ' '))
        forms.add(cl.replace('_', ''))
        # CamelCase split: "CustomerName" → "customer name"
        camel = _re.sub(r'([a-z])([A-Z])', r'\1 \2', col_name).lower()
        forms.add(camel)
        forms.add(camel.replace(' ', '_'))
        forms.add(camel.replace(' ', ''))
        forms.discard('')
        return list(forms)

    def _col_match(self, col_name: str, query_lower: str) -> bool:
        """Check if a column name (in any form) appears in the query.
        Uses word-boundary matching to avoid false positives like
        'order' in 'in order to' or 'status' in 'what is the status'.
        """
        import re as _re
        # Common English words that happen to be column names — skip them
        _common_words = {
            'id', 'name', 'type', 'date', 'time', 'status', 'state',
            'order', 'product', 'item', 'price', 'amount', 'value',
            'note', 'description', 'comment', 'code', 'number',
            'count', 'total', 'size', 'level', 'class', 'group',
        }
        for form in self._col_expand(col_name):
            if len(form) < 3:
                continue
            # Single common word → require strong data context nearby
            if form in _common_words:
                ctx_pat = (
                    r'(?:the file|this file|the data|this data|uploaded|csv|excel|'
                    r'the table|this table|column|row|'
                    r'sum of|average of|total of|count of|max of|min of|'
                    r'group by|filter by|sort by|order by|'
                    r'calculate|compute|add\s+\d|multiply|divide|percent)'
                    r'.*\b' + _re.escape(form) + r'\b'
                )
                if _re.search(ctx_pat, query_lower):
                    return True
                continue
            # Multi-word or uncommon column name — word boundary match
            if _re.search(r'\b' + _re.escape(form) + r'\b', query_lower):
                return True
        return False

    def _is_data_question(self, text: str) -> bool:
        """Decide if the user question is about the uploaded tabular data."""
        import re as _re
        q = text.lower().strip()

        # --- 1) Check if any column name from the uploaded file appears in the question ---
        if self.uploaded_file_path:
            try:
                df = self._load_tabular_df(self.uploaded_file_path)
                for col in df.columns:
                    if self._col_match(col, q):
                        return True
            except Exception:
                pass

        # --- 2) Check for explicit data / file references ---
        fname = (self.uploaded_file_name or "").lower()
        if fname:
            base = os.path.splitext(fname)[0].replace('_', ' ').replace('-', ' ').lower()
            if base and base in q:
                return True
        if _re.search(r'\b(the file|uploaded file|this file|the data|this data|the table|this table|the csv|this csv|the spreadsheet|this spreadsheet|the excel)\b', q):
            return True

        # --- 3) Data-analysis keywords (only strong indicators) ---
        # These must be phrases/combos unlikely in casual conversation.
        strong_data_kw = (
            r'\b(sum of|total of|average of|avg of|mean of'
            r'|count of|how many rows|how much .* in the'
            r'|group by|grouped by|for each .* (show|get|find)'
            r'|order by|sort by|filter by|where .* (=|>|<|is)'
            r'|schema|describe .* (table|file|data|columns)'
            r'|head|preview|sample rows'
            r'|histogram|distribution of|chart .* (column|data)'
            r'|unique values|distinct values|null values|missing values'
            r'|add \d+.*% to|calculat\w* .* column|comput\w* .* column'
            r'|new column|multiply .* by|divide .* by'
            r'|top \d+|bottom \d+|highest .* (in|from)|lowest .* (in|from)'
            r'|median of|percentile|standard deviation'
            r'|rows in|number of rows|row count)\b'
        )
        if _re.search(strong_data_kw, q):
            return True

        return False

    def _model_generate_sql(self, conn, df, question: str, all_cols: list, total_rows: int) -> str | None:
        """Ask the model to write a DuckDB SQL query, execute it, return formatted result."""
        import re as _re
        if self.model is None:
            return None
        try:
            col_desc = ", ".join(f'"{c}" ({df[c].dtype})' for c in all_cols)
            sample_rows = conn.execute("SELECT * FROM data_table LIMIT 3").df().to_string(index=False)
            sql_prompt = (
                f"Table: data_table  ({total_rows} rows)\n"
                f"Columns: {col_desc}\n"
                f"Sample data:\n{sample_rows}\n\n"
                f"User request: {question}\n\n"
                "Write a single DuckDB SQL SELECT query to fulfill this request.\n"
                "Rules:\n"
                "- Use double quotes for column names, e.g. \"Total_Amount\"\n"
                "- For computed columns use AS with a clear alias\n"
                "- For percentage: multiply by 0.10 for 10%, etc.\n"
                "- Only SELECT queries allowed (no INSERT/UPDATE/DELETE)\n"
                "- Return ONLY the raw SQL, no explanation, no markdown fences.\n"
                "SQL:"
            )
            sql_resp = self.model.create_completion(
                sql_prompt,
                max_tokens=400,
                temperature=0.1,
                stop=["</s>", "<|im_end|>", "<|end|>", "\n\n", "```"],
            )
            raw_sql = (sql_resp.get("choices", [{}])[0]
                       .get("text", "").strip())
            # Clean up markdown fences
            raw_sql = _re.sub(r'^```\w*\n?', '', raw_sql)
            raw_sql = _re.sub(r'\n?```$', '', raw_sql).strip()
            # Strip non-SQL prefix
            if not raw_sql.upper().startswith("SELECT"):
                sel_match = _re.search(r'(SELECT\s.+)', raw_sql, _re.IGNORECASE | _re.DOTALL)
                if sel_match:
                    raw_sql = sel_match.group(1).strip()
            # Remove trailing semicolons
            raw_sql = raw_sql.rstrip(';').strip()
            # Safety: only SELECT
            if not raw_sql or not raw_sql.upper().startswith("SELECT"):
                return None
            # Block dangerous keywords
            sql_upper = raw_sql.upper()
            for forbidden in ("DROP", "DELETE", "INSERT", "UPDATE", "ALTER", "CREATE", "TRUNCATE"):
                if forbidden in sql_upper:
                    return None
            print(f"[TABULAR] Model SQL: {raw_sql}")
            res = conn.execute(raw_sql).df()
            lines = []
            if res.empty:
                lines.append("**Result:** 0 rows matched the query.")
            else:
                lines.append(f"**Result** ({len(res)} rows):")
                lines.append(self._fmt_table(res))
            return "\n".join(lines)
        except Exception as sql_err:
            print(f"[TABULAR] Model SQL failed: {sql_err}")
            # Retry with simplified prompt
            try:
                simple_prompt = (
                    f"Table: data_table\nColumns: {col_desc}\n\n"
                    f"Write a DuckDB SQL SELECT for: {question}\n"
                    f"Use double-quoted column names. Return only SQL.\nSQL: SELECT"
                )
                simple_resp = self.model.create_completion(
                    simple_prompt, max_tokens=300, temperature=0.05,
                    stop=["</s>", "<|im_end|>", "<|end|>", "\n\n"],
                )
                retry_sql = "SELECT " + (simple_resp.get("choices", [{}])[0].get("text", "").strip())
                retry_sql = retry_sql.rstrip(';').strip()
                retry_sql = _re.sub(r'```.*', '', retry_sql).strip()
                if retry_sql and retry_sql.upper().startswith("SELECT"):
                    sql_upper2 = retry_sql.upper()
                    safe = True
                    for forbidden in ("DROP", "DELETE", "INSERT", "UPDATE", "ALTER", "CREATE", "TRUNCATE"):
                        if forbidden in sql_upper2:
                            safe = False
                    if safe:
                        print(f"[TABULAR] Model SQL retry: {retry_sql}")
                        res2 = conn.execute(retry_sql).df()
                        if not res2.empty:
                            return f"**Result** ({len(res2)} rows):\n" + self._fmt_table(res2)
            except Exception as retry_err:
                print(f"[TABULAR] Model SQL retry also failed: {retry_err}")
            return None

    def _precompute_tabular_answer(self, path: str, question: str) -> str:
        """Universal DuckDB query engine — always returns computed results."""
        try:
            import importlib
            import re as _re
            duckdb = importlib.import_module("duckdb")
            df = self._load_tabular_df(path)

            conn = duckdb.connect(database=":memory:")
            conn.register("data_table", df)

            q = question.lower().strip()
            total_rows = len(df)
            numeric_cols = [c for c in df.columns if df[c].dtype.kind in ('i', 'f')]
            text_cols = [c for c in df.columns if df[c].dtype.kind in ('O', 'U', 'S')]
            date_cols = [c for c in df.columns if df[c].dtype.kind in ('M',)]
            all_cols = list(df.columns)
            lines: list[str] = []

            # ── 1) Schema / describe / structure ──────────────────────
            _has_schema_kw = _re.search(r'\b(schema|describe|structure|columns?|fields?|info|dtype|dtypes|head|preview)\b', q)
            _has_transform_kw = _re.search(
                r'\b(add|create|calculat|comput|deriv|new\s+column|multiply|divide|subtract|'
                r'percentage|percent|%|ratio|convert|transform|modify|update|rename)\b', q)
            if _has_schema_kw and not _has_transform_kw:
                lines.append(f"**File:** {os.path.basename(path)}")
                lines.append(f"**Rows:** {total_rows:,}  |  **Columns:** {len(all_cols)}")
                col_info = []
                for c in all_cols:
                    dtype_str = str(df[c].dtype)
                    nulls = int(df[c].isna().sum())
                    uniq = int(df[c].nunique())
                    col_info.append([c, dtype_str, f"{uniq:,}", f"{nulls:,}"])
                import pandas as pd
                info_df = pd.DataFrame(col_info, columns=["Column", "Type", "Unique", "Nulls"])
                lines.append(self._fmt_table(info_df))
                if _re.search(r'\b(head|preview|sample)\b', q):
                    sample = conn.execute("SELECT * FROM data_table LIMIT 10").df()
                    lines.append("\n**Sample rows (first 10):**")
                    lines.append(self._fmt_table(sample))
                conn.close()
                return "\n".join(lines)

            # ── 1b) Direct computed-column patterns (regex, no model needed) ──
            # "add 10% to Total_Amount as/called/new Net_Amount"
            _pct_match = _re.search(
                r'(?:add|apply|calculat\w*|comput\w*)\s+'
                r'(\d+(?:\.\d+)?)\s*%\s*(?:to|of|on)\s+'
                r'(\w+)',
                q,
            )
            if _pct_match:
                pct_val = float(_pct_match.group(1))
                src_token = _pct_match.group(2)
                # find the real column
                src_col = None
                for c in all_cols:
                    for form in self._col_expand(c):
                        if form == src_token.lower() or src_token.lower() in form or form in src_token.lower():
                            src_col = c
                            break
                    if src_col:
                        break
                if src_col and src_col in numeric_cols:
                    # Try to find the alias name from the query
                    alias_match = _re.search(
                        r'(?:(?:new|as|called|named|into|column)\s+)(\w+)\s*(?:column)?$',
                        q,
                    )
                    if not alias_match:
                        alias_match = _re.search(r'(?:new|as|called|named|into|column)\s+(\w+)', q)
                    alias = alias_match.group(1) if alias_match else f"{src_col}_plus_{int(pct_val)}pct"
                    # Clean alias — don't use 'column' as the alias
                    if alias.lower() == 'column':
                        alias = f"{src_col}_plus_{int(pct_val)}pct"
                    multiplier = 1.0 + pct_val / 100.0
                    sql = f'SELECT *, "{src_col}" * {multiplier} AS "{alias}" FROM data_table'
                    print(f"[TABULAR] Regex-computed SQL: {sql}")
                    try:
                        res = conn.execute(sql).df()
                        lines.append(f"**Result** ({len(res)} rows):")
                        lines.append(self._fmt_table(res))
                        conn.close()
                        return "\n".join(lines)
                    except Exception as e:
                        print(f"[TABULAR] Regex-computed SQL failed: {e}")
                        lines.clear()

            # "multiply/divide column by N" or "column * N"
            _arith_match = _re.search(
                r'(?:multiply|divide)\s+(\w+)\s+(?:by|with)\s+(\d+(?:\.\d+)?)', q
            )
            if not _arith_match:
                _arith_match = _re.search(
                    r'(\w+)\s*[\*x×]\s*(\d+(?:\.\d+)?)', q
                )
            if _arith_match:
                col_token = _arith_match.group(1)
                factor = float(_arith_match.group(2))
                op = '*' if 'divide' not in q else '/'
                src_col = None
                for c in all_cols:
                    for form in self._col_expand(c):
                        if form == col_token.lower() or col_token.lower() in form or form in col_token.lower():
                            src_col = c
                            break
                    if src_col:
                        break
                if src_col and src_col in numeric_cols:
                    alias = f"{src_col}_{'div' if op == '/' else 'mult'}_{int(factor)}"
                    sql = f'SELECT *, "{src_col}" {op} {factor} AS "{alias}" FROM data_table'
                    print(f"[TABULAR] Regex-arith SQL: {sql}")
                    try:
                        res = conn.execute(sql).df()
                        lines.append(f"**Result** ({len(res)} rows):")
                        lines.append(self._fmt_table(res))
                        conn.close()
                        return "\n".join(lines)
                    except Exception as e:
                        print(f"[TABULAR] Regex-arith SQL failed: {e}")
                        lines.clear()

            # ── 1c) Dynamic / transformation requests → model SQL ─────
            _transform_kw = _re.search(
                r'\b(add|create|calculat|comput|deriv|new\s+column|rename|'
                r'multiply|divide|subtract|percentage|percent|%|ratio|'
                r'convert|cast|transform|modify|update|change|replace|'
                r'concat|merge|split|extract|format|round|ceil|floor|'
                r'case when|if\s+.*\s+then|pivot|unpivot|cross\s*tab|'
                r'rank|dense_rank|row_number|window|over\s*\(|partition|'
                r'join|union|intersect|except|subquer|'
                r'date_part|year|month|day|hour|minute|second|'
                r'regex|pattern|like\s+\'|substring|trim|upper|lower|length)\b',
                q,
            )
            if _transform_kw and self.model is not None:
                sql_result = self._model_generate_sql(conn, df, question, all_cols, total_rows)
                if sql_result:
                    conn.close()
                    return sql_result

            # ── Detect referenced columns ─────────────────────────────
            group_col = None
            agg_col = None
            filter_col = None
            filter_val = None

            # Best-match column finder: returns the column whose expanded
            # form is longest-match in the query, avoiding false positives.
            # Also checks if query words appear IN the column form (bidirectional).
            def _find_col(candidates, query):
                """Find best matching column. Uses word-boundary matching and
                prefers longer form matches to avoid false positives."""
                best = None
                best_len = 0
                qw = query.lower().strip()
                for c in candidates:
                    for form in self._col_expand(c):
                        if len(form) < 3:
                            continue
                        # Must match as a whole word (word boundary)
                        if _re.search(r'\b' + _re.escape(form) + r'\b', qw) and len(form) > best_len:
                            best = c
                            best_len = len(form)
                return best

            # Detect grouping col: "by X", "of X", "per X", "for each X", "each X"
            grp_match = _re.search(
                r'\b(?:by|of|per|for\s+each|each|group\s*by)\s+'
                r'(\w[\w\s]*?)\s*(?:and|or|for|in|with|where|from|limit|top|order|asc|desc|tabular|table|form|format|$)',
                q,
            )
            if grp_match:
                target = grp_match.group(1).strip()
                group_col = _find_col(text_cols + all_cols, target)

            # If "total of X of Y" pattern — X likely agg, Y likely group
            of_of_match = _re.search(
                r'\b(?:total|sum|average|avg|mean)\s+(?:of\s+)?(\w[\w\s]*?)\s+(?:of|by|per|for)\s+(\w[\w\s]*?)(?:\s|$)',
                q,
            )
            if of_of_match:
                part1 = of_of_match.group(1).strip()
                part2 = of_of_match.group(2).strip()
                a = _find_col(numeric_cols, part1)
                g = _find_col(text_cols + all_cols, part2)
                if a:
                    agg_col = a
                if g:
                    group_col = g
                # If part1 matched a text col and part2 matched numeric, swap
                if not a and not g:
                    a2 = _find_col(numeric_cols, part2)
                    g2 = _find_col(text_cols + all_cols, part1)
                    if a2:
                        agg_col = a2
                    if g2:
                        group_col = g2

            # Direct column mention fallback
            if not group_col:
                group_col = _find_col(text_cols, q)
            if not agg_col:
                agg_col = _find_col(numeric_cols, q)
            if not agg_col and numeric_cols:
                agg_col = numeric_cols[0]

            # ── 2) Filter / where / find ──────────────────────────────
            filter_match = _re.search(r'\b(?:where|filter|find|show|for|with)\s+(\w[\w\s]*?)\s*(?:=|==|is|equals?|like)\s*["\']?(\w[\w\s]*?)["\']?\s*$', q)
            if filter_match:
                fcol_search = filter_match.group(1).strip()
                filter_val = filter_match.group(2).strip()
                filter_col = _find_col(all_cols, fcol_search)

            if filter_col and filter_val:
                sql = f"SELECT * FROM data_table WHERE LOWER(CAST(\"{filter_col}\" AS VARCHAR)) LIKE '%{filter_val.lower()}%' LIMIT 100"
                try:
                    res = conn.execute(sql).df()
                    lines.append(f"**Filtered rows** where {filter_col} matches '{filter_val}' ({len(res)} rows shown, max 100):")
                    lines.append(self._fmt_table(res))
                    conn.close()
                    return "\n".join(lines)
                except Exception as _e:
                    print(f"[TABULAR] Branch 2 failed: {_e}, falling through to model SQL")
                    lines.clear()

            # ── 3) Top-N / bottom-N ───────────────────────────────────
            top_match = _re.search(r'\b(top|bottom|best|worst|highest|lowest|largest|smallest)\s*(\d+)?\b', q)
            if top_match:
                direction = top_match.group(1).lower()
                n = int(top_match.group(2)) if top_match.group(2) else 10
                n = min(n, 1000)
                is_desc = direction in ('top', 'best', 'highest', 'largest')
                order = "DESC" if is_desc else "ASC"
                # Try to detect target column from "by X" or "in X" context
                by_match = _re.search(r'\b(?:by|in|of|for)\s+(\w[\w\s]*?)(?:\s*$)', q)
                sort_col = None
                if by_match:
                    sort_col = _find_col(numeric_cols, by_match.group(1).strip())
                if not sort_col:
                    # Also try finding numeric column mentioned anywhere in query
                    sort_col = _find_col(numeric_cols, q)
                if not sort_col:
                    sort_col = agg_col or (numeric_cols[0] if numeric_cols else all_cols[0])
                sql = f'SELECT * FROM data_table ORDER BY CAST(\"{sort_col}\" AS DOUBLE) {order} NULLS LAST LIMIT {n}'
                try:
                    res = conn.execute(sql).df()
                    lines.append(f"**{direction.title()} {n}** rows by {sort_col}:")
                    lines.append(self._fmt_table(res))
                    conn.close()
                    return "\n".join(lines)
                except Exception as _e:
                    print(f"[TABULAR] Branch 3 failed: {_e}, falling through to model SQL")
                    lines.clear()

            # ── 4) Correlation / relationship ─────────────────────────
            if _re.search(r'\b(correlat|relationship|relation)\b', q) and len(numeric_cols) >= 2:
                try:
                    import pandas as pd
                    corr = df[numeric_cols].corr()
                    lines.append("**Correlation matrix** (numeric columns):")
                    lines.append(self._fmt_table(corr.reset_index().rename(columns={"index": "Column"})))
                    conn.close()
                    return "\n".join(lines)
                except Exception as _e:
                    print(f"[TABULAR] Branch 4 failed: {_e}, falling through to model SQL")
                    lines.clear()

            # ── 5) Distinct / unique values ───────────────────────────
            if _re.search(r'\b(distinct|unique|categories|values)\b', q):
                target_col = group_col or (text_cols[0] if text_cols else all_cols[0])
                sql = f'SELECT DISTINCT \"{target_col}\" AS value, COUNT(*) AS count FROM data_table GROUP BY \"{target_col}\" ORDER BY count DESC LIMIT 100'
                try:
                    res = conn.execute(sql).df()
                    lines.append(f"**Distinct values** in {target_col} ({len(res)} shown, max 100):")
                    lines.append(self._fmt_table(res))
                    conn.close()
                    return "\n".join(lines)
                except Exception as _e:
                    print(f"[TABULAR] Branch 5 failed: {_e}, falling through to model SQL")
                    lines.clear()

            # ── 6) Count / how many ───────────────────────────────────
            if _re.search(r'\b(count|how many|number of|total\s+(?:rows?|records?|entries|items?))\b', q) and not _re.search(r'\b(sum|average|avg|mean|total\s+(?:amount|price|sales|revenue|value|quantity))\b', q):
                try:
                    if group_col:
                        sql = f'SELECT \"{group_col}\", COUNT(*) AS count FROM data_table GROUP BY \"{group_col}\" ORDER BY count DESC'
                        res = conn.execute(sql).df()
                        lines.append(f"**Count** by {group_col}:")
                        lines.append(self._fmt_table(res))
                        lines.append(f"\n**Total rows:** {total_rows:,}")
                    else:
                        lines.append(f"**Total rows:** {total_rows:,}")
                    conn.close()
                    return "\n".join(lines)
                except Exception as _e:
                    print(f"[TABULAR] Branch 6 failed: {_e}, falling through to model SQL")
                    lines.clear()

            # ── 7) Group-by aggregation (sum/avg/min/max) ─────────────
            if group_col and agg_col:
                try:
                    agg_type = "SUM"
                    if _re.search(r'\b(average|avg|mean)\b', q):
                        agg_type = "AVG"
                    elif _re.search(r'\b(max|maximum|highest)\b', q):
                        agg_type = "MAX"
                    elif _re.search(r'\b(min|minimum|lowest)\b', q):
                        agg_type = "MIN"
                    sql = (
                        f'SELECT \"{group_col}\", '
                        f'{agg_type}(CAST(\"{agg_col}\" AS DOUBLE)) AS {agg_type.lower()}_{agg_col}, '
                        f'COUNT(*) AS count '
                        f'FROM data_table GROUP BY \"{group_col}\" '
                        f'ORDER BY {agg_type.lower()}_{agg_col} DESC'
                    )
                    res = conn.execute(sql).df()
                    val_col = f"{agg_type.lower()}_{agg_col}"
                    lines.append(f"**{agg_type}({agg_col})** by {group_col}:")
                    lines.append(self._fmt_table(res))
                    lines.append(f"\n**Grand total ({agg_type}):** {res[val_col].sum():,.2f}")
                    conn.close()
                    return "\n".join(lines)
                except Exception as _e:
                    print(f"[TABULAR] Branch 7 failed: {_e}, falling through to model SQL")
                    lines.clear()

            # ── 8) Model-generated SQL (handles any question) ─────────
            if self.model is not None:
                sql_result = self._model_generate_sql(conn, df, question, all_cols, total_rows)
                if sql_result:
                    conn.close()
                    return sql_result

            # ── 9) General numeric summary (last resort) ──────────────
            if numeric_cols:
                import pandas as pd
                parts = []
                for nc in numeric_cols:
                    try:
                        sql = (
                            f'SELECT '
                            f'COUNT(\"{nc}\") AS count, '
                            f'SUM(CAST(\"{nc}\" AS DOUBLE)) AS sum, '
                            f'AVG(CAST(\"{nc}\" AS DOUBLE)) AS avg, '
                            f'MIN(CAST(\"{nc}\" AS DOUBLE)) AS min, '
                            f'MAX(CAST(\"{nc}\" AS DOUBLE)) AS max '
                            f'FROM data_table'
                        )
                        r = conn.execute(sql).df()
                        parts.append([
                            nc,
                            f"{r['count'].iloc[0]:,}",
                            f"{r['sum'].iloc[0]:,.2f}",
                            f"{r['avg'].iloc[0]:,.2f}",
                            f"{r['min'].iloc[0]:,.2f}",
                            f"{r['max'].iloc[0]:,.2f}",
                        ])
                    except Exception:
                        pass
                if parts:
                    stats_df = pd.DataFrame(parts, columns=["Column", "Count", "Sum", "Avg", "Min", "Max"])
                    lines.append(f"**Numeric summary** for {os.path.basename(path)} ({total_rows:,} rows):")
                    lines.append(self._fmt_table(stats_df))

            # If we still have nothing, provide a basic overview
            if not lines:
                lines.append(f"**File:** {os.path.basename(path)}")
                lines.append(f"**Total rows:** {total_rows:,}  |  **Columns:** {len(all_cols)}")
                lines.append(f"**Columns:** {', '.join(all_cols[:30])}")
                if text_cols:
                    first_text = text_cols[0]
                    sql = f'SELECT \"{first_text}\", COUNT(*) AS count FROM data_table GROUP BY \"{first_text}\" ORDER BY count DESC LIMIT 20'
                    try:
                        res = conn.execute(sql).df()
                        lines.append(f"\n**Distribution of {first_text}:**")
                        lines.append(self._fmt_table(res))
                    except Exception:
                        pass

            conn.close()
            return "\n".join(lines)
        except Exception as e:
            print(f"[TABULAR] Precompute error: {e}")
            return ""

    def _run_tabular_chat_analysis(self, text: str, request_options: dict | None = None):
        self.generation_in_progress = True
        self.stop_generation_flag = False
        self._emit("generation_start", None)
        self._status("Analyzing tabular file...")

        try:
            if not self.uploaded_file_path or not self.uploaded_file_name:
                raise RuntimeError("No uploaded tabular file is available")

            # Step 1: Always compute the answer with DuckDB (handles ANY row count)
            precomputed = self._precompute_tabular_answer(self.uploaded_file_path, text)

            if precomputed:
                # DuckDB produced correct data — this is the ONLY source of truth.
                display = precomputed
            else:
                # DuckDB returned nothing — run structural analysis as fallback
                analysis = self._analyze_tabular_files(
                    [{"name": self.uploaded_file_name, "path": self.uploaded_file_path}],
                    text,
                    mode="chat",
                )
                display = analysis.get("summary_text", "Could not analyze the file. Please try a more specific question.")
                warning_text = analysis.get("warning", "")
                if warning_text:
                    display += f"\n\nWarnings:\n{warning_text}"

            self._emit("generation_done", {"text": display})
            self.message_history.append({"role": "assistant", "content": display})
            self.chats[self.current_chat_id] = self.message_history
            self._save_current_chat()
            self._auto_export_response(
                self.current_chat_id,
                len(self.message_history) - 1,
                (request_options or {}).get("export_format"),
            )
            self._status("Ready")
        except Exception as e:
            self._emit("generation_error", {"error": str(e)})
            self._status(f"Tabular analysis error: {e}")
        finally:
            self.generation_in_progress = False

    def _apply_selected_tts_voice(self, engine):
        """Apply selected TTS voice from settings if it exists on this system.
        Returns the selected voice id string (may start with 'piper:')."""
        selected = str((self.app_settings or {}).get("tts_voice_id", "")).strip()
        if not selected:
            return selected
        if selected.startswith("piper:"):
            return selected  # Piper voices not handled by pyttsx3
        try:
            voices = engine.getProperty("voices") or []
            for v in voices:
                vid = str(getattr(v, "id", "") or "")
                if vid == selected:
                    engine.setProperty("voice", vid)
                    return selected
        except Exception:
            pass
        return selected

    def _get_piper_onnx_path(self, voice_id: str):
        """Resolve a piper:<folder_name> voice id to its .onnx file path."""
        folder_name = voice_id.removeprefix("piper:").strip()
        if not folder_name:
            return None
        piper_dir = app_data_path(os.path.join("models", "voices", "piper"))
        voice_folder = os.path.join(piper_dir, folder_name)
        if not os.path.isdir(voice_folder):
            return None
        onnx_files = [f for f in os.listdir(voice_folder) if f.endswith(".onnx")]
        if not onnx_files:
            return None
        return os.path.join(voice_folder, onnx_files[0])

    def _piper_speak(self, text: str):
        """Speak text using a downloaded Piper voice. Stoppable via _tts_stop_event."""
        selected = str((self.app_settings or {}).get("tts_voice_id", "")).strip()
        onnx_path = self._get_piper_onnx_path(selected)
        if not onnx_path:
            raise FileNotFoundError(f"Piper voice model not found for: {selected}")

        import wave, tempfile
        from piper import PiperVoice

        voice = PiperVoice.load(onnx_path)
        clean = re.sub(r"[#*`_~\[\]()]", "", text)
        clean = re.sub(r"\n+", ". ", clean).strip()[:2000]

        # Synthesize to a temp WAV
        tmp_path = os.path.join(tempfile.gettempdir(), f"piper_speak_{os.getpid()}.wav")
        with wave.open(tmp_path, "wb") as wav_file:
            voice.synthesize_wav(clean, wav_file)

        if self._tts_stop_event.is_set():
            return

        self._piper_tmp_path = tmp_path
        try:
            import winsound
            # Play asynchronously so we can stop it
            winsound.PlaySound(tmp_path, winsound.SND_FILENAME | winsound.SND_ASYNC)
            # Wait for playback to finish or stop signal
            # Estimate duration from file size (16-bit mono 22050 Hz typical)
            try:
                file_size = os.path.getsize(tmp_path)
                duration_sec = max(1.0, file_size / (22050 * 2))  # rough estimate
            except Exception:
                duration_sec = 30.0
            # Poll for stop event
            elapsed = 0.0
            interval = 0.1
            while elapsed < duration_sec + 2.0:
                if self._tts_stop_event.wait(timeout=interval):
                    # Stop requested — kill playback
                    winsound.PlaySound(None, winsound.SND_PURGE)
                    return
                elapsed += interval
        finally:
            self._piper_tmp_path = None
            try:
                os.remove(tmp_path)
            except Exception:
                pass

    def _piper_save_wav(self, text: str, output_path: str):
        """Save Piper TTS synthesis directly to a WAV file."""
        selected = str((self.app_settings or {}).get("tts_voice_id", "")).strip()
        onnx_path = self._get_piper_onnx_path(selected)
        if not onnx_path:
            raise FileNotFoundError(f"Piper voice model not found for: {selected}")

        import wave
        from piper import PiperVoice

        voice = PiperVoice.load(onnx_path)
        clean = re.sub(r"[#*`_~\[\]()]", "", text)
        clean = re.sub(r"\n+", ". ", clean).strip()[:12000]

        with wave.open(output_path, "wb") as wav_file:
            voice.synthesize_wav(clean, wav_file)

    def list_tts_voices(self):
        """Return available pyttsx3 voices + downloaded Piper voices for settings."""
        selected = str((self.app_settings or {}).get("tts_voice_id", "")).strip()
        payload = []

        # ── System (pyttsx3 / SAPI) voices ──────────────────────
        try:
            import pyttsx3
            engine = pyttsx3.init()
            voices = engine.getProperty("voices") or []
            for i, v in enumerate(voices):
                vid = str(getattr(v, "id", "") or "")
                name = str(getattr(v, "name", "") or "").strip() or f"Voice {i + 1}"
                langs = getattr(v, "languages", None)
                lang_text = ""
                if isinstance(langs, (list, tuple)) and langs:
                    try:
                        lang_text = ", ".join(
                            str(x.decode("utf-8", errors="ignore") if isinstance(x, (bytes, bytearray)) else x)
                            for x in langs
                        )
                    except Exception:
                        lang_text = ""
                label = name if not lang_text else f"{name} ({lang_text})"
                payload.append({"id": vid, "name": name, "label": label, "selected": bool(selected and vid == selected)})
            try:
                engine.stop()
            except Exception:
                pass
        except Exception:
            pass

        # ── Downloaded Piper voices (models/voices/piper/*/*.onnx) ──
        try:
            piper_dir = app_data_path(os.path.join("models", "voices", "piper"))
            if os.path.isdir(piper_dir):
                for folder_name in sorted(os.listdir(piper_dir)):
                    voice_folder = os.path.join(piper_dir, folder_name)
                    if not os.path.isdir(voice_folder):
                        continue
                    # Look for an .onnx file inside
                    onnx_files = [f for f in os.listdir(voice_folder) if f.endswith(".onnx")]
                    if not onnx_files:
                        continue
                    onnx_path = os.path.join(voice_folder, onnx_files[0])
                    vid = f"piper:{folder_name}"
                    # Derive friendly name from folder:  en_US-lessac-medium → English US - Lessac (Piper)
                    pretty = folder_name.replace("-", " ").replace("_", " ").title()
                    label = f"🔊 {pretty} (Piper)"
                    payload.append({
                        "id": vid,
                        "name": folder_name,
                        "label": label,
                        "piper": True,
                        "onnx_path": onnx_path,
                        "selected": bool(selected and vid == selected),
                    })
        except Exception:
            pass

        return {"ok": True, "voices": payload, "selected_voice_id": selected}

    def _free_piper_voice_catalog(self) -> list[dict]:
        """Static free Piper voice catalog (Hugging Face-hosted)."""
        _B = "https://huggingface.co/rhasspy/piper-voices/resolve/main"
        def _v(vid, name, lang, lang2, voice, quality="medium"):
            prefix = f"{lang}/{lang2}/{voice}/{quality}/{lang2}-{voice}-{quality}"
            return {"id": vid, "name": name, "language": lang2, "quality": quality,
                    "onnx_url": f"{_B}/{prefix}.onnx?download=true",
                    "json_url": f"{_B}/{prefix}.onnx.json?download=true"}
        return [
            # ── English US ─────────────────────────────────
            _v("en_US-lessac-medium",   "English US - Lessac",        "en", "en_US", "lessac"),
            _v("en_US-amy-medium",      "English US - Amy",           "en", "en_US", "amy"),
            _v("en_US-bryce-medium",    "English US - Bryce",         "en", "en_US", "bryce"),
            _v("en_US-danny-low",       "English US - Danny (low)",   "en", "en_US", "danny", "low"),
            _v("en_US-hfc_female-medium","English US - HFC Female",   "en", "en_US", "hfc_female"),
            _v("en_US-hfc_male-medium", "English US - HFC Male",      "en", "en_US", "hfc_male"),
            _v("en_US-joe-medium",      "English US - Joe",           "en", "en_US", "joe"),
            _v("en_US-john-medium",     "English US - John",          "en", "en_US", "john"),
            _v("en_US-kathleen-low",    "English US - Kathleen (low)","en", "en_US", "kathleen", "low"),
            _v("en_US-kristin-medium",  "English US - Kristin",       "en", "en_US", "kristin"),
            _v("en_US-kusal-medium",    "English US - Kusal",         "en", "en_US", "kusal"),
            _v("en_US-ljspeech-medium", "English US - LJSpeech",      "en", "en_US", "ljspeech"),
            _v("en_US-norman-medium",   "English US - Norman",        "en", "en_US", "norman"),
            _v("en_US-ryan-medium",     "English US - Ryan",          "en", "en_US", "ryan"),
            # ── English UK ─────────────────────────────────
            _v("en_GB-alan-medium",     "English UK - Alan",          "en", "en_GB", "alan"),
            _v("en_GB-alba-medium",     "English UK - Alba",          "en", "en_GB", "alba"),
            _v("en_GB-aru-medium",      "English UK - Aru",           "en", "en_GB", "aru"),
            _v("en_GB-cori-medium",     "English UK - Cori",          "en", "en_GB", "cori"),
            _v("en_GB-jenny_dioco-medium","English UK - Jenny Dioco", "en", "en_GB", "jenny_dioco"),
            _v("en_GB-northern_english_male-medium","English UK - Northern Male","en","en_GB","northern_english_male"),
            _v("en_GB-southern_english_female-medium","English UK - Southern Female","en","en_GB","southern_english_female"),
            # ── German ─────────────────────────────────────
            _v("de_DE-thorsten-medium", "German - Thorsten",          "de", "de_DE", "thorsten"),
            _v("de_DE-thorsten_emotional-medium","German - Thorsten Emotional","de","de_DE","thorsten_emotional"),
            _v("de_DE-eva_k-medium",    "German - Eva K",             "de", "de_DE", "eva_k"),
            _v("de_DE-karlsson-low",    "German - Karlsson (low)",    "de", "de_DE", "karlsson", "low"),
            _v("de_DE-kerstin-low",     "German - Kerstin (low)",     "de", "de_DE", "kerstin", "low"),
            _v("de_DE-ramona-low",      "German - Ramona (low)",      "de", "de_DE", "ramona", "low"),
            _v("de_DE-pavoque-low",     "German - Pavoque (low)",     "de", "de_DE", "pavoque", "low"),
            # ── French ─────────────────────────────────────
            _v("fr_FR-siwis-medium",    "French - Siwis",             "fr", "fr_FR", "siwis"),
            _v("fr_FR-gilles-low",      "French - Gilles (low)",      "fr", "fr_FR", "gilles", "low"),
            _v("fr_FR-tom-medium",      "French - Tom",               "fr", "fr_FR", "tom"),
            # ── Spanish ────────────────────────────────────
            _v("es_ES-carlfm-medium",   "Spanish - Carlfm",           "es", "es_ES", "carlfm"),
            _v("es_ES-davefx-medium",   "Spanish - Davefx",           "es", "es_ES", "davefx"),
            _v("es_ES-sharvard-medium", "Spanish - Sharvard",         "es", "es_ES", "sharvard"),
            # ── Hindi ──────────────────────────────────────
            _v("hi_IN-pratham-medium",  "Hindi India - Pratham",       "hi", "hi_IN", "pratham"),
            # ── Tamil ──────────────────────────────────────
            _v("ta_IN-kani-medium",     "Tamil India - Kani",          "ta", "ta_IN", "kani"),
        ]

    def list_free_piper_voices(self):
        """Expose free downloadable Piper voices for UI."""
        voices = self._free_piper_voice_catalog()
        return {"ok": True, "voices": voices}

    def download_free_piper_voice(self, voice_id: str):
        """Download selected free Piper voice files to local app folder."""
        import urllib.request

        selected = None
        for item in self._free_piper_voice_catalog():
            if str(item.get("id", "")) == str(voice_id or ""):
                selected = item
                break
        if not selected:
            return {"error": "Unknown Piper voice id"}

        safe_id = re.sub(r"[^A-Za-z0-9_.-]+", "_", selected["id"]).strip("._") or "voice"
        base_dir = app_data_path(os.path.join("models", "voices", "piper"))
        voice_dir = os.path.join(base_dir, safe_id)
        os.makedirs(voice_dir, exist_ok=True)

        onnx_path = os.path.join(voice_dir, f"{safe_id}.onnx")
        json_path = os.path.join(voice_dir, f"{safe_id}.onnx.json")

        def _download(url: str, path: str):
            tmp = path + ".part"
            with urllib.request.urlopen(url, timeout=300) as resp, open(tmp, "wb") as out:
                while True:
                    chunk = resp.read(1024 * 128)
                    if not chunk:
                        break
                    out.write(chunk)
            os.replace(tmp, path)

        try:
            self._status(f"Downloading Piper voice: {selected['name']}")
            _download(selected["onnx_url"], onnx_path)
            _download(selected["json_url"], json_path)
            self._status(f"Downloaded Piper voice to {voice_dir}")
            return {
                "ok": True,
                "voice_id": safe_id,
                "folder": voice_dir,
                "onnx_path": onnx_path,
                "json_path": json_path,
            }
        except Exception as e:
            return {"error": f"Piper download failed: {e}"}

    def speak_text(self, text: str):
        """Speak text via pyttsx3 or Piper TTS in background thread."""
        if not text or not text.strip():
            return {"error": "Empty text"}

        selected = str((self.app_settings or {}).get("tts_voice_id", "")).strip()
        use_piper = selected.startswith("piper:")

        def _tts_worker(content: str):
            self._tts_stop_event.clear()
            try:
                with self._tts_lock:
                    self._tts_active = True
                self._emit("tts_started", None)

                if use_piper:
                    self._piper_speak(content)
                else:
                    engine = None
                    try:
                        import pyttsx3
                        engine = pyttsx3.init()
                        engine.setProperty("rate", 170)
                        self._apply_selected_tts_voice(engine)
                        clean = re.sub(r"[#*`_~\[\]()]", "", content)
                        clean = re.sub(r"\n+", ". ", clean)
                        with self._tts_lock:
                            self._tts_engine = engine
                        engine.say(clean[:2000])
                        engine.runAndWait()
                    finally:
                        try:
                            if engine is not None:
                                engine.stop()
                        except Exception:
                            pass
            except Exception as e:
                self._emit("tts_error", {"error": str(e)})
            finally:
                with self._tts_lock:
                    self._tts_engine = None
                    self._tts_active = False
                self._emit("tts_done", None)

        threading.Thread(target=_tts_worker, args=(text,), daemon=True).start()
        return {"ok": True}

    def stop_speaking(self):
        """Stop active text-to-speech playback if running."""
        # Signal Piper playback to stop
        self._tts_stop_event.set()

        # Also try to purge winsound in case Piper is playing
        try:
            import winsound
            winsound.PlaySound(None, winsound.SND_PURGE)
        except Exception:
            pass

        # Stop pyttsx3 engine if active
        engine = None
        with self._tts_lock:
            engine = self._tts_engine
        if engine is not None:
            try:
                engine.stop()
            except Exception:
                pass

        self._emit("tts_stopped", None)
        return {"ok": True, "stopped": True}

    def export_assistant_message_wav(self, chat_id: str, message_index: int):
        """Export one assistant message from chat to a WAV file."""
        try:
            if not chat_id or chat_id not in self.chats:
                return {"error": "No chat to export"}

            msgs = self.chats.get(chat_id, [])
            if message_index < 0 or message_index >= len(msgs):
                return {"error": "Message index out of range"}

            msg = msgs[message_index]
            if msg.get("role") != "assistant":
                return {"error": "Only assistant messages can be exported to WAV"}

            raw_text = str(msg.get("content", "")).strip()
            if not raw_text:
                return {"error": "Assistant message is empty"}

            # Light cleanup for TTS synthesis.
            clean = re.sub(r"[#*`_~\[\]()]", "", raw_text)
            clean = re.sub(r"\n+", ". ", clean).strip()

            exports_dir = app_data_path("exports")
            os.makedirs(exports_dir, exist_ok=True)
            timestamp = time.strftime("%Y%m%d_%H%M%S")
            base = f"{self._sanitize_export_name(chat_id)}_assistant_{message_index + 1}_{timestamp}"
            path = os.path.join(exports_dir, f"{base}.wav")

            selected = str((self.app_settings or {}).get("tts_voice_id", "")).strip()
            if selected.startswith("piper:"):
                self._piper_save_wav(raw_text, path)
            else:
                try:
                    import pyttsx3
                except Exception:
                    return {"error": "pyttsx3 is not available for WAV export"}
                engine = pyttsx3.init()
                self._apply_selected_tts_voice(engine)
                engine.save_to_file(clean[:12000], path)
                engine.runAndWait()

            # Ensure file exists before returning success.
            if not os.path.exists(path):
                return {"error": "WAV export failed"}

            self._emit("export_ready", {"path": path, "format": "wav", "chat_id": chat_id})
            self._status(f"Exported WAV to {path}")
            return {"ok": True, "path": path}
        except Exception as e:
            return {"error": f"WAV export failed: {e}"}

    # ── Generation ─────────────────────────────────────────────

    def send_message(self, text: str):
        """Process user message. Returns immediately; streams via events."""
        if not self._has_full_access():
            return {"error": "Trial expired. Enter passkey to activate full access."}
        text = text.strip()
        if not text:
            return {"error": "Empty message"}

        # Parse @rag_name mention → activate that RAG database for this message
        text, mentioned_rag = self._resolve_rag_mention(text)
        if mentioned_rag:
            self.current_rag_database = mentioned_rag
            if self.current_chat_id:
                self.chat_rag_settings[self.current_chat_id] = mentioned_rag
            self._status(f"RAG: {mentioned_rag}")

        request_options = self._parse_prompt_options(text)

        if not self.current_chat_id:
            self.new_chat()

        # Add user message
        msg = {"role": "user", "content": text}
        self.message_history.append(msg)
        self.chats[self.current_chat_id] = self.message_history
        self._emit("message_added", {"role": "user", "content": text})
        self._save_current_chat()

        # Plugin command short-circuit (e.g. /hello)
        plugin_out = self._run_plugin_command(text)
        if plugin_out.get("handled"):
            return {"ok": True, "plugin_command": True, "plugin_result": plugin_out.get("result")}

        # ── Direct file-export short-circuit ──────────────────────
        # When the user has a file uploaded and the intent is purely
        # "export this to csv/excel/docx", skip the LLM entirely and
        # produce the file immediately.
        export_fmt = (request_options or {}).get("export_format")
        if export_fmt and self.uploaded_file_path:
            def _direct_export():
                try:
                    self._status(f"Exporting to {export_fmt}...")
                    path = self._export_uploaded_file_as(export_fmt, self.current_chat_id or "export")
                    if path:
                        fname = os.path.basename(path)
                        reply = f"✅ Exported to **{export_fmt.upper()}**: `{fname}`\n\nFile saved at: `{path}`"
                        self._emit("message_added", {"role": "assistant", "content": reply})
                        self.message_history.append({"role": "assistant", "content": reply})
                        self.chats[self.current_chat_id] = self.message_history
                        self._save_current_chat()
                        self._emit("export_ready", {"path": path, "format": export_fmt,
                                                     "chat_id": self.current_chat_id})
                        self._status(f"Exported to {path}")
                    else:
                        # Could not export directly — fall back to normal LLM flow
                        self._generate(text, request_options)
                except Exception as exc:
                    self._emit("generation_error", {"error": str(exc)})
                    self._status(f"Export error: {exc}")
            threading.Thread(target=_direct_export, daemon=True).start()
            return {"ok": True}

        # Route
        if self.attached_image:
            image_path = self.attached_image
            self.attached_image = None
            self._status("Processing image...")
            threading.Thread(target=self._generate_with_image,
                             args=(text, image_path, request_options), daemon=True).start()
        elif self.web_search_enabled:
            self._status("Searching web...")
            threading.Thread(target=self._web_search_and_respond,
                             args=(text, request_options), daemon=True).start()
        elif self.uploaded_file_path and self._is_tabular_file(self.uploaded_file_name or "") and self._is_data_question(text):
            self._status("Analyzing uploaded table...")
            threading.Thread(target=self._run_tabular_chat_analysis,
                             args=(text, request_options), daemon=True).start()
        elif self.model is not None:
            self._status("AI is thinking...")
            threading.Thread(target=self._generate,
                             args=(text, request_options), daemon=True).start()
        else:
            self._emit("message_added",
                       {"role": "assistant",
                        "content": "⚠ No model loaded. Load a model or enable Web Search."})
            self._status("No model loaded")
        return {"ok": True}

    def edit_user_message(self, message_index: int, new_text: str):
        """Edit an existing user message, truncate following history, and regenerate."""
        if not self.current_chat_id:
            return {"error": "No active chat"}
        text = (new_text or "").strip()
        if not text:
            return {"error": "Edited message is empty"}

        try:
            idx = int(message_index)
        except Exception:
            return {"error": "Invalid message index"}

        if idx < 0 or idx >= len(self.message_history):
            return {"error": "Message index out of range"}
        if self.message_history[idx].get("role") != "user":
            return {"error": "Only user messages can be edited"}

        # Desktop parity: update message then truncate conversation after it.
        self.message_history[idx]["content"] = text
        self.message_history = self.message_history[: idx + 1]
        self.chats[self.current_chat_id] = self.message_history
        self._save_current_chat()

        plugin_out = self._run_plugin_command(text)
        if plugin_out.get("handled"):
            return {
                "ok": True,
                "messages": self.message_history,
                "regenerating": False,
                "chat_id": self.current_chat_id,
                "plugin_command": True,
                "plugin_result": plugin_out.get("result"),
            }

        regenerating = False
        if self.attached_image:
            regenerating = True
            image_path = self.attached_image
            self.attached_image = None
            threading.Thread(target=self._generate_with_image,
                             args=(text, image_path, self._parse_prompt_options(text)), daemon=True).start()
        elif self.web_search_enabled:
            regenerating = True
            threading.Thread(target=self._web_search_and_respond,
                             args=(text, self._parse_prompt_options(text)), daemon=True).start()
        elif self.model is not None:
            regenerating = True
            threading.Thread(target=self._generate,
                             args=(text, self._parse_prompt_options(text)), daemon=True).start()
        else:
            assistant_msg = {
                "role": "assistant",
                "content": "⚠ No model loaded. Load a model or enable Web Search.",
            }
            self.message_history.append(assistant_msg)
            self.chats[self.current_chat_id] = self.message_history
            self._save_current_chat()

        return {
            "ok": True,
            "messages": self.message_history,
            "regenerating": regenerating,
            "chat_id": self.current_chat_id,
        }

    def _generate(self, text: str, request_options: dict | None = None):
        self.generation_in_progress = True
        self.stop_generation_flag = False
        self._emit("generation_start", None)
        self._status("AI is thinking...")

        try:
            # ── 1) Gather context ─────────────────────────────────────
            user_system = self.chat_system_prompts.get(self.current_chat_id, "")
            system_prompt = user_system if user_system else self._DEFAULT_SYSTEM

            conversation_msgs = self._get_conversation_messages()
            rag_context = ""
            rag_sources = []
            file_context = self._get_relevant_uploaded_chunk(text)

            # RAG retrieval
            if self.current_rag_database and self.rag_manager:
                try:
                    results = self.rag_manager.retrieve(
                        self.current_rag_database, text, k=5)
                    if results:
                        chunks = [r[0] for r in results]
                        rag_sources = self._extract_rag_sources(chunks)
                        rag_context = "Reference context:\n" + "\n---\n".join(chunks)
                except Exception:
                    pass

            # Inject tool definitions so the model can use /code for exports,
            # computations, etc.  Skip only when RAG context is active AND the
            # user is NOT asking for an export/file operation (to save tokens).
            tool_defs = ""
            export_fmt = (request_options or {}).get("export_format")
            user_wants_action = export_fmt is not None or re.search(
                r"\b(export|save|convert|extract|download)\b.{0,20}\b(csv|xlsx|excel|file|pdf)\b",
                text, re.IGNORECASE,
            )
            if self.actual_n_ctx >= 2048 and (
                not rag_context or user_wants_action
            ):
                tool_defs = self._build_tool_definitions()
                if tool_defs:
                    system_prompt = system_prompt + "\n\n" + tool_defs

            # ── 2) Build extra context block ──────────────────────────
            extra_parts = []
            if rag_context:
                extra_parts.append(rag_context)
            if file_context:
                file_label = self.uploaded_file_name or "uploaded file"
                if len(self.uploaded_files) > 1:
                    file_label = ", ".join(f["name"] for f in self.uploaded_files)
                extra_parts.append(
                    f"Uploaded file(s) ({file_label}):\n{file_context}\n"
                    "Use this file context only when relevant to the question."
                )
            format_instruction = self._get_response_format_instruction(request_options)
            if format_instruction:
                extra_parts.append(f"Formatting: {format_instruction}")
            workflow_guidance = self._build_chat_execution_guidance(text)
            if workflow_guidance:
                extra_parts.append(f"Execution guidance:\n{workflow_guidance}")
            extra_context = "\n\n".join(extra_parts)

            # ── 3) Build prompt with proper chat template ─────────────
            prompt = self._build_chat_prompt(
                system=system_prompt,
                messages=conversation_msgs,
                user_text=text,
                extra_context=extra_context,
            )

            # ── 4) Token budget ───────────────────────────────────────
            try:
                prompt_tokens = len(self.model.tokenize(prompt.encode("utf-8")))
            except Exception:
                prompt_tokens = len(prompt) // 4
            available = max(64, self.actual_n_ctx - prompt_tokens - 64)
            max_tokens = available

            # Respect user-set max response tokens limit
            # -1 = "Max" (use full available context), 0 = Auto, >0 = fixed cap
            user_max = self.app_settings.get("max_response_tokens")
            if user_max and int(user_max) == -1:
                pass  # Max — use all available tokens
            elif user_max and int(user_max) > 0:
                max_tokens = min(max_tokens, int(user_max))

            self._emit("context_usage", {
                "prompt_tokens": prompt_tokens,
                "effective_window": self.actual_n_ctx,
                "max_tokens": max_tokens,
            })

            # ── 5) Adaptive sampling based on task ────────────────────
            task = self._detect_task_type(text)
            params = self._get_adaptive_params(task)
            stop_tokens = self._get_stop_tokens()
            # Add tool_call stop so model halts when it wants to call a tool
            if tool_defs and "</tool_call>" not in stop_tokens:
                stop_tokens = stop_tokens + ["</tool_call>"]

            response = ""
            token_count = 0
            was_stopped = False

            # ── 6) Stream inference ───────────────────────────────────
            with self.model_lock:
                if self.model is None:
                    raise RuntimeError("Model not loaded")

                stream = self.model(
                    prompt,
                    max_tokens=max_tokens,
                    temperature=params["temperature"],
                    top_p=params["top_p"],
                    top_k=params.get("top_k", 50),
                    repeat_penalty=params["repeat_penalty"],
                    stream=True,
                    stop=stop_tokens,
                )

                for chunk in stream:
                    if self.stop_generation_flag:
                        was_stopped = True
                        break
                    token = chunk["choices"][0].get("text", "")
                    response += token
                    token_count += 1
                    # Emit display every 3 tokens (preserve think blocks separately)
                    if token_count % 3 == 0:
                        think_blocks = _RE_THINK.findall(response)
                        display = _RE_THINK.sub("", response).strip()
                        display = re.sub(r"</?tool_call[^>]*>", "", display).strip()
                        # Include incomplete (in-progress) think block
                        incomplete = _RE_THINK_INCOMPLETE.search(response)
                        thinking_text = incomplete.group(0) if incomplete else ""
                        self._emit("generation_token", {"text": display, "think": think_blocks, "thinking": thinking_text})
                        # Update context usage
                        self._emit("context_usage", {
                            "prompt_tokens": prompt_tokens + token_count,
                            "effective_window": self.actual_n_ctx,
                            "max_tokens": max_tokens,
                        })

            # ── 7) Finalize output ────────────────────────────────────
            think_blocks = _RE_THINK.findall(response)
            display = _RE_THINK.sub("", response).strip()
            # Strip incomplete think tags at the end
            display = _RE_THINK_INCOMPLETE.sub("", display).strip()
            # Strip any malformed tool_call tags from display
            display = re.sub(r"</?tool_call[^>]*>", "", display).strip()
            display = re.sub(r"\[multimodal\]", "", display).strip()
            if not display:
                display = response.strip()

            if was_stopped:
                self._emit("generation_stopped", {"text": display})
                self._status("Generation stopped")
                return

            # ── 7b) Tool-call detection & execution ───────────────────
            tool_call = self._parse_tool_call(response + "</tool_call>") if tool_defs else None
            # Also check for MCP tools even without local plugins
            if not tool_call and hasattr(self, 'mcp_manager') and self.mcp_manager.get_all_tools():
                tool_call = self._parse_tool_call(response + "</tool_call>")
            if tool_call:
                cmd, args_str = tool_call

                # ── Permission gate for dangerous tools ────────────────
                is_mcp = isinstance(cmd, str) and cmd.startswith("mcp:")
                if is_mcp or cmd not in self._SAFE_TOOLS:
                    desc = self.plugin_commands.get(cmd, {}).get("description", cmd)
                    self._emit("tool_permission_request", {
                        "command": cmd,
                        "args": args_str,
                        "description": desc,
                    })
                    self._tool_permission_event.clear()
                    self._tool_permission_granted = False
                    self._status(f"Waiting for permission to run {cmd}...")
                    # Wait up to 60 s for user response
                    self._tool_permission_event.wait(timeout=60)
                    if not self._tool_permission_granted:
                        display = f"⚠️ Tool `{cmd}` was not approved. Operation cancelled."
                        self._emit("generation_token", {"text": display})
                        self._emit("generation_done", {"text": display})
                        self.message_history.append({"role": "assistant", "content": display})
                        self.chats[self.current_chat_id] = self.message_history
                        self._save_current_chat()
                        return

                self._status(f"Running tool: {cmd}...")
                self._emit("generation_token", {"text": f"🔧 Running `{cmd}`..."})

                tool_result = self._execute_tool_call(cmd, args_str)

                # Strip the tool_call block from what the user sees
                visible_before = self._TOOL_CALL_RE.sub("", display).strip()
                # Remove incomplete trailing <tool_call>... from display
                visible_before = re.sub(r"<tool_call>.*", "", visible_before, flags=re.DOTALL).strip()

                # Build follow-up prompt: ask model to summarize the tool result
                followup_msgs = list(conversation_msgs)  # copy
                followup_msgs.append({"role": "user", "content": text})
                followup_msgs.append({"role": "assistant", "content":
                    f"I used the `{cmd}` tool. Here is the result:\n\n{tool_result}"})
                followup_prompt = self._build_chat_prompt(
                    system=system_prompt,
                    messages=followup_msgs,
                    user_text="Now present this tool result to the user in a clear, readable format. "
                              "Use markdown. Be concise.",
                    extra_context="",
                )
                # Generate follow-up summary
                followup_response = ""
                with self.model_lock:
                    if self.model is not None:
                        try:
                            fup_tokens = min(512, max(64, self.actual_n_ctx - len(followup_prompt) // 4 - 64))
                            # Remove </tool_call> from stop tokens for the follow-up
                            fup_stop = [s for s in stop_tokens if s != "</tool_call>"]
                            fup_stream = self.model(
                                followup_prompt,
                                max_tokens=fup_tokens,
                                temperature=0.3,
                                top_p=0.9,
                                repeat_penalty=1.1,
                                stream=True,
                                stop=fup_stop,
                            )
                            for chunk in fup_stream:
                                if self.stop_generation_flag:
                                    break
                                token = chunk["choices"][0].get("text", "")
                                followup_response += token
                                fup_display = _RE_THINK.sub("", followup_response).strip()
                                self._emit("generation_token", {"text": fup_display})
                        except Exception as e:
                            print(f"[TOOLS] Follow-up generation error: {e}")

                followup_clean = _RE_THINK.sub("", followup_response).strip()
                followup_clean = _RE_THINK_INCOMPLETE.sub("", followup_clean).strip()

                # Final display: tool result (or AI summary if model produced one)
                if followup_clean:
                    display = followup_clean
                else:
                    display = tool_result

            # Build structured doc sources for card display
            doc_sources = []
            if rag_context and self.current_rag_database:
                if rag_sources:
                    for src in rag_sources:
                        doc_sources.append({"name": src["name"], "type": "rag", "snippet": src.get("snippet", ""), "db": self.current_rag_database})
                else:
                    doc_sources.append({"name": self.current_rag_database, "type": "rag", "snippet": "", "db": self.current_rag_database})
            if file_context and self.uploaded_file_name:
                snippet = (file_context[:120].replace("\n", " ").strip())
                if len(snippet) > 117:
                    snippet = snippet[:117] + "..."
                doc_sources.append({"name": self.uploaded_file_name, "type": "file", "snippet": snippet})

            if doc_sources:
                self._emit("doc_sources", {"sources": doc_sources})

            self._emit("generation_done", {"text": display, "think": think_blocks, "doc_sources": doc_sources if doc_sources else None})

            # Save
            self.message_history.append(
                {"role": "assistant", "content": display})
            self.chats[self.current_chat_id] = self.message_history
            self._save_current_chat()
            self._auto_export_response(
                self.current_chat_id,
                len(self.message_history) - 1,
                (request_options or {}).get("export_format"),
            )

        except Exception as e:
            self._emit("generation_error", {"error": str(e)})
            self._status(f"Generation error: {e}")
        finally:
            self.generation_in_progress = False
            self._status("Ready")

    def stop_generation(self):
        """Stop current generation.

        Sets the flag immediately so all generation loops notice it on their
        next iteration.  The ``generation_stopped`` event is emitted here ONLY
        when nothing is actively running — streaming chat emits the event
        itself once the token loop exits, and agent_chat is a synchronous call
        whose JS finally-block resets the UI when it returns.
        Firing the event prematurely would reset the UI while the backend
        thread is still running, making stop appear to do nothing.
        """
        self.stop_generation_flag = True
        self._status("Stopping generation...")
        if not self.generation_in_progress:
            # Nothing blocking: emit now so the UI resets immediately
            self._emit("generation_stopped", {"text": ""})
        return {"ok": True}

    # ── Prompt engineering helpers ─────────────────────────────

    _DEFAULT_SYSTEM = (
        "You are a helpful, accurate, and concise AI assistant. "
        "Follow these guidelines:\n"
        "- Answer directly and stay on topic. Do not repeat the question.\n"
        "- Use markdown formatting: **bold** for emphasis, `code` for code terms, "
        "bullet lists for multiple points, and code blocks with language tags for code.\n"
        "- For code: write clean, idiomatic, production-quality code with brief explanations.\n"
        "- For math/logic: show your reasoning step-by-step.\n"
        "- For factual questions: be precise. If uncertain, say so.\n"
        "- Be concise. Prefer short paragraphs over long walls of text.\n"
        "- Never fabricate URLs, citations, or data."
    )

    def _detect_model_family(self) -> str:
        """Detect model family from filename for chat template selection."""
        name = Path(self.model_path or "").stem.lower()
        if "gemma" in name:
            # Gemma 4 uses different chat tokens than Gemma 3 and earlier
            if "gemma-4" in name or "gemma_4" in name or "gemma4" in name:
                return "gemma4"
            return "gemma"
        if "qwen" in name:
            return "qwen"
        if "llama" in name or "meta" in name:
            return "llama"
        if "mistral" in name or "mixtral" in name:
            return "mistral"
        if "phi" in name:
            return "phi"
        return "generic"

    def _get_stop_tokens(self) -> list[str]:
        """Return model-family-specific stop tokens."""
        family = self._detect_model_family()
        base = ["</s>", "<|im_end|>", "<|end|>", "<|eot_id|>"]
        if family == "gemma4":
            return base + ["<turn|>", "<|turn>", "<channel|>"]
        if family == "gemma":
            return base + ["<end_of_turn>", "<start_of_turn>"]
        if family == "qwen":
            return base + ["<|endoftext|>"]
        if family == "llama":
            return base + ["<|eot_id|>"]
        if family == "phi":
            return base + ["<|assistant|>", "<|user|>"]
        return base + ["User:", "\nUser:"]

    def _build_chat_prompt(self, system: str, messages: list[dict],
                           user_text: str, extra_context: str = "") -> str:
        """Build a properly formatted chat prompt using the right template."""
        family = self._detect_model_family()

        # Clean messages — strip think tags, limit content length
        cleaned = []
        for m in messages:
            content = _RE_THINK.sub("", m["content"]).strip()
            # Truncate very long messages to save context
            if len(content) > 800:
                content = content[:780] + "..."
            cleaned.append({"role": m["role"], "content": content})

        if family == "gemma4":
            return self._build_gemma4_prompt(system, cleaned, user_text, extra_context)
        elif family == "gemma":
            return self._build_gemma_prompt(system, cleaned, user_text, extra_context)
        elif family == "qwen":
            return self._build_chatml_prompt(system, cleaned, user_text, extra_context)
        else:
            return self._build_chatml_prompt(system, cleaned, user_text, extra_context)

    def _build_gemma_prompt(self, system: str, messages: list, user_text: str, extra: str) -> str:
        parts = []
        # Gemma 3 and earlier: <start_of_turn> / <end_of_turn>
        if system:
            parts.append(f"<start_of_turn>user\nSystem instructions: {system}<end_of_turn>")
        if extra:
            parts.append(f"<start_of_turn>user\n{extra}<end_of_turn>")
        for m in messages:
            role = "user" if m["role"] == "user" else "model"
            parts.append(f"<start_of_turn>{role}\n{m['content']}<end_of_turn>")
        parts.append(f"<start_of_turn>user\n{user_text}<end_of_turn>")
        parts.append("<start_of_turn>model\n")
        return "\n".join(parts)

    def _build_gemma4_prompt(self, system: str, messages: list, user_text: str, extra: str) -> str:
        parts = []
        # Gemma 4: <|turn> / <turn|> with native system role support
        sys_content = system
        if extra:
            sys_content += f"\n\n{extra}" if sys_content else extra
        if sys_content:
            parts.append(f"<|turn>system\n{sys_content}<turn|>")
        for m in messages:
            role = "user" if m["role"] == "user" else "model"
            parts.append(f"<|turn>{role}\n{m['content']}<turn|>")
        parts.append(f"<|turn>user\n{user_text}<turn|>")
        parts.append("<|turn>model\n")
        return "\n".join(parts)

    def _build_chatml_prompt(self, system: str, messages: list, user_text: str, extra: str) -> str:
        parts = []
        # ChatML format used by Qwen, Llama, Mistral, Phi, and most modern models
        sys_content = system
        if extra:
            sys_content += f"\n\n{extra}" if sys_content else extra
        if sys_content:
            parts.append(f"<|im_start|>system\n{sys_content}<|im_end|>")
        for m in messages:
            role = m["role"]
            parts.append(f"<|im_start|>{role}\n{m['content']}<|im_end|>")
        parts.append(f"<|im_start|>user\n{user_text}<|im_end|>")
        parts.append("<|im_start|>assistant\n")
        return "\n".join(parts)

    def _detect_task_type(self, text: str) -> str:
        """Classify user intent for adaptive sampling."""
        q = text.lower()
        if re.search(r'\b(write|code|function|class|implement|debug|fix|refactor|script|program|api|html|css|sql|python|java|javascript)\b', q):
            return "code"
        if re.search(r'\b(create|story|poem|essay|write me|creative|imagine|generate|brainstorm)\b', q):
            return "creative"
        if re.search(r'\b(explain|what is|define|how does|why|difference between|compare|summarize|summary)\b', q):
            return "explain"
        if re.search(r'\b(calculate|solve|math|equation|formula|proof|logic)\b', q):
            return "reasoning"
        return "general"

    def _get_adaptive_params(self, task: str) -> dict:
        """Return optimal sampling parameters for the task type."""
        base_temp = self.model_config.get("temperature", 0.25)
        params = {
            "code":      {"temperature": min(base_temp, 0.15), "top_p": 0.85, "repeat_penalty": 1.05, "top_k": 40},
            "reasoning": {"temperature": min(base_temp, 0.10), "top_p": 0.80, "repeat_penalty": 1.05, "top_k": 30},
            "explain":   {"temperature": max(base_temp, 0.30), "top_p": 0.90, "repeat_penalty": 1.10, "top_k": 50},
            "creative":  {"temperature": max(base_temp, 0.60), "top_p": 0.95, "repeat_penalty": 1.15, "top_k": 60},
            "general":   {"temperature": base_temp, "top_p": 0.88, "repeat_penalty": 1.10, "top_k": 50},
        }
        return params.get(task, params["general"])

    def _get_conversation_context(self, max_messages=10):
        msgs = self.message_history[:-1]  # exclude the current user msg
        recent = msgs[-max_messages:]
        lines = []
        for m in recent:
            role = "User" if m["role"] == "user" else "Assistant"
            content = _RE_THINK.sub("", m["content"]).strip()
            lines.append(f"{role}: {content}")
        return "\n".join(lines)

    def _get_conversation_messages(self, max_messages: int = 12) -> list[dict]:
        """Return recent conversation as structured messages for chat template."""
        msgs = self.message_history[:-1]  # exclude the current user msg (appended before _generate)
        recent = msgs[-max_messages:]
        result = []
        for m in recent:
            content = _RE_THINK.sub("", m["content"]).strip()
            if content:
                result.append({"role": m["role"], "content": content})
        return result

    def _get_relevant_uploaded_chunk(self, query: str, max_chars: int = 0):
        """Retrieve the most relevant portion of uploaded content for the query.
        Searches across all uploaded files when multiple files are present.
        """
        # Collect all content sources
        all_content = self.uploaded_content
        all_pages = self.uploaded_pages
        file_sources = []

        if len(self.uploaded_files) > 1:
            # Multi-file: gather content from all files
            for uf in self.uploaded_files:
                uf_content = uf.get("content")
                uf_pages = uf.get("pages")
                if uf_pages:
                    for p in uf_pages:
                        file_sources.append((uf["name"], p))
                elif uf_content:
                    file_sources.append((uf["name"], uf_content))

        if not all_content and not file_sources:
            return ""

        # Adaptive max_chars: use ~30% of model context window if not specified
        if max_chars <= 0:
            max_chars = max(6000, int(self.actual_n_ctx * 3.5 * 0.30))

        query_terms = [t.lower() for t in re.findall(r"\b\w{3,}\b", query)]

        # Multi-file search
        if file_sources:
            scored = []
            for fname, text in file_sources:
                low = text.lower()
                score = sum(1 for t in query_terms if t in low)
                q_lower = query.lower()
                words = q_lower.split()
                for w in range(len(words) - 1):
                    bigram = words[w] + " " + words[w + 1]
                    if bigram in low:
                        score += 2
                scored.append((score, fname, text))

            scored.sort(key=lambda x: x[0], reverse=True)
            selected = []
            total = 0
            for _score, fname, text in scored:
                header = f"[{fname}]\n"
                chunk = header + text
                if total + len(chunk) + 2 > max_chars:
                    remaining = max_chars - total
                    if remaining > 200:
                        selected.append(chunk[:remaining])
                    break
                selected.append(chunk)
                total += len(chunk) + 2
            if selected:
                return "\n\n".join(selected)

        # Single file - page-level search
        pages = all_pages
        if pages and len(pages) > 1:
            scored_pages = []
            for i, page in enumerate(pages):
                low = page.lower()
                # Score: count of matched terms + bonus for exact phrase fragments
                score = sum(1 for t in query_terms if t in low)
                # Bonus for multi-word phrase matches
                q_lower = query.lower()
                words = q_lower.split()
                for w in range(len(words) - 1):
                    bigram = words[w] + " " + words[w + 1]
                    if bigram in low:
                        score += 2
                if score > 0:
                    scored_pages.append((score, i, page))

            if scored_pages:
                scored_pages.sort(key=lambda x: x[0], reverse=True)
                selected = []
                total = 0
                for _score, _idx, page_text in scored_pages:
                    if total + len(page_text) + 2 > max_chars:
                        # Add partial page if space remains
                        remaining = max_chars - total
                        if remaining > 200:
                            selected.append(page_text[:remaining])
                        break
                    selected.append(page_text)
                    total += len(page_text) + 2
                return "\n\n".join(selected)

            # No keyword match — return first + last pages as context
            first_pages = "\n\n".join(pages[:3])
            if len(first_pages) > max_chars:
                return first_pages[:max_chars]
            return first_pages

        # Fallback: line-level search for non-paged content
        lines = [ln.strip() for ln in self.uploaded_content.splitlines() if ln.strip()]
        if not lines:
            return self.uploaded_content[:max_chars]

        scored = []
        for ln in lines:
            low = ln.lower()
            score = sum(1 for t in query_terms if t in low)
            if score > 0:
                scored.append((score, ln))

        if not scored:
            return self.uploaded_content[:max_chars]

        scored.sort(key=lambda it: (it[0], len(it[1])), reverse=True)
        selected = []
        total = 0
        for _score, line in scored[:30]:
            if total + len(line) + 1 > max_chars:
                break
            selected.append(line)
            total += len(line) + 1
        return "\n".join(selected)

    # ── Web search ─────────────────────────────────────────────

    def toggle_web_search(self):
        self.web_search_enabled = not self.web_search_enabled
        return {"enabled": self.web_search_enabled}

    def _web_search_and_respond(self, query: str, request_options: dict | None = None):
        self.generation_in_progress = True
        self.stop_generation_flag = False
        self._emit("generation_start", None)

        # Step 1 — AI query refinement
        self._status("Refining search query...")
        refined_query = self._refine_search_query(query)
        self._emit("web_search_start", {"query": refined_query})
        self._status("Searching web...")

        try:
            context, sources = self._search_web(refined_query)
            # Emit sources for card rendering
            if sources:
                self._emit("web_sources", {"sources": sources})
            if self.stop_generation_flag:
                self._emit("generation_stopped", {"text": ""})
                self._status("Generation stopped")
                return

            # Step 2 — Deepening: try original query & alternative phrasing
            if len(sources) < 4 and not self.stop_generation_flag:
                self._status("Deepening research...")
                seen_urls = {s["url"] for s in sources}

                # Try original (un-refined) query
                _, src2 = self._search_web(query)
                for s in (src2 or []):
                    if s["url"] not in seen_urls:
                        sources.append(s)
                        seen_urls.add(s["url"])

                # If still thin, try a broader variant
                if len(sources) < 4 and self.model is not None:
                    broader = self._refine_search_query(
                        f"What is {query}? overview explanation")
                    _, src3 = self._search_web(broader)
                    for s in (src3 or []):
                        if s["url"] not in seen_urls:
                            sources.append(s)
                            seen_urls.add(s["url"])

                if len(sources) > 0:
                    # Rebuild context with merged sources
                    n_ctx = getattr(self, "actual_n_ctx", 2048)
                    per_src = max(400, int(n_ctx * 3.5 * 0.30) // max(1, len(sources)))
                    lines = []
                    for i, s in enumerate(sources, 1):
                        snippet = self._truncate_at_sentence(
                            s.get("snippet", ""), 350)
                        entry = f"[{i}] {s['title']}: {snippet}"
                        if s.get("excerpt"):
                            exc = self._truncate_at_sentence(
                                s["excerpt"], per_src)
                            entry += f"\n  Content: {exc}"
                        lines.append(entry)
                    context = "\n".join(lines)
                    self._emit("web_sources", {"sources": sources})

            if self.model is not None:
                self._status("Generating answer...")
                self._generate_from_web_context(context, sources, query, request_options)
            else:
                if self.stop_generation_flag:
                    self._emit("generation_stopped", {"text": ""})
                    self._status("Generation stopped")
                    return
                # No model: just show raw results
                display = context
                self._emit("generation_done", {"text": display, "web_sources": sources})
                self.message_history.append(
                    {"role": "assistant", "content": display})
                self.chats[self.current_chat_id] = self.message_history
                self._save_current_chat()
                self._auto_export_response(
                    self.current_chat_id,
                    len(self.message_history) - 1,
                    (request_options or {}).get("export_format"),
                )
                self._status("Ready")
        except Exception as e:
            self._emit("generation_error", {"error": str(e)})
            self._status(f"Web search error: {e}")
            print(f"[WEB] Search+respond error: {type(e).__name__}: {e}")
        finally:
            self.generation_in_progress = False

    def _refine_search_query(self, user_query: str) -> str:
        """Use the loaded model to rewrite a casual user question into a
        better web-search query.  Returns the original query on any failure."""
        if self.model is None:
            return user_query
        try:
            prompt = self._build_chat_prompt(
                system=(
                    "You are a search-query optimizer. "
                    "Rewrite the user's question into a concise, keyword-rich "
                    "web search query (max 12 words). "
                    "Output ONLY the search query, nothing else."
                ),
                messages=[],
                user_text=user_query,
            )
            with self.model_lock:
                if self.model is None:
                    return user_query
                r = self.model.create_completion(
                    prompt, max_tokens=40, temperature=0.1,
                    stop=["\n", "<|im_end|>", "<end_of_turn>"])
            refined = (r.get("choices", [{}])[0].get("text", "") or "").strip()
            refined = _RE_THINK.sub("", refined).strip().strip('"').strip("'")
            if 3 <= len(refined) <= 120:
                return refined
        except Exception:
            pass
        return user_query

    def _search_web(self, query: str, num_results: int = 8) -> tuple:
        """Return (context_str, sources_list).

        Priority: Brave API (if key set) → DuckDuckGo → Bing (fallback).
        If primary engine returns thin results (<4), merges with next engine
        for better coverage. Results cached per-session."""
        cache_key = query.strip().lower()
        if cache_key in self._web_search_cache:
            return self._web_search_cache[cache_key]

        primary_sources = []
        brave_key = (self.app_settings or {}).get("brave_api_key", "").strip()

        # Try Brave first
        if brave_key:
            result = self._search_brave(query, num_results, brave_key)
            if result[1]:
                primary_sources = result[1]

        # DDG — either as primary or supplementary
        if len(primary_sources) < 4:
            ddg_result = self._search_ddg(query, num_results)
            if ddg_result[1]:
                if not primary_sources:
                    primary_sources = ddg_result[1]
                else:
                    # Merge unique DDG results
                    seen = {s["url"] for s in primary_sources}
                    for s in ddg_result[1]:
                        if s["url"] not in seen:
                            primary_sources.append(s)
                            seen.add(s["url"])

        # Bing — as final fallback or supplementary
        if len(primary_sources) < 4:
            bing_result = self._search_bing(query, num_results)
            if bing_result[1]:
                if not primary_sources:
                    primary_sources = bing_result[1]
                else:
                    seen = {s["url"] for s in primary_sources}
                    for s in bing_result[1]:
                        if s["url"] not in seen:
                            primary_sources.append(s)
                            seen.add(s["url"])

        # Trim to num_results and build context
        primary_sources = primary_sources[:num_results]
        if primary_sources:
            result = self._build_search_context(primary_sources, query)
        else:
            result = ("No results found.", [])

        self._web_search_cache[cache_key] = result
        print(f"[WEB] Search complete: {len(primary_sources)} sources for '{query[:60]}'")
        return result

    # ── Individual search engine helpers ──────────────────────

    def _build_search_context(self, sources: list, query: str) -> tuple:
        """Fetch excerpts in parallel and build LLM context string.
        Dynamically scales excerpt size to use available context window."""
        # Dynamic excerpt limit based on model context
        n_ctx = getattr(self, "actual_n_ctx", 2048)
        # Use ~30% of context for web content, convert tokens→chars (~3.5 chars/token)
        max_web_chars = int(n_ctx * 3.5 * 0.30)
        per_source_chars = max(400, max_web_chars // max(1, len(sources)))

        # Parallel excerpt fetching — up to 8 workers
        need_excerpt = [s for s in sources if not s.get("excerpt")]
        if need_excerpt:
            def _fetch(s):
                try:
                    s["excerpt"] = self._fetch_web_excerpt(
                        s["url"], query, max_chars=per_source_chars)
                except Exception as e:
                    print(f"[WEB] Excerpt fetch failed for {s.get('url','?')}: "
                          f"{type(e).__name__}")
            with concurrent.futures.ThreadPoolExecutor(max_workers=8) as pool:
                pool.map(_fetch, need_excerpt)

        lines = []
        for i, s in enumerate(sources, 1):
            snippet = self._truncate_at_sentence(s.get("snippet", ""), 350)
            entry = f"[{i}] {s['title']}: {snippet}"
            if s.get("excerpt"):
                exc = self._truncate_at_sentence(s["excerpt"], per_source_chars)
                entry += f"\n  Content: {exc}"
            lines.append(entry)
        context = "\n".join(lines) if lines else "No results found."
        # Final cap to prevent context overflow
        if len(context) > max_web_chars:
            context = context[:max_web_chars].rsplit("\n", 1)[0]
        return context, sources

    def _search_brave(self, query: str, num_results: int, api_key: str) -> tuple:
        """Brave Search API — best quality, requires free API key."""
        try:
            import requests
            r = requests.get(
                "https://api.search.brave.com/res/v1/web/search",
                headers={
                    "Accept": "application/json",
                    "Accept-Encoding": "gzip",
                    "X-Subscription-Token": api_key,
                },
                params={"q": query, "count": num_results},
                timeout=10,
            )
            r.raise_for_status()
            data = r.json()
            results = data.get("web", {}).get("results", [])
            sources = []
            for item in results[:num_results]:
                snippet_raw = item.get("description", "")
                sources.append({
                    "title": item.get("title", ""),
                    "url": item.get("url", ""),
                    "snippet": self._truncate_at_sentence(snippet_raw, 350),
                    "excerpt": item.get("extra_snippets", [""])[0][:600]
                              if item.get("extra_snippets") else "",
                })
            return self._build_search_context(sources, query)
        except requests.exceptions.HTTPError as e:
            print(f"[WEB] Brave API HTTP error: {e.response.status_code}")
            return "Search error.", []
        except Exception as e:
            print(f"[WEB] Brave search failed: {type(e).__name__}: {e}")
            return "Search error.", []

    def _search_ddg(self, query: str, num_results: int) -> tuple:
        """DuckDuckGo search — prefers `duckduckgo_search` library,
        falls back to HTML scraping, then Lite page."""
        # Attempt 1: Use ddgs library (most reliable)
        try:
            try:
                from ddgs import DDGS
            except ImportError:
                from duckduckgo_search import DDGS
            with DDGS() as ddgs:
                raw = list(ddgs.text(query, max_results=num_results))
            sources = []
            seen = set()
            for item in raw:
                title = item.get("title", "")
                url = item.get("href", "") or item.get("link", "")
                snippet = item.get("body", "") or item.get("snippet", "")
                if title and url and snippet:
                    key = (title.lower(), url)
                    if key in seen:
                        continue
                    seen.add(key)
                    sources.append({
                        "title": title,
                        "url": url,
                        "snippet": self._truncate_at_sentence(
                            " ".join(snippet.split()), 350),
                        "excerpt": "",
                    })
            if sources:
                print(f"[WEB] DDG library returned {len(sources)} results")
                return self._build_search_context(sources, query)
        except ImportError:
            pass   # Library not installed — fall through to HTML
        except Exception as e:
            print(f"[WEB] DDG library error: {type(e).__name__}: {e}")

        # Attempt 2: HTML scraping
        try:
            from bs4 import BeautifulSoup
            import requests
            from urllib.parse import quote
            headers = {
                "User-Agent": (
                    "Mozilla/5.0 (Windows NT 10.0; Win64; x64) "
                    "AppleWebKit/537.36 (KHTML, like Gecko) "
                    "Chrome/125.0.0.0 Safari/537.36"
                ),
                "Accept": "text/html,application/xhtml+xml,application/xml;q=0.9,*/*;q=0.8",
                "Accept-Language": "en-US,en;q=0.5",
                "Referer": "https://duckduckgo.com/",
            }
            search_query = " ".join(query.split()[:12])
            session = requests.Session()
            try:
                session.get("https://duckduckgo.com/", headers=headers, timeout=5)
            except Exception:
                pass
            r = session.get(
                f"https://html.duckduckgo.com/html/?q={quote(search_query)}",
                headers=headers, timeout=10)
            soup = BeautifulSoup(r.text, "html.parser")
            results_div = soup.find_all("div", class_="result__body")
            # Fallback to DDG Lite if blocked
            if not results_div and r.status_code != 200:
                r2 = session.get(
                    f"https://lite.duckduckgo.com/lite/?q={quote(search_query)}",
                    headers=headers, timeout=10)
                soup2 = BeautifulSoup(r2.text, "html.parser")
                results_div = soup2.find_all("div", class_="result__body")
                if not results_div:
                    return self._parse_ddg_lite(soup2, query, num_results)
            sources = []
            seen = set()
            for div in results_div:
                title_tag = div.find("a", class_="result__a")
                snippet_tag = div.find(class_="result__snippet")
                title = title_tag.get_text(strip=True) if title_tag else ""
                link = self._resolve_search_result_url(
                    title_tag.get("href", "") if title_tag else "")
                snippet = snippet_tag.get_text(strip=True) if snippet_tag else ""
                if title and snippet:
                    key = (title.lower(), link)
                    if key in seen:
                        continue
                    seen.add(key)
                    sources.append({
                        "title": title, "url": link,
                        "snippet": self._truncate_at_sentence(
                            " ".join(snippet.split()), 350),
                        "excerpt": "",
                    })
                if len(sources) >= num_results:
                    break
            if sources:
                print(f"[WEB] DDG HTML returned {len(sources)} results")
            return self._build_search_context(sources, query)
        except Exception as e:
            print(f"[WEB] DDG HTML scraping failed: {type(e).__name__}: {e}")
            return "", []

    def _search_bing(self, query: str, num_results: int) -> tuple:
        """Bing HTML scraping — no API key required, used as fallback."""
        try:
            from bs4 import BeautifulSoup
            import requests
            from urllib.parse import quote
            headers = {
                "User-Agent": (
                    "Mozilla/5.0 (Windows NT 10.0; Win64; x64) "
                    "AppleWebKit/537.36 (KHTML, like Gecko) "
                    "Chrome/125.0.0.0 Safari/537.36"
                ),
                "Accept-Language": "en-US,en;q=0.9",
            }
            r = requests.get(
                f"https://www.bing.com/search?q={quote(query)}&count={num_results}",
                headers=headers, timeout=10)
            soup = BeautifulSoup(r.text, "html.parser")
            sources = []
            seen = set()
            for li in soup.select("li.b_algo"):
                h2 = li.find("h2")
                a = h2.find("a") if h2 else None
                cap = li.find("div", class_="b_caption") or li.find("p")
                title = a.get_text(strip=True) if a else ""
                link = a.get("href", "") if a else ""
                snippet = cap.get_text(strip=True) if cap else ""
                if title and link and snippet:
                    key = (title.lower(), link)
                    if key in seen:
                        continue
                    seen.add(key)
                    sources.append({
                        "title": title, "url": link,
                        "snippet": self._truncate_at_sentence(
                            " ".join(snippet.split()), 350),
                        "excerpt": "",
                    })
                if len(sources) >= num_results:
                    break
            if sources:
                print(f"[WEB] Bing returned {len(sources)} results")
            return self._build_search_context(sources, query)
        except Exception as e:
            print(f"[WEB] Bing search failed: {type(e).__name__}: {e}")
            return f"Search error: {type(e).__name__}", []

    def _generate_from_web_context(self, context: str, sources: list,
                                    query: str,
                                    request_options: dict | None = None):
        format_instruction = self._get_response_format_instruction(request_options)
        # Build numbered source list for the prompt
        src_labels = ""
        for i, s in enumerate(sources, 1):
            src_labels += f"  [{i}] {s['title']}\n"
        web_system = (
            "You are a thorough AI research assistant. "
            "IMPORTANT: Answer ONLY based on the search results provided below. "
            "Do NOT use prior knowledge or make up any information not found in the results. "
            "If the search results do not contain enough information to answer, explicitly say so instead of guessing. "
            "Synthesize the search results into a comprehensive, well-structured answer. "
            "Use paragraphs, bullet points, or numbered lists to organize your response clearly. "
            "ALWAYS add citation numbers like [1], [2] after every factual claim to show which source it came from. "
            "If sources disagree, mention both perspectives and cite each. "
            "If the results are insufficient, say what you found and what's still unclear.\n"
            "Sources:\n" + src_labels +
            "Do NOT simply list the search results. Synthesize them into a coherent answer."
        )
        # Dynamic context cap — use up to 40% of context window for web content
        max_web_chars = int(self.actual_n_ctx * 3.5 * 0.40)
        trimmed = context[:max(4000, max_web_chars)]
        extra = f"Search results:\n{trimmed}"
        if format_instruction:
            extra += f"\n\nFormatting: {format_instruction}"
        prompt = self._build_chat_prompt(
            system=web_system,
            messages=[],
            user_text=query,
            extra_context=extra,
        )
        try:
            prompt_tokens = len(self.model.tokenize(prompt.encode("utf-8")))
        except Exception:
            prompt_tokens = len(prompt) // 4
        available = max(64, self.actual_n_ctx - prompt_tokens - 64)
        max_tokens = available
        stop_tokens = self._get_stop_tokens()

        response = ""
        count = 0
        was_stopped = False
        
        # Lock during inference to prevent model unload
        with self.model_lock:
            if self.model is None:
                display = "Model not available for web search processing."
                self._emit("generation_error", {"error": display})
                self._status("Generation error: " + display)
                return
            
            stream = self.model(
                prompt, max_tokens=max_tokens, temperature=0.3,
                top_p=0.92, stream=True, stop=stop_tokens)

            for chunk in stream:
                if self.stop_generation_flag:
                    was_stopped = True
                    break
                token = chunk["choices"][0].get("text", "")
                response += token
                count += 1
                if count % 3 == 0:
                    display = _RE_THINK.sub("", response).strip()
                    self._emit("generation_token", {"text": display})

        display = _RE_THINK.sub("", response).strip()
        if was_stopped:
            self._emit("generation_stopped", {"text": display})
            self._status("Generation stopped")
            return

        if not display:
            # Fallback when the model returns empty/near-empty output.
            if context and not context.startswith(("Search error:", "No results found.")):
                top_lines = [ln.strip() for ln in context.splitlines() if ln.strip()][:3]
                display = (
                    "I found web results but could not generate a full summary. "
                    "Here are the top snippets:\n" + "\n".join(top_lines)
                )
            elif context:
                display = context
            else:
                display = "No web results available for this query."

        self._emit("generation_done", {"text": display, "web_sources": sources})
        self.message_history.append(
            {"role": "assistant", "content": display})
        self.chats[self.current_chat_id] = self.message_history
        self._save_current_chat()
        self._auto_export_response(
            self.current_chat_id,
            len(self.message_history) - 1,
            (request_options or {}).get("export_format"),
        )
        self._status("Ready")

    # ── Compare models ────────────────────────────────────────

    def compare_models(self, model_a_label: str, model_b_label: str, prompt_text: str):
        """Run side-by-side comparison using two models."""
        if not prompt_text or not prompt_text.strip():
            return {"error": "Prompt is empty"}
        path_a = self.model_map.get(model_a_label)
        path_b = self.model_map.get(model_b_label)
        if not path_a or not path_b:
            return {"error": "Invalid model selection"}

        original_model = self.model
        original_path = self.model_path
        result = {"model_a": model_a_label, "model_b": model_b_label,
                  "response_a": "", "response_b": ""}

        try:
            for key, path in (("response_a", path_a), ("response_b", path_b)):
                temp_model = Llama(
                    model_path=path,
                    n_threads=self.config["n_threads"],
                    n_gpu_layers=self.config["n_gpu_layers"],
                    n_ctx=2048,
                    verbose=False,
                )
                out = temp_model(
                    f"Question: {prompt_text}\n\nAnswer:",
                    max_tokens=512,
                    temperature=0.25,
                    top_p=0.9,
                    stream=False,
                )
                result[key] = out["choices"][0]["text"].strip() if out.get("choices") else "No response"
                del temp_model
                gc.collect()
        except Exception as e:
            return {"error": str(e)}
        finally:
            # keep current loaded model untouched
            self.model = original_model
            self.model_path = original_path

        return {"ok": True, **result}

    # ── Image support ─────────────────────────────────────────

    def get_vision_status(self):
        """Return whether multimodal assets are available."""
        if not self.model_path:
            return {"ready": False, "reason": "No model selected"}
        model_dir = os.path.dirname(self.model_path)
        try:
            for f in os.listdir(model_dir):
                lower = f.lower()
                if "mmproj" in lower or "clip" in lower:
                    return {"ready": True, "clip": os.path.join(model_dir, f)}
        except Exception:
            pass
        return {
            "ready": False,
            "reason": "Missing mmproj/clip file for multimodal model",
        }

    def open_image_dialog(self):
        """Open native image picker and store selected path."""
        try:
            import webview
            if not webview.windows:
                return {"error": "Window not ready"}
            paths = webview.windows[0].create_file_dialog(
                webview.FileDialog.OPEN,
                allow_multiple=False,
                file_types=("Images (*.png;*.jpg;*.jpeg;*.webp;*.bmp;*.gif)",),
            )
            if not paths:
                self.attached_image = None
                return {"selected": False}
            self.attached_image = paths[0]
            return {"selected": True, "path": self.attached_image,
                    "name": os.path.basename(self.attached_image)}
        except Exception as e:
            return {"error": str(e)}

    def clear_attached_image(self):
        self.attached_image = None
        return {"ok": True}

    def get_attached_image_preview(self):
        """Return base64 data URI for the currently attached image (for inline preview)."""
        if not self.attached_image or not os.path.isfile(self.attached_image):
            return {"preview": None}
        try:
            import base64
            ext = os.path.splitext(self.attached_image)[1].lower()
            mime_map = {".png": "image/png", ".jpg": "image/jpeg", ".jpeg": "image/jpeg",
                        ".webp": "image/webp", ".gif": "image/gif", ".bmp": "image/bmp"}
            mime = mime_map.get(ext, "image/png")
            with open(self.attached_image, "rb") as f:
                data = base64.b64encode(f.read()).decode("ascii")
            return {"preview": f"data:{mime};base64,{data}", "name": os.path.basename(self.attached_image)}
        except Exception:
            return {"preview": None}

    def open_document_dialog(self):
        """Open native document picker and extract text for context use.
        Supports selecting multiple files at once.
        Replaces any previously uploaded files (use drag-drop to add more)."""
        try:
            import webview
            if not webview.windows:
                return {"error": "Window not ready"}
            paths = webview.windows[0].create_file_dialog(
                webview.FileDialog.OPEN,
                allow_multiple=True,
                file_types=(
                    "All supported (*.txt;*.md;*.pdf;*.csv;*.docx;*.xlsx;*.xls;*.json;*.xml;*.html;*.htm;*.yaml;*.yml;*.log;*.tsv;*.sql;*.py;*.js;*.css;*.ini;*.cfg;*.toml;*.rtf)",
                ),
            )
            if not paths:
                return {"selected": False}

            # Clear previous files — dialog upload replaces, drag-drop appends
            self.uploaded_files = []
            self.uploaded_content = None
            self.uploaded_file_name = None
            self.uploaded_file_path = None
            self.uploaded_pages = None

            results = []
            has_async = False
            for path in paths:
                r = self._process_single_upload(path)
                results.append(r)
                if r.get("processing"):
                    has_async = True

            # Return summary for frontend
            names = [r["name"] for r in results if r.get("name")]
            return {
                "selected": True,
                "files": results,
                "name": names[0] if len(names) == 1 else f"{len(names)} files",
                "chars": sum(r.get("chars", 0) for r in results),
                "processing": has_async,
                "multi": len(results) > 1,
            }
        except Exception as e:
            return {"error": str(e)}

    def _process_single_upload(self, path: str) -> dict:
        """Process a single file upload: extract content, add to uploaded_files list."""
        name = os.path.basename(path)

        # Avoid duplicates
        for existing in self.uploaded_files:
            if existing["name"] == name:
                return {"name": name, "chars": len(existing.get("content") or ""), "duplicate": True}

        file_entry = {"name": name, "path": path, "content": None, "pages": None}

        if self._is_tabular_file(name):
            preview = self._build_tabular_preview_fast(path)
            file_entry["content"] = preview
            self.uploaded_files.append(file_entry)
            self._sync_primary_file()
            return {"name": name, "chars": len(preview or "")}

        ext = os.path.splitext(path)[1].lower()

        if ext == ".pdf":
            self.uploaded_files.append(file_entry)
            self._sync_primary_file()
            self._status(f"Processing PDF: {name}...")
            threading.Thread(
                target=self._load_pdf_background_multi, args=(path, name), daemon=True
            ).start()
            return {"name": name, "chars": 0, "processing": True}

        extracted = self._extract_text_from_file(path)
        if not extracted:
            return {"name": name, "error": "Could not extract text"}

        file_entry["content"] = extracted
        self.uploaded_files.append(file_entry)
        self._sync_primary_file()
        return {"name": name, "chars": len(extracted)}

    def _sync_primary_file(self):
        """Keep legacy single-file state in sync with uploaded_files list.
        Primary = last uploaded file (most recent is most relevant)."""
        if self.uploaded_files:
            primary = self.uploaded_files[-1]
            self.uploaded_file_name = primary["name"]
            self.uploaded_file_path = primary["path"]
            self.uploaded_content = primary.get("content")
            self.uploaded_pages = primary.get("pages")
        else:
            self.uploaded_file_name = None
            self.uploaded_file_path = None
            self.uploaded_content = None
            self.uploaded_pages = None

    def _load_pdf_background_multi(self, path: str, name: str):
        """Extract PDF text in background for multi-file upload."""
        try:
            extracted, page_list = self._extract_pdf_full(path=path)
            if not extracted:
                self._emit("file_upload_done", {
                    "name": name,
                    "error": "Could not extract readable text from PDF",
                })
                return
            # Update the entry in uploaded_files
            for f in self.uploaded_files:
                if f["name"] == name and f["path"] == path:
                    f["content"] = extracted
                    f["pages"] = page_list
                    break
            self._sync_primary_file()
            self._emit("file_upload_done", {
                "name": name,
                "chars": len(extracted),
                "pages": len(page_list),
            })
            self._status("Ready")
        except Exception as e:
            self._emit("file_upload_done", {"name": name, "error": str(e)})
            self._status("Ready")

    def upload_file_from_data(self, file_name: str, data_base64: str):
        """Receive a file from drag-drop (base64-encoded), save to temp, and process."""
        if not file_name or not data_base64:
            return {"error": "Missing file name or data"}

        # Avoid duplicates
        for existing in self.uploaded_files:
            if existing["name"] == file_name:
                return {"name": file_name, "duplicate": True, "chars": len(existing.get("content") or "")}

        temp_dir = os.path.join(app_data_path(), "agent_temp")
        os.makedirs(temp_dir, exist_ok=True)
        safe_name = re.sub(r'[<>:"/\\|?*]', '_', file_name)
        dest = os.path.join(temp_dir, safe_name)

        try:
            import base64
            raw = base64.b64decode(data_base64)
            with open(dest, "wb") as fh:
                fh.write(raw)
        except Exception as e:
            return {"error": f"Failed to save file: {e}"}

        return self._process_single_upload(dest)

    def get_uploaded_files(self):
        """Return list of currently uploaded file names for the frontend."""
        return {"files": [f["name"] for f in self.uploaded_files]}

    def _load_pdf_background(self, path: str):
        """Extract PDF text (with OCR) in background, emit event when done."""
        name = os.path.basename(path)
        try:
            extracted, page_list = self._extract_pdf_full(path=path)
            if not extracted:
                self._emit("file_upload_done", {
                    "error": "Could not extract readable text from PDF"
                })
                return
            self.uploaded_content = extracted
            self.uploaded_pages = page_list
            # Update uploaded_files entry if present
            for f in self.uploaded_files:
                if f["path"] == path:
                    f["content"] = extracted
                    f["pages"] = page_list
                    break
            char_count = len(extracted)
            page_count = len(page_list)
            self._emit("file_upload_done", {
                "name": self.uploaded_file_name or name,
                "chars": char_count,
                "pages": page_count,
            })
            self._status("Ready")
        except Exception as e:
            self._emit("file_upload_done", {"error": str(e)})
            self._status("Ready")

    def clear_uploaded_document(self, file_name: str | None = None):
        """Clear uploaded files. If file_name given, remove only that file."""
        if file_name:
            self.uploaded_files = [f for f in self.uploaded_files if f["name"] != file_name]
            if self.uploaded_file_name == file_name:
                if self.uploaded_files:
                    primary = self.uploaded_files[-1]
                    self.uploaded_file_name = primary["name"]
                    self.uploaded_file_path = primary["path"]
                    self.uploaded_content = primary.get("content")
                    self.uploaded_pages = primary.get("pages")
                else:
                    self.uploaded_file_name = None
                    self.uploaded_file_path = None
                    self.uploaded_content = None
                    self.uploaded_pages = None
            return {"ok": True, "files": [f["name"] for f in self.uploaded_files]}
        self.uploaded_content = None
        self.uploaded_file_name = None
        self.uploaded_file_path = None
        self.uploaded_pages = None
        self.uploaded_files = []
        return {"ok": True, "files": []}

    def _cleanup_agent_temp(self):
        """Delete all files in agent_temp/ to free disk space after processing."""
        temp_dir = os.path.join(app_data_path(), "agent_temp")
        if not os.path.isdir(temp_dir):
            return
        removed = 0
        for fname in os.listdir(temp_dir):
            fpath = os.path.join(temp_dir, fname)
            try:
                if os.path.isfile(fpath):
                    os.remove(fpath)
                    removed += 1
            except Exception:
                pass
        if removed:
            print(f"[CLEANUP] Removed {removed} temp file(s) from agent_temp/")

    def save_agent_temp_file(self, file_name: str, content: str):
        """Save agent file content to a temp directory and return the path."""
        try:
            import tempfile
            temp_dir = os.path.join(app_data_path(), "agent_temp")
            os.makedirs(temp_dir, exist_ok=True)
            # Sanitize filename
            safe_name = re.sub(r'[^\w.\-]', '_', file_name)
            path = os.path.join(temp_dir, safe_name)
            with open(path, 'w', encoding='utf-8', errors='ignore') as f:
                f.write(content)
            return {"ok": True, "path": path}
        except Exception as e:
            return {"error": str(e)}

    def _build_tabular_preview_fast(self, path: str) -> str:
        """Lightweight preview — reads only a few rows for instant UI feedback."""
        try:
            import pandas as pd
            ext = os.path.splitext(path)[1].lower()
            name = os.path.basename(path)
            if ext == ".csv":
                # Read just first 5 rows for preview, count total rows separately
                preview_df = pd.read_csv(path, nrows=5, low_memory=False)
                # Count total rows without loading all data
                with open(path, 'r', encoding='utf-8', errors='ignore') as f:
                    total_rows = sum(1 for _ in f) - 1  # subtract header
            elif ext in (".xlsx", ".xls"):
                full_df = pd.read_excel(path, sheet_name=0)
                total_rows = len(full_df)
                preview_df = full_df.head(5)
            else:
                return f"Tabular file attached: {name}"

            cols = list(preview_df.columns)
            numeric_cols = [c for c in cols if preview_df[c].dtype.kind in ('i', 'f')]
            lines = [
                f"Tabular file: {name}",
                f"Rows: {total_rows:,}  |  Columns: {len(cols)}",
                f"Column names: {', '.join(cols[:30])}",
            ]
            if numeric_cols:
                lines.append(f"Numeric columns: {', '.join(numeric_cols[:20])}")
            lines.append(f"\nAsk me questions about this data — I'll compute accurate answers using SQL.")
            return "\n".join(lines)
        except Exception as e:
            return f"Tabular file attached: {os.path.basename(path)} (preview error: {e})"

    def _build_tabular_preview_for_chat(self, path: str) -> str:
        try:
            analysis = self._analyze_tabular_files(
                [{"name": os.path.basename(path), "path": path}],
                "Provide structural preview",
                mode="chat",
            )
            if analysis.get("ok"):
                return (analysis.get("summary_text", "") or "Tabular file attached.")[:4000]
        except Exception:
            pass
        return f"Tabular file attached: {os.path.basename(path)}"

    def _extract_text_from_file(self, path: str) -> str:
        ext = os.path.splitext(path)[1].lower()
        try:
            # Plain-text / code / config files — read directly
            if ext in (".txt", ".md", ".log", ".json", ".xml", ".html", ".htm",
                        ".yaml", ".yml", ".sql", ".py", ".js", ".css",
                        ".ini", ".cfg", ".toml", ".tsv", ".rtf"):
                with open(path, "r", encoding="utf-8", errors="ignore") as f:
                    return f.read()
            if ext == ".pdf":
                text, _pages = self._extract_pdf_full(path=path)
                return text
            if ext == ".csv":
                try:
                    import pandas as pd
                    df = pd.read_csv(path)
                    return df.to_string(index=False)
                except Exception:
                    with open(path, "r", encoding="utf-8", errors="ignore") as f:
                        return f.read()
            if ext in (".xlsx", ".xls"):
                import pandas as pd
                sheets = pd.read_excel(path, sheet_name=None)
                parts = []
                for name, frame in sheets.items():
                    parts.append(f"Sheet: {name}")
                    parts.append(frame.to_string(index=False))
                return "\n\n".join(parts)
            if ext == ".docx":
                from docx import Document
                doc = Document(path)
                return "\n".join(p.text for p in doc.paragraphs if p.text.strip())
            # Fallback: try reading as text
            with open(path, "r", encoding="utf-8", errors="ignore") as f:
                return f.read()
        except Exception:
            return ""
        return ""

    def _extract_pdf_full(self, path: str = "", raw_bytes: bytes = b"") -> tuple:
        """Extract ALL text from a PDF — text pages + OCR for scanned pages.
        Accepts a file *path* OR *raw_bytes* (at least one required).
        Returns (full_text, page_list) so callers can use pages without
        mutating self.uploaded_pages.  The chat-mode caller still writes
        self.uploaded_pages for backward compat.
        Uses parallel OCR workers for speed.
        """
        cache_key = ""
        try:
            if path and os.path.isfile(path):
                cache_key = self._pdf_hash_for_path(path)
            elif raw_bytes:
                cache_key = self._pdf_hash_for_bytes(raw_bytes)
        except Exception:
            cache_key = ""

        if cache_key:
            cached = self._load_pdf_cached_extract(cache_key)
            if cached:
                self._status("PDF ready (cached)")
                return cached

        try:
            import fitz  # PyMuPDF
        except ImportError:
            return "", []

        from concurrent.futures import ThreadPoolExecutor, as_completed

        pages: list[str] = []
        ocr_tasks: list[tuple[int, bytes]] = []  # (page_idx, png_bytes)

        try:
            if path:
                doc = fitz.open(path)
            elif raw_bytes:
                doc = fitz.open(stream=raw_bytes, filetype="pdf")
            else:
                return "", []
        except Exception:
            return "", []

        total_pages = len(doc)
        self._status(f"Reading PDF ({total_pages} pages)...")

        # Pass 1: extract text pages, collect scanned pages for OCR
        for i, page in enumerate(doc):
            page_text = (page.get_text("text") or "").strip()
            if len(page_text) > 30:
                # Readable text page
                pages.append(f"[Page {i+1}]\n{page_text}")
            else:
                # Scanned / image-only page — queue for OCR
                try:
                    pix = page.get_pixmap(dpi=200, alpha=False)
                    img_bytes = pix.tobytes("png")
                    ocr_tasks.append((i, img_bytes))
                    pages.append(None)  # placeholder
                except Exception:
                    pages.append(None)
            if (i + 1) % 25 == 0:
                self._status(f"Reading PDF... {i+1}/{total_pages} pages")

        doc.close()

        # Pass 2: parallel OCR for scanned pages
        if ocr_tasks:
            self._status(f"OCR processing {len(ocr_tasks)} scanned pages...")

            def _ocr_one(item):
                idx, img_bytes = item
                text = self._ocr_image_bytes(img_bytes)
                return idx, text

            workers = min(4, len(ocr_tasks))  # cap at 4 threads
            done_count = 0
            with ThreadPoolExecutor(max_workers=workers) as pool:
                futures = {pool.submit(_ocr_one, t): t[0] for t in ocr_tasks}
                for fut in as_completed(futures):
                    done_count += 1
                    idx, text = fut.result()
                    if text:
                        pages[idx] = f"[Page {idx+1} — OCR]\n{text}"
                    if done_count % 10 == 0:
                        self._status(f"OCR... {done_count}/{len(ocr_tasks)} pages done")

        # Build final page list (drop empty pages)
        final_pages = [p for p in pages if p]

        full_text = "\n\n".join(final_pages)
        if cache_key and full_text:
            self._save_pdf_cached_extract(cache_key, full_text, final_pages)
        text_pages = sum(1 for p in pages if p and "[Page" in p and "OCR" not in p)
        ocr_pages = sum(1 for p in pages if p and "OCR" in p)
        empty_pages = total_pages - text_pages - ocr_pages
        self._status(f"PDF ready: {text_pages} text + {ocr_pages} OCR + {empty_pages} empty pages")
        return full_text, final_pages

    def _generate_with_image(self, text: str, image_path: str,
                             request_options: dict | None = None):
        self.generation_in_progress = True
        self.stop_generation_flag = False
        self._emit("generation_start", None)
        try:
            # Lock during inference to prevent model unload
            with self.model_lock:
                if self.model is None:
                    self._emit("generation_error", {"error": "No model loaded"})
                    return

                if self.stop_generation_flag:
                    self._emit("generation_stopped", {"text": ""})
                    self._status("Generation stopped")
                    return

                result = self._try_multimodal_generate(text, image_path)
                if result:
                    final_text = result.strip()
                else:
                    vision_status = self.get_vision_status()
                    ocr_text = ""
                    try:
                        with open(image_path, "rb") as f:
                            ocr_text = self._ocr_image_bytes(f.read())
                    except Exception:
                        ocr_text = ""

                    has_ocr = bool((ocr_text or "").strip())
                    note = (
                        f"📷 Image attached: {os.path.basename(image_path)}\n\n"
                        "⚠ Vision chat is unavailable with current assets. "
                        f"{vision_status.get('reason', 'Missing multimodal projector file')}\n"
                    )
                    if has_ocr:
                        note += "Using local Tesseract OCR as fallback.\n\n"
                    else:
                        note += "No OCR text could be extracted from this image.\n\n"

                    web_context = ""
                    web_sources = []
                    if self.web_search_enabled:
                        try:
                            base_q = text.strip()
                            if has_ocr:
                                ocr_hint = " ".join((ocr_text or "").split())[:220]
                                if ocr_hint:
                                    base_q = f"{base_q} {ocr_hint}".strip()
                            refined = self._refine_search_query(base_q)
                            web_context, web_sources = self._search_web(refined)
                            if web_sources:
                                self._emit("web_sources", {"sources": web_sources})
                        except Exception:
                            web_context, web_sources = "", []

                    format_instruction = self._get_response_format_instruction(request_options)
                    system = (
                        "You are an assistant answering a question about an attached image. "
                        "If OCR text is provided, rely on it as the primary evidence. "
                        "If web results are provided, use them to cross-check facts and cite with [1], [2] where used. "
                        "If OCR is weak or missing, clearly say what is uncertain."
                    )

                    extra_parts = []
                    if has_ocr:
                        extra_parts.append("OCR text from image:\n" + (ocr_text or "")[:5000])
                    if web_context:
                        extra_parts.append("Web matches:\n" + web_context[:3000])
                    if format_instruction:
                        extra_parts.append(f"Formatting: {format_instruction}")
                    extra_context = "\n\n".join(extra_parts)

                    prompt = self._build_chat_prompt(
                        system=system,
                        messages=[],
                        user_text=text,
                        extra_context=extra_context,
                    )
                    stop_tokens = self._get_stop_tokens()
                    response_obj = self.model(
                        prompt,
                        max_tokens=min(420, max(128, self.actual_n_ctx // 3)),
                        temperature=0.15,
                        top_p=0.9,
                        stream=False,
                        stop=stop_tokens,
                    )
                    fallback = (response_obj["choices"][0].get("text", "") or "").strip()
                    final_text = note + (fallback if fallback else "I could not produce a reliable answer.")

                if self.stop_generation_flag:
                    self._emit("generation_stopped", {"text": ""})
                    self._status("Generation stopped")
                    return

            self._emit("generation_done", {"text": final_text})
            self.message_history.append({"role": "assistant", "content": final_text})
            self.chats[self.current_chat_id] = self.message_history
            self._save_current_chat()
            self._auto_export_response(
                self.current_chat_id,
                len(self.message_history) - 1,
                (request_options or {}).get("export_format"),
            )
        except Exception as e:
            self._emit("generation_error", {"error": str(e)})
        finally:
            self.generation_in_progress = False

    @staticmethod
    def _make_clip_handler(clip_path: str):
        """Create the appropriate multimodal chat handler."""
        try:
            from llama_cpp.llama_chat_format import MiniCPMv26ChatHandler
            return MiniCPMv26ChatHandler(clip_model_path=clip_path)
        except ImportError:
            pass
        try:
            from llama_cpp.llama_chat_format import Llava16ChatHandler
            return Llava16ChatHandler(clip_model_path=clip_path)
        except ImportError:
            pass
        from llama_cpp.llama_chat_format import Llava15ChatHandler
        return Llava15ChatHandler(clip_model_path=clip_path)

    def _try_multimodal_generate(self, prompt_text: str, image_path: str):
        try:
            # Vision must be loaded at model-init time via chat_handler
            if not hasattr(self.model, 'chat_handler') and \
               not getattr(self.model, '_chat_handler', None):
                # Check if model was loaded with a handler
                handler = getattr(self.model, 'chat_handler', None)
                if handler is None:
                    return None

            import base64
            with open(image_path, "rb") as f:
                img_b64 = base64.b64encode(f.read()).decode("utf-8")

            ext = os.path.splitext(image_path)[1].lower().lstrip(".")
            mime = {
                "jpg": "image/jpeg",
                "jpeg": "image/jpeg",
                "png": "image/png",
                "gif": "image/gif",
                "webp": "image/webp",
                "bmp": "image/bmp",
            }.get(ext, "image/png")

            result = self.model.create_chat_completion(
                messages=[{
                    "role": "user",
                    "content": [
                        {
                            "type": "image_url",
                            "image_url": {"url": f"data:{mime};base64,{img_b64}"},
                        },
                        {"type": "text", "text": prompt_text},
                    ],
                }],
                max_tokens=512,
            )
            return result["choices"][0]["message"]["content"]
        except Exception as e:
            print(f"[Vision] generation failed: {e}")
            return None

    # ── RAG management ─────────────────────────────────────────

    def get_rag_databases(self):
        """Return list of RAG databases."""
        if not self.rag_manager:
            return []
        names = self.rag_manager.list_databases()
        result = []
        for name in names:
            info = self.rag_manager.get_database_info(name) or {}
            md_path = self.rag_manager.get_markdown_path(name)
            result.append({"name": name,
                           "chunks": info.get("num_chunks", 0),
                           "markdown_path": md_path,
                           "selected": name == self.current_rag_database})
        return result

    def select_rag(self, db_name: str):
        """Toggle RAG database selection for current chat."""
        if self.current_rag_database == db_name:
            self.current_rag_database = None
        else:
            self.current_rag_database = db_name
        if self.current_chat_id:
            if self.current_rag_database:
                self.chat_rag_settings[self.current_chat_id] = self.current_rag_database
            else:
                self.chat_rag_settings.pop(self.current_chat_id, None)
        if self.current_rag_database:
            self._status(f"RAG selected: {self.current_rag_database}")
        else:
            self._status("RAG deselected")
        return {"selected": self.current_rag_database}

    def _rag_progress_cb(self, rag_name, file_name, file_index, total_files):
        """Emit progress event for RAG creation."""
        pct = int((file_index / total_files) * 100) if total_files else 0
        self._emit("rag_progress", {
            "name": rag_name,
            "file": file_name,
            "file_index": file_index,
            "total_files": total_files,
            "percent": pct,
        })

    def create_rag_from_folder(self, folder_path: str, name: str,
                                chunk_size: int = 512, overlap: int = 100):
        """Create RAG database from folder path (blocking)."""
        if not self.rag_manager:
            return {"error": "RAG not initialized"}
        try:
            self._status(f"Creating RAG: {name}...")
            cb = lambda fname, idx, total: self._rag_progress_cb(name, fname, idx, total)
            self.rag_manager.create_from_folder(folder_path, name,
                                                 chunk_size=chunk_size,
                                                 chunk_overlap=overlap,
                                                 progress_callback=cb)
            self._emit("rag_progress", {"name": name, "percent": 100, "done": True})
            info = self.rag_manager.get_database_info(name)
            md_path = self.rag_manager.get_markdown_path(name)
            self._status(f"RAG '{name}' created ({info.get('num_chunks', 0)} chunks)")
            return {"ok": True, "name": name,
                    "chunks": info.get("num_chunks", 0),
                    "markdown_path": md_path}
        except Exception as e:
            self._emit("rag_progress", {"name": name, "percent": 0, "done": True, "error": True})
            self._status(f"RAG create error: {e}")
            return {"error": str(e)}

    def delete_rag(self, db_name: str):
        """Delete RAG database."""
        if not self.rag_manager:
            return {"error": "RAG not initialized"}
        try:
            self.rag_manager.delete_database(db_name)
            if self.current_rag_database == db_name:
                self.current_rag_database = None
            self._status(f"Deleted RAG: {db_name}")
            return {"ok": True}
        except Exception as e:
            self._status(f"Delete RAG error: {e}")
            return {"error": str(e)}

    def create_rag_from_url(self, url: str, name: str,
                            chunk_size: int = 512, overlap: int = 100):
        """Create RAG database from URL content."""
        if not self.rag_manager:
            return {"error": "RAG not initialized"}
        try:
            self._emit("rag_progress", {"name": name, "file": "Fetching URL...", "percent": 10})
            self._status(f"Importing RAG from URL: {name}...")
            ok = self.rag_manager.create_from_url(
                url, name, chunk_size=chunk_size, chunk_overlap=overlap)
            if not ok:
                self._emit("rag_progress", {"name": name, "percent": 0, "done": True, "error": True})
                return {"error": "Failed to build RAG from URL"}
            self._emit("rag_progress", {"name": name, "percent": 100, "done": True})
            info = self.rag_manager.get_database_info(name) or {}
            md_path = self.rag_manager.get_markdown_path(name)
            self._status(f"RAG '{name}' imported ({info.get('num_chunks', 0)} chunks)")
            return {"ok": True, "name": name, "chunks": info.get("num_chunks", 0), "markdown_path": md_path}
        except Exception as e:
            self._emit("rag_progress", {"name": name, "percent": 0, "done": True, "error": True})
            self._status(f"RAG URL import error: {e}")
            return {"error": str(e)}

    def reindex_rag(self, db_name: str):
        """Reindex an existing RAG database."""
        if not self.rag_manager:
            return {"error": "RAG not initialized"}
        try:
            self._status(f"Reindexing RAG: {db_name}...")
            cb = lambda fname, idx, total: self._rag_progress_cb(db_name, fname, idx, total)
            ok = self.rag_manager.reindex_database(db_name, progress_callback=cb)
            self._emit("rag_progress", {"name": db_name, "percent": 100, "done": True})
            if ok:
                self._status(f"Reindexed RAG: {db_name}")
            return {"ok": bool(ok)}
        except Exception as e:
            self._status(f"Reindex error: {e}")
            return {"error": str(e)}

    # ── Settings ───────────────────────────────────────────────

    def get_app_info(self):
        """Return system + app info for settings display."""
        return {
            "system_ram": self.system_ram,
            "gpu": self.gpu_info,
            "config": self.config,
            "theme": self.current_theme,
            "web_search": self.web_search_enabled,
            "model_loaded": self.model is not None,
            "model_name": Path(self.model_path).name if self.model_path else None,
            "n_ctx": self.actual_n_ctx,
            "activation": self._get_activation_status(),
        }

    def set_theme(self, theme: str):
        normalized = str(theme or "Dark").strip()
        if not normalized:
            normalized = "Dark"
        normalized = normalized[0].upper() + normalized[1:].lower()

        self.current_theme = normalized
        self.app_settings["theme"] = normalized
        self.chat_db.set_kv("app_settings", self.app_settings)
        return {"ok": True}

    def list_plugins(self):
        """List discovered plugins."""
        try:
            return self.plugin_manager.list_plugins()
        except Exception as e:
            return [{"name": "error", "info": str(e), "loaded": False}]

    def get_plugin_commands(self):
        """Return list of registered plugin commands for frontend autocomplete."""
        cmds = []
        for cmd, entry in self.plugin_commands.items():
            cmds.append({
                "command": cmd,
                "description": entry.get("description", ""),
                "plugin": entry.get("plugin", ""),
            })
        return cmds

    # ── Tool permission gate (called from JS) ─────────────────

    def approve_tool_execution(self):
        """User approved a dangerous tool call."""
        self._tool_permission_granted = True
        self._tool_permission_event.set()

    def deny_tool_execution(self):
        """User denied a dangerous tool call."""
        self._tool_permission_granted = False
        self._tool_permission_event.set()

    def reload_plugins(self):
        """Reload all plugins."""
        try:
            for name in list(self.plugin_manager.loaded.keys()):
                self.plugin_manager.unload(name, self)
            self.plugin_manager.loaded.clear()
            self.plugin_manager.load_all(self)
            return {"ok": True, "plugins": self.plugin_manager.list_plugins()}
        except Exception as e:
            return {"error": str(e)}

    # ── Create Plugin with AI ──────────────────────────────────

    _PLUGIN_SYSTEM_PROMPT = """You are an expert Python developer writing plugins for SIMPLE_AI, \
a local AI chat desktop application built with llama-cpp-python and pywebview.

## Plugin structure
A plugin is a single .py file saved in the plugins/ directory. Required entry point:

```python
PLUGIN_INFO = "One-line description of plugin"

def register(app):
    \"\"\"Called when plugin loads. app is the Bridge instance.\"\"\"
    pass

def unregister(app):  # optional
    \"\"\"Called when plugin is unloaded.\"\"\"
    pass
```

## What `app` exposes
- `app.message_history` — list[dict] like [{"role":"user","content":"..."}, ...]
- `app.current_chat_id` — str ID of the active chat
- `app.model` — loaded llama-cpp-python Llama model (None when no model loaded)
- `app._emit(event_name, data)` — push event to the frontend JS; data must be JSON-serialisable
- `app._status(text)` — update the status bar
- `app.app_settings` — dict of user settings (keys: theme, font_size, brave_api_key, …)
- `app.plugin_manager` — PluginManager instance
- `app.register_plugin_command('/name', handler, plugin_name='...', description='...')`
- `app.unregister_plugin_command('/name')`
- `app.execute_python_snippet(code, timeout=8)` — restricted Python runner for plugin commands
- `app._DEFAULT_SYSTEM` — default system prompt string (can be overridden)
- `app.send_message(text)` — send a message programmatically (wraps the full generation pipeline)

## Common patterns

### Add a custom chat command (recommended)
```python
PLUGIN_INFO = "Adds /roll command to roll dice"
import random, re
ROLL_RE = re.compile(r'^/roll\\s*(?:(\\d+)d)?(\\d+)$', re.I)

def handle_roll(app, text):
    m = ROLL_RE.match(text.strip())
    if not m:
        return {"handled": False}
    n = int(m.group(1) or 1)
    sides = int(m.group(2))
    rolls = [random.randint(1, sides) for _ in range(n)]
    return {"content": f"Rolled {n}d{sides}: {rolls} = {sum(rolls)}"}

def register(app):
    app.register_plugin_command('/roll', handle_roll,
                                plugin_name='roll_plugin',
                                description='Roll dice using /roll NdM')

def unregister(app):
    app.unregister_plugin_command('/roll')
```

### Background task / status ticker
```python
PLUGIN_INFO = "Shows a live clock in the status bar"
import time, threading
_running = False
def register(app):
    global _running; _running = True
    def tick():
        while _running:
            app._status(time.strftime("🕐 %H:%M:%S"))
            time.sleep(1)
    threading.Thread(target=tick, daemon=True).start()
def unregister(app):
    global _running; _running = False
```

### Startup configuration
```python
PLUGIN_INFO = "Forces concise replies"
def register(app):
    app._DEFAULT_SYSTEM = "You are a helpful assistant. Always be brief and use bullet points."
```

## Rules
- Always define PLUGIN_INFO as a plain string at module level
- register(app) is required; unregister(app) is optional
- Do NOT use blocking I/O or input() — the app is event-driven
- Put all imports at the top
- Use try/except around risky code
- Output ONLY raw Python code — no markdown fences, no prose before or after the code

Write a plugin that fulfils the following requirement:
"""

    def generate_plugin_with_ai(self, name: str, description: str):
        """Stream AI-generated plugin code. Emits plugin_gen_token / plugin_gen_done / plugin_gen_error."""
        if self.model is None:
            self._emit("plugin_gen_error", {"error": "No model loaded. Please load a model first."})
            return {"error": "No model loaded"}
        if self.generation_in_progress:
            self._emit("plugin_gen_error", {"error": "Generation already in progress."})
            return {"error": "Busy"}

        def _run():
            self.generation_in_progress = True
            self.stop_generation_flag = False
            try:
                system = self._PLUGIN_SYSTEM_PROMPT
                user_text = f"Plugin name: {name.strip()}\n\nRequirement: {description.strip()}"

                prompt = self._build_chat_prompt(
                    system=system,
                    messages=[],
                    user_text=user_text,
                    extra_context="",
                )

                try:
                    prompt_tokens = len(self.model.tokenize(prompt.encode("utf-8")))
                except Exception:
                    prompt_tokens = len(prompt) // 4
                max_tokens = max(256, self.actual_n_ctx - prompt_tokens - 64)

                response = ""
                token_count = 0

                with self.model_lock:
                    if self.model is None:
                        raise RuntimeError("Model not loaded")
                    stream = self.model(
                        prompt,
                        max_tokens=max_tokens,
                        temperature=0.2,
                        top_p=0.95,
                        top_k=40,
                        repeat_penalty=1.1,
                        stream=True,
                        stop=self._get_stop_tokens(),
                    )
                    for chunk in stream:
                        if self.stop_generation_flag:
                            break
                        token = chunk["choices"][0].get("text", "")
                        response += token
                        token_count += 1
                        if token_count % 3 == 0:
                            display = _RE_THINK.sub("", response).strip()
                            self._emit("plugin_gen_token", {"text": display})

                code = _RE_THINK.sub("", response).strip()
                code = _RE_THINK_INCOMPLETE.sub("", code).strip()
                # Strip accidental markdown fences the model might add
                import re as _re
                code = _re.sub(r'^```(?:python)?\s*', '', code, flags=_re.MULTILINE)
                code = _re.sub(r'^```\s*$', '', code, flags=_re.MULTILINE)
                code = code.strip()
                self._emit("plugin_gen_done", {"code": code})

            except Exception as e:
                self._emit("plugin_gen_error", {"error": str(e)})
            finally:
                self.generation_in_progress = False

        threading.Thread(target=_run, daemon=True).start()
        return {"ok": True}

    def test_plugin_code(self, name: str, code: str):
        """Syntax-check and dry-run the plugin code in an isolated subprocess."""
        import tempfile, sys
        import subprocess as _sub
        try:
            # First check syntax locally (fast, no subprocess needed)
            compile(code, f"{name}.py", "exec")
        except SyntaxError as e:
            return {"error": f"Syntax error: {e}"}

        test_script = f"""
import sys, types, traceback

code = {repr(code)}
name = {repr(name)}

class MockApp:
    message_history = []
    current_chat_id = "test-chat"
    model = None
    app_settings = {{}}
    _DEFAULT_SYSTEM = "You are a helpful assistant."
    _events = []
    plugin_commands = {{}}
    def _emit(self, event, data=None):
        print(f"[emit] {{event}}: {{data}}")
    def _status(self, text):
        print(f"[status] {{text}}")
    def send_message(self, text, *a, **kw):
        print(f"[send_message] {{text}}")
    def register_plugin_command(self, command, handler, plugin_name='', description=''):
        self.plugin_commands[str(command).lower()] = handler
        print(f"[register_plugin_command] {{command}}")
        return {{"ok": True}}
    def unregister_plugin_command(self, command):
        self.plugin_commands.pop(str(command).lower(), None)
        print(f"[unregister_plugin_command] {{command}}")
        return {{"ok": True}}

app = MockApp()
try:
    import importlib.util, tempfile, os
    with tempfile.NamedTemporaryFile(suffix='.py', delete=False, mode='w', encoding='utf-8') as f:
        f.write(code)
        tmp = f.name
    spec = importlib.util.spec_from_file_location(name, tmp)
    mod = importlib.util.module_from_spec(spec)
    spec.loader.exec_module(mod)
    if hasattr(mod, 'register'):
        mod.register(app)
        print("[ok] register(app) called successfully")
        if hasattr(mod, 'unregister'):
            mod.unregister(app)
            print("[ok] unregister(app) called successfully")
    else:
        print("[warn] No register() function found")
    os.unlink(tmp)
except Exception:
    traceback.print_exc()
    sys.exit(1)
"""
        try:
            result = _sub.run(
                [sys.executable, "-c", test_script],
                capture_output=True, text=True, timeout=15
            )
            output = result.stdout.strip()
            errors = result.stderr.strip()
            if result.returncode != 0:
                return {"error": errors or "Test failed with no output"}
            return {"ok": True, "output": output + ("\n" + errors if errors else "")}
        except _sub.TimeoutExpired:
            return {"error": "Test timed out (15 s)"}
        except Exception as e:
            return {"error": str(e)}

    def save_generated_plugin(self, name: str, code: str):
        """Save generated plugin code to the plugins directory and reload."""
        safe_name = self._sanitize_plugin_name(name)
        path = os.path.join(self.plugin_manager.plugins_dir, safe_name + ".py")
        try:
            with open(path, "w", encoding="utf-8") as f:
                f.write(code)
            # Reload just that plugin
            self.plugin_manager.loaded.pop(safe_name, None)
            self.plugin_manager.load_all(self)
            return {"ok": True, "path": path, "name": safe_name}
        except Exception as e:
            return {"error": str(e)}

    def _sanitize_plugin_name(self, name: str) -> str:
        import re as _re
        safe_name = _re.sub(r'[^a-zA-Z0-9_]', '_', str(name or "").strip())
        return safe_name or "custom_plugin"

    def get_plugin_code(self, name: str):
        """Return plugin source code by plugin name (without .py)."""
        safe_name = self._sanitize_plugin_name(name)
        path = os.path.join(self.plugin_manager.plugins_dir, safe_name + ".py")
        if not os.path.isfile(path):
            return {"error": f"Plugin not found: {safe_name}"}
        try:
            with open(path, "r", encoding="utf-8") as f:
                code = f.read()
            return {"ok": True, "name": safe_name, "code": code}
        except Exception as e:
            return {"error": str(e)}

    def update_plugin_code(self, name: str, code: str):
        """Update an existing plugin file and reload plugins."""
        safe_name = self._sanitize_plugin_name(name)
        path = os.path.join(self.plugin_manager.plugins_dir, safe_name + ".py")
        if not os.path.isfile(path):
            return {"error": f"Plugin not found: {safe_name}"}
        try:
            with open(path, "w", encoding="utf-8") as f:
                f.write(str(code or ""))
            self.reload_plugins()
            return {"ok": True, "name": safe_name, "path": path}
        except Exception as e:
            return {"error": str(e)}

    def delete_plugin(self, name: str):
        """Delete plugin file and unload/reload plugin state."""
        safe_name = self._sanitize_plugin_name(name)
        path = os.path.join(self.plugin_manager.plugins_dir, safe_name + ".py")
        if not os.path.isfile(path):
            return {"error": f"Plugin not found: {safe_name}"}
        try:
            # Best-effort unload to run unregister hooks.
            self.plugin_manager.unload(safe_name, self)
            if safe_name in self.plugin_manager.loaded:
                self.plugin_manager.loaded.pop(safe_name, None)
            if isinstance(self.plugin_commands, dict):
                self.plugin_commands = {
                    cmd: meta for cmd, meta in self.plugin_commands.items()
                    if meta.get("plugin") != safe_name
                }
            os.remove(path)
            self.reload_plugins()
            return {"ok": True, "name": safe_name}
        except Exception as e:
            return {"error": str(e)}

    def get_app_settings(self):
        return dict(self.app_settings)

    def get_ai_generation_debug(self):
        """Return current AI generation debug toggle state."""
        return {"enabled": self._is_ai_debug_enabled()}

    def set_ai_generation_debug(self, enabled: bool):
        """Enable or disable verbose AI generation diagnostics."""
        self.app_settings["debug_ai_generation"] = bool(enabled)
        self.chat_db.set_kv("app_settings", self.app_settings)
        return {"ok": True, "enabled": bool(enabled)}

    def save_app_settings(self, settings: dict):
        self.app_settings.update(settings)
        # Keep current_theme in sync when theme is saved via settings payload.
        if "theme" in settings:
            normalized = str(settings.get("theme") or "Dark").strip()
            if normalized:
                self.current_theme = normalized[0].upper() + normalized[1:].lower()
        self.chat_db.set_kv("app_settings", self.app_settings)
        if self.model is not None and "temperature" in settings:
            self.model_config["temperature"] = settings["temperature"]
        return {"ok": True}

    def get_system_prompt(self):
        if not self.current_chat_id:
            return ""
        return self.chat_system_prompts.get(self.current_chat_id, "")

    def set_system_prompt(self, prompt: str):
        if not self.current_chat_id:
            return {"error": "No active chat"}
        if prompt.strip():
            self.chat_system_prompts[self.current_chat_id] = prompt.strip()
            self.chat_db.set_meta(self.current_chat_id, "system_prompt", prompt.strip())
        else:
            self.chat_system_prompts.pop(self.current_chat_id, None)
            self.chat_db.delete_meta(self.current_chat_id, "system_prompt")
        return {"ok": True}

    # ── System monitor ─────────────────────────────────────────

    def start_monitor(self):
        """Start system monitor thread, emits 'system_stats' events."""
        def loop():
            while True:
                cpu = psutil.cpu_percent()
                ram = psutil.virtual_memory().percent
                gpu_text = ""
                gpu_type = self.gpu_info.get("type", "CPU")
                if gpu_type == "NVIDIA":
                    try:
                        r = subprocess.run(
                            ["nvidia-smi", "--query-gpu=utilization.gpu",
                             "--format=csv,noheader,nounits"],
                            capture_output=True, text=True, timeout=3)
                        if r.returncode == 0 and r.stdout.strip():
                            gpu_text = f" | GPU: {r.stdout.strip()}%"
                    except Exception:
                        pass
                elif gpu_type == "APPLE_METAL":
                    gpu_text = " | Metal"
                elif gpu_type == "AMD":
                    gpu_text = " | Vulkan"
                self._emit("system_stats",
                           {"text": f"CPU: {cpu}% | RAM: {ram}%{gpu_text}"})
                time.sleep(3)
        threading.Thread(target=loop, daemon=True).start()

    # ── Persistence helpers ────────────────────────────────────

    def rate_message(self, chat_id: str, message_index: int, rating):
        """Store a thumbs-up/down rating for a specific message."""
        key = f"rating_{message_index}"
        if rating:
            self.chat_db.set_meta(chat_id, key, str(rating))
        else:
            self.chat_db.set_meta(chat_id, key, "")
        return {"ok": True}

    # ── MCP Server Management ──────────────────────────────────

    def get_mcp_servers(self):
        """Return list of configured MCP servers and their status."""
        return self.mcp_manager.get_server_list()

    def add_mcp_server(self, config):
        """Add a new MCP server configuration."""
        return self.mcp_manager.add_server(config)

    def update_mcp_server(self, old_name, config):
        """Update an existing MCP server configuration."""
        return self.mcp_manager.update_server(old_name, config)

    def remove_mcp_server(self, name):
        """Remove an MCP server."""
        return self.mcp_manager.remove_server(name)

    def connect_mcp_server(self, name):
        """Connect to an MCP server and discover its tools."""
        return self.mcp_manager.connect_server(name)

    def disconnect_mcp_server(self, name):
        """Disconnect from an MCP server."""
        return self.mcp_manager.disconnect_server(name)

    def get_mcp_tools(self):
        """Return all tools from all connected MCP servers."""
        return self.mcp_manager.get_all_tools()

    def _save_current_chat(self):
        if self.current_chat_id and self.current_chat_id in self.chats:
            self.chat_db.save_chat(
                self.current_chat_id, self.chats[self.current_chat_id])
            self._chats_dirty = True

    def save_all(self):
        """Manual save-all trigger."""
        self.chat_db.save_all_chats(self.chats)
        return {"ok": True}

    # ── HuggingFace downloader ─────────────────────────────────

    def search_hf_models(self, query: str):
        """Search HuggingFace for GGUF models with accurate sizes and smart sorting."""
        try:
            from urllib.parse import quote
            import requests
            from concurrent.futures import ThreadPoolExecutor, as_completed

            url = (f"https://huggingface.co/api/models"
                   f"?search={quote(query)}&filter=gguf"
                   f"&sort=downloads&direction=-1&limit=20&full=true")
            resp = requests.get(url, timeout=15)
            resp.raise_for_status()
            models = resp.json()
        except Exception as e:
            return {"error": str(e)}

        avail_ram = psutil.virtual_memory().available / (1024**3)
        vram = self.gpu_info.get("vram", 0)
        gpu_type = self.gpu_info.get("type", "CPU")
        if gpu_type == "APPLE_METAL":
            max_gb = round(avail_ram * 0.75, 1)
        elif gpu_type in ("NVIDIA", "AMD") and vram >= 2:
            max_gb = round((vram + min(avail_ram * 0.25, 3.0)) * 0.9, 1)
        else:
            max_gb = round(min(avail_ram, self.system_ram) * 0.65, 1)

        # Fetch file sizes from tree API in parallel (siblings API no longer has sizes)
        def fetch_tree(model_id):
            try:
                tree_url = f"https://huggingface.co/api/models/{model_id}/tree/main?recursive=true"
                r = requests.get(tree_url, timeout=10)
                if r.status_code == 200:
                    return {f.get("path", ""): f.get("size") or (f.get("lfs") or {}).get("size")
                            for f in r.json() if f.get("path", "").lower().endswith(".gguf")}
            except Exception:
                pass
            return {}

        model_ids = [m.get("modelId") or m.get("id", "") for m in models]
        size_maps = {}
        with ThreadPoolExecutor(max_workers=8) as pool:
            futures = {pool.submit(fetch_tree, mid): mid for mid in model_ids}
            for fut in as_completed(futures):
                size_maps[futures[fut]] = fut.result()

        rows = []
        for m in models:
            model_id = m.get("modelId") or m.get("id", "")
            tree_sizes = size_maps.get(model_id, {})

            for sib in m.get("siblings", []):
                fname = sib.get("rfilename", "")
                if not fname.lower().endswith(".gguf"):
                    continue
                # Skip multi-part shard files (keep only single .gguf or first shard)
                fname_lower = fname.lower()
                if "-00002-of-" in fname_lower or "-00003-of-" in fname_lower:
                    continue

                size_bytes = tree_sizes.get(fname)
                if size_bytes is None:
                    lfs = sib.get("lfs")
                    if isinstance(lfs, dict):
                        size_bytes = lfs.get("size")
                if size_bytes is None:
                    size_bytes = sib.get("size")
                size_gb = round(size_bytes / (1024**3), 2) if size_bytes else None
                compatible = size_gb is not None and size_gb <= max_gb
                rows.append({
                    "model_id": model_id, "filename": fname,
                    "size_gb": size_gb,
                    "compatible": compatible,
                    "unknown": size_gb is None,
                })

        # Smart sort: compatible first (smallest first), then too-large, then unknown
        rows.sort(key=lambda r: (
            0 if r["compatible"] else (2 if r["unknown"] else 1),
            r["size_gb"] if r["size_gb"] is not None else 9999,
        ))
        return {"results": rows, "max_gb": max_gb, "system_ram": round(self.system_ram, 1)}

    def download_hf_model(self, model_id: str, filename: str):
        """Start background download of HF model."""
        threading.Thread(target=self._download_hf_thread,
                         args=(model_id, filename), daemon=True).start()
        return {"status": "downloading"}

    def _download_hf_thread(self, model_id, filename):
        import requests
        # Use only the base filename (strip subfolder paths like BF16/)
        base_name = filename.split("/")[-1] if "/" in filename else filename
        dest = Path(app_data_path("models")) / base_name
        try:
            url = f"https://huggingface.co/{model_id}/resolve/main/{filename}"
            self._emit("download_status",
                       {"text": f"Connecting to {model_id}...", "progress": 0})
            with requests.get(url, stream=True, timeout=60) as r:
                r.raise_for_status()
                total = int(r.headers.get("Content-Length", 0))
                downloaded = 0
                with open(dest, "wb") as f:
                    for chunk in r.iter_content(chunk_size=1024 * 1024):
                        if chunk:
                            f.write(chunk)
                            downloaded += len(chunk)
                            if total > 0:
                                pct = round(downloaded / total * 100)
                                mb = downloaded // (1024 * 1024)
                                total_mb = total // (1024 * 1024)
                                self._emit("download_status", {
                                    "text": f"{filename}: {mb}/{total_mb} MB",
                                    "progress": pct})
            self._emit("download_done", {"filename": filename})
        except Exception as e:
            self._emit("download_error", {"error": str(e)})
            if dest.exists():
                try:
                    dest.unlink()
                except Exception:
                    pass

    # ── Per-model settings ─────────────────────────────────────

    def get_per_model_settings(self):
        """Get per-model settings (temperature, n_ctx, n_threads, etc).
        
        These settings are user-configurable from the UI:
        - Model Settings button (⚙) in the top panel
        - Settings persist per model in chat_db
        """
        if not self.model_path:
            return {}
        fn = Path(self.model_path).name
        return self.model_configs.get(fn, {})

    def save_per_model_settings(self, settings: dict):
        """Save per-model settings (called from UI when user configures model).
        
        Supported keys:
        - n_ctx: context window (integer, or -1 for "Max")
        - temperature: sampling temperature (0.0-2.0)
        - n_threads: CPU threads to use
        """
        if not self.model_path:
            return {"error": "No model selected"}
        fn = Path(self.model_path).name
        old = self.model_configs.get(fn, {})
        n_ctx_changed = settings.get("n_ctx") != old.get("n_ctx")
        self.model_configs[fn] = settings
        self.chat_db.set_kv("model_configs", self.model_configs)
        if self.model is not None and "temperature" in settings:
            self.model_config["temperature"] = settings["temperature"]
        return {"ok": True, "n_ctx_changed": n_ctx_changed}

    # ── JSONL Job Queue with Checkpointing ─────────────────────────────────

    def run_jsonl_queue(
        self,
        jsonl: str = "",
        jsonl_path: str = "",
        checkpoint_dir: str = "",
        checkpoint_every: int = 12,
    ) -> dict:
        """Process a JSONL job queue one task at a time, saving checkpoints.

        Each JSONL line must be a JSON object with at minimum:
          - ``job_id``  (str)  — unique identifier; auto-assigned if missing.
          - ``type``    (str)  — ``"chat"`` (default) or ``"process_files"``.

        For ``type = "chat"``:
          - ``text``          (str)
          - ``role`` / ``task`` / ``steps``  (str, optional)
          - ``output_format`` (str, optional)

        For ``type = "process_files"``:
          - ``instructions``  (str)
          - ``files``         (list of file dicts)
          - ``output_format`` (str, optional, default ``"excel"``)

        A checkpoint file (``{checkpoint_dir}/{queue_id}_checkpoint.jsonl``)
        stores one completed-result JSON per line so the queue can resume
        after interruption without re-running finished jobs.
        """
        # ── Resolve JSONL content ─────────────────────────────────────────
        if jsonl_path:
            try:
                with open(jsonl_path, "r", encoding="utf-8") as _f:
                    jsonl = _f.read()
            except Exception as exc:
                return {"error": f"Cannot read JSONL file: {exc}"}

        if not jsonl or not jsonl.strip():
            return {"error": "No JSONL content provided"}

        # ── Parse jobs ────────────────────────────────────────────────────
        jobs = []
        parse_errors = []
        for _i, _raw in enumerate(jsonl.splitlines(), 1):
            _line = _raw.strip()
            if not _line:
                continue
            try:
                _job = json.loads(_line)
                if not isinstance(_job, dict):
                    raise ValueError("Each JSONL line must be a JSON object")
                if not _job.get("job_id"):
                    _job["job_id"] = f"job_{_i:04d}"
                jobs.append(_job)
            except Exception as exc:
                parse_errors.append({"line": _i, "error": str(exc), "raw": _raw[:120]})

        if not jobs:
            return {"error": "No valid jobs found in JSONL", "parse_errors": parse_errors}

        # ── Checkpoint setup ──────────────────────────────────────────────
        if not checkpoint_dir:
            checkpoint_dir = os.path.join(app_data_path(), "jsonl_checkpoints")
        os.makedirs(checkpoint_dir, exist_ok=True)

        # Stable queue ID from a hash of the raw JSONL content
        queue_id = hashlib.sha1(jsonl.encode("utf-8")).hexdigest()[:12]
        checkpoint_file = os.path.join(checkpoint_dir, f"{queue_id}_checkpoint.jsonl")

        # Load already-completed jobs from checkpoint (skip on resume)
        completed: dict = {}
        if os.path.exists(checkpoint_file):
            try:
                with open(checkpoint_file, "r", encoding="utf-8") as _cf:
                    for _cl in _cf:
                        _cl = _cl.strip()
                        if _cl:
                            try:
                                _entry = json.loads(_cl)
                                completed[_entry["job_id"]] = _entry
                            except Exception:
                                pass
                print(f"[JSONL-QUEUE] Resuming — {len(completed)} jobs already done in checkpoint")
            except Exception as exc:
                print(f"[JSONL-QUEUE] Warning: could not read checkpoint: {exc}")

        def _append_checkpoint(entry: dict):
            try:
                with open(checkpoint_file, "a", encoding="utf-8") as _cf:
                    _cf.write(json.dumps(entry, ensure_ascii=False) + "\n")
            except Exception as exc:
                print(f"[JSONL-QUEUE] Warning: checkpoint write failed: {exc}")

        # ── Run pending jobs ──────────────────────────────────────────────
        results = list(completed.values())
        newly_done = 0
        total = len(jobs)

        for idx, job in enumerate(jobs):
            job_id = job["job_id"]

            if job_id in completed:
                print(f"[JSONL-QUEUE] Skipping completed: {job_id}")
                continue

            job_type = str(job.get("type", "chat")).lower()
            print(f"[JSONL-QUEUE] [{idx + 1}/{total}] {job_type}: {job_id}")
            _t0 = time.time()

            try:
                if job_type == "process_files":
                    raw = self.process_files_with_ai(
                        files=job.get("files", []),
                        instructions=str(job.get("instructions", "")),
                        output_format=str(job.get("output_format", "excel")),
                    )
                else:
                    raw = self.agent_chat(
                        text=str(job.get("text", "")),
                        role=str(job.get("role", "")),
                        task=str(job.get("task", "")),
                        steps=str(job.get("steps", "")),
                        output_format=str(job.get("output_format", "none")),
                    )
                entry = {
                    "job_id": job_id,
                    "ok": not bool(raw.get("error")),
                    "elapsed_s": round(time.time() - _t0, 2),
                    "result": raw,
                }
            except Exception as exc:
                entry = {
                    "job_id": job_id,
                    "ok": False,
                    "elapsed_s": round(time.time() - _t0, 2),
                    "error": str(exc),
                }

            results.append(entry)
            newly_done += 1
            _append_checkpoint(entry)

            if checkpoint_every > 0 and newly_done % checkpoint_every == 0:
                print(f"[JSONL-QUEUE] Checkpoint saved at {newly_done} completed jobs ({checkpoint_file})")

            if self.stop_generation_flag:
                print(f"[JSONL-QUEUE] Stopped by user after {newly_done} jobs")
                break

        done_count = sum(1 for r in results if r.get("ok"))
        failed_count = sum(1 for r in results if not r.get("ok"))

        print(f"[JSONL-QUEUE] Done: {done_count} succeeded, {failed_count} failed — {checkpoint_file}")
        return {
            "ok": True,
            "total": total,
            "completed": len(results),
            "succeeded": done_count,
            "failed": failed_count,
            "stopped": bool(self.stop_generation_flag),
            "checkpoint_file": checkpoint_file,
            "parse_errors": parse_errors,
            "results": results,
        }
