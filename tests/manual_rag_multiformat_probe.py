import json
import tempfile
from collections import Counter
from pathlib import Path
import sys

import fitz
import pandas as pd
from PIL import Image, ImageDraw
from docx import Document

WORKSPACE = Path(__file__).resolve().parents[1]
if str(WORKSPACE) not in sys.path:
    sys.path.insert(0, str(WORKSPACE))

from app_core.rag_manager import RAGManager


def make_text_pdf(path: Path):
    doc = fitz.open()
    page = doc.new_page()
    page.insert_text(
        (72, 72),
        "Invoice PDF-101. Vendor Northwind Traders. Total amount 4500 USD. Payment due 2027-06-15.",
    )
    doc.save(str(path))
    doc.close()


def make_scanned_pdf(path: Path, image_path: Path):
    image = Image.new("RGB", (1200, 300), color="white")
    draw = ImageDraw.Draw(image)
    draw.text((40, 120), "SCANNED REF SCN-777 OWNER LINA", fill="black")
    image.save(image_path)

    doc = fitz.open()
    page = doc.new_page(width=1200, height=300)
    page.insert_image(fitz.Rect(0, 0, 1200, 300), filename=str(image_path))
    doc.save(str(path))
    doc.close()


def make_table_pdf(path: Path):
    doc = fitz.open()
    page = doc.new_page()
    page.insert_text((60, 60), "Quarterly Sales Table")
    page.insert_text((60, 95), "Quarter | Region | Revenue")
    page.insert_text((60, 120), "Q1 | East | 120000")
    page.insert_text((60, 145), "Q2 | West | 98000")
    page.insert_text((60, 170), "Q3 | North | 111000")
    page.insert_text((60, 195), "Q4 | South | 127500")
    doc.save(str(path))
    doc.close()


def make_docx(path: Path):
    doc = Document()
    doc.add_paragraph("Project Phoenix owner is Arjun Patel.")
    doc.add_paragraph("Deadline: 2027-11-30.")
    doc.save(str(path))


def make_csv(path: Path):
    df = pd.DataFrame(
        [
            {"product": "Widget", "units": 42, "region": "East"},
            {"product": "Gadget", "units": 17, "region": "West"},
        ]
    )
    df.to_csv(path, index=False)


def make_txt(path: Path):
    path.write_text("Support ticket TCK-909 status Escalated. Priority High.", encoding="utf-8")


def make_jpeg(path: Path):
    image = Image.new("RGB", (600, 200), color="white")
    draw = ImageDraw.Draw(image)
    draw.text((20, 80), "JPG NOTE J-55 COLOR BLUE", fill="black")
    image.save(path, format="JPEG")


def run_probe():
    with tempfile.TemporaryDirectory() as temp_dir:
        base = Path(temp_dir)
        docs = base / "docs"
        docs.mkdir(parents=True, exist_ok=True)

        files = {
            "text_pdf": docs / "invoice_text.pdf",
            "scanned_pdf": docs / "invoice_scanned.pdf",
            "table_pdf": docs / "sales_table.pdf",
            "scanned_img_src": docs / "scan_source.png",
            "docx": docs / "project.docx",
            "csv": docs / "sales.csv",
            "txt": docs / "ticket.txt",
            "jpeg": docs / "photo_note.jpg",
        }

        make_text_pdf(files["text_pdf"])
        make_scanned_pdf(files["scanned_pdf"], files["scanned_img_src"])
        make_table_pdf(files["table_pdf"])
        make_docx(files["docx"])
        make_csv(files["csv"])
        make_txt(files["txt"])
        make_jpeg(files["jpeg"])

        rag_store = base / "rag_store"
        manager = RAGManager(base_directory=str(rag_store))

        extracted = manager._extract_texts_from_folder(str(docs))
        extracted_names = [name for name, _text in extracted]
        extracted_preview = {name: text[:180] for name, text in extracted}

        db = manager.create_from_folder(str(docs), "multiformat_probe", chunk_size=260, chunk_overlap=40)

        db_dir = rag_store / "multiformat_probe"
        storage_files = sorted([p.name for p in db_dir.glob("*")])

        source_counts = Counter([m.get("source", "") for m in db.metadata])

        queries = {
            "text_pdf": {
                "query": "What is the invoice total amount and vendor?",
                "expect": ["4500", "Northwind"],
            },
            "scanned_pdf": {
                "query": "What is the scanned reference and owner?",
                "expect": ["SCN-777", "LINA"],
            },
            "docx": {
                "query": "Who owns Project Phoenix and what is the deadline?",
                "expect": ["Arjun", "2027-11-30"],
            },
            "table_pdf": {
                "query": "In the quarterly sales table, what is Q4 South revenue?",
                "expect": ["Q4", "South", "127500"],
            },
            "csv": {
                "query": "How many units of Widget were sold and in which region?",
                "expect": ["Widget", "42", "East"],
            },
            "txt": {
                "query": "What is the support ticket id and status?",
                "expect": ["TCK-909", "Escalated"],
            },
            "jpeg": {
                "query": "What jpg note id is present?",
                "expect": ["J-55"],
            },
        }

        retrieval = {}
        for key, cfg in queries.items():
            rows = manager.retrieve("multiformat_probe", cfg["query"], k=3)
            top_text = "\n".join([chunk for chunk, _score in rows])
            hit_all = all(token.lower() in top_text.lower() for token in cfg["expect"])
            retrieval[key] = {
                "query": cfg["query"],
                "expected_tokens": cfg["expect"],
                "top_hits": [
                    {
                        "score": round(float(score), 4),
                        "preview": chunk[:180],
                    }
                    for chunk, score in rows
                ],
                "expected_tokens_found": hit_all,
            }

        result = {
            "input_files": [p.name for p in docs.iterdir() if p.is_file()],
            "supported_extensions_in_rag_manager": [
                ".txt", ".pdf", ".csv", ".xlsx", ".xls", ".docx", ".pptx", ".ppt",
                ".jpg", ".jpeg", ".png", ".bmp", ".webp", ".tif", ".tiff"
            ],
            "extracted_file_names": extracted_names,
            "extracted_preview": extracted_preview,
            "jpeg_indexed": "photo_note.jpg" in extracted_names,
            "scanned_pdf_has_extracted_text": "invoice_scanned.pdf" in extracted_names,
            "db_name": db.name,
            "total_chunks": len(db.chunks),
            "source_chunk_counts": dict(source_counts),
            "stored_files": storage_files,
            "retrieval_quality": retrieval,
        }

        print(json.dumps(result, indent=2, ensure_ascii=False))


if __name__ == "__main__":
    run_probe()
